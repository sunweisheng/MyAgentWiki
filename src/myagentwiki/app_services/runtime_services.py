from __future__ import annotations

import csv
import hashlib
import json
import mimetypes
import re
import shlex
import ssl
import subprocess
import struct
import tempfile
import urllib.error
import urllib.request
import zipfile
from pathlib import Path
from datetime import datetime, timezone
from urllib.parse import unquote, urlparse
from xml.etree import ElementTree as ET

try:
    import docx
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    docx = None

try:
    import openpyxl
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    openpyxl = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    PdfReader = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    Image = None

from ..runtime_env import build_doctor_payload, command_exists, load_simple_yaml

AUTOMATION_STRATEGIES = {"safe_auto", "agent_assisted"}
WORKSPACE_SCHEMA_VERSION = "v1"
MARKDOWN_IMAGE_PATTERN = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<target>[^)\s]+)(?:\s+\"[^\"]*\")?\)")
DOCX_NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}
XLSX_NAMESPACES = {
    "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}
OLE_HEADER_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def run_doctor_service(root: Path) -> dict:
    return build_doctor_payload(root)


def run_bootstrap_service(
    *,
    root: Path,
    python_executable: str,
    extras: list[str],
    dry_run: bool,
) -> tuple[int, dict]:
    install_command = [python_executable, "-m", "pip", "install", "-e"]
    if extras:
        install_command.append(f"{str(root)}[{','.join(extras)}]")
    else:
        install_command.append(str(root))

    doctor_payload = build_doctor_payload(root)

    if dry_run:
        payload = {
            "action": "dry_run",
            "install_command": install_command,
            "project_root": str(root),
            "requested_extras": extras,
            "doctor_summary": doctor_payload["summary"],
        }
        return 0, payload

    completed = subprocess.run(
        install_command,
        check=False,
        capture_output=True,
        text=True,
    )
    payload = {
        "action": "install",
        "install_command": install_command,
        "requested_extras": extras,
        "returncode": completed.returncode,
        "stdout": completed.stdout.strip(),
        "stderr": completed.stderr.strip(),
        "doctor_summary": doctor_payload["summary"],
    }
    return completed.returncode, payload


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def ensure_directory(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def path_is_within_root(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def resolve_workspace_path(target: Path, configured_path: str) -> Path:
    path = Path(configured_path).expanduser()
    if path.is_absolute():
        return path
    return (target / path).resolve()


def load_workspace_config(target: Path) -> dict:
    return load_simple_yaml(target / "config" / "project.yml")


def raw_assets_dir_for_workspace(target: Path, raw_dir: Path | None = None) -> Path:
    resolved_raw_dir = raw_dir or resolve_workspace_raw_dir(target)
    return (resolved_raw_dir.parent / "assets").resolve()


def resolve_workspace_raw_dir(target: Path) -> Path:
    config = load_workspace_config(target)
    raw_dir = resolve_workspace_path(target, config["paths"]["raw"])
    if raw_dir.name != "raw":
        raise ValueError(f"Workspace raw directory must be named 'raw': {raw_dir}")
    if raw_dir.parent != target.parent:
        raise ValueError(
            f"Workspace raw directory must be a sibling of the workspace: raw={raw_dir} target={target}"
        )
    return raw_dir


def resolve_source_record_path(target: Path, source_path: str) -> Path:
    path = Path(source_path).expanduser()
    if path.is_absolute():
        return path
    return (target / path).resolve()


def ensure_path_within_raw_root(path: Path, raw_root: Path, *, purpose: str) -> Path:
    resolved_path = path.resolve()
    resolved_root = raw_root.resolve()
    if not path_is_within_root(resolved_path, resolved_root):
        raise ValueError(
            f"{purpose} must stay within raw directory: path={resolved_path} raw={resolved_root}"
        )
    return resolved_path


def coerce_int(value, default: int) -> int:
    if isinstance(value, int):
        return value
    if isinstance(value, str):
        digits = value.strip()
        if digits.isdigit():
            return int(digits)
    return default


def coerce_float(value, default: float) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        stripped = value.strip()
        try:
            return float(stripped)
        except ValueError:
            return default
    return default


def normalize_command_config(value) -> list[str]:
    if isinstance(value, str):
        stripped = value.strip()
        return shlex.split(stripped) if stripped else []
    if isinstance(value, list):
        normalized = []
        for item in value:
            text = str(item).strip()
            if text:
                normalized.append(text)
        return normalized
    return []


def load_automation_target_config(config: dict, target_name: str) -> dict:
    automation_config = config.get("automation", {})
    if not isinstance(automation_config, dict):
        automation_config = {}

    target_config = automation_config.get(target_name, {})
    if not isinstance(target_config, dict):
        target_config = {}

    inherited_strategy = str(automation_config.get("mode", "safe_auto")).strip() or "safe_auto"
    strategy = str(target_config.get("strategy", inherited_strategy)).strip() or inherited_strategy
    if strategy not in AUTOMATION_STRATEGIES:
        strategy = "safe_auto"

    command = normalize_command_config(target_config.get("command", []))
    timeout_seconds = max(coerce_int(target_config.get("timeout_seconds", 45), 45), 5)
    min_confidence = min(max(coerce_float(target_config.get("min_confidence", 0.8), 0.8), 0.0), 1.0)
    return {
        "strategy": strategy,
        "command": command,
        "timeout_seconds": timeout_seconds,
        "min_confidence": min_confidence,
        "enabled": strategy == "agent_assisted" and bool(command),
    }


def run_json_automation_command(
    target: Path,
    command: list[str],
    payload: dict,
    timeout_seconds: int,
) -> dict | None:
    if not command:
        return None
    try:
        completed = subprocess.run(
            command,
            cwd=target,
            input=json.dumps(payload, ensure_ascii=False),
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None

    if completed.returncode != 0:
        return None

    stdout = (completed.stdout or "").strip()
    if not stdout:
        return None

    try:
        parsed = json.loads(stdout)
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def normalize_text_content(source_type: str, raw_text: str) -> str:
    normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalized.split("\n")]
    normalized = "\n".join(lines).strip()
    if source_type == "plain_text":
        return normalized + "\n"
    return normalized + "\n"


def normalize_markdown_or_text_record(
    target: Path,
    source_record: dict,
    *,
    allow_insecure_downloads: bool = True,
) -> dict:
    source_type = source_record["source_type"]
    raw_path = resolve_source_record_path(target, source_record["source_path"])
    ensure_path_within_raw_root(raw_path, resolve_workspace_raw_dir(target), purpose="Source record")
    raw_text = raw_path.read_text(encoding="utf-8")
    if source_type == "markdown":
        normalized_text, metadata = enrich_markdown_with_embedded_images(
            target=target,
            source_record=source_record,
            raw_path=raw_path,
            raw_text=raw_text,
            allow_insecure_downloads=allow_insecure_downloads,
        )
    else:
        normalized_text = normalize_text_content(source_type, raw_text)
        metadata = {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good",
            "warnings": [],
            "location_map": {
                "type": "line_map",
                "source_path": source_record["source_path"],
            },
        }

    normalized_rel_path = Path("normalized") / f"{source_record['source_id']}.md"
    normalized_abs_path = target / normalized_rel_path
    normalized_abs_path.write_text(normalized_text, encoding="utf-8")

    normalized_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    line_count = len(normalized_text.splitlines()) or 1
    title = raw_path.stem
    location_map = dict(metadata.get("location_map", {}))
    if location_map.get("type") == "line_map" and "normalized_line_range" not in location_map:
        location_map["normalized_line_range"] = f"1-{line_count}"

    return {
        "source_id": source_record["source_id"],
        "source_type": source_type,
        "source_path": source_record["source_path"],
        "normalized_path": str(normalized_rel_path),
        "title": title,
        "language": "unknown",
        "content_format": "markdown",
        "raw_hash": source_record["source_hash"],
        "normalized_hash": normalized_hash,
        "normalizer_version": "normalize_v1",
        "document_kind": "note",
        "structure_quality": "unknown",
        "chunk_strategy_hint": "heading_first",
        "extraction_method": metadata.get("extraction_method", "python_only"),
        "extraction_quality": metadata.get("extraction_quality", "good"),
        "warnings": metadata.get("warnings", []),
        "location_map": location_map,
        "updated_at": utc_now_iso(),
    }


def convert_pdf_to_markdown(raw_path: Path) -> tuple[str, dict]:
    if PdfReader is None:
        return convert_pdf_to_markdown_fallback(raw_path)

    try:
        reader = PdfReader(str(raw_path))
        parts: list[str] = [f"# {raw_path.stem}"]
        page_map: list[dict] = []
        warnings: list[str] = []

        for page_index, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            cleaned = extracted.strip()
            parts.append(f"\n## 第 {page_index} 页\n")
            if cleaned:
                parts.append(cleaned)
            else:
                parts.append("_本页未提取到文本_")
                warnings.append(f"page_{page_index}_empty_text")
            page_map.append({
                "page": page_index,
                "char_count": len(cleaned),
            })

        markdown = "\n\n".join(parts).strip() + "\n"
        metadata = {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "partial" if warnings else "good",
            "warnings": warnings,
            "location_map": {
                "type": "pdf_page_map",
                "pages": page_map,
                "source_path": str(raw_path),
            },
        }
        return markdown, metadata
    except Exception:
        return convert_pdf_to_markdown_fallback(raw_path)


def pdf_count_pages_from_bytes(pdf_bytes: bytes) -> int:
    matches = re.findall(rb"/Type\s*/Page\b", pdf_bytes)
    return max(len(matches), 0)


def pdf_try_inflate_stream(stream_bytes: bytes) -> bytes:
    try:
        import zlib
        return zlib.decompress(stream_bytes)
    except Exception:
        return stream_bytes


def decode_pdf_literal_string(raw: bytes) -> str:
    text = raw.replace(b"\\(", b"(").replace(b"\\)", b")").replace(b"\\\\", b"\\")
    try:
        return text.decode("utf-8")
    except UnicodeDecodeError:
        return text.decode("latin-1", errors="ignore")


def pdf_extract_text_snippets_from_bytes(pdf_bytes: bytes) -> list[str]:
    snippets: list[str] = []

    object_matches = re.finditer(rb"(\d+\s+\d+\s+obj.*?endobj)", pdf_bytes, flags=re.DOTALL)
    for match in object_matches:
        object_bytes = match.group(1)
        stream_match = re.search(rb"stream\r?\n(.*?)\r?\nendstream", object_bytes, flags=re.DOTALL)
        if not stream_match:
            continue
        stream_bytes = stream_match.group(1)
        if b"/FlateDecode" in object_bytes:
            stream_bytes = pdf_try_inflate_stream(stream_bytes)

        for block in re.findall(rb"BT(.*?)ET", stream_bytes, flags=re.DOTALL):
            pieces = re.findall(rb"\(([^()]*)\)", block)
            if not pieces:
                continue
            combined = "".join(decode_pdf_literal_string(piece) for piece in pieces).strip()
            combined = re.sub(r"\s+", " ", combined)
            if len(combined) >= 4 and combined not in snippets:
                snippets.append(combined)

    return snippets[:20]


def convert_pdf_to_markdown_fallback(raw_path: Path) -> tuple[str, dict]:
    pdf_bytes = raw_path.read_bytes()
    page_count = pdf_count_pages_from_bytes(pdf_bytes)
    snippets = pdf_extract_text_snippets_from_bytes(pdf_bytes)
    warnings: list[str] = []

    parts = [f"# {raw_path.stem}", "", f"- 估计页数: {page_count or 'unknown'}"]
    if snippets:
        parts.extend(["", "## 提取文本片段"])
        for index, snippet in enumerate(snippets, start=1):
            parts.append(f"{index}. {snippet}")
    else:
        warnings.append("pdf_fallback_no_text")
        parts.extend([
            "",
            "> 当前环境未启用 `pypdf`，且标准库 fallback 未提取到正文文本。",
        ])

    markdown = "\n".join(parts).strip() + "\n"
    pages = [{"page": index + 1, "char_count": None} for index in range(page_count)] if page_count else []
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "partial",
        "warnings": warnings if warnings else ["pdf_fallback_used"],
        "location_map": {
            "type": "pdf_page_map",
            "pages": pages,
            "source_path": str(raw_path),
        },
    }


def iter_docx_blocks(document) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name:
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(digits) if digits else 1
            level = min(max(level, 1), 6)
            blocks.append(("heading", "#" * level + " " + text))
        else:
            blocks.append(("paragraph", text))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") or " " for cell in row.cells]
            rows.append(values)
        if not rows:
            continue
        blocks.append(("table_title", f"## 表格 {table_index}"))
        header = rows[0]
        divider = ["---"] * len(header)
        table_lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(divider) + " |",
        ]
        for row in rows[1:]:
            padded = row + [" "] * (len(header) - len(row))
            table_lines.append("| " + " | ".join(padded[: len(header)]) + " |")
        blocks.append(("table", "\n".join(table_lines)))
    return blocks


def docx_style_map_from_archive(archive: zipfile.ZipFile) -> dict[str, str]:
    try:
        styles_xml = archive.read("word/styles.xml")
    except KeyError:
        return {}
    root = ET.fromstring(styles_xml)
    style_map: dict[str, str] = {}
    for style in root.findall("w:style", DOCX_NAMESPACES):
        style_id = style.get(f"{{{DOCX_NAMESPACES['w']}}}styleId")
        name_node = style.find("w:name", DOCX_NAMESPACES)
        if not style_id:
            continue
        style_map[style_id] = name_node.get(f"{{{DOCX_NAMESPACES['w']}}}val", style_id) if name_node is not None else style_id
    return style_map


def docx_paragraph_text(node: ET.Element) -> str:
    texts = [item.text or "" for item in node.findall(".//w:t", DOCX_NAMESPACES)]
    return "".join(texts).strip()


def docx_heading_level_from_style(style_name: str, style_id: str) -> int | None:
    candidates = f"{style_name} {style_id}".lower()
    if "heading" not in candidates:
        return None
    digits = "".join(ch for ch in candidates if ch.isdigit())
    if not digits:
        return 1
    return min(max(int(digits[0]), 1), 6)


def docx_table_to_markdown(node: ET.Element) -> str:
    rows: list[list[str]] = []
    for row in node.findall("w:tr", DOCX_NAMESPACES):
        values = []
        for cell in row.findall("w:tc", DOCX_NAMESPACES):
            cell_texts = []
            for paragraph in cell.findall(".//w:p", DOCX_NAMESPACES):
                text = docx_paragraph_text(paragraph)
                if text:
                    cell_texts.append(text)
            values.append(" ".join(cell_texts).strip() or " ")
        if rows or any(value.strip() for value in values):
            rows.append(values)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized_rows = [row + [" "] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    divider = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_docx_to_markdown_fallback(raw_path: Path) -> tuple[str, dict]:
    parts = [f"# {raw_path.stem}"]
    warnings: list[str] = []

    with zipfile.ZipFile(raw_path) as archive:
        style_map = docx_style_map_from_archive(archive)
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    body = root.find("w:body", DOCX_NAMESPACES)
    if body is None:
        warnings.append("docx_missing_body")
        markdown = "\n".join(parts).strip() + "\n"
        return markdown, {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "partial",
            "warnings": warnings,
            "location_map": {
                "type": "line_map",
                "normalized_line_range": "1-1",
                "source_path": str(raw_path),
            },
        }

    table_index = 0
    for child in list(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = docx_paragraph_text(child)
            if not text:
                continue
            style_node = child.find("w:pPr/w:pStyle", DOCX_NAMESPACES)
            style_id = style_node.get(f"{{{DOCX_NAMESPACES['w']}}}val", "") if style_node is not None else ""
            style_name = style_map.get(style_id, style_id)
            heading_level = docx_heading_level_from_style(style_name, style_id)
            if heading_level is not None:
                parts.append("#" * heading_level + " " + text)
            else:
                parts.append(text)
        elif tag == "tbl":
            table_index += 1
            table_markdown = docx_table_to_markdown(child)
            if table_markdown:
                parts.append(f"## 表格 {table_index}")
                parts.append(table_markdown)

    markdown = "\n\n".join(part for part in parts if part.strip()).strip() + "\n"
    line_count = len(markdown.splitlines()) or 1
    metadata = {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if len(parts) > 1 else "partial",
        "warnings": warnings if warnings else [],
        "location_map": {
            "type": "line_map",
            "normalized_line_range": f"1-{line_count}",
            "source_path": str(raw_path),
        },
    }
    return markdown, metadata


def convert_docx_to_markdown(raw_path: Path) -> tuple[str, dict]:
    if docx is not None:
        document = docx.Document(str(raw_path))
        blocks = iter_docx_blocks(document)
        parts = [f"# {raw_path.stem}"]
        parts.extend(block for _, block in blocks)
        markdown = "\n\n".join(part for part in parts if part.strip()).strip() + "\n"
        line_count = len(markdown.splitlines()) or 1
        return markdown, {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good" if blocks else "partial",
            "warnings": [] if blocks else ["docx_no_visible_blocks"],
            "location_map": {
                "type": "line_map",
                "normalized_line_range": f"1-{line_count}",
                "source_path": str(raw_path),
            },
        }
    return convert_docx_to_markdown_fallback(raw_path)


def is_probably_ole_document(raw_path: Path) -> bool:
    with raw_path.open("rb") as fh:
        return fh.read(len(OLE_HEADER_MAGIC)) == OLE_HEADER_MAGIC


def normalize_binary_snippet_text(text: str) -> str:
    cleaned = []
    for char in text:
        if char in {"\n", "\r", "\t"}:
            cleaned.append(" ")
            continue
        category = ord(char)
        if char.isprintable() or 0x4E00 <= category <= 0x9FFF:
            cleaned.append(char)
    normalized = "".join(cleaned)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def binary_text_candidate_is_meaningful(text: str, min_length: int = 8) -> bool:
    if len(text) < min_length:
        return False
    interesting_chars = [char for char in text if char.isalnum() or "\u4e00" <= char <= "\u9fff"]
    if len(interesting_chars) < max(4, min_length // 2):
        return False
    return True


def extract_printable_ascii_snippets(binary_bytes: bytes, min_length: int = 8) -> list[str]:
    snippets: list[str] = []
    for match in re.finditer(rb"[\x20-\x7e]{8,}", binary_bytes):
        text = normalize_binary_snippet_text(match.group(0).decode("utf-8", errors="ignore"))
        if binary_text_candidate_is_meaningful(text, min_length=min_length):
            snippets.append(text)
    return snippets


def extract_utf16_text_snippets(binary_bytes: bytes, min_length: int = 8) -> list[str]:
    snippets: list[str] = []
    pattern = re.compile(r"[0-9A-Za-z\u4e00-\u9fff][0-9A-Za-z\u4e00-\u9fff\s，。、“”‘’；：？！,.!?:;()（）\-_/]{7,}")
    for offset in (0, 1):
        if len(binary_bytes) <= offset + 2:
            continue
        decoded = binary_bytes[offset:].decode("utf-16le", errors="ignore")
        for match in pattern.finditer(decoded):
            text = normalize_binary_snippet_text(match.group(0))
            if binary_text_candidate_is_meaningful(text, min_length=min_length):
                snippets.append(text)
    return snippets


def extract_text_snippets_from_binary_document(
    binary_bytes: bytes,
    limit: int = 20,
    min_length: int = 8,
) -> list[str]:
    ordered_candidates = [
        *extract_printable_ascii_snippets(binary_bytes, min_length=min_length),
        *extract_utf16_text_snippets(binary_bytes, min_length=min_length),
    ]
    snippets: list[str] = []
    seen: set[str] = set()
    for candidate in ordered_candidates:
        dedupe_key = candidate.lower()
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        snippets.append(candidate)
        if len(snippets) >= limit:
            break
    return snippets


def convert_legacy_doc_to_markdown(raw_path: Path) -> tuple[str, dict]:
    binary_bytes = raw_path.read_bytes()
    snippets = extract_text_snippets_from_binary_document(binary_bytes)
    is_ole = is_probably_ole_document(raw_path)
    parts = [
        f"# {raw_path.stem}",
        "",
        "## 文档信息 / Document Info",
        "",
        "- 原始格式: `.doc`",
        f"- 文件大小: {len(binary_bytes)} bytes",
        f"- OLE 容器: `{is_ole}`",
    ]
    if snippets:
        parts.extend(["", "## 提取文本片段 / Extracted Snippets", ""])
        for index, snippet in enumerate(snippets, start=1):
            parts.append(f"{index}. {snippet}")
    else:
        parts.extend([
            "",
            "## 提取文本片段 / Extracted Snippets",
            "",
            "> 当前纯 Python fallback 未提取到可读正文片段。",
        ])
    markdown = "\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "partial" if snippets else "poor",
        "warnings": ["legacy_doc_binary_fallback"] if snippets else ["legacy_doc_no_text_snippets"],
        "location_map": {
            "type": "binary_snippet_map",
            "source_path": str(raw_path),
            "snippet_count": len(snippets),
            "is_ole_container": is_ole,
        },
    }


def image_size_from_binary(raw_path: Path) -> tuple[int | None, int | None, str | None]:
    with raw_path.open("rb") as fh:
        header = fh.read(64)
    if header.startswith(b"\x89PNG\r\n\x1a\n") and len(header) >= 24:
        width, height = struct.unpack(">II", header[16:24])
        return width, height, "PNG"
    if header.startswith(b"\xff\xd8"):
        with raw_path.open("rb") as fh:
            fh.read(2)
            while True:
                marker_start = fh.read(1)
                if not marker_start:
                    break
                if marker_start != b"\xff":
                    continue
                marker = fh.read(1)
                while marker == b"\xff":
                    marker = fh.read(1)
                if marker in {b"\xc0", b"\xc1", b"\xc2", b"\xc3", b"\xc5", b"\xc6", b"\xc7", b"\xc9", b"\xca", b"\xcb", b"\xcd", b"\xce", b"\xcf"}:
                    _segment_length = struct.unpack(">H", fh.read(2))[0]
                    _precision = fh.read(1)
                    height, width = struct.unpack(">HH", fh.read(4))
                    return width, height, "JPEG"
                if marker in {b"\xd8", b"\xd9"}:
                    continue
                segment_length_data = fh.read(2)
                if len(segment_length_data) < 2:
                    break
                segment_length = struct.unpack(">H", segment_length_data)[0]
                fh.seek(segment_length - 2, 1)
    if header[:4] == b"RIFF" and header[8:12] == b"WEBP":
        if header[12:16] == b"VP8X" and len(header) >= 30:
            width = 1 + int.from_bytes(header[24:27], "little")
            height = 1 + int.from_bytes(header[27:30], "little")
            return width, height, "WEBP"
    return None, None, None


def worksheet_to_markdown(sheet) -> str:
    rows = []
    for row in sheet.iter_rows(values_only=True):
        values = ["" if value is None else str(value).replace("\n", " ") for value in row]
        if any(value != "" for value in values):
            rows.append(values)
    if not rows:
        return "_此工作表无可见数据_"
    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    divider = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def xlsx_shared_strings_from_archive(archive: zipfile.ZipFile) -> list[str]:
    try:
        xml_data = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    root = ET.fromstring(xml_data)
    values: list[str] = []
    for item in root.findall("main:si", XLSX_NAMESPACES):
        texts = [node.text or "" for node in item.findall(".//main:t", XLSX_NAMESPACES)]
        values.append("".join(texts))
    return values


def xlsx_sheet_names_from_workbook(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
    rel_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    rel_map = {}
    for rel in rel_root.findall("rel:Relationship", XLSX_NAMESPACES):
        rel_id = rel.get("Id")
        target = rel.get("Target")
        if rel_id and target:
            rel_map[rel_id] = target
    sheets: list[tuple[str, str]] = []
    for sheet in workbook_root.findall("main:sheets/main:sheet", XLSX_NAMESPACES):
        name = sheet.get("name", "Sheet")
        rel_id = sheet.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        target = rel_map.get(rel_id, "")
        if target:
            normalized_target = target.lstrip("/")
            sheet_path = normalized_target if normalized_target.startswith("xl/") else f"xl/{normalized_target}"
            sheets.append((name, sheet_path))
    return sheets


def column_letters_to_index(value: str) -> int:
    index = 0
    for char in value:
        if not char.isalpha():
            break
        index = index * 26 + (ord(char.upper()) - ord("A") + 1)
    return max(index - 1, 0)


def xlsx_cell_value(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.get("t")
    value_node = cell.find("main:v", XLSX_NAMESPACES)
    inline_node = cell.find("main:is", XLSX_NAMESPACES)
    if inline_node is not None:
        texts = [node.text or "" for node in inline_node.findall(".//main:t", XLSX_NAMESPACES)]
        return "".join(texts).strip()
    if value_node is None or value_node.text is None:
        return ""
    raw_value = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value)]
        except (ValueError, IndexError):
            return raw_value
    return raw_value


def xlsx_rows_from_sheet_xml(archive: zipfile.ZipFile, sheet_path: str, shared_strings: list[str]) -> list[list[str]]:
    root = ET.fromstring(archive.read(sheet_path))
    rows: list[list[str]] = []
    for row in root.findall("main:sheetData/main:row", XLSX_NAMESPACES):
        row_values: list[str] = []
        for cell in row.findall("main:c", XLSX_NAMESPACES):
            ref = cell.get("r", "")
            col_index = column_letters_to_index(ref)
            while len(row_values) < col_index:
                row_values.append("")
            value = xlsx_cell_value(cell, shared_strings).replace("\n", " ").strip()
            row_values.append(value)
        if row_values and any(value != "" for value in row_values):
            rows.append(row_values)
    return rows


def rows_to_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return "_此工作表无可见数据_"
    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    divider = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_xlsx_to_markdown_fallback(raw_path: Path) -> tuple[str, dict]:
    with zipfile.ZipFile(raw_path) as archive:
        shared_strings = xlsx_shared_strings_from_archive(archive)
        sheets = xlsx_sheet_names_from_workbook(archive)
        parts = [f"# {raw_path.stem}"]
        sheet_map: list[dict] = []
        for sheet_name, sheet_path in sheets:
            rows = xlsx_rows_from_sheet_xml(archive, sheet_path, shared_strings)
            parts.append(f"## 工作表: {sheet_name}")
            parts.append(rows_to_markdown_table(rows))
            sheet_map.append({
                "sheet_name": sheet_name,
                "row_count": len(rows),
                "sheet_path": sheet_path,
            })
    markdown = "\n\n".join(part for part in parts if part.strip()).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if sheet_map else "partial",
        "warnings": [] if sheet_map else ["xlsx_no_visible_sheets"],
        "location_map": {
            "type": "sheet_map",
            "sheets": sheet_map,
            "source_path": str(raw_path),
        },
    }


def convert_legacy_xls_to_markdown(raw_path: Path) -> tuple[str, dict]:
    binary_bytes = raw_path.read_bytes()
    snippets = extract_text_snippets_from_binary_document(binary_bytes)
    is_ole = is_probably_ole_document(raw_path)
    parts = [
        f"# {raw_path.stem}",
        "",
        "## 工作簿信息 / Workbook Info",
        "",
        "- 原始格式: `.xls`",
        f"- 文件大小: {len(binary_bytes)} bytes",
        f"- OLE 容器: `{is_ole}`",
    ]
    if snippets:
        parts.extend(["", "## 可见文本片段 / Visible Text Snippets", ""])
        for index, snippet in enumerate(snippets, start=1):
            parts.append(f"{index}. {snippet}")
    else:
        parts.extend([
            "",
            "## 可见文本片段 / Visible Text Snippets",
            "",
            "> 当前纯 Python fallback 未提取到可读工作簿文本。",
        ])
    markdown = "\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "partial" if snippets else "poor",
        "warnings": ["legacy_xls_binary_fallback"] if snippets else ["legacy_xls_no_text_snippets"],
        "location_map": {
            "type": "binary_snippet_map",
            "source_path": str(raw_path),
            "snippet_count": len(snippets),
            "is_ole_container": is_ole,
        },
    }


def convert_spreadsheet_to_markdown(raw_path: Path) -> tuple[str, dict]:
    if raw_path.suffix.lower() == ".csv":
        return convert_csv_to_markdown(raw_path)
    if raw_path.suffix.lower() == ".xls":
        return convert_legacy_xls_to_markdown(raw_path)
    if openpyxl is None:
        return convert_xlsx_to_markdown_fallback(raw_path)
    workbook = openpyxl.load_workbook(str(raw_path), data_only=False)
    parts = [f"# {raw_path.stem}"]
    sheet_map: list[dict] = []
    for sheet in workbook.worksheets:
        parts.append(f"\n## 工作表: {sheet.title}\n")
        parts.append(worksheet_to_markdown(sheet))
        sheet_map.append({
            "sheet_name": sheet.title,
            "max_row": sheet.max_row,
            "max_column": sheet.max_column,
        })
    markdown = "\n\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if workbook.worksheets else "partial",
        "warnings": [] if workbook.worksheets else ["spreadsheet_no_worksheets"],
        "location_map": {
            "type": "sheet_map",
            "sheets": sheet_map,
            "source_path": str(raw_path),
        },
    }


def convert_csv_to_markdown(raw_path: Path) -> tuple[str, dict]:
    rows: list[list[str]] = []
    with raw_path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.reader(fh)
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append([cell.strip() for cell in row])
    parts = [f"# {raw_path.stem}", "", "## 工作表: CSV", ""]
    if rows:
        width = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (width - len(row)) for row in rows]
        header = normalized_rows[0]
        divider = ["---"] * width
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join(divider) + " |")
        for row in normalized_rows[1:]:
            parts.append("| " + " | ".join(row) + " |")
    else:
        parts.append("_CSV 文件没有可见数据_")
    markdown = "\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if rows else "partial",
        "warnings": [] if rows else ["csv_no_visible_rows"],
        "location_map": {
            "type": "sheet_map",
            "sheets": [{"sheet_name": "CSV", "row_count": len(rows)}],
            "source_path": str(raw_path),
        },
    }


def ocr_text_is_meaningful(text: str, min_length: int = 12) -> bool:
    normalized = normalize_binary_snippet_text(text)
    if len(normalized) < min_length:
        return False
    interesting_chars = [char for char in normalized if char.isalnum() or "\u4e00" <= char <= "\u9fff"]
    return len(interesting_chars) >= max(6, min_length // 2)


def image_understanding_text_is_meaningful(text: str, min_length: int = 16) -> bool:
    normalized = normalize_binary_snippet_text(text)
    if len(normalized) < min_length:
        return False
    signal_chars = [char for char in normalized if char.isalnum() or "\u4e00" <= char <= "\u9fff"]
    return len(signal_chars) >= max(8, min_length // 2)


def sanitize_asset_filename(value: str, default_stem: str = "asset") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = cleaned.strip("._")
    return cleaned or default_stem


def markdown_image_target_candidates(target_value: str) -> list[str]:
    normalized = target_value.strip()
    if not normalized:
        return []
    if normalized.startswith("<") and normalized.endswith(">"):
        normalized = normalized[1:-1].strip()
    normalized = normalized.replace("\\)", ")").replace("\\(", "(")
    return [normalized]


def markdown_image_asset_extension(*, parsed_url, source_name: str, content_type: str | None = None) -> str:
    candidate = Path(unquote(parsed_url.path or source_name)).suffix.lower()
    if candidate:
        return candidate
    if content_type:
        guessed = mimetypes.guess_extension(content_type.split(";", 1)[0].strip())
        if guessed:
            return guessed.lower()
    fallback = Path(source_name).suffix.lower()
    return fallback or ".bin"


def build_markdown_asset_path(assets_dir: Path, source_record: dict, image_index: int, source_name: str) -> Path:
    stem = sanitize_asset_filename(Path(source_name).stem or f"image_{image_index}", default_stem=f"image_{image_index}")
    suffix = Path(source_name).suffix.lower() or ".bin"
    relative_dir = Path(source_record["source_id"])
    return assets_dir / relative_dir / f"{image_index:03d}_{stem}{suffix}"


def is_certificate_verification_error(exc: Exception) -> bool:
    if isinstance(exc, ssl.SSLCertVerificationError):
        return True
    if isinstance(exc, urllib.error.URLError):
        reason = exc.reason
        if isinstance(reason, ssl.SSLCertVerificationError):
            return True
        if isinstance(reason, ssl.SSLError):
            return "CERTIFICATE_VERIFY_FAILED" in str(reason)
    return "CERTIFICATE_VERIFY_FAILED" in str(exc)


def download_url_with_optional_insecure_retry(
    target_value: str,
    *,
    allow_insecure_downloads: bool,
) -> tuple[bytes, str | None, str, list[str]]:
    try:
        with urllib.request.urlopen(target_value, timeout=20) as response:
            return response.read(), response.headers.get("Content-Type"), "verified", []
    except Exception as exc:
        if not allow_insecure_downloads or not is_certificate_verification_error(exc):
            raise
    insecure_context = ssl._create_unverified_context()
    with urllib.request.urlopen(target_value, timeout=20, context=insecure_context) as response:
        return (
            response.read(),
            response.headers.get("Content-Type"),
            "insecure_retry",
            ["markdown_remote_image_download_used_insecure_retry"],
        )


def download_markdown_image_to_assets(
    *,
    target: Path,
    source_record: dict,
    raw_dir: Path,
    image_index: int,
    target_value: str,
    allow_insecure_downloads: bool = True,
) -> dict:
    assets_dir = raw_assets_dir_for_workspace(target, raw_dir)
    parsed = urlparse(target_value)
    source_name = Path(unquote(parsed.path)).name or f"image_{image_index}"
    content, content_type, download_mode, warnings = download_url_with_optional_insecure_retry(
        target_value,
        allow_insecure_downloads=allow_insecure_downloads,
    )
    suffix = markdown_image_asset_extension(parsed_url=parsed, source_name=source_name, content_type=content_type)
    asset_path = build_markdown_asset_path(
        assets_dir,
        source_record,
        image_index,
        f"{Path(source_name).stem}{suffix}",
    )
    ensure_directory(asset_path.parent)
    asset_path.write_bytes(content)
    return {
        "storage_kind": "downloaded",
        "asset_path": asset_path,
        "asset_hash": hashlib.sha256(content).hexdigest(),
        "content_type": content_type,
        "download_mode": download_mode,
        "warnings": warnings,
    }


def resolve_markdown_local_image_path(raw_path: Path, target_value: str, raw_dir: Path) -> Path:
    parsed = urlparse(target_value)
    if parsed.scheme and parsed.scheme != "file":
        raise ValueError(f"Unsupported local image reference scheme: {target_value}")
    if parsed.scheme == "file":
        candidate = Path(unquote(parsed.path)).expanduser()
    else:
        candidate = (raw_path.parent / unquote(target_value)).expanduser()
    return ensure_path_within_raw_root(candidate, raw_dir, purpose="Markdown image reference")


def convert_markdown_embedded_image_to_section(
    *,
    target: Path,
    asset_path: Path,
    asset_label: str,
    alt_text: str,
    image_context: dict | None = None,
) -> tuple[list[str], dict]:
    image_markdown, image_metadata = convert_image_to_markdown(
        asset_path,
        target=target,
        image_context=image_context,
    )
    body_lines = [line.rstrip() for line in image_markdown.splitlines()]
    if body_lines and body_lines[0].startswith("# "):
        body_lines = body_lines[1:]
        if body_lines and not body_lines[0]:
            body_lines = body_lines[1:]
    lines = [
        f"## 内嵌图片 {asset_label}",
        "",
        f"- alt: {alt_text or '(empty)'}",
        f"- asset_path: {asset_path}",
        f"- extraction_quality: {image_metadata.get('extraction_quality', 'partial')}",
    ]
    lines.extend(body_lines if body_lines else ["> 图片存在，但当前未生成附加文本内容。"])
    return lines, image_metadata


def enrich_markdown_with_embedded_images(
    *,
    target: Path,
    source_record: dict,
    raw_path: Path,
    raw_text: str,
    allow_insecure_downloads: bool = True,
) -> tuple[str, dict]:
    raw_dir = resolve_workspace_raw_dir(target)
    assets_dir = raw_assets_dir_for_workspace(target, raw_dir)
    ensure_directory(assets_dir)
    warnings: list[str] = []
    image_records: list[dict] = []
    section_lines: list[str] = []
    extraction_quality = "good"
    used_downloads = False
    used_ocr = False
    matches = list(MARKDOWN_IMAGE_PATTERN.finditer(raw_text))
    if not matches:
        return normalize_text_content("markdown", raw_text), {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good",
            "warnings": [],
            "location_map": {
                "type": "line_map",
                "source_path": source_record["source_path"],
            },
        }
    for image_index, match in enumerate(matches, start=1):
        alt_text = (match.group("alt") or "").strip()
        target_value = (match.group("target") or "").strip()
        candidate_values = markdown_image_target_candidates(target_value)
        if not candidate_values:
            warnings.append(f"markdown_image_invalid_target:{image_index}")
            extraction_quality = "partial"
            section_lines.extend([
                f"## 内嵌图片 image_{image_index}",
                "",
                f"- alt: {alt_text or '(empty)'}",
                "> 源文件包含图片，但图片地址为空或无法解析。",
                "",
            ])
            continue
        candidate = candidate_values[0]
        parsed = urlparse(candidate)
        try:
            if parsed.scheme in {"http", "https"}:
                asset_result = download_markdown_image_to_assets(
                    target=target,
                    source_record=source_record,
                    raw_dir=raw_dir,
                    image_index=image_index,
                    target_value=candidate,
                    allow_insecure_downloads=allow_insecure_downloads,
                )
                used_downloads = True
                asset_path = asset_result["asset_path"]
            else:
                asset_path = resolve_markdown_local_image_path(raw_path, candidate, raw_dir)
                asset_result = {
                    "storage_kind": "local_raw",
                    "asset_path": asset_path,
                    "asset_hash": file_sha256(asset_path),
                    "content_type": None,
                    "download_mode": "local_raw",
                    "warnings": [],
                }
            section_block, image_metadata = convert_markdown_embedded_image_to_section(
                target=target,
                asset_path=asset_path,
                asset_label=f"image_{image_index}",
                alt_text=alt_text,
                image_context={
                    "markdown_source_path": source_record["source_path"],
                    "image_index": image_index,
                    "image_alt": alt_text,
                    "image_target": candidate,
                },
            )
            used_ocr = used_ocr or bool(image_metadata.get("location_map", {}).get("ocr", {}).get("used"))
            image_quality = image_metadata.get("extraction_quality", "partial")
            if image_quality in {"failed", "poor", "partial"}:
                extraction_quality = "partial"
            warnings.extend(image_metadata.get("warnings", []))
            warnings.extend(asset_result.get("warnings", []))
            section_lines.extend(section_block)
            section_lines.append("")
            image_records.append({
                "index": image_index,
                "alt": alt_text,
                "target": candidate,
                "storage_kind": asset_result["storage_kind"],
                "asset_path": str(asset_path),
                "asset_hash": asset_result["asset_hash"],
                "content_type": asset_result.get("content_type"),
                "download_mode": asset_result.get("download_mode"),
                "image_metadata": image_metadata.get("location_map", {}),
            })
        except Exception as exc:
            warnings.append(f"markdown_image_conversion_failed:{image_index}:{type(exc).__name__}")
            extraction_quality = "partial"
            section_lines.extend([
                f"## 内嵌图片 image_{image_index}",
                "",
                f"- alt: {alt_text or '(empty)'}",
                f"- target: {candidate}",
                "> 源文件包含图片，但内容暂时无法转换为文本。",
                "",
            ])
            image_records.append({
                "index": image_index,
                "alt": alt_text,
                "target": candidate,
                "storage_kind": "failed",
                "error": f"{type(exc).__name__}: {exc}",
            })
    normalized_markdown = normalize_text_content("markdown", raw_text)
    if section_lines:
        normalized_markdown = (
            normalized_markdown.rstrip()
            + "\n\n## 内嵌图片内容 / Embedded Image Content\n\n"
            + "\n".join(section_lines).strip()
            + "\n"
        )
    extraction_method = "python_only"
    if used_downloads:
        extraction_method += "+remote_assets"
    if used_ocr:
        extraction_method += "+tesseract"
    return normalized_markdown, {
        "content_format": "markdown",
        "extraction_method": extraction_method,
        "extraction_quality": extraction_quality,
        "warnings": sorted(set(warnings)),
        "location_map": {
            "type": "markdown_with_embedded_images",
            "source_path": source_record["source_path"],
            "image_count": len(image_records),
            "images": image_records,
        },
    }


def normalize_agent_assisted_image_result(hook_result: dict | None, min_confidence: float) -> tuple[str, dict]:
    if not isinstance(hook_result, dict):
        return "", {
            "used": False,
            "ok": False,
            "quality": "unavailable",
            "warnings": [],
            "confidence": 0.0,
            "reason": "image_to_text_unavailable",
            "summary": "",
        }
    confidence = coerce_float(hook_result.get("confidence", 0.0), 0.0)
    reason = str(hook_result.get("reason", "")).strip() or "image_to_text_agent_result"
    warnings = [str(item).strip() for item in hook_result.get("warnings", []) if str(item).strip()]
    if confidence < min_confidence:
        return "", {
            "used": True,
            "ok": False,
            "quality": "low_confidence",
            "warnings": sorted(set([*warnings, "image_to_text_low_confidence"])),
            "confidence": confidence,
            "reason": reason,
            "summary": normalize_binary_snippet_text(str(hook_result.get("summary", ""))),
        }
    extracted_text = normalize_binary_snippet_text(str(hook_result.get("extracted_text") or hook_result.get("text") or ""))
    summary = normalize_binary_snippet_text(str(hook_result.get("summary", "")))
    combined_text = extracted_text
    if summary and summary not in combined_text:
        combined_text = f"摘要: {summary}\n\n{extracted_text}".strip() if extracted_text else summary
    if not combined_text:
        return "", {
            "used": True,
            "ok": False,
            "quality": "empty",
            "warnings": sorted(set([*warnings, "image_to_text_no_text"])),
            "confidence": confidence,
            "reason": reason,
            "summary": summary,
        }
    quality = "good" if image_understanding_text_is_meaningful(combined_text) else "partial"
    if quality == "partial":
        warnings.append("image_to_text_low_signal")
    return combined_text, {
        "used": True,
        "ok": True,
        "quality": quality,
        "warnings": sorted(set(warnings)),
        "confidence": confidence,
        "reason": reason,
        "summary": summary,
    }


def run_agent_assisted_image_to_text(
    target: Path | None,
    raw_path: Path,
    image_context: dict | None = None,
) -> tuple[str, dict]:
    if target is None:
        return "", {
            "used": False,
            "ok": False,
            "quality": "disabled",
            "warnings": [],
            "confidence": 0.0,
            "reason": "image_to_text_target_missing",
            "summary": "",
        }
    automation_config = load_automation_target_config(load_workspace_config(target), "image_to_text")
    if not automation_config.get("enabled"):
        return "", {
            "used": False,
            "ok": False,
            "quality": "disabled",
            "warnings": [],
            "confidence": 0.0,
            "reason": "image_to_text_agent_disabled",
            "summary": "",
        }
    payload = {
        "task": "describe_image",
        "image_path": str(raw_path.resolve()),
        "image_name": raw_path.name,
        "image_context": image_context or {},
    }
    hook_result = run_json_automation_command(
        target=target,
        command=automation_config.get("command", []),
        payload=payload,
        timeout_seconds=automation_config.get("timeout_seconds", 45),
    )
    return normalize_agent_assisted_image_result(hook_result, automation_config.get("min_confidence", 0.8))


def run_tesseract_ocr(raw_path: Path) -> tuple[str, dict]:
    if not command_exists("tesseract"):
        return "", {
            "used": False,
            "ok": False,
            "quality": "missing",
            "warnings": ["tesseract_missing"],
            "details": "tesseract command not found in PATH",
        }
    completed = subprocess.run(
        ["tesseract", str(raw_path), "stdout", "--psm", "3"],
        check=False,
        capture_output=True,
        text=True,
    )
    raw_text = (completed.stdout or "").strip()
    normalized_text = normalize_binary_snippet_text(raw_text)
    if completed.returncode != 0:
        stderr_text = (completed.stderr or "").strip() or "unknown_tesseract_error"
        return "", {
            "used": True,
            "ok": False,
            "quality": "failed",
            "warnings": ["tesseract_ocr_failed"],
            "details": stderr_text,
        }
    if not normalized_text:
        return "", {
            "used": True,
            "ok": True,
            "quality": "empty",
            "warnings": ["tesseract_ocr_no_text"],
            "details": (completed.stderr or "").strip(),
        }
    quality = "good" if ocr_text_is_meaningful(normalized_text) else "partial"
    warnings: list[str] = []
    if quality == "partial":
        warnings.append("tesseract_ocr_low_signal")
    return normalized_text, {
        "used": True,
        "ok": True,
        "quality": quality,
        "warnings": warnings,
        "details": (completed.stderr or "").strip(),
    }


def convert_image_to_markdown(
    raw_path: Path,
    *,
    target: Path | None = None,
    image_context: dict | None = None,
) -> tuple[str, dict]:
    warnings: list[str] = []
    stat = raw_path.stat()
    metadata_lines = [
        f"# {raw_path.stem}",
        "",
        f"- 文件名: {raw_path.name}",
        f"- 文件大小: {stat.st_size} bytes",
    ]
    location_map = {
        "type": "image_metadata",
        "source_path": str(raw_path),
    }
    width = None
    height = None
    image_format = None
    image_mode = "unknown"
    exif = {}
    if Image is None:
        warnings.append("pillow_missing")
        width, height, image_format = image_size_from_binary(raw_path)
    else:
        try:
            with Image.open(raw_path) as image:
                width = image.width
                height = image.height
                image_mode = image.mode
                image_format = image.format or "unknown"
                if hasattr(image, "getexif"):
                    exif_data = image.getexif()
                    if exif_data:
                        exif = {str(key): str(value) for key, value in exif_data.items()}
        except Exception:
            warnings.append("pillow_image_open_failed")
            width, height, image_format = image_size_from_binary(raw_path)
            image_mode = "unknown"
    metadata_lines.extend([
        f"- 尺寸: {width}x{height}" if width and height else "- 尺寸: unknown",
        f"- 模式: {image_mode}",
        f"- 格式: {image_format}" if image_format else "- 格式: unknown",
    ])
    if exif:
        metadata_lines.append("")
        metadata_lines.append("## EXIF")
        for key, value in sorted(exif.items()):
            metadata_lines.append(f"- {key}: {value}")
    ocr_text, ocr_result = run_tesseract_ocr(raw_path)
    warnings.extend(ocr_result.get("warnings", []))
    llm_text = ""
    llm_result = {
        "used": False,
        "ok": False,
        "quality": "disabled",
        "warnings": [],
        "confidence": 0.0,
        "reason": "image_to_text_not_attempted",
        "summary": "",
    }
    should_try_llm = not ocr_text or ocr_result.get("quality") in {"missing", "failed", "partial", "empty"}
    if should_try_llm:
        llm_text, llm_result = run_agent_assisted_image_to_text(
            target=target,
            raw_path=raw_path,
            image_context=image_context,
        )
        warnings.extend(llm_result.get("warnings", []))
    extraction_quality = "partial"
    combined_sections: list[str] = []
    if ocr_text:
        combined_sections.extend(["## OCR 文本 / OCR Text", "", ocr_text])
        extraction_quality = "good" if ocr_result.get("quality") == "good" else "partial"
    if llm_text:
        combined_sections.extend(["## LLM 图片理解 / LLM Image Understanding", "", llm_text])
        if llm_result.get("quality") == "good":
            extraction_quality = "good"
    if combined_sections:
        metadata_lines.extend(["", *combined_sections])
    else:
        if ocr_result.get("quality") == "failed":
            metadata_lines.extend(["", "> tesseract 已安装，但本次 OCR 执行失败，当前仅保留图片元数据。"])
        elif ocr_result.get("quality") == "missing":
            metadata_lines.extend(["", "> 当前环境未检测到 tesseract，图片仅生成元数据级 normalized 文档。"])
        else:
            metadata_lines.extend(["", "> 本次 OCR 未提取到稳定正文，当前保留图片元数据供后续人工或 Agent 处理。"])
        if llm_result.get("used") and not llm_text:
            metadata_lines.extend(["", "> 已尝试使用 LLM 识别图片内容，但当前未得到可稳定落盘的文本结果。"])
    markdown = "\n".join(metadata_lines).strip() + "\n"
    location_map["image"] = {
        "has_exif": bool(exif),
        "width": width,
        "height": height,
        "mode": image_mode,
        "format": image_format,
    }
    location_map["ocr"] = {
        "used": ocr_result.get("used", False),
        "ok": ocr_result.get("ok", False),
        "quality": ocr_result.get("quality"),
        "char_count": len(ocr_text),
    }
    location_map["llm_image_understanding"] = {
        "used": llm_result.get("used", False),
        "ok": llm_result.get("ok", False),
        "quality": llm_result.get("quality"),
        "char_count": len(llm_text),
        "confidence": llm_result.get("confidence", 0.0),
        "reason": llm_result.get("reason"),
        "summary": llm_result.get("summary", ""),
    }
    extraction_method = "python_only"
    if ocr_result.get("used"):
        extraction_method += "+tesseract"
    if llm_result.get("used"):
        extraction_method += "+agent_assisted"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": extraction_method,
        "extraction_quality": extraction_quality,
        "warnings": warnings if warnings else ([] if (ocr_text or llm_text) else ["image_metadata_only"]),
        "location_map": location_map,
    }


def convert_unknown_source_to_placeholder(raw_path: Path, source_type: str) -> tuple[str, dict]:
    markdown = (
        f"# {raw_path.stem}\n\n"
        f"- source_type: {source_type}\n"
        f"- source_path: {raw_path}\n\n"
        "> 当前版本暂不支持该格式的自动标准化，请等待后续转换器或使用 Agent 辅助处理。\n"
    )
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "poor",
        "warnings": [f"unsupported_source_type:{source_type}"],
        "location_map": {
            "type": "placeholder",
            "source_path": str(raw_path),
        },
    }


def convert_source_to_normalized_markdown(raw_path: Path, source_type: str) -> tuple[str, dict]:
    if source_type in {"markdown", "plain_text"}:
        raw_text = raw_path.read_text(encoding="utf-8")
        return normalize_text_content(source_type, raw_text), {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good",
            "warnings": [],
        }
    if source_type == "pdf":
        return convert_pdf_to_markdown(raw_path)
    if source_type == "docx":
        return convert_docx_to_markdown(raw_path)
    if source_type == "doc":
        return convert_legacy_doc_to_markdown(raw_path)
    if source_type == "spreadsheet":
        return convert_spreadsheet_to_markdown(raw_path)
    if source_type == "image":
        return convert_image_to_markdown(raw_path)
    return convert_unknown_source_to_placeholder(raw_path, source_type)


def build_failed_conversion_placeholder(raw_path: Path, source_type: str, exc: Exception) -> tuple[str, dict]:
    markdown = (
        f"# {raw_path.stem}\n\n"
        f"- source_type: {source_type}\n"
        f"- source_path: {raw_path}\n"
        f"- converter_error: {type(exc).__name__}: {exc}\n\n"
        "> 当前版本在标准化该文件时失败，已生成占位文档等待后续修复或 Agent 辅助处理。\n"
    )
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "failed",
        "warnings": [f"converter_error:{type(exc).__name__}", str(exc)],
        "location_map": {
            "type": "conversion_error",
            "source_path": str(raw_path),
        },
    }


def normalize_source_record(
    target: Path,
    source_record: dict,
    *,
    allow_insecure_downloads: bool = True,
) -> dict | None:
    source_type = source_record["source_type"]
    if source_type in {"markdown", "plain_text"}:
        return normalize_markdown_or_text_record(
            target,
            source_record,
            allow_insecure_downloads=allow_insecure_downloads,
        )
    raw_path = resolve_source_record_path(target, source_record["source_path"])
    ensure_path_within_raw_root(raw_path, resolve_workspace_raw_dir(target), purpose="Source record")
    try:
        normalized_text, metadata = convert_source_to_normalized_markdown(raw_path, source_type)
    except Exception as exc:
        normalized_text, metadata = build_failed_conversion_placeholder(raw_path, source_type, exc)
    normalized_text = normalize_text_content("markdown", normalized_text)
    normalized_rel_path = Path("normalized") / f"{source_record['source_id']}.md"
    normalized_abs_path = target / normalized_rel_path
    normalized_abs_path.write_text(normalized_text, encoding="utf-8")
    normalized_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    title = raw_path.stem
    location_map = dict(metadata.get("location_map", {}))
    location_map["source_path"] = source_record["source_path"]
    if location_map.get("type") == "line_map" and "normalized_line_range" not in location_map:
        line_count = len(normalized_text.splitlines()) or 1
        location_map["normalized_line_range"] = f"1-{line_count}"
    return {
        "source_id": source_record["source_id"],
        "source_type": source_type,
        "source_path": source_record["source_path"],
        "normalized_path": str(normalized_rel_path),
        "title": title,
        "language": "unknown",
        "content_format": metadata.get("content_format", "markdown"),
        "raw_hash": source_record["source_hash"],
        "normalized_hash": normalized_hash,
        "normalizer_version": "normalize_v1",
        "document_kind": "note",
        "structure_quality": "unknown",
        "chunk_strategy_hint": "heading_first",
        "extraction_method": metadata.get("extraction_method", "python_only"),
        "extraction_quality": metadata.get("extraction_quality", "partial"),
        "warnings": metadata.get("warnings", []),
        "location_map": location_map,
        "updated_at": utc_now_iso(),
    }
