from __future__ import annotations

import argparse
import copy
import json
import importlib.util
import os
import shutil
import subprocess
import sys
import hashlib
import re
import csv
import zipfile
import struct
import zlib
import math
import difflib
import shlex
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from string import Template
import ast
import tomllib
import tempfile
import fcntl
import mimetypes
import ssl
import urllib.error
import urllib.request
from urllib.parse import quote
from urllib.parse import unquote
from urllib.parse import urlparse
from xml.etree import ElementTree as ET

from .semantic import (
    SemanticTaskConfig,
    build_semantic_decision_id,
    fingerprint_payload,
    item_type_for_task,
    normalize_string_list,
    normalize_semantic_hook_decision,
    semantic_batches_dir,
)

try:
    import docx
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    docx = None

try:
    import openpyxl
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    openpyxl = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    PdfReader = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover - 依赖缺失时走降级逻辑
    Image = None


ROOT_MARKERS = ("pyproject.toml", ".git")
DEFAULT_CHUNK_TARGET_TOKENS = 1000
DEFAULT_CHUNK_MAX_TOKENS = 1600
DEFAULT_CHUNK_MIN_TOKENS = 200
MAX_FILENAME_COMPONENT_BYTES = 240
FILENAME_HASH_LENGTH = 12
QUERY_READING_DEPTH_LIMITS = {
    "standard": {
        "claim_limit": 3,
        "chunk_limit": 2,
    },
    "deep": {
        "claim_limit": 6,
        "chunk_limit": 5,
    },
}
ALIAS_INDEX_REL_PATH = Path("indexes") / "aliases.json"
PAGE_ALIAS_OVERRIDES_REL_PATH = Path("state") / "page_alias_overrides.json"
PAGE_ALIAS_OVERRIDES_LOCK_REL_PATH = Path("state") / ".page_alias_overrides.lock"
STRUCTURE_BLOCKS_REL_PATH = Path("state") / "structure_blocks.jsonl"
EVIDENCE_BLOCKS_REL_PATH = Path("state") / "evidence_blocks.jsonl"
KNOWLEDGE_UNITS_REL_PATH = Path("state") / "knowledge_units.jsonl"
SEMANTIC_DECISIONS_REL_PATH = Path("state") / "semantic_decisions.jsonl"
MARKDOWN_IMAGE_PATTERN = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<target>[^)\s]+)(?:\s+\"[^\"]*\")?\)")
NEGATION_MARKERS = ("不", "不是", "没有", "无法", "不能", "未", "无", "禁止", "不要", "not ", "no ", "never ", "cannot ")
PACKAGE_IMPORT_ALIASES = {
    "python-docx": "docx",
    "pillow": "PIL",
}
PAGE_RENDER_TARGETS = {
    "readable_concept": {
        "page_types": {"concept"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": "readable_concept",
    },
    "guide": {
        "page_types": {"guide"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": None,
    },
    "example": {
        "page_types": {"example"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": None,
    },
    "topic": {
        "page_types": {"topic"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": None,
    },
    "reference": {
        "page_types": {"reference"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": None,
    },
    "timeline": {
        "page_types": {"timeline"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": None,
    },
    "overview": {
        "page_types": {"overview"},
        "rebuild_strategy": "review_affected_pages",
        "grounding_checker": None,
    },
    "qa_note": {
        "page_types": {"qa-note"},
        "rebuild_strategy": "none",
        "grounding_checker": None,
    },
    "concept_update": {
        "page_types": {"concept-update"},
        "rebuild_strategy": "none",
        "grounding_checker": None,
    },
}
QUERY_FIELD_WEIGHTS = {
    "title": 5.0,
    "aliases": 4.0,
    "hierarchy": 3.5,
    "summary": 3.0,
    "headings": 2.5,
    "body": 1.0,
    "claim_text": 2.0,
    "source_refs": 0.5,
}
QUERY_PAGE_TYPE_WEIGHTS = {
    "overview": 1.25,
    "concept": 1.22,
    "topic": 1.08,
    "guide": 1.05,
    "example": 0.95,
    "reference": 1.04,
    "timeline": 1.02,
    "entity": 1.10,
    "source-summary": 1.00,
    "qa": 0.95,
    "draft": 0.70,
}
QUERY_PAGE_STATUS_WEIGHTS = {
    "stable": 1.10,
    "draft": 0.80,
    "disputed": 0.90,
    "outdated": 0.60,
    # 设计文档里没有 needs_review，这里把它视为比 draft 更需要谨慎的状态。
    "needs_review": 0.75,
}
QUERY_BM25_K1 = 1.5
QUERY_BM25_B = 0.75
QUERY_EXACT_MATCH_MAX_BOOST = 1.35
QUERY_HEADING_BLACKLIST = {
    "原文概览 / source overview",
    "核心观点 / key points",
    "知识声明 / claims",
    "证据切块 / chunks",
    "后续建议 / next steps",
    "概念摘要 / concept summary",
    "核心陈述 / canonical claim",
    "支撑声明 / supporting claims",
    "来源页面 / source pages",
    "来源证据 / source evidence",
    "审核提示 / review notes",
    "工作区综述 / workspace overview",
    "稳定概念 / stable concepts",
    "来源覆盖 / source coverage",
    "维护状态 / maintenance",
}
SEARCH_PAGES_INDEX_REL_PATH = Path("indexes") / "search_pages.jsonl"
SEARCH_PAGES_INDEX_VERSION = "search_pages_v2"
ALIAS_INDEX_VERSION = "aliases_v1"
QUERY_ANSWER_HANDOFF_CONTRACT_VERSION = "query_answer_handoff/v1"
REVIEW_AUTO_HANDOFF_CONTRACT_VERSION = "review_auto_handoff/v1"
ANSWER_READY_OUTPUT_VERSION = "answer_ready_query/v1"
AUTOMATION_STRATEGIES = {"safe_auto", "agent_assisted"}
SEMANTIC_TASK_NAMES = ("document_analysis", "claim_candidate_quality", "claim_role", "page_intent", "page_route")
WORKSPACE_SCHEMA_VERSION = "v1"
WORKSPACE_MIN_SUPPORTED_SCHEMA_VERSION = "v1"
WORKSPACE_SCHEMA_VERSION_ORDER = ("v1",)
QUERY_INTENT_MARKERS = {
    "overview": (
        "概览", "概况", "总览", "总述", "整体", "全局", "框架", "脉络", "overview",
        "主要讲什么", "主要内容", "整体内容", "有哪些主题",
    ),
    "definition": (
        "是什么", "什么是", "定义", "是指", "指什么", "介绍一下", "what is", "define",
    ),
    "compare": (
        "区别", "对比", "比较", "差异", "vs", "versus", "compare",
    ),
    "timeline": (
        "时间线", "演变", "历史", "历程", "timeline",
    ),
    "reference": (
        "参数", "清单", "列表", "faq", "FAQ", "参考", "规则", "字段", "配置项", "reference",
    ),
    "how_to": (
        "如何", "怎么", "步骤", "做法", "实践", "how to", "tutorial",
    ),
    "evidence": (
        "来源", "证据", "出处", "引用", "依据", "为什么", "source", "evidence", "trace",
    ),
}
QUERY_INTENT_FIELD_MULTIPLIERS = {
    "lookup": {},
    "overview": {
        "title": 1.15,
        "summary": 1.20,
        "headings": 1.15,
        "body": 1.08,
        "claim_text": 1.05,
    },
    "definition": {
        "title": 1.15,
        "summary": 1.15,
        "aliases": 1.10,
        "hierarchy": 1.08,
    },
    "compare": {
        "claim_text": 1.15,
        "body": 1.10,
        "headings": 1.05,
        "hierarchy": 1.08,
    },
    "timeline": {
        "body": 1.10,
        "summary": 1.05,
        "source_refs": 1.10,
    },
    "reference": {
        "title": 1.10,
        "headings": 1.12,
        "body": 1.08,
        "claim_text": 1.06,
        "hierarchy": 1.10,
    },
    "how_to": {
        "headings": 1.15,
        "body": 1.15,
        "claim_text": 1.10,
        "hierarchy": 1.12,
    },
    "evidence": {
        "source_refs": 1.80,
        "claim_text": 1.20,
        "body": 1.05,
        "hierarchy": 1.08,
    },
}
CLAIM_DEPENDENT_PREFIXES = (
    "旨在",
    "以便",
    "从而",
    "从而让",
    "从而使",
    "但",
    "具体细节",
    "而不是",
    "并且",
    "同时",
    "以及",
    "其中",
    "例如",
    "比如",
)
CLAIM_META_PREFIXES = (
    "这是一",
    "这是",
    "本文件",
    "本文",
    "这份",
    "这个",
    "该文",
    "该文件",
)
CLAIM_STANDALONE_PREDICATE_MARKERS = (
    "是",
    "不是",
    "意味着",
    "需要",
    "应该",
    "可以",
    "能够",
    "会",
    "能",
    "支持",
    "保留",
    "保持",
    "维护",
    "构成",
    "属于",
    "记录",
    "标注",
    "更新",
    "提取",
    "整合",
    "构建",
)


@dataclass
class CommandResult:
    # 所有命令统一返回这一层结构，方便后面既能打印给人看，也能转成 JSON 给 Agent 消费。
    exit_code: int = 0
    payload: dict | None = None
    message: str | None = None


def find_project_root(start: Path | None = None) -> Path:
    # 从当前文件或给定路径一路向上找项目根目录。
    # 这里优先依赖 pyproject.toml 和 .git 这样的标记文件，避免把临时目录误判成项目根。
    current = (start or Path(__file__)).resolve()
    for path in [current, *current.parents]:
        if all((path / marker).exists() or marker == ".git" and (path / marker).exists() for marker in ROOT_MARKERS):
            return path
    for path in [current, *current.parents]:
        if (path / "pyproject.toml").exists():
            return path
    raise FileNotFoundError("Could not locate project root from current path.")


def parse_yaml_scalar(raw: str):
    # 当前项目里的 YAML 结构比较简单，这里做一个轻量标量解析器就够用了。
    # 目标不是支持完整 YAML 语法，而是稳定读取我们自己维护的配置文件。
    value = raw.strip()
    if value == "":
        return {}
    if value in {"true", "false"}:
        return value == "true"
    if value in {"[]", "{}"}:
        return ast.literal_eval(value)
    if value.startswith(("[", "{", "'", '"')):
        return ast.literal_eval(value)
    return value


def load_simple_yaml(path: Path) -> dict:
    # 这里实现的是“够用版 YAML 读取器”：
    # 只处理字典嵌套、简单标量、缩进列表，不引入额外依赖。
    root: dict = {}
    stack: list[tuple[int, object, dict | None, str | None]] = [(-1, root, None, None)]

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        if not raw_line.strip() or raw_line.lstrip().startswith("#"):
            continue

        # 通过缩进层级维护一个栈，逐步构造嵌套字典结构。
        indent = len(raw_line) - len(raw_line.lstrip(" "))
        while len(stack) > 1 and indent <= stack[-1][0]:
            stack.pop()

        stripped = raw_line.strip()
        current_indent, parent, owner, owner_key = stack[-1]

        if stripped.startswith("- "):
            if isinstance(parent, dict):
                # 空字典节点如果下一层第一条就是 "- item"，就把它视为列表容器。
                if parent == {} and owner is not None and owner_key is not None:
                    parent = []
                    owner[owner_key] = parent
                    stack[-1] = (current_indent, parent, owner, owner_key)
                else:
                    raise ValueError(f"Invalid YAML list item placement in {path}: {raw_line}")
            if not isinstance(parent, list):
                raise ValueError(f"Invalid YAML list item placement in {path}: {raw_line}")
            parent.append(parse_yaml_scalar(stripped[2:]))
            continue

        key, sep, rest = stripped.partition(":")
        if not sep:
            continue

        if rest.strip() == "":
            node: dict | list = {}
            parent[key] = node
            stack.append((indent, node, parent, key))
        else:
            parent[key] = parse_yaml_scalar(rest)

    return root


def load_runtime_manifest(root: Path) -> dict:
    # runtime_manifest 是运行环境的统一来源，doctor/bootstrap 都会依赖它。
    manifest_path = root / "config" / "runtime_manifest.yml"
    return load_simple_yaml(manifest_path)


def load_project_metadata(root: Path) -> dict:
    # pyproject.toml 负责项目名、版本、依赖和 CLI 入口等元信息。
    with (root / "pyproject.toml").open("rb") as fh:
        return tomllib.load(fh)


def print_result(result: CommandResult, as_json: bool = False) -> int:
    # 所有命令统一走这里输出，避免每个命令自己决定打印格式。
    if as_json:
        print(json.dumps(result.payload or {}, ensure_ascii=False, indent=2))
        return result.exit_code

    if result.message:
        print(result.message)
    elif result.payload is not None:
        print(json.dumps(result.payload, ensure_ascii=False, indent=2))
    return result.exit_code


def build_workspace_summary(target_dir: Path, raw_dir: Path | None = None) -> dict:
    # 给上层 Agent 和人类读者一份“可以直接复述”的路径摘要，避免只剩目录名。
    lint_report_path = target_dir / "reports" / "lint" / "lint_latest.md"
    schema_version = None
    schema_guard = {
        "status": "unknown",
        "expected_schema_version": WORKSPACE_SCHEMA_VERSION,
        "minimum_supported_schema_version": WORKSPACE_MIN_SUPPORTED_SCHEMA_VERSION,
    }
    config_path = target_dir / "config" / "project.yml"
    if config_path.exists():
        workspace_config = load_simple_yaml(config_path)
        workspace_block = workspace_config.get("workspace", {})
        if not isinstance(workspace_block, dict):
            workspace_block = {}
        schema_version = str(workspace_block.get("schema_version", "")).strip() or None
        schema_guard["status"] = workspace_schema_guard_status(schema_version).replace("_schema_version", "")
    summary = {
        "workspace_dir": str(target_dir),
        "workspace_name": target_dir.name,
        "entry_page_path": str(target_dir / "wiki" / "index.md"),
        "wiki_log_path": str(target_dir / "wiki" / "log.md"),
        "lint_report_path": str(lint_report_path),
        "lint_report_exists": lint_report_path.exists(),
        "schema_version": schema_version,
        "schema_guard": schema_guard,
    }
    if raw_dir is not None:
        summary["raw_dir"] = str(raw_dir)
    return summary


def render_workspace_summary_message(
    action_label: str,
    target_dir: Path,
    raw_dir: Path | None = None,
    extra_lines: list[str] | None = None,
) -> str:
    # 纯文本模式也显式带绝对路径，减少 UI 把链接文案压缩成目录名后的歧义。
    summary = build_workspace_summary(target_dir, raw_dir)
    lines = [
        f"{action_label}",
        f"Workspace: {summary['workspace_dir']}",
    ]
    if raw_dir is not None:
        lines.append(f"Raw sibling: {summary['raw_dir']}")
    lines.extend(
        [
            f"Entry page: {summary['entry_page_path']}",
            (
                f"Lint report: {summary['lint_report_path']}"
                if summary["lint_report_exists"]
                else f"Lint report: {summary['lint_report_path']} (will be created after the first lint run)"
            ),
        ]
    )
    schema_guard = summary.get("schema_guard", {})
    if schema_guard.get("status") == "unsupported":
        lines.append(
            "Schema guard: "
            f"workspace_schema={summary.get('schema_version')}, "
            f"supported={schema_guard.get('minimum_supported_schema_version')}..{schema_guard.get('expected_schema_version')}"
        )
    elif summary.get("schema_version"):
        lines.append(f"Schema version: {summary.get('schema_version')}")
    if extra_lines:
        lines.extend(line for line in extra_lines if line)
    return "\n".join(lines)


def compare_python_version(spec: str) -> bool:
    # 目前只需要支持 >= 这种最常见的版本约束，先不把比较器做复杂。
    if not spec.startswith(">="):
        return True
    required = tuple(int(part) for part in spec[2:].split("."))
    current = sys.version_info[: len(required)]
    return current >= required


def run_check_command(command: list[str]) -> tuple[bool, str]:
    # doctor 检查系统工具时统一走这个函数。
    # 成功时返回命令输出，失败时返回异常信息，尽量保留可诊断线索。
    try:
        completed = subprocess.run(
            command,
            check=True,
            capture_output=True,
            text=True,
        )
        output = (completed.stdout or completed.stderr).strip()
        return True, output
    except (OSError, subprocess.CalledProcessError) as exc:
        return False, str(exc)


def command_exists(name: str) -> bool:
    # 有些增强能力只需要知道系统命令是否在 PATH 中，不必真的跑一遍。
    return shutil.which(name) is not None


def normalize_package_name(name: str) -> str:
    # Python 包名和 import 名有时会用 - / _ 混用，这里先做最常见的归一化。
    if name in PACKAGE_IMPORT_ALIASES:
        return PACKAGE_IMPORT_ALIASES[name]
    return name.replace("-", "_")


def check_python_package_installed(name: str) -> bool:
    # 通过 importlib 检测包是否可导入，避免直接 import 触发副作用。
    return importlib.util.find_spec(normalize_package_name(name)) is not None


def collect_python_package_report(manifest: dict) -> dict:
    # 把 Python 包依赖拆成 required / optional 两层，doctor 输出时会更容易读。
    required_packages = manifest.get("python_packages", {}).get("required", [])
    optional_groups = manifest.get("python_packages", {}).get("optional", {})

    required_report = {
        package: {"ok": check_python_package_installed(package)}
        for package in required_packages
    }
    optional_report = {
        group: {
            package: {"ok": check_python_package_installed(package)}
            for package in packages
        }
        for group, packages in optional_groups.items()
    }
    return {
        "required": required_report,
        "optional": optional_report,
    }


def build_doctor_payload(root: Path) -> dict:
    # doctor 的核心逻辑集中在这里：
    # 读取项目元信息、读取运行时清单、检查 Python / git / 可选系统工具 / Python 包。
    manifest = load_runtime_manifest(root)
    project_meta = load_project_metadata(root)

    python_spec = manifest["runtime"]["python"]["version"]
    python_ok = compare_python_version(python_spec)
    git_command = manifest["runtime"]["git"]["check"]["command"]
    git_ok, git_output = run_check_command(git_command)

    required = {
        "python": {
            "required_version": python_spec,
            "current_version": ".".join(map(str, sys.version_info[:3])),
            "ok": python_ok,
        },
        "git": {
            "ok": git_ok,
            "details": git_output,
        },
    }

    optional_tools: dict[str, dict] = {}
    current_platform = sys.platform
    platform_label = (
        "windows" if current_platform.startswith("win")
        else "macos" if current_platform == "darwin"
        else "linux"
    )
    # 平台标签单独算出来，后面 Windows/macOS/Linux 的提示逻辑都会复用。
    for tool_name, tool_data in manifest.get("system_tools", {}).get("optional", {}).items():
        ok, details = run_check_command(tool_data["check"]["command"])
        optional_tools[tool_name] = {
            "ok": ok,
            "purpose": tool_data["purpose"],
            "fallback": tool_data["fallback"],
            "supported_platforms": tool_data["supported_platforms"],
            "supported_on_current_platform": platform_label in tool_data["supported_platforms"],
            "details": details,
        }

    python_packages = collect_python_package_report(manifest)
    required_python_packages_ok = all(
        item["ok"] for item in python_packages["required"].values()
    )
    bootstrap_examples = {
        "windows": [
            r"py -3.12 -m venv .venv",
            r".venv\Scripts\python -m pip install -U pip",
            r".venv\Scripts\python -m myagentwiki bootstrap --extra dev",
        ],
        "macos": [
            "python3.12 -m venv .venv",
            ".venv/bin/python -m pip install -U pip",
            ".venv/bin/python -m myagentwiki bootstrap --extra dev",
        ],
        "linux": [
            "python3.12 -m venv .venv",
            ".venv/bin/python -m pip install -U pip",
            ".venv/bin/python -m myagentwiki bootstrap --extra dev",
        ],
    }
    return {
        "project": {
            "name": project_meta["project"]["name"],
            "version": project_meta["project"]["version"],
            "root": str(root),
        },
        "platform": {
            "os_name": os.name,
            "sys_platform": sys.platform,
            "platform_label": platform_label,
        },
        "required": required,
        "python_packages": python_packages,
        "optional": optional_tools,
        "summary": {
            "required_ok": all(item["ok"] for item in required.values()) and required_python_packages_ok,
            "missing_required_python_packages": [
                name for name, item in python_packages["required"].items() if not item["ok"]
            ],
            "optional_missing": [name for name, item in optional_tools.items() if not item["ok"]],
        },
        "bootstrap_guidance": {
            "recommended_shell_examples": bootstrap_examples.get(platform_label, bootstrap_examples["linux"]),
            "notes": [
                "核心流程不依赖 shell 专属语法，Windows / macOS / Linux 都应可执行。",
                "可选系统工具缺失时，优先走 Python 降级路径，再视需要交给 Agent 补强。",
            ],
        },
    }


def command_doctor(args: argparse.Namespace) -> CommandResult:
    # command_* 函数负责组装业务结果；真正的公共逻辑尽量往辅助函数里收。
    root = find_project_root()
    payload = build_doctor_payload(root)
    return CommandResult(payload=payload, message="MyAgentWiki doctor report ready.")


def command_bootstrap(args: argparse.Namespace) -> CommandResult:
    # bootstrap 当前先聚焦 Python 依赖安装，不擅自处理系统级软件。
    # 命令本身保持跨平台：直接调用当前 Python 解释器，不依赖 bash/zsh 专属语法。
    root = find_project_root()
    install_command = [sys.executable, "-m", "pip", "install", "-e"]
    if args.extras:
        # extras 允许用户按需装增强能力，比如 pdf / dev。
        install_command.append(f"{str(root)}[{','.join(args.extras)}]")
    else:
        install_command.append(str(root))

    doctor_payload = build_doctor_payload(root)

    if args.dry_run:
        # dry-run 先把“将做什么”讲清楚，避免一上来就改环境。
        payload = {
            "action": "dry_run",
            "install_command": install_command,
            "project_root": str(root),
            "requested_extras": args.extras,
            "doctor_summary": doctor_payload["summary"],
        }
        return CommandResult(payload=payload, message="Bootstrap dry run complete.")

    completed = subprocess.run(
        install_command,
        check=False,
        capture_output=True,
        text=True,
    )
    # pip 的 stdout / stderr 原样带回去，方便排查安装失败的真实原因。
    payload = {
        "action": "install",
        "install_command": install_command,
        "requested_extras": args.extras,
        "returncode": completed.returncode,
        "stdout": completed.stdout.strip(),
        "stderr": completed.stderr.strip(),
        "doctor_summary": doctor_payload["summary"],
    }
    return CommandResult(
        exit_code=completed.returncode,
        payload=payload,
        message="Bootstrap finished." if completed.returncode == 0 else "Bootstrap finished with errors.",
    )


def render_template(template_path: Path, context: dict[str, str]) -> str:
    # 用户工程初始化时，大部分文本文件都通过模板渲染生成。
    template = Template(template_path.read_text(encoding="utf-8"))
    return template.safe_substitute(context)


def ensure_clean_target(target: Path) -> None:
    # init 不应覆盖已有非空目录，否则非常容易误伤用户自己的文件。
    if target.exists() and any(target.iterdir()):
        raise FileExistsError(f"Target directory already exists and is not empty: {target}")


def ensure_directory(path: Path) -> None:
    # 原始资料目录与工作区子目录都统一走这里，避免各处重复 mkdir 参数。
    path.mkdir(parents=True, exist_ok=True)


def baseline_git_paths(target: Path) -> list[str]:
    # 基线提交只纳入 MyAgentWiki 自己生成的骨架与状态文件；
    # 外部 raw/ 不属于工作区仓库，基线里也不应试图追踪它。
    candidates = [
        ".gitignore",
        "AGENTS.md",
        "CLAUDE.md",
        "config/project.yml",
        "config/runtime_manifest.yml",
        "indexes/aliases.json",
        str(SEARCH_PAGES_INDEX_REL_PATH),
        "reports/lint/lint_latest.md",
        "state/claims.jsonl",
        "state/chunks.jsonl",
        str(EVIDENCE_BLOCKS_REL_PATH),
        "state/error_log.jsonl",
        "state/ingest_state.jsonl",
        str(KNOWLEDGE_UNITS_REL_PATH),
        "state/normalized.jsonl",
        "state/pages.jsonl",
        "state/reviews.jsonl",
        str(SEMANTIC_DECISIONS_REL_PATH),
        "state/sources.jsonl",
        str(STRUCTURE_BLOCKS_REL_PATH),
        "wiki/index.md",
        "wiki/log.md",
    ]
    return [path for path in candidates if (target / path).exists()]


def git_init_and_commit(target: Path) -> list[str]:
    # 初始化工作区时自动建一个 Git 基线，方便后续所有自动化改动都可回滚。
    steps: list[str] = []
    # 这里故意拆成三步记录，后面如果你要把这些步骤展示到日志或界面，会更直观。
    subprocess.run(["git", "init"], cwd=target, check=True, capture_output=True, text=True)
    steps.append("git init")
    tracked_paths = baseline_git_paths(target)
    subprocess.run(
        ["git", "add", "--", *tracked_paths],
        cwd=target,
        check=True,
        capture_output=True,
        text=True,
    )
    steps.append("git add whitelist")
    # 用固定的本地身份提交第一次基线，避免依赖用户电脑上是否已经配置 git 用户名邮箱。
    subprocess.run(
        [
            "git",
            "-c",
            "user.name=MyAgentWiki",
            "-c",
            "user.email=myagentwiki@local",
            "commit",
            "-m",
            "init: bootstrap MyAgentWiki workspace",
        ],
        cwd=target,
        check=True,
        capture_output=True,
        text=True,
    )
    steps.append("git commit")
    return steps


def utc_now_iso() -> str:
    # 统一使用 UTC 时间戳，避免不同机器和时区写出来的状态难对齐。
    return datetime.now(timezone.utc).isoformat()


def atomic_write_text(path: Path, text: str, encoding: str = "utf-8") -> None:
    # 关键状态文件和页面文件统一走原子写：
    # 先写同目录临时文件，再 replace 到目标路径，尽量避免留下半截文件。
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        "w",
        encoding=encoding,
        dir=path.parent,
        prefix=f".{path.name}.tmp-",
        delete=False,
    ) as fh:
        fh.write(text)
        temp_path = Path(fh.name)
    temp_path.replace(path)


def write_jsonl(path: Path, records: list[dict]) -> None:
    # 初始化占位文件时统一覆盖写入，保证 JSONL 文件总是处于可读状态。
    # JSONL 的优点是“每行一个 JSON 对象”，增量追加和排查问题都很方便。
    lines = [json.dumps(record, ensure_ascii=False) for record in records]
    text = "\n".join(lines)
    if lines:
        # 末尾补一个换行，后续 append 时不用担心和上一条黏在一起。
        text += "\n"
    atomic_write_text(path, text, encoding="utf-8")


def append_jsonl(path: Path, record: dict) -> None:
    # 运行中的状态和来源登记采用 append-only，尽量减少覆盖写坏文件的风险。
    with path.open("a", encoding="utf-8") as fh:
        fh.write(json.dumps(record, ensure_ascii=False))
        fh.write("\n")


def load_jsonl(path: Path) -> list[dict]:
    # JSONL 读取保持极简：逐行解析，后续状态恢复和 lint 都会复用。
    if not path.exists():
        return []
    records: list[dict] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        # 每一行都应该是一个独立 JSON 对象，所以这里不需要做复杂的流式解析。
        records.append(json.loads(line))
    return records


def load_json(path: Path) -> dict:
    # claim / review 单文件读取统一走这里，避免各处重复写编码与 JSON 解析逻辑。
    return json.loads(path.read_text(encoding="utf-8"))


def file_sha256(path: Path) -> str:
    # 文件去重和 source_id 稳定性都依赖内容哈希，因此统一走分块读取。
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def sanitize_source_key(value: str) -> str:
    # 把路径名压成适合放进 source_id 的安全片段。
    # 例如 "topic-a/Note 01" 会变成 "topic_a_note_01"。
    cleaned = []
    for char in value.lower():
        if char.isalnum():
            cleaned.append(char)
        else:
            cleaned.append("_")
    compact = "".join(cleaned)
    # 连续下划线压缩一下，避免生成过于难读的 ID。
    while "__" in compact:
        compact = compact.replace("__", "_")
    return compact.strip("_") or "source"


def truncate_utf8_text(value: str, max_bytes: int) -> str:
    if max_bytes <= 0:
        return ""
    encoded = value.encode("utf-8")
    if len(encoded) <= max_bytes:
        return value
    truncated = encoded[:max_bytes]
    while truncated:
        try:
            return truncated.decode("utf-8").rstrip(" ._-")
        except UnicodeDecodeError:
            truncated = truncated[:-1]
    return ""


def stabilize_filename_component(
    value: str,
    *,
    max_bytes: int = MAX_FILENAME_COMPONENT_BYTES,
    separator: str = "__",
) -> str:
    cleaned = value.strip(" .")
    if not cleaned:
        return ""
    if len(cleaned.encode("utf-8")) <= max_bytes:
        return cleaned

    digest = hashlib.sha256(cleaned.encode("utf-8")).hexdigest()[:FILENAME_HASH_LENGTH]
    suffix = f"{separator}{digest}"
    prefix_budget = max(max_bytes - len(suffix.encode("utf-8")), 1)
    prefix = truncate_utf8_text(cleaned, prefix_budget)
    if not prefix:
        return digest
    return f"{prefix}{suffix}"


def build_source_id(raw_root: Path, file_path: Path, source_hash: str) -> str:
    # source_id 不能只看文件名 stem，否则 raw 子目录里同名文件会互相挤占命名空间。
    relative = file_path.relative_to(raw_root).with_suffix("")
    relative_key = sanitize_source_key(relative.as_posix())
    return f"src_{relative_key}_{source_hash[:12]}"


def build_source_version_group(raw_root: Path, file_path: Path) -> str:
    # version_group 用来表达“同一路径来源的多次版本演进”。
    # 它不跟随内容 hash 变化，便于后面做路径级更新和版本链追踪。
    relative = file_path.relative_to(raw_root).with_suffix("")
    relative_key = sanitize_source_key(relative.as_posix())
    return f"vgrp_{relative_key}"


def source_path_to_raw_relative(source_path: str) -> str:
    # source_path 可能是 raw/topic/a.md，也可能是 ../raw/topic/a.md；
    # 这里统一裁成“相对 raw 根目录的路径”，供 version_group 与展示逻辑复用。
    path = Path(source_path)
    parts = list(path.parts)
    if "raw" in parts:
        raw_index = parts.index("raw")
        relative = Path(*parts[raw_index + 1:]) if raw_index + 1 < len(parts) else Path()
    else:
        relative = path
    return relative.with_suffix("").as_posix().lstrip("./")


def build_source_version_group_from_source_path(source_path: str) -> str:
    # sources.jsonl 里保存的是可回到 raw 的路径，例如 raw/topic/a.md 或 ../raw/topic/a.md。
    # 这里补一个从已存记录反推 version_group 的帮助函数。
    raw_relative = source_path_to_raw_relative(source_path)
    relative_key = sanitize_source_key(raw_relative)
    return f"vgrp_{relative_key}"


def build_latest_source_record_by_path(records: list[dict]) -> dict[str, dict]:
    # 同一路径可能被重复导入，这里统一选“最近导入”的那条。
    latest_by_path: dict[str, dict] = {}
    for record in records:
        source_path = record.get("source_path")
        if not source_path:
            continue
        current = latest_by_path.get(source_path)
        if current is None or record.get("imported_at", "") >= current.get("imported_at", ""):
            latest_by_path[source_path] = record
    return latest_by_path


def path_is_within_root(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def collect_files(root: Path) -> list[Path]:
    # 递归遍历 raw 下所有文件，允许用户按主题、来源、年份自由分子目录管理原始资料。
    return sorted(
        path
        for path in root.rglob("*")
        if (
            path.is_file()
            and path_is_within_root(path, root)
            and not any(part.startswith(".") for part in path.relative_to(root).parts)
        )
    )


def infer_source_type(path: Path) -> str:
    # 这里只做最基础的后缀判断，后面可以再升级成 MIME / 魔数检测。
    suffix = path.suffix.lower()
    mapping = {
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "plain_text",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".xlsx": "spreadsheet",
        ".xls": "spreadsheet",
        ".csv": "spreadsheet",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".webp": "image",
    }
    return mapping.get(suffix, "unknown")


def load_workspace_config(target: Path) -> dict:
    # 工作区自己的配置放在 config/project.yml，后续 chunk/query 也会继续依赖它。
    config_path = target / "config" / "project.yml"
    return load_simple_yaml(config_path)


def workspace_schema_version_rank(schema_version: str | None) -> int | None:
    normalized = str(schema_version or "").strip()
    if not normalized:
        return None
    try:
        return WORKSPACE_SCHEMA_VERSION_ORDER.index(normalized)
    except ValueError:
        return None


def workspace_schema_guard_status(schema_version: str | None) -> str:
    normalized = str(schema_version or "").strip() or None
    if normalized is None:
        return "missing_schema_version"

    current_rank = workspace_schema_version_rank(normalized)
    minimum_rank = workspace_schema_version_rank(WORKSPACE_MIN_SUPPORTED_SCHEMA_VERSION)
    expected_rank = workspace_schema_version_rank(WORKSPACE_SCHEMA_VERSION)
    if current_rank is None or minimum_rank is None or expected_rank is None:
        return "unsupported"
    if minimum_rank <= current_rank <= expected_rank:
        return "supported"
    return "unsupported"


def normalize_optional_cli_string(value) -> str | None:
    if value is None:
        return None
    normalized = str(value).strip()
    if not normalized or normalized.lower() == "none":
        return None
    return normalized


def update_workspace_schema_version(target: Path, schema_version: str) -> None:
    config_path = target / "config" / "project.yml"
    if not config_path.exists():
        raise ValueError(f"Missing config/project.yml in {target}.")
    config_text = config_path.read_text(encoding="utf-8")
    updated_text, replacement_count = re.subn(
        r'(^\s*schema_version:\s*")[^"]*(".*$)',
        rf'\g<1>{schema_version}\2',
        config_text,
        count=1,
        flags=re.MULTILINE,
    )
    if replacement_count == 0:
        raise ValueError("Could not update workspace.schema_version in config/project.yml.")
    atomic_write_text(config_path, updated_text, encoding="utf-8")


def workspace_schema_guard_payload(target: Path) -> dict:
    config_path = target / "config" / "project.yml"
    if not config_path.exists():
        return {
            "status": "missing_config",
            "workspace_schema_version": None,
            "expected_schema_version": WORKSPACE_SCHEMA_VERSION,
            "minimum_supported_schema_version": WORKSPACE_MIN_SUPPORTED_SCHEMA_VERSION,
        }
    config = load_simple_yaml(config_path)
    workspace_block = config.get("workspace", {})
    if not isinstance(workspace_block, dict):
        workspace_block = {}
    schema_version = str(workspace_block.get("schema_version", "")).strip() or None
    status = workspace_schema_guard_status(schema_version)
    return {
        "status": status,
        "workspace_schema_version": schema_version,
        "expected_schema_version": WORKSPACE_SCHEMA_VERSION,
        "minimum_supported_schema_version": WORKSPACE_MIN_SUPPORTED_SCHEMA_VERSION,
    }


def ensure_workspace_schema_supported(target: Path) -> None:
    payload = workspace_schema_guard_payload(target)
    if payload["status"] == "supported":
        return
    if payload["status"] == "missing_config":
        raise ValueError(
            f"Workspace schema guard failed: missing config/project.yml in {target}. "
            "Re-initialize the workspace or point the command at a valid workspace."
        )
    if payload["status"] == "missing_schema_version":
        raise ValueError(
            "Workspace schema guard failed: workspace.schema_version is missing in config/project.yml. "
            "Re-initialize the workspace or update config/project.yml to the current scaffold."
        )
    raise ValueError(
        "Workspace schema guard failed: "
        f"workspace.schema_version={payload['workspace_schema_version']} is not supported by this CLI "
        f"(supported={payload['minimum_supported_schema_version']}..{payload['expected_schema_version']}). "
        "Re-initialize the workspace with the current CLI before attempting mutating commands."
    )


def coerce_int(value, default: int) -> int:
    if isinstance(value, int):
        return value
    if isinstance(value, str):
        digits = value.strip()
        if digits.isdigit():
            return int(digits)
    return default


def coerce_float(value, default: float) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        stripped = value.strip()
        try:
            return float(stripped)
        except ValueError:
            return default
    return default


def normalize_command_config(value) -> list[str]:
    if isinstance(value, str):
        stripped = value.strip()
        return shlex.split(stripped) if stripped else []
    if isinstance(value, list):
        normalized = []
        for item in value:
            text = str(item).strip()
            if text:
                normalized.append(text)
        return normalized
    return []


def load_automation_target_config(config: dict, target_name: str) -> dict:
    automation_config = config.get("automation", {})
    if not isinstance(automation_config, dict):
        automation_config = {}

    target_config = automation_config.get(target_name, {})
    if not isinstance(target_config, dict):
        target_config = {}

    inherited_strategy = str(automation_config.get("mode", "safe_auto")).strip() or "safe_auto"
    strategy = str(target_config.get("strategy", inherited_strategy)).strip() or inherited_strategy
    if strategy not in AUTOMATION_STRATEGIES:
        strategy = "safe_auto"

    command = normalize_command_config(target_config.get("command", []))
    timeout_seconds = max(coerce_int(target_config.get("timeout_seconds", 45), 45), 5)
    min_confidence = min(max(coerce_float(target_config.get("min_confidence", 0.8), 0.8), 0.0), 1.0)
    return {
        "strategy": strategy,
        "command": command,
        "timeout_seconds": timeout_seconds,
        "min_confidence": min_confidence,
        "enabled": strategy == "agent_assisted" and bool(command),
    }


def load_post_ingest_review_auto_config(config: dict) -> dict:
    automation_config = config.get("automation", {})
    if not isinstance(automation_config, dict):
        automation_config = {}

    post_ingest_config = automation_config.get("post_ingest", {})
    if not isinstance(post_ingest_config, dict):
        post_ingest_config = {}

    review_auto_enabled = post_ingest_config.get("review_auto", True)
    return {
        "review_auto": bool(review_auto_enabled),
    }


def load_semantic_task_config(config: dict, task_name: str) -> SemanticTaskConfig:
    semantic_config = config.get("semantic", {})
    if not isinstance(semantic_config, dict):
        semantic_config = {}

    scheduler_config = semantic_config.get("batch_scheduler", {})
    if not isinstance(scheduler_config, dict):
        scheduler_config = {}

    task_config = semantic_config.get(task_name, {})
    if not isinstance(task_config, dict):
        task_config = {}

    strategy = str(task_config.get("strategy", "agent_assisted")).strip() or "agent_assisted"
    if strategy not in AUTOMATION_STRATEGIES:
        strategy = "agent_assisted"

    command = normalize_command_config(task_config.get("command", []))
    timeout_seconds = max(coerce_int(task_config.get("timeout_seconds", 45), 45), 5)
    min_confidence = min(max(coerce_float(task_config.get("min_confidence", 0.75), 0.75), 0.0), 1.0)
    batch_size = max(
        coerce_int(
            task_config.get("batch_size", scheduler_config.get("default_batch_size", 12)),
            12,
        ),
        1,
    )
    model_key = str(task_config.get("model_key", "local-default")).strip() or "local-default"
    prompt_version = str(task_config.get("prompt_version", "v1")).strip() or "v1"
    schema_version = str(task_config.get("schema_version", "v1")).strip() or "v1"
    enabled = strategy == "agent_assisted" and bool(command)
    return SemanticTaskConfig(
        task_name=task_name,
        strategy=strategy,
        command=command,
        timeout_seconds=timeout_seconds,
        min_confidence=min_confidence,
        batch_size=batch_size,
        enabled=enabled,
        model_key=model_key,
        prompt_version=prompt_version,
        schema_version=schema_version,
    )


def run_json_automation_command(
    target: Path,
    command: list[str],
    payload: dict,
    timeout_seconds: int,
) -> dict | None:
    if not command:
        return None
    try:
        completed = subprocess.run(
            command,
            cwd=target,
            input=json.dumps(payload, ensure_ascii=False),
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None

    if completed.returncode != 0:
        return None

    stdout = (completed.stdout or "").strip()
    if not stdout:
        return None

    try:
        parsed = json.loads(stdout)
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def semantic_structure_records_by_id(target: Path) -> tuple[dict[str, dict], dict[str, dict]]:
    evidence_blocks = {
        str(record.get("evidence_block_id", "")).strip(): record
        for record in load_jsonl(target / EVIDENCE_BLOCKS_REL_PATH)
        if str(record.get("evidence_block_id", "")).strip()
    }
    knowledge_units = {
        str(record.get("knowledge_unit_id", "")).strip(): record
        for record in load_jsonl(target / KNOWLEDGE_UNITS_REL_PATH)
        if str(record.get("knowledge_unit_id", "")).strip()
    }
    return evidence_blocks, knowledge_units


def sorted_counter_dict(counter: Counter[str], limit: int = 12) -> dict[str, int]:
    return {
        key: count
        for key, count in sorted(counter.items(), key=lambda item: (-item[1], item[0]))[:limit]
        if key
    }


def first_non_empty_string(values: list[object]) -> str:
    for value in values:
        text = str(value or "").strip()
        if text:
            return text
    return ""


def claim_structure_context(
    claim_record: dict,
    evidence_blocks_by_id: dict[str, dict],
    knowledge_units_by_id: dict[str, dict],
) -> dict:
    knowledge_unit_ids = normalize_string_list(claim_record.get("knowledge_unit_ids"))
    evidence_block_ids = normalize_string_list(claim_record.get("evidence_block_ids"))
    source_refs = claim_record.get("source_refs", [])
    if not isinstance(source_refs, list):
        source_refs = []

    for source_ref in source_refs:
        if not isinstance(source_ref, dict):
            continue
        for knowledge_unit_id in normalize_string_list(source_ref.get("knowledge_unit_id")):
            if knowledge_unit_id not in knowledge_unit_ids:
                knowledge_unit_ids.append(knowledge_unit_id)
        for evidence_block_id in normalize_string_list(source_ref.get("evidence_block_ids")):
            if evidence_block_id not in evidence_block_ids:
                evidence_block_ids.append(evidence_block_id)

    knowledge_units = [
        knowledge_units_by_id[unit_id]
        for unit_id in knowledge_unit_ids
        if unit_id in knowledge_units_by_id
    ]
    evidence_blocks = [
        evidence_blocks_by_id[evidence_id]
        for evidence_id in evidence_block_ids
        if evidence_id in evidence_blocks_by_id
    ]

    section_path_parts: list[str] = []
    section_title = ""
    parent_section_path = ""
    heading_level = 0
    for source_ref in source_refs:
        if not isinstance(source_ref, dict):
            continue
        if not section_path_parts:
            section_path_parts = normalize_string_list(source_ref.get("section_path_parts"))
        if not section_title:
            section_title = str(source_ref.get("section_title", "")).strip()
        if not parent_section_path:
            parent_section_path = str(source_ref.get("parent_section_path", "")).strip()
        if not heading_level:
            heading_level = coerce_int(source_ref.get("heading_level", 0), 0)

    if not section_path_parts:
        for evidence_block in evidence_blocks:
            section_path_parts = normalize_string_list(evidence_block.get("section_path_parts"))
            if section_path_parts:
                break
    if not section_title and section_path_parts:
        section_title = section_path_parts[-1]
    if not parent_section_path and len(section_path_parts) > 1:
        parent_section_path = " > ".join(section_path_parts[:-1])
    if not heading_level and section_path_parts:
        heading_level = len(section_path_parts)

    content_tag_counter: Counter[str] = Counter()
    semantic_feature_counter: Counter[str] = Counter()
    semantic_feature_strength_counter: Counter[str] = Counter()
    unit_kind_counter: Counter[str] = Counter()
    evidence_kind_counter: Counter[str] = Counter()
    metadata_key_counter: Counter[str] = Counter()
    local_headings: list[str] = []
    seen_content_tag_sources: set[tuple[str, str]] = set()

    for knowledge_unit in knowledge_units:
        unit_kind = str(knowledge_unit.get("unit_kind", "")).strip()
        if unit_kind:
            unit_kind_counter[unit_kind] += 1
        local_heading = str(knowledge_unit.get("local_heading", "") or "").strip()
        if local_heading and local_heading not in local_headings:
            local_headings.append(local_heading)
        metadata = knowledge_unit.get("metadata", {})
        if isinstance(metadata, dict):
            for key in metadata:
                if str(key).strip():
                    metadata_key_counter[str(key).strip()] += 1
        projection = knowledge_unit.get("semantic_projection", {})
        if isinstance(projection, dict):
            source_key = first_non_empty_string(
                normalize_string_list(knowledge_unit.get("evidence_block_ids"))
                + [knowledge_unit.get("knowledge_unit_id", "")]
            )
            for tag in normalize_string_list(projection.get("content_tags")):
                tag_key = (source_key or f"knowledge_unit:{len(seen_content_tag_sources)}", tag)
                if tag_key not in seen_content_tag_sources:
                    seen_content_tag_sources.add(tag_key)
                    content_tag_counter[tag] += 1
            for feature in projection.get("semantic_features", []) or []:
                if not isinstance(feature, dict):
                    continue
                tag = str(feature.get("tag", "")).strip()
                strength = str(feature.get("strength", "")).strip()
                if tag:
                    semantic_feature_counter[tag] += 1
                if strength:
                    semantic_feature_strength_counter[strength] += 1

    for evidence_block in evidence_blocks:
        block_kind = str(evidence_block.get("block_kind", "")).strip()
        if block_kind:
            evidence_kind_counter[block_kind] += 1
        local_heading = str(evidence_block.get("local_heading", "") or "").strip()
        if local_heading and local_heading not in local_headings:
            local_headings.append(local_heading)
        metadata = evidence_block.get("metadata", {})
        if isinstance(metadata, dict):
            for key in metadata:
                if str(key).strip():
                    metadata_key_counter[str(key).strip()] += 1
        source_key = str(evidence_block.get("evidence_block_id", "")).strip()
        for tag in normalize_string_list(evidence_block.get("content_tags")):
            tag_key = (source_key or f"evidence_block:{len(seen_content_tag_sources)}", tag)
            if tag_key not in seen_content_tag_sources:
                seen_content_tag_sources.add(tag_key)
                content_tag_counter[tag] += 1
        for feature in evidence_block.get("semantic_features", []) or []:
            if not isinstance(feature, dict):
                continue
            tag = str(feature.get("tag", "")).strip()
            strength = str(feature.get("strength", "")).strip()
            if tag:
                semantic_feature_counter[tag] += 1
            if strength:
                semantic_feature_strength_counter[strength] += 1

    return {
        "section_path_parts": section_path_parts,
        "section_title": section_title,
        "parent_section_path": parent_section_path,
        "heading_level": heading_level,
        "local_headings": local_headings[:5],
        "unit_kind_counts": sorted_counter_dict(unit_kind_counter),
        "evidence_block_kind_counts": sorted_counter_dict(evidence_kind_counter),
        "content_tag_counts": sorted_counter_dict(content_tag_counter),
        "semantic_feature_counts": sorted_counter_dict(semantic_feature_counter),
        "semantic_feature_strength_counts": sorted_counter_dict(semantic_feature_strength_counter),
        "metadata_key_counts": sorted_counter_dict(metadata_key_counter),
        "source_ref_count": len(source_refs),
        "knowledge_unit_ids": knowledge_unit_ids[:8],
        "evidence_block_ids": evidence_block_ids[:12],
    }


def page_intent_group_context(grouped_claims: list[dict]) -> dict:
    role_counter: Counter[str] = Counter()
    hint_counter: Counter[str] = Counter()
    content_tag_counter: Counter[str] = Counter()
    unit_kind_counter: Counter[str] = Counter()
    evidence_kind_counter: Counter[str] = Counter()
    semantic_feature_counter: Counter[str] = Counter()
    section_counter: Counter[str] = Counter()
    local_headings: list[str] = []

    for claim_record in grouped_claims:
        role = claim_knowledge_role(claim_record)
        if role:
            role_counter[role] += 1
        hint_counter.update(claim_page_intent_hints(claim_record))

        context = claim_record.get("structure_context", {})
        if not isinstance(context, dict):
            context = {}
        content_tag_counter.update(dict(context.get("content_tag_counts", {}) or {}))
        unit_kind_counter.update(dict(context.get("unit_kind_counts", {}) or {}))
        evidence_kind_counter.update(dict(context.get("evidence_block_kind_counts", {}) or {}))
        semantic_feature_counter.update(dict(context.get("semantic_feature_counts", {}) or {}))
        section_path = " > ".join(normalize_string_list(context.get("section_path_parts")))
        if section_path:
            section_counter[section_path] += 1
        for heading in normalize_string_list(context.get("local_headings")):
            if heading not in local_headings:
                local_headings.append(heading)

    return {
        "knowledge_role_counts": sorted_counter_dict(role_counter),
        "page_intent_hint_counts": sorted_counter_dict(hint_counter),
        "content_tag_counts": sorted_counter_dict(content_tag_counter),
        "unit_kind_counts": sorted_counter_dict(unit_kind_counter),
        "evidence_block_kind_counts": sorted_counter_dict(evidence_kind_counter),
        "semantic_feature_counts": sorted_counter_dict(semantic_feature_counter),
        "section_path_counts": sorted_counter_dict(section_counter),
        "representative_local_headings": local_headings[:8],
    }


def collect_semantic_task_items(target: Path, task_name: str) -> list[dict]:
    if task_name == "document_analysis":
        records = load_jsonl(target / "state" / "normalized.jsonl")
        items = []
        for record in records:
            source_id = str(record.get("source_id", "")).strip()
            normalized_path = str(record.get("normalized_path", "")).strip()
            if not source_id or not normalized_path:
                continue
            items.append(
                {
                    "item_id": source_id,
                    "source_id": source_id,
                    "normalized_path": normalized_path,
                    "title": record.get("title", ""),
                    "extraction_quality": record.get("extraction_quality"),
                }
            )
        return items

    if task_name == "claim_candidate_quality":
        records = load_jsonl(target / "state" / "claims.jsonl")
        items = []
        for record in records:
            claim_id = str(record.get("claim_id", "")).strip()
            text = str(record.get("text", "")).strip()
            if (
                not claim_id
                or not text
                or record.get("lifecycle_status", "active") != "active"
                or not claim_candidate_has_short_gray_zone(text)
            ):
                continue
            cleaned_text = clean_claim_candidate_text(text)
            natural_char_count = len([
                char for char in cleaned_text
                if char.isalnum() or "\u4e00" <= char <= "\u9fff"
            ])
            items.append(
                {
                    "item_id": claim_id,
                    "claim_id": claim_id,
                    "text": text,
                    "cleaned_text": cleaned_text,
                    "claim_type": record.get("claim_type"),
                    "natural_char_count": natural_char_count,
                    "source_ids": record.get("source_ids", []),
                    "source_refs": record.get("source_refs", []),
                }
            )
        return items

    if task_name == "claim_role":
        records = load_jsonl(target / "state" / "claims.jsonl")
        evidence_blocks_by_id, knowledge_units_by_id = semantic_structure_records_by_id(target)
        items = []
        for record in records:
            claim_id = str(record.get("claim_id", "")).strip()
            if not claim_id or record.get("lifecycle_status", "active") != "active":
                continue
            items.append(
                {
                    "item_id": claim_id,
                    "claim_id": claim_id,
                    "text": record.get("text", ""),
                    "claim_type": record.get("claim_type"),
                    "quality_label": record.get("quality_label"),
                    "quality_reason": record.get("quality_reason"),
                    "quality_safe_auto_ready": record.get("quality_safe_auto_ready"),
                    "source_ids": record.get("source_ids", []),
                    "source_refs": record.get("source_refs", []),
                    "structure_context": claim_structure_context(
                        record,
                        evidence_blocks_by_id=evidence_blocks_by_id,
                        knowledge_units_by_id=knowledge_units_by_id,
                    ),
                }
            )
        return items

    if task_name == "page_intent":
        records = load_jsonl(target / "state" / "claims.jsonl")
        evidence_blocks_by_id, knowledge_units_by_id = semantic_structure_records_by_id(target)
        groups: dict[str, list[dict]] = {}
        for record in records:
            if record.get("lifecycle_status", "active") != "active":
                continue
            bucket_key = build_concept_group_key(record)
            if not bucket_key:
                continue
            enriched_record = dict(record)
            enriched_record["structure_context"] = claim_structure_context(
                record,
                evidence_blocks_by_id=evidence_blocks_by_id,
                knowledge_units_by_id=knowledge_units_by_id,
            )
            groups.setdefault(bucket_key, []).append(enriched_record)

        items = []
        for bucket_key, grouped_claims in sorted(groups.items()):
            items.append(build_page_intent_item_payload(bucket_key, grouped_claims))
        return items

    raise KeyError(f"Unsupported semantic task: {task_name}")


def chunk_semantic_items(items: list[dict], batch_size: int) -> list[list[dict]]:
    if batch_size <= 0:
        batch_size = 1
    return [items[index:index + batch_size] for index in range(0, len(items), batch_size)]


def normalize_semantic_batch_results(
    task_name: str,
    hook_result: dict,
    batch_items: list[dict],
    config: SemanticTaskConfig,
) -> tuple[list[dict], list[dict]]:
    decisions = hook_result.get("decisions", [])
    if not isinstance(decisions, list):
        return [], []

    item_map = {str(item.get("item_id")): item for item in batch_items if item.get("item_id")}
    normalized: list[dict] = []
    skipped: list[dict] = []
    for decision in decisions:
        if not isinstance(decision, dict):
            continue
        item_id = str(decision.get("item_id", "")).strip()
        if not item_id or item_id not in item_map:
            continue
        confidence = coerce_float(decision.get("confidence", 0.0), 0.0)
        if confidence < config.min_confidence:
            skipped.append({
                "item_id": item_id,
                "decision_status": "rejected",
                "reason_code": str(decision.get("reason_code", "")).strip() or "semantic_batch_low_confidence",
                "confidence": confidence,
                "risk_flags": ["semantic_decision_low_confidence"],
            })
            continue
        normalized_decision = normalize_semantic_hook_decision(task_name, decision)
        if normalized_decision["decision_status"] != "accepted":
            skipped.append({
                "item_id": item_id,
                "decision_status": normalized_decision["decision_status"],
                "reason_code": str(decision.get("reason_code", "")).strip() or "semantic_batch_result",
                "confidence": confidence,
                "risk_flags": normalized_decision["risk_flags"],
                "abstain_reason": normalized_decision["abstain_reason"],
                "missing_fields": normalized_decision["missing_fields"],
            })
            continue
        item_payload = item_map[item_id]
        input_fingerprint = fingerprint_payload(
            task_name=task_name,
            item_payloads=[item_payload],
            prompt_version=config.prompt_version,
            schema_version=config.schema_version,
        )
        normalized.append(
            {
                "decision_id": build_semantic_decision_id(task_name, input_fingerprint),
                "task_type": task_name,
                "item_type": item_type_for_task(task_name),
                "item_ids": [item_id],
                "decision": normalized_decision["decision"],
                "decision_status": normalized_decision["decision_status"],
                "confidence": confidence,
                "reason_code": str(decision.get("reason_code", "")).strip() or "semantic_batch_result",
                "risk_flags": normalized_decision["risk_flags"],
                "supporting_ids": normalized_decision["supporting_ids"],
                "abstain_reason": normalized_decision["abstain_reason"],
                "prompt_version": config.prompt_version,
                "model_key": config.model_key,
                "schema_version": config.schema_version,
                "input_fingerprint": input_fingerprint,
                "created_at": utc_now_iso(),
                "superseded_by": [],
            }
        )
    return normalized, skipped


def run_semantic_batch_task(
    target: Path,
    task_name: str,
    dry_run: bool = False,
) -> dict:
    config = load_semantic_task_config(load_workspace_config(target), task_name)
    items = collect_semantic_task_items(target, task_name)
    existing_records = load_semantic_decisions(target)
    existing_by_fingerprint = build_latest_semantic_decisions_by_fingerprint(existing_records)

    cache_hits = 0
    pending_batches: list[tuple[list[dict], list[str]]] = []
    for batch_items in chunk_semantic_items(items, config.batch_size):
        pending_items = []
        cached_ids = []
        for item in batch_items:
            input_fingerprint = fingerprint_payload(
                task_name=task_name,
                item_payloads=[item],
                prompt_version=config.prompt_version,
                schema_version=config.schema_version,
            )
            if input_fingerprint in existing_by_fingerprint:
                cache_hits += 1
                cached_ids.append(str(item.get("item_id")))
            else:
                pending_items.append(item)
        if pending_items:
            pending_batches.append((pending_items, cached_ids))

    written_decisions: list[dict] = []
    batch_reports = []
    ensure_directory(semantic_batches_dir(target))

    for batch_index, (batch_items, cached_ids) in enumerate(pending_batches, start=1):
        payload = {
            "task": f"review_{task_name}_batch",
            "task_name": task_name,
            "prompt_version": config.prompt_version,
            "schema_version": config.schema_version,
            "items": batch_items,
        }
        hook_result = run_json_automation_command(
            target=target,
            command=config.command,
            payload=payload,
            timeout_seconds=config.timeout_seconds,
        ) if config.enabled else None
        normalized_results, skipped_results = normalize_semantic_batch_results(task_name, hook_result or {}, batch_items, config)

        batch_report = {
            "task_name": task_name,
            "batch_index": batch_index,
            "item_ids": [str(item.get("item_id")) for item in batch_items],
            "cached_item_ids": cached_ids,
            "decision_count": len(normalized_results),
            "skipped_decision_count": len(skipped_results),
            "skipped_decisions": skipped_results,
            "created_at": utc_now_iso(),
        }
        write_json(
            semantic_batches_dir(target) / f"{task_name}_batch_{batch_index:04d}.json",
            batch_report,
        )
        batch_reports.append(batch_report)
        written_decisions.extend(normalized_results)

    if written_decisions and not dry_run:
        for record in written_decisions:
            append_jsonl(semantic_decisions_path(target), record)

    return {
        "task_name": task_name,
        "workspace_summary": build_workspace_summary(target),
        "summary": {
            "item_count": len(items),
            "cache_hits": cache_hits,
            "pending_batch_count": len(pending_batches),
            "written_decision_count": 0 if dry_run else len(written_decisions),
            "dry_run": dry_run,
        },
        "config": {
            "strategy": config.strategy,
            "batch_size": config.batch_size,
            "model_key": config.model_key,
            "prompt_version": config.prompt_version,
            "schema_version": config.schema_version,
            "enabled": config.enabled,
        },
        "batch_reports": batch_reports,
        "decisions": written_decisions,
    }


def apply_document_analysis_decisions_to_normalized_records(
    target: Path,
    normalized_records: list[dict],
    task_config: SemanticTaskConfig,
) -> list[dict]:
    latest_decisions = build_latest_semantic_decisions_by_fingerprint(load_semantic_decisions(target))
    normalized_by_source_id = {record["source_id"]: dict(record) for record in normalized_records}
    changed_records: list[dict] = []

    for record in normalized_records:
        item_payload = {
            "item_id": record["source_id"],
            "source_id": record["source_id"],
            "normalized_path": record.get("normalized_path", ""),
            "title": record.get("title", ""),
            "extraction_quality": record.get("extraction_quality"),
        }
        fingerprint = fingerprint_payload(
            task_name="document_analysis",
            item_payloads=[item_payload],
            prompt_version=task_config.prompt_version,
            schema_version=task_config.schema_version,
        )
        decision_record = latest_decisions.get(fingerprint)
        if decision_record is None:
            continue
        decision = decision_record.get("decision", {})
        if not isinstance(decision, dict):
            continue

        updated_record = dict(normalized_by_source_id[record["source_id"]])
        updated_record["document_kind"] = decision.get("document_kind", updated_record.get("document_kind", "note"))
        updated_record["structure_quality"] = decision.get("structure_quality", updated_record.get("structure_quality", "unknown"))
        updated_record["chunk_strategy_hint"] = decision.get("chunk_strategy_hint", updated_record.get("chunk_strategy_hint", "heading_first"))
        normalized_by_source_id[record["source_id"]] = updated_record
        changed_records.append(updated_record)

    if changed_records:
        ordered_records = []
        for record in normalized_records:
            ordered_records.append(normalized_by_source_id[record["source_id"]])
        write_jsonl(target / "state" / "normalized.jsonl", ordered_records)
        return ordered_records
    return normalized_records


def apply_claim_role_decisions_to_claim_records(
    target: Path,
    claim_records: list[dict],
    task_config: SemanticTaskConfig,
) -> list[dict]:
    latest_decisions = build_latest_semantic_decisions_by_fingerprint(load_semantic_decisions(target))
    claims_by_id = {record["claim_id"]: dict(record) for record in claim_records}
    evidence_blocks_by_id, knowledge_units_by_id = semantic_structure_records_by_id(target)
    changed = False

    for record in claim_records:
        item_payload = {
            "item_id": record["claim_id"],
            "claim_id": record["claim_id"],
            "text": record.get("text", ""),
            "claim_type": record.get("claim_type"),
            "quality_label": record.get("quality_label"),
            "quality_reason": record.get("quality_reason"),
            "quality_safe_auto_ready": record.get("quality_safe_auto_ready"),
            "source_ids": record.get("source_ids", []),
            "source_refs": record.get("source_refs", []),
            "structure_context": claim_structure_context(
                record,
                evidence_blocks_by_id=evidence_blocks_by_id,
                knowledge_units_by_id=knowledge_units_by_id,
            ),
        }
        fingerprint = fingerprint_payload(
            task_name="claim_role",
            item_payloads=[item_payload],
            prompt_version=task_config.prompt_version,
            schema_version=task_config.schema_version,
        )
        decision_record = latest_decisions.get(fingerprint)
        if decision_record is None:
            continue
        decision = decision_record.get("decision", {})
        if not isinstance(decision, dict):
            continue

        updated = dict(claims_by_id[record["claim_id"]])
        updated["knowledge_role"] = decision.get("knowledge_role", updated.get("knowledge_role"))
        updated["page_intent_hints"] = list(decision.get("page_intent_hints", updated.get("page_intent_hints", [])) or [])
        updated["concept_candidate_score"] = coerce_float(
            decision.get("concept_candidate_score", updated.get("concept_candidate_score", 0.0)),
            coerce_float(updated.get("concept_candidate_score", 0.0), 0.0),
        )
        append_unique(updated.setdefault("semantic_decision_ids", []), decision_record["decision_id"])
        updated = sync_claim_semantic_projection(updated)
        updated["updated_at"] = utc_now_iso()
        claims_by_id[record["claim_id"]] = updated
        changed = True

    ordered_records = []
    for record in claim_records:
        ordered_records.append(claims_by_id[record["claim_id"]])

    if changed:
        write_jsonl(target / "state" / "claims.jsonl", ordered_records)
        for record in ordered_records:
            write_claim_file(target, record)
    return ordered_records


def apply_claim_candidate_quality_decisions_to_claim_records(
    target: Path,
    claim_records: list[dict],
    task_config: SemanticTaskConfig,
) -> tuple[list[dict], set[str], set[str]]:
    latest_decisions = build_latest_semantic_decisions_by_fingerprint(load_semantic_decisions(target))
    live_claims_by_id = {
        record["claim_id"]: dict(record)
        for record in claim_records
        if is_live_claim_record(record)
    }
    historical_claims_by_id = {
        record["claim_id"]: dict(record)
        for record in claim_records
        if not is_live_claim_record(record)
    }
    live_reviews_by_id, historical_reviews_by_id, _ = load_review_state_maps(target)
    changed = False
    archived_claim_ids: set[str] = set()
    affected_review_ids: set[str] = set()

    for record in claim_records:
        if not is_live_claim_record(record):
            continue
        if not claim_candidate_has_short_gray_zone(record.get("text", "")):
            continue
        item_payload = {
            "item_id": record["claim_id"],
            "claim_id": record["claim_id"],
            "text": record.get("text", ""),
            "cleaned_text": clean_claim_candidate_text(record.get("text", "")),
            "claim_type": record.get("claim_type"),
            "natural_char_count": len([
                char for char in clean_claim_candidate_text(record.get("text", ""))
                if char.isalnum() or "\u4e00" <= char <= "\u9fff"
            ]),
            "source_ids": record.get("source_ids", []),
            "source_refs": record.get("source_refs", []),
        }
        fingerprint = fingerprint_payload(
            task_name="claim_candidate_quality",
            item_payloads=[item_payload],
            prompt_version=task_config.prompt_version,
            schema_version=task_config.schema_version,
        )
        decision_record = latest_decisions.get(fingerprint)
        if decision_record is None:
            continue
        decision = decision_record.get("decision", {})
        if not isinstance(decision, dict):
            continue

        claim_id = record["claim_id"]
        updated = dict(live_claims_by_id[claim_id])
        updated["quality_label"] = str(decision.get("quality_label", "")).strip() or updated.get("quality_label")
        updated["quality_reason"] = str(decision.get("reason", "")).strip() or updated.get("quality_reason")
        updated["quality_confidence"] = coerce_float(
            decision_record.get("confidence", updated.get("quality_confidence", 0.0)),
            coerce_float(updated.get("quality_confidence", 0.0), 0.0),
        )
        quality_review_required = decision.get("review_required")
        if quality_review_required is not None:
            updated["quality_review_required"] = bool(quality_review_required)
        quality_safe_auto_ready = decision.get("safe_auto_ready")
        if quality_safe_auto_ready is not None:
            updated["quality_safe_auto_ready"] = bool(quality_safe_auto_ready)
        updated["quality_decision_source"] = "semantic_batch"
        append_unique(updated.setdefault("semantic_decision_ids", []), decision_record["decision_id"])
        updated = sync_claim_semantic_projection(updated)
        updated["updated_at"] = utc_now_iso()

        quality_label = str(updated.get("quality_label") or "").strip().lower()
        if quality_label in {"noise", "title_shell"}:
            archived_claim_ids.add(claim_id)
            archived_record = archive_live_claim(
                claim_record=updated,
                live_claims_by_id=live_claims_by_id,
                historical_claims_by_id=historical_claims_by_id,
            )
            affected_review_ids.update(
                purge_deleted_claims_from_reviews(
                    reviews_by_id=live_reviews_by_id,
                    historical_reviews_by_id=historical_reviews_by_id,
                    deleted_claim_ids={claim_id},
                )[0]
            )
            live_claims_by_id.pop(claim_id, None)
            historical_claims_by_id[archived_record["claim_id"]] = archived_record
            changed = True
            continue

        if updated.get("quality_review_required") and updated.get("status") == "draft":
            updated["status"] = "needs_review"
            updated["review_reason"] = "claim_quality_requires_human_review"
        live_claims_by_id[claim_id] = updated
        changed = True

    ordered_records = build_ordered_claim_state_records(
        live_claims_by_id=live_claims_by_id,
        historical_claims_by_id=historical_claims_by_id,
    )
    if changed:
        write_jsonl(target / "state" / "claims.jsonl", ordered_records)
        for record in ordered_records:
            write_claim_file(target, record)
        review_state_records = build_ordered_review_state_records(
            live_reviews_by_id=live_reviews_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )
        write_jsonl(target / "state" / "reviews.jsonl", review_state_records)
        for review_record in review_state_records:
            write_review_file(target, review_record)
        cleanup_superseded_record_files(
            target=target,
            historical_claims_by_id=historical_claims_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )
    return ordered_records, archived_claim_ids, affected_review_ids


SPECIALIZED_PAGE_INTENTS = {"guide", "example", "reference", "timeline"}
PAGE_INTENT_ROLE_SIGNALS = {
    "guide": {"procedure"},
    "example": {"example"},
}
PAGE_INTENT_CONTENT_TAG_SIGNALS = {
    "guide": {"procedural_language"},
    "example": {"cases"},
    "reference": {"rules"},
    "timeline": {"temporal_language"},
}
PAGE_INTENT_BLOCK_SIGNALS = {
    "reference": {"table_row", "metadata_line"},
    "example": {"code_example"},
}


def counter_value(counter_payload: dict, key: str) -> int:
    if not isinstance(counter_payload, dict):
        return 0
    return coerce_int(counter_payload.get(key, 0), 0)


def sum_counter_values(counter_payload: dict, keys: set[str]) -> int:
    return sum(counter_value(counter_payload, key) for key in keys)


def page_intent_signal_counts(page_intent: str, item_payload: dict) -> dict[str, int]:
    group_context = item_payload.get("group_context", {})
    if not isinstance(group_context, dict):
        group_context = {}
    return {
        "hint_count": counter_value(group_context.get("page_intent_hint_counts", {}), page_intent),
        "role_count": sum_counter_values(
            group_context.get("knowledge_role_counts", {}),
            PAGE_INTENT_ROLE_SIGNALS.get(page_intent, set()),
        ),
        "content_tag_count": sum_counter_values(
            group_context.get("content_tag_counts", {}),
            PAGE_INTENT_CONTENT_TAG_SIGNALS.get(page_intent, set()),
        ),
        "block_count": sum_counter_values(
            group_context.get("evidence_block_kind_counts", {}),
            PAGE_INTENT_BLOCK_SIGNALS.get(page_intent, set()),
        ),
    }


def downgrade_specialized_page_intent(grouped_claims: list[dict]) -> str:
    return "concept" if should_generate_concept_page(grouped_claims) else "topic"


def page_intent_has_enough_group_evidence(page_intent: str, item_payload: dict) -> bool:
    if page_intent not in SPECIALIZED_PAGE_INTENTS:
        return True
    counts = page_intent_signal_counts(page_intent, item_payload)
    multi_signal_count = counts["hint_count"] + counts["role_count"] + counts["content_tag_count"]
    return multi_signal_count >= 2 or counts["block_count"] >= 1


def validate_page_intent_candidate(
    page_intent: str,
    grouped_claims: list[dict],
    item_payload: dict,
    decision_record: dict | None,
    route_reason: str,
) -> tuple[str, str]:
    normalized_intent = str(page_intent or "").strip().lower() or "topic"
    if normalized_intent not in SPECIALIZED_PAGE_INTENTS:
        return normalized_intent, route_reason

    counts = page_intent_signal_counts(normalized_intent, item_payload)
    if page_intent_has_enough_group_evidence(normalized_intent, item_payload):
        return normalized_intent, route_reason

    risk_flags = []
    decision_content_tags: list[str] = []
    if isinstance(decision_record, dict):
        risk_flags = normalize_string_list(decision_record.get("risk_flags"))
        decision = decision_record.get("decision", {})
        if isinstance(decision, dict):
            decision_content_tags = normalize_string_list(decision.get("content_tags"))
    expected_content_tags = PAGE_INTENT_CONTENT_TAG_SIGNALS.get(normalized_intent, set())
    decision_tag_signal = bool(expected_content_tags & set(decision_content_tags))
    source_is_strong = (
        decision_record is not None
        and "strong_" in str(route_reason)
        and not any("ambiguous" in flag for flag in risk_flags)
    )
    has_any_signal = any(count > 0 for count in counts.values())
    if source_is_strong and (has_any_signal or decision_tag_signal):
        return normalized_intent, route_reason

    downgraded_intent = downgrade_specialized_page_intent(grouped_claims)
    return downgraded_intent, f"page_intent_validation_downgraded_{normalized_intent}_insufficient_group_evidence"


def choose_bucket_page_intent(grouped_claims: list[dict]) -> str:
    if not grouped_claims:
        return "reject"
    item_payload = build_page_intent_item_payload("heuristic_bucket", grouped_claims)
    hint_counts: Counter[str] = Counter()
    for claim_record in grouped_claims:
        for hint in claim_page_intent_hints(claim_record):
            normalized_hint = str(hint).strip().lower()
            if normalized_hint:
                hint_counts[normalized_hint] += 1
    for preferred in ("reject", "timeline", "reference", "guide", "example", "concept", "topic"):
        if not hint_counts.get(preferred):
            continue
        if preferred in SPECIALIZED_PAGE_INTENTS and not page_intent_has_enough_group_evidence(preferred, item_payload):
            continue
        if preferred == "reject" and hint_counts[preferred] < len(grouped_claims):
            continue
        return preferred
    return "concept" if should_generate_concept_page(grouped_claims) else "topic"


def build_page_intent_item_payload(bucket_key: str, grouped_claims: list[dict]) -> dict:
    ordered_claims = sorted(
        grouped_claims,
        key=lambda item: str(item.get("claim_id", "")).strip(),
    )
    claim_ids = [
        str(item.get("claim_id", "")).strip()
        for item in ordered_claims
        if str(item.get("claim_id", "")).strip()
    ]
    preview_texts = [
        str(item.get("text", "")).strip()
        for item in ordered_claims[:5]
        if str(item.get("text", "")).strip()
    ]
    claim_semantics = []
    for item in ordered_claims:
        claim_id = str(item.get("claim_id", "")).strip()
        if not claim_id:
            continue
        claim_semantics.append(
            {
                "claim_id": claim_id,
                "knowledge_role": claim_knowledge_role(item),
                "page_intent_hints": claim_page_intent_hints(item),
                "concept_candidate_score": claim_concept_candidate_score(item),
            }
        )
    return {
        "item_id": bucket_key,
        "bucket_key": bucket_key,
        "claim_ids": claim_ids,
        "claim_texts": preview_texts,
        "claim_count": len(claim_ids),
        "claim_semantics": claim_semantics,
        "group_context": page_intent_group_context(ordered_claims),
    }


def apply_page_intent_decisions_to_claim_groups(
    target: Path,
    concept_claim_groups: dict[str, list[dict]],
    task_config: SemanticTaskConfig,
) -> dict[str, dict]:
    latest_decisions = build_latest_semantic_decisions_by_fingerprint(load_semantic_decisions(target))
    evidence_blocks_by_id, knowledge_units_by_id = semantic_structure_records_by_id(target)
    page_routes: dict[str, dict] = {}
    new_route_decisions: list[dict] = []

    for bucket_key, grouped_claims in concept_claim_groups.items():
        enriched_grouped_claims = []
        for record in grouped_claims:
            enriched_record = dict(record)
            enriched_record["structure_context"] = claim_structure_context(
                record,
                evidence_blocks_by_id=evidence_blocks_by_id,
                knowledge_units_by_id=knowledge_units_by_id,
            )
            enriched_grouped_claims.append(enriched_record)
        item_payload = build_page_intent_item_payload(bucket_key, enriched_grouped_claims)
        fingerprint = fingerprint_payload(
            task_name="page_intent",
            item_payloads=[item_payload],
            prompt_version=task_config.prompt_version,
            schema_version=task_config.schema_version,
        )
        decision_record = latest_decisions.get(fingerprint)
        page_intent = ""
        route_reason = "heuristic_page_intent"
        source_decision_id = None
        if decision_record is not None:
            decision = decision_record.get("decision", {})
            if isinstance(decision, dict):
                page_intent = str(decision.get("page_intent", "")).strip().lower()
                if page_intent:
                    route_reason = str(decision_record.get("reason_code", "")).strip() or "semantic_page_intent"
                    source_decision_id = decision_record.get("decision_id")
        if not page_intent:
            page_intent = choose_bucket_page_intent(grouped_claims)

        original_page_intent = page_intent
        page_intent, route_reason = validate_page_intent_candidate(
            page_intent=page_intent,
            grouped_claims=grouped_claims,
            item_payload=item_payload,
            decision_record=decision_record,
            route_reason=route_reason,
        )
        if page_intent == "topic" and should_generate_concept_page(grouped_claims):
            page_intent = "concept"
            route_reason = "topic_promoted_to_concept_by_claim_group"

        route_payload = {
            "item_id": bucket_key,
            "bucket_key": bucket_key,
            "claim_ids": item_payload["claim_ids"],
            "source_page_intent": original_page_intent,
            "page_intent": page_intent,
            "route_target": page_intent,
            "route_reason": route_reason,
            "source_decision_id": source_decision_id,
            "supporting_unit_ids": sorted({
                unit_id
                for claim_record in grouped_claims
                for unit_id in claim_record.get("knowledge_unit_ids", [])
            }),
            "rejected_alternatives": [
                candidate
                for candidate in ("concept", "guide", "example", "topic", "reference", "timeline")
                if candidate != page_intent
            ],
        }
        route_fingerprint = fingerprint_payload(
            task_name="page_route",
            item_payloads=[route_payload],
            prompt_version=task_config.prompt_version,
            schema_version=task_config.schema_version,
        )
        existing_route_decision = latest_decisions.get(route_fingerprint)
        if existing_route_decision is None:
            existing_route_decision = {
                "decision_id": build_semantic_decision_id("page_route", route_fingerprint),
                "task_type": "page_route",
                "item_type": item_type_for_task("page_route"),
                "item_ids": [bucket_key],
                "decision": route_payload,
                "confidence": 1.0 if source_decision_id else 0.75,
                "reason_code": route_reason,
                "prompt_version": task_config.prompt_version,
                "model_key": task_config.model_key,
                "schema_version": task_config.schema_version,
                "input_fingerprint": route_fingerprint,
                "created_at": utc_now_iso(),
                "superseded_by": [],
            }
            new_route_decisions.append(existing_route_decision)

        page_routes[bucket_key] = {
            "page_intent": page_intent,
            "semantic_decision_id": existing_route_decision["decision_id"],
            "route_reason": route_reason,
            "route_target": page_intent,
            "source_decision_id": source_decision_id,
            "supporting_unit_ids": route_payload["supporting_unit_ids"],
            "rejected_alternatives": route_payload["rejected_alternatives"],
        }

    for record in new_route_decisions:
        append_jsonl(semantic_decisions_path(target), record)

    return page_routes


def preferred_page_intent_for_claim_group(
    grouped_claims: list[dict],
    page_intent: str,
) -> str:
    if page_intent == "topic" and should_generate_concept_page(grouped_claims):
        return "concept"
    return page_intent


def page_route_for_bucket(page_routes_by_bucket: dict[str, dict], bucket_key: str) -> dict:
    route = dict(page_routes_by_bucket.get(bucket_key) or {})
    route.setdefault("page_intent", "topic")
    route.setdefault("route_target", route["page_intent"])
    route.setdefault("semantic_decision_id", None)
    route.setdefault("route_reason", "missing_page_route_fallback")
    route.setdefault("source_decision_id", None)
    route.setdefault("supporting_unit_ids", [])
    route.setdefault("rejected_alternatives", [])
    return route


def apply_page_route_to_page_record(page_record: dict, page_route: dict) -> dict:
    updated = dict(page_record)
    decision_ids = list(updated.get("semantic_decision_ids", []) or [])
    semantic_decision_id = page_route.get("semantic_decision_id")
    if semantic_decision_id:
        append_unique(decision_ids, semantic_decision_id)
    source_decision_id = page_route.get("source_decision_id")
    if source_decision_id:
        append_unique(decision_ids, source_decision_id)
    updated["semantic_decision_ids"] = decision_ids
    updated["page_route"] = {
        "page_intent": page_route.get("page_intent"),
        "route_target": page_route.get("route_target"),
        "route_reason": page_route.get("route_reason"),
        "semantic_decision_id": semantic_decision_id,
        "source_decision_id": source_decision_id,
        "supporting_unit_ids": list(page_route.get("supporting_unit_ids", []) or []),
        "rejected_alternatives": list(page_route.get("rejected_alternatives", []) or []),
    }
    return updated


def supported_page_render_targets() -> tuple[str, ...]:
    return tuple(PAGE_RENDER_TARGETS.keys())


def page_record_render_target(page_record: dict) -> str | None:
    explicit_target = page_record.get("render_target")
    if explicit_target in PAGE_RENDER_TARGETS:
        return explicit_target

    page_type = page_record.get("type")
    for render_target, spec in PAGE_RENDER_TARGETS.items():
        if page_type in spec.get("page_types", set()):
            return render_target
    return None


def page_record_matches_render_target(page_record: dict, render_target: str) -> bool:
    return page_record_render_target(page_record) == render_target


def live_pages_for_render_target(page_records: list[dict], render_target: str) -> list[dict]:
    return [
        record
        for record in filter_live_page_records(page_records)
        if page_record_matches_render_target(record, render_target)
    ]


def load_page_render_config(config: dict, render_target: str) -> dict:
    if render_target not in PAGE_RENDER_TARGETS:
        raise KeyError(f"Unknown render target: {render_target}")
    rendering_config = config.get("rendering", {})
    target_config = (
        rendering_config.get(render_target, {})
        if isinstance(rendering_config, dict)
        else {}
    )
    if not isinstance(target_config, dict):
        target_config = {}

    mode = str(target_config.get("mode", "llm_assisted")).strip() or "llm_assisted"
    if mode not in {"deterministic", "llm_assisted"}:
        mode = "llm_assisted"

    command = normalize_command_config(target_config.get("command", []))
    timeout_seconds = coerce_int(target_config.get("timeout_seconds", 20), 20)
    timeout_seconds = max(timeout_seconds, 5)
    return {
        "render_target": render_target,
        "mode": mode,
        "command": command,
        "timeout_seconds": timeout_seconds,
    }


def load_readable_concept_render_config(config: dict) -> dict:
    return load_page_render_config(config, "readable_concept")


def resolve_workspace_path(target: Path, configured_path: str) -> Path:
    # config 里的路径既可能是相对工作区，也可能是绝对路径。
    path = Path(configured_path).expanduser()
    if path.is_absolute():
        return path
    return (target / path).resolve()


def raw_assets_dir_for_workspace(target: Path, raw_dir: Path | None = None) -> Path:
    resolved_raw_dir = raw_dir or resolve_workspace_raw_dir(target)
    return (resolved_raw_dir.parent / "assets").resolve()


def resolve_workspace_raw_dir(target: Path) -> Path:
    config = load_workspace_config(target)
    raw_dir = resolve_workspace_path(target, config["paths"]["raw"])
    if raw_dir.name != "raw":
        raise ValueError(f"Workspace raw directory must be named 'raw': {raw_dir}")
    if raw_dir.parent != target.parent:
        raise ValueError(
            f"Workspace raw directory must be a sibling of the workspace: raw={raw_dir} target={target}"
        )
    return raw_dir


def resolve_source_record_path(target: Path, source_path: str) -> Path:
    # source_path 默认按“相对工作区可访问路径”解释，这样 ../raw/... 也能稳定解析。
    path = Path(source_path).expanduser()
    if path.is_absolute():
        return path
    return (target / path).resolve()


def ensure_path_within_raw_root(path: Path, raw_root: Path, *, purpose: str) -> Path:
    resolved_path = path.resolve()
    resolved_root = raw_root.resolve()
    if not path_is_within_root(resolved_path, resolved_root):
        raise ValueError(
            f"{purpose} must stay within raw directory: path={resolved_path} raw={resolved_root}"
        )
    return resolved_path


def alias_index_path(target: Path) -> Path:
    # alias registry 是工作区级派生索引，和 search index 一样放在 indexes/ 下。
    return target / ALIAS_INDEX_REL_PATH


def page_alias_overrides_path(target: Path) -> Path:
    # 人工对页面 alias 的修订单独存一层覆盖，避免被后续自动页面重建直接抹掉。
    return target / PAGE_ALIAS_OVERRIDES_REL_PATH


def page_alias_overrides_lock_path(target: Path) -> Path:
    # review-apply 可能被多个 Agent/进程同时触发，覆盖层更新要串行化。
    return target / PAGE_ALIAS_OVERRIDES_LOCK_REL_PATH


def semantic_decisions_path(target: Path) -> Path:
    return target / SEMANTIC_DECISIONS_REL_PATH


def load_semantic_decisions(target: Path) -> list[dict]:
    return load_jsonl(semantic_decisions_path(target))


def build_latest_semantic_decisions_by_fingerprint(records: list[dict]) -> dict[str, dict]:
    latest: dict[str, dict] = {}
    for record in records:
        fingerprint = str(record.get("input_fingerprint", "")).strip()
        if not fingerprint:
            continue
        current = latest.get(fingerprint)
        if current is None or record.get("created_at", "") >= current.get("created_at", ""):
            latest[fingerprint] = record
    return latest


def normalize_alias_value(text: str) -> str:
    # alias / canonical 查询归一化尽量沿用 claim 文本清洗逻辑，
    # 这样页面标题、别名、查询词之间更容易对齐。
    return normalize_claim_text(text)


def load_alias_index(target: Path) -> dict:
    path = alias_index_path(target)
    if not path.exists():
        return {
            "index_version": ALIAS_INDEX_VERSION,
            "updated_at": None,
            "canonical_map": {},
            "alias_map": {},
            "conflicts": [],
        }
    return load_json(path)


def load_page_alias_overrides(target: Path) -> dict:
    path = page_alias_overrides_path(target)
    if not path.exists():
        return {"page_aliases": {}, "accepted_conflicts": []}
    payload = load_json(path)
    payload.setdefault("page_aliases", {})
    payload.setdefault("accepted_conflicts", [])
    return payload


def write_page_alias_overrides(target: Path, payload: dict) -> None:
    payload.setdefault("page_aliases", {})
    payload.setdefault("accepted_conflicts", [])
    write_json(page_alias_overrides_path(target), payload)


@dataclass
class FileLockHandle:
    path: Path
    file_handle: object


def acquire_file_lock(path: Path) -> FileLockHandle:
    path.parent.mkdir(parents=True, exist_ok=True)
    file_handle = path.open("w", encoding="utf-8")
    fcntl.flock(file_handle.fileno(), fcntl.LOCK_EX)
    return FileLockHandle(path=path, file_handle=file_handle)


def release_file_lock(lock_handle: FileLockHandle | None) -> None:
    if lock_handle is None:
        return
    fcntl.flock(lock_handle.file_handle.fileno(), fcntl.LOCK_UN)
    lock_handle.file_handle.close()


def apply_page_alias_overrides_payload(page_record: dict, overrides: dict) -> dict:
    page_aliases = overrides.get("page_aliases", {})
    override = page_aliases.get(page_record.get("page_id"), {})
    if not override:
        return page_record

    updated_record = dict(page_record)
    if "aliases" in override:
        updated_record["aliases"] = sorted(set(override.get("aliases", [])))
    if "title" in override and override.get("title"):
        updated_record["title"] = override["title"]
    updated_record["updated"] = utc_now_iso()
    return updated_record


def load_live_page_aliases_by_id(target: Path) -> dict[str, list[str]]:
    # alias 覆盖层只记录“人工最终想保留的页面 alias 集合”，
    # 但在第一次人工处理前，覆盖层里往往还没有任何内容。
    # 这里补一层从 pages.jsonl 读取当前 live alias 的快照，
    # 让 assign/remove 基于“页面现状”增删，而不是误把别名列表清成只剩人工刚操作的那一项。
    pages_path = target / "state" / "pages.jsonl"
    if not pages_path.exists():
        return {}

    aliases_by_page_id: dict[str, list[str]] = {}
    for record in load_jsonl(pages_path):
        record = ensure_page_lifecycle_defaults(record)
        if not is_live_page_record(record):
            continue
        aliases_by_page_id[record["page_id"]] = sorted(set(record.get("aliases", [])))
    return aliases_by_page_id


def remove_alias_from_overrides(
    target: Path,
    page_ids: list[str],
    alias_value: str,
) -> dict:
    # 从指定页面的人工 alias 覆盖层里移除某个 alias。
    overrides = load_page_alias_overrides(target)
    page_aliases = overrides.setdefault("page_aliases", {})
    live_aliases_by_page_id = load_live_page_aliases_by_id(target)
    normalized_alias = normalize_alias_value(alias_value)

    for page_id in page_ids:
        page_override = page_aliases.setdefault(page_id, {})
        aliases = sorted(set(page_override.get("aliases", live_aliases_by_page_id.get(page_id, []))))
        aliases = [alias for alias in aliases if normalize_alias_value(alias) != normalized_alias]
        page_override["aliases"] = aliases

    write_page_alias_overrides(target, overrides)
    return overrides


def apply_alias_override_action(
    overrides: dict,
    live_aliases_by_page_id: dict[str, list[str]],
    candidate_page_ids: list[str],
    primary_page_id: str,
    alias_value: str,
    action: str,
) -> dict:
    updated_overrides = copy.deepcopy(overrides)
    page_aliases = updated_overrides.setdefault("page_aliases", {})
    normalized_alias = normalize_alias_value(alias_value)

    for page_id in candidate_page_ids:
        page_override = page_aliases.setdefault(page_id, {})
        aliases = sorted(set(page_override.get("aliases", live_aliases_by_page_id.get(page_id, []))))
        aliases = [alias for alias in aliases if normalize_alias_value(alias) != normalized_alias]
        if action == "assign_alias" and page_id == primary_page_id and alias_value not in aliases:
            aliases.append(alias_value)
        page_override["aliases"] = sorted(set(aliases))
    return updated_overrides


def accepted_alias_conflict_signature(alias_value: str, canonical_ids: list[str]) -> str:
    normalized_alias = normalize_alias_value(alias_value)
    canonical_part = "|".join(sorted(str(item).strip() for item in canonical_ids if str(item).strip()))
    return f"{normalized_alias}::{canonical_part}"


def build_accepted_alias_conflict_signatures(overrides: dict) -> set[str]:
    accepted = overrides.get("accepted_conflicts", [])
    signatures: set[str] = set()
    for item in accepted:
        if isinstance(item, str) and item.strip():
            signatures.add(item.strip())
            continue
        if not isinstance(item, dict):
            continue
        alias_value = str(item.get("alias", "")).strip()
        canonical_ids = [str(value).strip() for value in item.get("canonical_ids", []) if str(value).strip()]
        if not alias_value or not canonical_ids:
            continue
        signatures.add(accepted_alias_conflict_signature(alias_value, canonical_ids))
    return signatures


def persist_accepted_alias_conflict(
    overrides: dict,
    alias_value: str,
    canonical_ids: list[str],
) -> dict:
    updated_overrides = copy.deepcopy(overrides)
    accepted = [
        item for item in updated_overrides.get("accepted_conflicts", [])
        if isinstance(item, dict) or (isinstance(item, str) and item.strip())
    ]
    normalized_canonical_ids = sorted({
        str(item).strip()
        for item in canonical_ids
        if str(item).strip()
    })
    signature = accepted_alias_conflict_signature(alias_value, normalized_canonical_ids)

    filtered_accepted: list[dict | str] = []
    for item in accepted:
        if isinstance(item, str):
            if item.strip() != signature:
                filtered_accepted.append(item)
            continue
        existing_signature = accepted_alias_conflict_signature(
            str(item.get("alias", "")).strip(),
            [str(value).strip() for value in item.get("canonical_ids", []) if str(value).strip()],
        )
        if existing_signature != signature:
            filtered_accepted.append(item)

    filtered_accepted.append({
        "alias": alias_value,
        "canonical_ids": normalized_canonical_ids,
        "accepted_at": utc_now_iso(),
    })
    updated_overrides["accepted_conflicts"] = filtered_accepted
    return updated_overrides


def clear_accepted_alias_conflict(
    overrides: dict,
    alias_value: str,
    canonical_ids: list[str],
) -> dict:
    updated_overrides = copy.deepcopy(overrides)
    signature = accepted_alias_conflict_signature(alias_value, canonical_ids)
    filtered_accepted: list[dict | str] = []
    for item in updated_overrides.get("accepted_conflicts", []):
        if isinstance(item, str):
            if item.strip() != signature:
                filtered_accepted.append(item)
            continue
        existing_signature = accepted_alias_conflict_signature(
            str(item.get("alias", "")).strip(),
            [str(value).strip() for value in item.get("canonical_ids", []) if str(value).strip()],
        )
        if existing_signature != signature:
            filtered_accepted.append(item)
    updated_overrides["accepted_conflicts"] = filtered_accepted
    return updated_overrides


def unresolved_alias_conflicts(alias_index: dict) -> list[dict]:
    return [
        conflict
        for conflict in alias_index.get("conflicts", [])
        if not conflict.get("accepted")
    ]


def update_page_alias_overrides_with_lock(
    target: Path,
    updater,
) -> dict:
    lock_handle = acquire_file_lock(page_alias_overrides_lock_path(target))
    try:
        overrides = load_page_alias_overrides(target)
        updated_overrides = updater(overrides)
        write_page_alias_overrides(target, updated_overrides)
        return updated_overrides
    finally:
        release_file_lock(lock_handle)


def build_alias_index(page_records: list[dict], accepted_conflict_signatures: set[str] | None = None) -> dict:
    # alias registry 统一记录 canonical_id、title、aliases 的双向映射关系。
    # query、lint、Agent 约定都依赖它，避免各自维护一份别名世界观。
    canonical_map: dict[str, dict] = {}
    alias_map: dict[str, list[dict]] = {}
    accepted_conflict_signatures = accepted_conflict_signatures or set()
    live_page_records = filter_live_page_records(page_records)
    pages_by_canonical_id: dict[str, list[dict]] = {}
    title_owners_by_alias: dict[str, list[dict]] = {}
    noisy_title_alias_values = {
        normalize_alias_value("一句话总结"),
        normalize_alias_value("注意"),
    }

    def canonical_page_rank_key(page_record: dict) -> tuple:
        page_type = page_record.get("type", "")
        page_status = page_record.get("status", "")
        return (
            1 if page_type == "concept" else 0,
            1 if page_status == "stable" else 0,
            QUERY_PAGE_TYPE_WEIGHTS.get(page_type, QUERY_PAGE_TYPE_WEIGHTS["draft"]),
            QUERY_PAGE_STATUS_WEIGHTS.get(page_status, QUERY_PAGE_STATUS_WEIGHTS["draft"]),
            len(page_record.get("claim_ids", [])),
        )

    for page_record in live_page_records:
        canonical_id = page_record.get("canonical_id") or page_record.get("page_id")
        pages_by_canonical_id.setdefault(canonical_id, []).append(page_record)
        normalized_title = normalize_alias_value(page_record.get("title", ""))
        if normalized_title:
            title_owners_by_alias.setdefault(normalized_title, []).append(page_record)

    def should_register_title_alias(page_record: dict, normalized_title: str) -> bool:
        if normalized_title in noisy_title_alias_values:
            owners = title_owners_by_alias.get(normalized_title, [])
            return len({
                owner.get("canonical_id") or owner.get("page_id")
                for owner in owners
            }) <= 1
        # source-summary 的标题常常只是原文文件名或章节名，
        # 如果它与概念/综述页重名，再把它注册成 alias 只会制造噪声和伪冲突。
        # 这里保留 source-summary 的正文/标题检索能力，但在 alias registry 里更保守。
        if page_record.get("type") != "source-summary":
            return True
        owners = title_owners_by_alias.get(normalized_title, [])
        if len(owners) <= 1:
            return True
        return not any(
            owner.get("type") in {"concept", "overview", "entity"}
            and owner.get("canonical_id") != page_record.get("canonical_id")
            for owner in owners
        )

    for canonical_id, grouped_pages in pages_by_canonical_id.items():
        representative_page = max(grouped_pages, key=canonical_page_rank_key)
        combined_aliases = sorted({
            alias
            for page_record in grouped_pages
            for alias in page_record.get("aliases", [])
            if alias
        })
        canonical_map[canonical_id] = {
            "canonical_id": canonical_id,
            "page_id": representative_page.get("page_id"),
            "title": representative_page.get("title", ""),
            "page_path": representative_page.get("page_path", ""),
            "type": representative_page.get("type", ""),
            "status": representative_page.get("status", ""),
            "aliases": combined_aliases,
        }

    for page_record in live_page_records:
        page_id = page_record.get("page_id")
        canonical_id = page_record.get("canonical_id") or page_id
        title = page_record.get("title", "")
        page_path = page_record.get("page_path", "")
        page_type = page_record.get("type", "")
        page_status = page_record.get("status", "")
        normalized_title = normalize_alias_value(title)
        candidates = [canonical_id, *page_record.get("aliases", [])]
        if normalized_title and should_register_title_alias(page_record, normalized_title):
            candidates.insert(0, title)
        seen_keys: set[str] = set()
        for candidate in candidates:
            normalized_candidate = normalize_alias_value(candidate)
            if not normalized_candidate or normalized_candidate in seen_keys:
                continue
            seen_keys.add(normalized_candidate)
            alias_map.setdefault(normalized_candidate, []).append({
                "canonical_id": canonical_id,
                "page_id": page_id,
                "title": title,
                "page_path": page_path,
                "type": page_type,
                "status": page_status,
                "matched_from": candidate,
            })

    conflicts = []
    for alias_key, matches in sorted(alias_map.items()):
        canonical_ids = sorted({item["canonical_id"] for item in matches})
        if len(canonical_ids) <= 1:
            continue
        signature = accepted_alias_conflict_signature(alias_key, canonical_ids)
        conflicts.append({
            "alias": alias_key,
            "canonical_ids": canonical_ids,
            "page_ids": sorted({item["page_id"] for item in matches}),
            "accepted": signature in accepted_conflict_signatures,
        })

    return {
        "index_version": ALIAS_INDEX_VERSION,
        "updated_at": utc_now_iso(),
        "canonical_map": canonical_map,
        "alias_map": alias_map,
        "conflicts": conflicts,
    }


def write_alias_index(target: Path, page_records: list[dict]) -> dict:
    overrides = load_page_alias_overrides(target)
    alias_index = build_alias_index(
        page_records,
        accepted_conflict_signatures=build_accepted_alias_conflict_signatures(overrides),
    )
    write_json(alias_index_path(target), alias_index)
    return alias_index


def apply_page_alias_overrides(target: Path, page_record: dict) -> dict:
    # 自动页面重建前先叠加人工 alias 覆盖层。
    overrides = load_page_alias_overrides(target)
    return apply_page_alias_overrides_payload(page_record, overrides)


def apply_page_alias_overrides_to_records(target: Path, page_records: list[dict]) -> list[dict]:
    overrides = load_page_alias_overrides(target)
    return [apply_page_alias_overrides_payload(record, overrides) for record in page_records]


def build_alias_conflict_reviews(
    alias_index: dict,
    existing_reviews: dict[str, dict],
) -> tuple[list[dict], list[str]]:
    # alias registry 里一旦出现“一词多义”的冲突，就应进入 review 队列而不是只留在 lint 里。
    created_reviews: list[dict] = []
    touched_review_ids: list[str] = []

    for conflict in alias_index.get("conflicts", []):
        if conflict.get("accepted"):
            continue
        canonical_ids = sorted(conflict.get("canonical_ids", []))
        page_ids = sorted(conflict.get("page_ids", []))
        review_record = build_review_record(
            kind="alias_conflict",
            candidate_claim_ids=[],
            reason="Detected alias key mapped to multiple canonical pages and requires manual disambiguation.",
            evidence=[{
                "alias": conflict.get("alias"),
                "canonical_ids": canonical_ids,
                "page_ids": page_ids,
            }],
            recommended_action="keep_both",
            signature_parts=[
                conflict.get("alias", ""),
                *canonical_ids,
            ],
        )
        review_record["candidate_page_ids"] = page_ids
        review_record["allowed_actions"] = ["keep_both", "edit_then_resume", "assign_alias", "remove_alias"]
        review_record["resume_from"] = "alias_registry"

        existing_review = existing_reviews.get(review_record["review_id"])
        if existing_review is not None:
            # 冲突仍然存在时，把 page_ids 刷新到最新集合即可。
            existing_review["candidate_page_ids"] = page_ids
            existing_review["evidence"] = review_record["evidence"]
            existing_review["reason"] = review_record["reason"]
            existing_review["recommended_action"] = review_record["recommended_action"]
            touched_review_ids.append(existing_review["review_id"])
            continue

        created_reviews.append(review_record)
        touched_review_ids.append(review_record["review_id"])

    return created_reviews, touched_review_ids


def archive_stale_alias_conflict_reviews(
    live_reviews_by_id: dict[str, dict],
    historical_reviews_by_id: dict[str, dict],
    active_alias_review_ids: set[str],
) -> set[str]:
    # alias 冲突一旦在当前 alias index 中消失，旧 review 就不该继续伪装成 active。
    # 这里把“仍为 alias_conflict、但已不在当前冲突集合里”的记录自动转入历史态。
    archived_review_ids: set[str] = set()
    for review_id, review_record in list(live_reviews_by_id.items()):
        if review_record.get("kind") != "alias_conflict":
            continue
        if review_id in active_alias_review_ids:
            continue
        archived_record = dict(review_record)
        archived_record["status"] = "resolved"
        archived_record["resolved_at"] = archived_record.get("resolved_at") or utc_now_iso()
        archived_record["lifecycle_status"] = "superseded"
        archived_record["archived_at"] = utc_now_iso()
        live_reviews_by_id.pop(review_id, None)
        historical_record = convert_review_record_to_historical(archived_record)
        historical_reviews_by_id[historical_record["review_id"]] = historical_record
        archived_review_ids.add(review_id)
    return archived_review_ids


def refresh_alias_conflict_reviews(
    target: Path,
    live_reviews_by_id: dict[str, dict],
    historical_reviews_by_id: dict[str, dict],
    page_records: list[dict] | None = None,
) -> tuple[dict, set[str], set[str]]:
    # 把 alias index 与 review 账本在一个入口里重新对齐：
    # 1. 基于当前 live pages 重建 alias index
    # 2. 刷新仍存在的 alias_conflict review
    # 3. 将已消失的 alias_conflict review 转成历史态
    if page_records is None:
        page_records = [
            ensure_page_lifecycle_defaults(record)
            for record in load_jsonl(target / "state" / "pages.jsonl")
        ]
    page_records = apply_page_alias_overrides_to_records(target, page_records)
    alias_index = write_alias_index(target, page_records)
    created_reviews, touched_review_ids = build_alias_conflict_reviews(alias_index, live_reviews_by_id)
    for review_record in created_reviews:
        live_reviews_by_id[review_record["review_id"]] = review_record
    archived_review_ids = archive_stale_alias_conflict_reviews(
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
        active_alias_review_ids=set(touched_review_ids),
    )
    return alias_index, set(touched_review_ids), archived_review_ids


def alias_index_matches_for_value(alias_index: dict, alias_value: str) -> list[dict]:
    normalized_alias = normalize_alias_value(alias_value)
    return list(alias_index.get("alias_map", {}).get(normalized_alias, []))


def replace_jsonl_record(path: Path, key_field: str, key_value: str, new_record: dict) -> None:
    # JSONL 天然适合追加，但“更新某条记录”就需要整文件重写一遍。
    # 这里先用最直白、最容易读懂的实现，后面数据量大了再考虑索引化。
    records = load_jsonl(path)
    replaced = False
    updated_records = []
    for record in records:
        if record.get(key_field) == key_value and not replaced:
            updated_records.append(new_record)
            replaced = True
        else:
            updated_records.append(record)
    if not replaced:
        updated_records.append(new_record)
    write_jsonl(path, updated_records)


def replace_jsonl_records_by_filter(path: Path, keep_predicate, replacement_records: list[dict]) -> None:
    # 有些 state 文件需要“替换同一来源的一组记录”，例如 chunks.jsonl。
    # 这里统一做成一个小工具，避免在主流程里反复手写整文件过滤逻辑。
    records = load_jsonl(path)
    kept_records = [record for record in records if keep_predicate(record)]
    write_jsonl(path, kept_records + replacement_records)


def replace_source_scoped_jsonl_records(path: Path, source_id: str, replacement_records: list[dict]) -> None:
    # V2 结构账本按 source_id 整体替换，保证重复 ingest 不会让结构记录膨胀。
    replace_jsonl_records_by_filter(
        path,
        keep_predicate=lambda record, source_id=source_id: record.get("source_id") != source_id,
        replacement_records=replacement_records,
    )


def normalize_text_content(source_type: str, raw_text: str) -> str:
    # 第一版先做最保守的规范化：统一换行、去掉尾部多余空白，并确保纯文本也能落成 Markdown。
    normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    # 行尾空白通常没有信息价值，反而会让 hash 和 diff 变得不稳定。
    lines = [line.rstrip() for line in normalized.split("\n")]
    normalized = "\n".join(lines).strip()
    if source_type == "plain_text":
        # 纯文本在 V1 里先直接当成 Markdown 文本存储，简化后续统一处理。
        return normalized + "\n"
    return normalized + "\n"


def normalize_markdown_or_text_record(
    target: Path,
    source_record: dict,
    *,
    allow_insecure_downloads: bool = True,
) -> dict:
    # Markdown 和纯文本是当前最稳定的一类输入，直接按文本规范化处理。
    source_type = source_record["source_type"]
    raw_path = resolve_source_record_path(target, source_record["source_path"])
    ensure_path_within_raw_root(raw_path, resolve_workspace_raw_dir(target), purpose="Source record")
    raw_text = raw_path.read_text(encoding="utf-8")
    if source_type == "markdown":
        normalized_text, metadata = enrich_markdown_with_embedded_images(
            target=target,
            source_record=source_record,
            raw_path=raw_path,
            raw_text=raw_text,
            allow_insecure_downloads=allow_insecure_downloads,
        )
    else:
        normalized_text = normalize_text_content(source_type, raw_text)
        metadata = {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good",
            "warnings": [],
            "location_map": {
                "type": "line_map",
                "source_path": source_record["source_path"],
            },
        }
    normalized_rel_path = Path("normalized") / f"{source_record['source_id']}.md"
    normalized_abs_path = target / normalized_rel_path
    normalized_abs_path.write_text(normalized_text, encoding="utf-8")

    normalized_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    line_count = len(normalized_text.splitlines()) or 1
    title = raw_path.stem
    location_map = dict(metadata.get("location_map", {}))
    if location_map.get("type") == "line_map" and "normalized_line_range" not in location_map:
        location_map["normalized_line_range"] = f"1-{line_count}"

    return {
        "source_id": source_record["source_id"],
        "source_type": source_type,
        "source_path": source_record["source_path"],
        "normalized_path": str(normalized_rel_path),
        "title": title,
        "language": "unknown",
        "content_format": "markdown",
        "raw_hash": source_record["source_hash"],
        "normalized_hash": normalized_hash,
        "normalizer_version": "normalize_v1",
        "document_kind": "note",
        "structure_quality": "unknown",
        "chunk_strategy_hint": "heading_first",
        "extraction_method": metadata.get("extraction_method", "python_only"),
        "extraction_quality": metadata.get("extraction_quality", "good"),
        "warnings": metadata.get("warnings", []),
        "location_map": location_map,
        "updated_at": utc_now_iso(),
    }


def convert_pdf_to_markdown(raw_path: Path) -> tuple[str, dict]:
    # PDF 先走纯 Python 文本提取，按页组织成 Markdown。
    if PdfReader is None:
        return convert_pdf_to_markdown_fallback(raw_path)

    try:
        reader = PdfReader(str(raw_path))
        parts: list[str] = [f"# {raw_path.stem}"]
        page_map: list[dict] = []
        warnings: list[str] = []

        for page_index, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            cleaned = extracted.strip()
            parts.append(f"\n## 第 {page_index} 页\n")
            if cleaned:
                parts.append(cleaned)
            else:
                parts.append("_本页未提取到文本_")
                warnings.append(f"page_{page_index}_empty_text")
            page_map.append({
                "page": page_index,
                "char_count": len(cleaned),
            })

        markdown = "\n\n".join(parts).strip() + "\n"
        metadata = {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "partial" if warnings else "good",
            "warnings": warnings,
            "location_map": {
                "type": "pdf_page_map",
                "pages": page_map,
                "source_path": str(raw_path),
            },
        }
        return markdown, metadata
    except Exception:
        # 主路径失败时退回低依赖 fallback，尽量别让 PDF 直接中断整批 ingest。
        return convert_pdf_to_markdown_fallback(raw_path)


def pdf_count_pages_from_bytes(pdf_bytes: bytes) -> int:
    # 这是一个很保守的页数估计：匹配 /Type /Page，避开 /Pages。
    matches = re.findall(rb"/Type\s*/Page\b", pdf_bytes)
    return max(len(matches), 0)


def pdf_try_inflate_stream(stream_bytes: bytes) -> bytes:
    try:
        return zlib.decompress(stream_bytes)
    except Exception:
        return stream_bytes


def decode_pdf_literal_string(raw: bytes) -> str:
    # PDF 文本操作符中的字符串常常放在 (...) 里，这里做一层轻量解码。
    text = raw.replace(b"\\(", b"(").replace(b"\\)", b")").replace(b"\\\\", b"\\")
    try:
        return text.decode("utf-8")
    except UnicodeDecodeError:
        return text.decode("latin-1", errors="ignore")


def pdf_extract_text_snippets_from_bytes(pdf_bytes: bytes) -> list[str]:
    # fallback 只求“尽量提取一点正文”，不追求完整版面恢复。
    snippets: list[str] = []

    object_matches = re.finditer(rb"(\d+\s+\d+\s+obj.*?endobj)", pdf_bytes, flags=re.DOTALL)
    for match in object_matches:
        object_bytes = match.group(1)
        stream_match = re.search(rb"stream\r?\n(.*?)\r?\nendstream", object_bytes, flags=re.DOTALL)
        if not stream_match:
            continue
        stream_bytes = stream_match.group(1)
        if b"/FlateDecode" in object_bytes:
            stream_bytes = pdf_try_inflate_stream(stream_bytes)

        for block in re.findall(rb"BT(.*?)ET", stream_bytes, flags=re.DOTALL):
            pieces = re.findall(rb"\(([^()]*)\)", block)
            if not pieces:
                continue
            combined = "".join(decode_pdf_literal_string(piece) for piece in pieces).strip()
            combined = re.sub(r"\s+", " ", combined)
            if len(combined) >= 4 and combined not in snippets:
                snippets.append(combined)

    return snippets[:20]


def convert_pdf_to_markdown_fallback(raw_path: Path) -> tuple[str, dict]:
    # 没有 pypdf 时，至少给 PDF 生成页数、可提取文本片段和占位信息。
    pdf_bytes = raw_path.read_bytes()
    page_count = pdf_count_pages_from_bytes(pdf_bytes)
    snippets = pdf_extract_text_snippets_from_bytes(pdf_bytes)
    warnings: list[str] = []

    parts = [f"# {raw_path.stem}", "", f"- 估计页数: {page_count or 'unknown'}"]
    if snippets:
        parts.extend(["", "## 提取文本片段"])
        for index, snippet in enumerate(snippets, start=1):
            parts.append(f"{index}. {snippet}")
    else:
        warnings.append("pdf_fallback_no_text")
        parts.extend([
            "",
            "> 当前环境未启用 `pypdf`，且标准库 fallback 未提取到正文文本。",
        ])

    markdown = "\n".join(parts).strip() + "\n"
    pages = [{"page": index + 1, "char_count": None} for index in range(page_count)] if page_count else []
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "partial",
        "warnings": warnings if warnings else ["pdf_fallback_used"],
        "location_map": {
            "type": "pdf_page_map",
            "pages": pages,
            "source_path": str(raw_path),
        },
    }


def iter_docx_blocks(document) -> list[tuple[str, str]]:
    # python-docx 的段落和表格分散在不同结构里。
    # V1 先分别收集，再按各自顺序做一个保守输出。
    blocks: list[tuple[str, str]] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name:
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(digits) if digits else 1
            level = min(max(level, 1), 6)
            blocks.append(("heading", "#" * level + " " + text))
        else:
            blocks.append(("paragraph", text))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") or " " for cell in row.cells]
            rows.append(values)
        if not rows:
            continue
        blocks.append(("table_title", f"## 表格 {table_index}"))
        header = rows[0]
        divider = ["---"] * len(header)
        table_lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(divider) + " |",
        ]
        for row in rows[1:]:
            padded = row + [" "] * (len(header) - len(row))
            table_lines.append("| " + " | ".join(padded[: len(header)]) + " |")
        blocks.append(("table", "\n".join(table_lines)))

    return blocks


DOCX_NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

XLSX_NAMESPACES = {
    "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}

OLE_HEADER_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def docx_style_map_from_archive(archive: zipfile.ZipFile) -> dict[str, str]:
    # 纯 Python 兜底解析 docx 时，需要先把 styleId 映射到更容易理解的样式名。
    try:
        styles_xml = archive.read("word/styles.xml")
    except KeyError:
        return {}

    root = ET.fromstring(styles_xml)
    style_map: dict[str, str] = {}
    for style in root.findall("w:style", DOCX_NAMESPACES):
        style_id = style.get(f"{{{DOCX_NAMESPACES['w']}}}styleId")
        name_node = style.find("w:name", DOCX_NAMESPACES)
        if not style_id:
            continue
        style_map[style_id] = name_node.get(f"{{{DOCX_NAMESPACES['w']}}}val", style_id) if name_node is not None else style_id
    return style_map


def docx_paragraph_text(node: ET.Element) -> str:
    # Word 段落里的文本经常被拆到多个 run / text 节点里，这里把它们拼回去。
    texts = [item.text or "" for item in node.findall(".//w:t", DOCX_NAMESPACES)]
    return "".join(texts).strip()


def docx_heading_level_from_style(style_name: str, style_id: str) -> int | None:
    # heading 样式名和 styleId 在不同文档里可能略有不同，这里做保守识别。
    candidates = f"{style_name} {style_id}".lower()
    if "heading" not in candidates:
        return None
    digits = "".join(ch for ch in candidates if ch.isdigit())
    if not digits:
        return 1
    return min(max(int(digits[0]), 1), 6)


def docx_table_to_markdown(node: ET.Element) -> str:
    # 纯 XML 路径下把 Word 表格转成 Markdown 表格。
    rows: list[list[str]] = []
    for row in node.findall("w:tr", DOCX_NAMESPACES):
        values = []
        for cell in row.findall("w:tc", DOCX_NAMESPACES):
            cell_texts = []
            for paragraph in cell.findall(".//w:p", DOCX_NAMESPACES):
                text = docx_paragraph_text(paragraph)
                if text:
                    cell_texts.append(text)
            values.append(" ".join(cell_texts).strip() or " ")
        if rows or any(value.strip() for value in values):
            rows.append(values)

    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized_rows = [row + [" "] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    divider = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_docx_to_markdown_fallback(raw_path: Path) -> tuple[str, dict]:
    # 这个 fallback 完全走标准库，避免因为 lxml / 二进制依赖不兼容把整条 docx 路堵死。
    parts = [f"# {raw_path.stem}"]
    warnings: list[str] = []

    with zipfile.ZipFile(raw_path) as archive:
        style_map = docx_style_map_from_archive(archive)
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    body = root.find("w:body", DOCX_NAMESPACES)
    if body is None:
        warnings.append("docx_missing_body")
        markdown = "\n".join(parts).strip() + "\n"
        return markdown, {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "partial",
            "warnings": warnings,
            "location_map": {
                "type": "line_map",
                "normalized_line_range": "1-1",
                "source_path": str(raw_path),
            },
        }

    table_index = 0
    for child in list(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = docx_paragraph_text(child)
            if not text:
                continue
            style_node = child.find("w:pPr/w:pStyle", DOCX_NAMESPACES)
            style_id = style_node.get(f"{{{DOCX_NAMESPACES['w']}}}val", "") if style_node is not None else ""
            style_name = style_map.get(style_id, style_id)
            heading_level = docx_heading_level_from_style(style_name, style_id)
            if heading_level is not None:
                parts.append("#" * heading_level + " " + text)
            else:
                parts.append(text)
        elif tag == "tbl":
            table_index += 1
            table_markdown = docx_table_to_markdown(child)
            if table_markdown:
                parts.append(f"## 表格 {table_index}")
                parts.append(table_markdown)

    markdown = "\n\n".join(part for part in parts if part.strip()).strip() + "\n"
    line_count = len(markdown.splitlines()) or 1
    metadata = {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if len(parts) > 1 else "partial",
        "warnings": warnings if warnings else [],
        "location_map": {
            "type": "line_map",
            "normalized_line_range": f"1-{line_count}",
            "source_path": str(raw_path),
        },
    }
    return markdown, metadata


def convert_docx_to_markdown(raw_path: Path) -> tuple[str, dict]:
    # Word 文档优先保留“标题 + 段落 + 表格”的主结构。
    if docx is not None:
        document = docx.Document(str(raw_path))
        blocks = iter_docx_blocks(document)
        parts = [f"# {raw_path.stem}"]
        parts.extend(block for _, block in blocks)
        markdown = "\n\n".join(part for part in parts if part.strip()).strip() + "\n"
        line_count = len(markdown.splitlines()) or 1

        metadata = {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good" if blocks else "partial",
            "warnings": [] if blocks else ["docx_no_visible_blocks"],
            "location_map": {
                "type": "line_map",
                "normalized_line_range": f"1-{line_count}",
                "source_path": str(raw_path),
            },
        }
        return markdown, metadata

    # 如果 python-docx 不可用，就退回到 zip+xml 路径。
    return convert_docx_to_markdown_fallback(raw_path)


def is_probably_ole_document(raw_path: Path) -> bool:
    # 老 Office 二进制格式通常基于 OLE Compound File。
    # 这里先做魔数级判断，用于给 fallback metadata 增加一点上下文。
    with raw_path.open("rb") as fh:
        return fh.read(len(OLE_HEADER_MAGIC)) == OLE_HEADER_MAGIC


def normalize_binary_snippet_text(text: str) -> str:
    # 二进制兜底抽出来的文本会混杂很多控制字符和奇怪空白，这里尽量压平。
    cleaned = []
    for char in text:
        if char in {"\n", "\r", "\t"}:
            cleaned.append(" ")
            continue
        category = ord(char)
        if char.isprintable() or 0x4E00 <= category <= 0x9FFF:
            cleaned.append(char)
    normalized = "".join(cleaned)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def binary_text_candidate_is_meaningful(text: str, min_length: int = 8) -> bool:
    # 兜底提取不要求完美，但至少要像“人会读的片段”，尽量别把纯噪声放进去。
    if len(text) < min_length:
        return False
    interesting_chars = [
        char for char in text
        if char.isalnum() or "\u4e00" <= char <= "\u9fff"
    ]
    if len(interesting_chars) < max(4, min_length // 2):
        return False
    return True


def extract_printable_ascii_snippets(binary_bytes: bytes, min_length: int = 8) -> list[str]:
    snippets: list[str] = []
    for match in re.finditer(rb"[\x20-\x7e]{8,}", binary_bytes):
        text = normalize_binary_snippet_text(match.group(0).decode("utf-8", errors="ignore"))
        if binary_text_candidate_is_meaningful(text, min_length=min_length):
            snippets.append(text)
    return snippets


def extract_utf16_text_snippets(binary_bytes: bytes, min_length: int = 8) -> list[str]:
    # 老 Office 二进制里经常能捞到 UTF-16LE 文本。
    # 为了适应可能的字节对齐偏移，这里从 0/1 两个 offset 都扫一遍。
    snippets: list[str] = []
    pattern = re.compile(r"[0-9A-Za-z\u4e00-\u9fff][0-9A-Za-z\u4e00-\u9fff\s，。、“”‘’；：？！,.!?:;()（）\-_/]{7,}")
    for offset in (0, 1):
        if len(binary_bytes) <= offset + 2:
            continue
        decoded = binary_bytes[offset:].decode("utf-16le", errors="ignore")
        for match in pattern.finditer(decoded):
            text = normalize_binary_snippet_text(match.group(0))
            if binary_text_candidate_is_meaningful(text, min_length=min_length):
                snippets.append(text)
    return snippets


def extract_text_snippets_from_binary_document(
    binary_bytes: bytes,
    limit: int = 20,
    min_length: int = 8,
) -> list[str]:
    # 这是老格式 Office 的兜底文本提取：
    # - 先扫 ASCII/UTF-8 可见串
    # - 再扫 UTF-16LE 候选串
    # 目标是尽量拿到“能提示内容主题”的片段，而不是高保真还原。
    ordered_candidates = [
        *extract_printable_ascii_snippets(binary_bytes, min_length=min_length),
        *extract_utf16_text_snippets(binary_bytes, min_length=min_length),
    ]

    snippets: list[str] = []
    seen: set[str] = set()
    for candidate in ordered_candidates:
        dedupe_key = candidate.lower()
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        snippets.append(candidate)
        if len(snippets) >= limit:
            break
    return snippets


def convert_legacy_doc_to_markdown(raw_path: Path) -> tuple[str, dict]:
    # `.doc` 在 V1 先不上复杂二进制解析器。
    # 先尽量给出“文件身份 + 可见文本片段 + 明确告警”，保证能进入后续链路。
    binary_bytes = raw_path.read_bytes()
    snippets = extract_text_snippets_from_binary_document(binary_bytes)
    is_ole = is_probably_ole_document(raw_path)

    parts = [
        f"# {raw_path.stem}",
        "",
        "## 文档信息 / Document Info",
        "",
        f"- 原始格式: `.doc`",
        f"- 文件大小: {len(binary_bytes)} bytes",
        f"- OLE 容器: `{is_ole}`",
    ]

    if snippets:
        parts.extend([
            "",
            "## 提取文本片段 / Extracted Snippets",
            "",
        ])
        for index, snippet in enumerate(snippets, start=1):
            parts.append(f"{index}. {snippet}")
    else:
        parts.extend([
            "",
            "## 提取文本片段 / Extracted Snippets",
            "",
            "> 当前纯 Python fallback 未提取到可读正文片段。",
        ])

    markdown = "\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "partial" if snippets else "poor",
        "warnings": ["legacy_doc_binary_fallback"] if snippets else ["legacy_doc_no_text_snippets"],
        "location_map": {
            "type": "binary_snippet_map",
            "source_path": str(raw_path),
            "snippet_count": len(snippets),
            "is_ole_container": is_ole,
        },
    }


def image_size_from_binary(raw_path: Path) -> tuple[int | None, int | None, str | None]:
    # 在 Pillow 不可用时，用文件头做一层轻量尺寸识别。
    with raw_path.open("rb") as fh:
        header = fh.read(64)

    if header.startswith(b"\x89PNG\r\n\x1a\n") and len(header) >= 24:
        width, height = struct.unpack(">II", header[16:24])
        return width, height, "PNG"

    if header.startswith(b"\xff\xd8"):
        with raw_path.open("rb") as fh:
            fh.read(2)
            while True:
                marker_start = fh.read(1)
                if not marker_start:
                    break
                if marker_start != b"\xff":
                    continue
                marker = fh.read(1)
                while marker == b"\xff":
                    marker = fh.read(1)
                if marker in {b"\xc0", b"\xc1", b"\xc2", b"\xc3", b"\xc5", b"\xc6", b"\xc7", b"\xc9", b"\xca", b"\xcb", b"\xcd", b"\xce", b"\xcf"}:
                    segment_length = struct.unpack(">H", fh.read(2))[0]
                    _precision = fh.read(1)
                    height, width = struct.unpack(">HH", fh.read(4))
                    return width, height, "JPEG"
                if marker in {b"\xd8", b"\xd9"}:
                    continue
                segment_length_data = fh.read(2)
                if len(segment_length_data) < 2:
                    break
                segment_length = struct.unpack(">H", segment_length_data)[0]
                fh.seek(segment_length - 2, 1)

    if header[:4] == b"RIFF" and header[8:12] == b"WEBP":
        if header[12:16] == b"VP8X" and len(header) >= 30:
            width = 1 + int.from_bytes(header[24:27], "little")
            height = 1 + int.from_bytes(header[27:30], "little")
            return width, height, "WEBP"

    return None, None, None


def worksheet_to_markdown(sheet) -> str:
    # Excel 每个 sheet 先保守转成一个 Markdown 表格。
    rows = []
    for row in sheet.iter_rows(values_only=True):
        values = ["" if value is None else str(value).replace("\n", " ") for value in row]
        if any(value != "" for value in values):
            rows.append(values)

    if not rows:
        return "_此工作表无可见数据_"

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    divider = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def xlsx_shared_strings_from_archive(archive: zipfile.ZipFile) -> list[str]:
    # xlsx 里的字符串常常集中存在 sharedStrings.xml，需要先整体解开。
    try:
        xml_data = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []

    root = ET.fromstring(xml_data)
    values: list[str] = []
    for item in root.findall("main:si", XLSX_NAMESPACES):
        texts = [node.text or "" for node in item.findall(".//main:t", XLSX_NAMESPACES)]
        values.append("".join(texts))
    return values


def xlsx_sheet_names_from_workbook(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    # 从 workbook 和关系文件里拿到 sheet 名和对应的 sheet xml 路径。
    workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
    rel_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))

    rel_map = {}
    for rel in rel_root.findall("rel:Relationship", XLSX_NAMESPACES):
        rel_id = rel.get("Id")
        target = rel.get("Target")
        if rel_id and target:
            rel_map[rel_id] = target

    sheets: list[tuple[str, str]] = []
    for sheet in workbook_root.findall("main:sheets/main:sheet", XLSX_NAMESPACES):
        name = sheet.get("name", "Sheet")
        rel_id = sheet.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        target = rel_map.get(rel_id, "")
        if target:
            normalized_target = target.lstrip("/")
            if normalized_target.startswith("xl/"):
                sheet_path = normalized_target
            else:
                sheet_path = f"xl/{normalized_target}"
            sheets.append((name, sheet_path))
    return sheets


def column_letters_to_index(value: str) -> int:
    index = 0
    for char in value:
        if not char.isalpha():
            break
        index = index * 26 + (ord(char.upper()) - ord("A") + 1)
    return max(index - 1, 0)


def xlsx_cell_value(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.get("t")
    value_node = cell.find("main:v", XLSX_NAMESPACES)
    inline_node = cell.find("main:is", XLSX_NAMESPACES)

    if inline_node is not None:
        texts = [node.text or "" for node in inline_node.findall(".//main:t", XLSX_NAMESPACES)]
        return "".join(texts).strip()
    if value_node is None or value_node.text is None:
        return ""
    raw_value = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value)]
        except (ValueError, IndexError):
            return raw_value
    return raw_value


def xlsx_rows_from_sheet_xml(archive: zipfile.ZipFile, sheet_path: str, shared_strings: list[str]) -> list[list[str]]:
    root = ET.fromstring(archive.read(sheet_path))
    rows: list[list[str]] = []

    for row in root.findall("main:sheetData/main:row", XLSX_NAMESPACES):
        row_values: list[str] = []
        last_index = -1
        for cell in row.findall("main:c", XLSX_NAMESPACES):
            ref = cell.get("r", "")
            col_index = column_letters_to_index(ref)
            while len(row_values) < col_index:
                row_values.append("")
            value = xlsx_cell_value(cell, shared_strings).replace("\n", " ").strip()
            row_values.append(value)
            last_index = col_index
        if row_values and any(value != "" for value in row_values):
            rows.append(row_values)

    return rows


def rows_to_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return "_此工作表无可见数据_"

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    divider = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_xlsx_to_markdown_fallback(raw_path: Path) -> tuple[str, dict]:
    # 没有 openpyxl 时，xlsx 仍然可以作为 zip+xml 做保守解析。
    with zipfile.ZipFile(raw_path) as archive:
        shared_strings = xlsx_shared_strings_from_archive(archive)
        sheets = xlsx_sheet_names_from_workbook(archive)

        parts = [f"# {raw_path.stem}"]
        sheet_map: list[dict] = []

        for sheet_name, sheet_path in sheets:
            rows = xlsx_rows_from_sheet_xml(archive, sheet_path, shared_strings)
            parts.append(f"## 工作表: {sheet_name}")
            parts.append(rows_to_markdown_table(rows))
            sheet_map.append({
                "sheet_name": sheet_name,
                "row_count": len(rows),
                "sheet_path": sheet_path,
            })

    markdown = "\n\n".join(part for part in parts if part.strip()).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if sheet_map else "partial",
        "warnings": [] if sheet_map else ["xlsx_no_visible_sheets"],
        "location_map": {
            "type": "sheet_map",
            "sheets": sheet_map,
            "source_path": str(raw_path),
        },
    }


def convert_legacy_xls_to_markdown(raw_path: Path) -> tuple[str, dict]:
    # `.xls` 也先走“可见文本片段 + 明确警告”的保守路径。
    # 这样至少能把工作表名、列名、公式痕迹或可读单元格内容捞出一部分。
    binary_bytes = raw_path.read_bytes()
    snippets = extract_text_snippets_from_binary_document(binary_bytes)
    is_ole = is_probably_ole_document(raw_path)

    parts = [
        f"# {raw_path.stem}",
        "",
        "## 工作簿信息 / Workbook Info",
        "",
        f"- 原始格式: `.xls`",
        f"- 文件大小: {len(binary_bytes)} bytes",
        f"- OLE 容器: `{is_ole}`",
    ]

    if snippets:
        parts.extend([
            "",
            "## 可见文本片段 / Visible Text Snippets",
            "",
        ])
        for index, snippet in enumerate(snippets, start=1):
            parts.append(f"{index}. {snippet}")
    else:
        parts.extend([
            "",
            "## 可见文本片段 / Visible Text Snippets",
            "",
            "> 当前纯 Python fallback 未提取到可读工作簿文本。",
        ])

    markdown = "\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "partial" if snippets else "poor",
        "warnings": ["legacy_xls_binary_fallback"] if snippets else ["legacy_xls_no_text_snippets"],
        "location_map": {
            "type": "binary_snippet_map",
            "source_path": str(raw_path),
            "snippet_count": len(snippets),
            "is_ole_container": is_ole,
        },
    }


def convert_spreadsheet_to_markdown(raw_path: Path) -> tuple[str, dict]:
    # 电子表格优先保留 sheet 结构、表头和主要单元格内容。
    if raw_path.suffix.lower() == ".csv":
        return convert_csv_to_markdown(raw_path)
    if raw_path.suffix.lower() == ".xls":
        return convert_legacy_xls_to_markdown(raw_path)
    if openpyxl is None:
        return convert_xlsx_to_markdown_fallback(raw_path)

    workbook = openpyxl.load_workbook(str(raw_path), data_only=False)
    parts = [f"# {raw_path.stem}"]
    sheet_map: list[dict] = []

    for sheet in workbook.worksheets:
        parts.append(f"\n## 工作表: {sheet.title}\n")
        parts.append(worksheet_to_markdown(sheet))
        sheet_map.append({
            "sheet_name": sheet.title,
            "max_row": sheet.max_row,
            "max_column": sheet.max_column,
        })

    markdown = "\n\n".join(parts).strip() + "\n"
    metadata = {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if workbook.worksheets else "partial",
        "warnings": [] if workbook.worksheets else ["spreadsheet_no_worksheets"],
        "location_map": {
            "type": "sheet_map",
            "sheets": sheet_map,
            "source_path": str(raw_path),
        },
    }
    return markdown, metadata


def convert_csv_to_markdown(raw_path: Path) -> tuple[str, dict]:
    # CSV 不依赖 openpyxl，直接走标准库就够了。
    rows: list[list[str]] = []
    with raw_path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.reader(fh)
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append([cell.strip() for cell in row])

    parts = [f"# {raw_path.stem}", "", "## 工作表: CSV", ""]
    if rows:
        width = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (width - len(row)) for row in rows]
        header = normalized_rows[0]
        divider = ["---"] * width
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join(divider) + " |")
        for row in normalized_rows[1:]:
            parts.append("| " + " | ".join(row) + " |")
    else:
        parts.append("_CSV 文件没有可见数据_")

    markdown = "\n".join(parts).strip() + "\n"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "good" if rows else "partial",
        "warnings": [] if rows else ["csv_no_visible_rows"],
        "location_map": {
            "type": "sheet_map",
            "sheets": [{"sheet_name": "CSV", "row_count": len(rows)}],
            "source_path": str(raw_path),
        },
    }


def ocr_text_is_meaningful(text: str, min_length: int = 12) -> bool:
    normalized = normalize_binary_snippet_text(text)
    if len(normalized) < min_length:
        return False
    interesting_chars = [
        char for char in normalized
        if char.isalnum() or "\u4e00" <= char <= "\u9fff"
    ]
    return len(interesting_chars) >= max(6, min_length // 2)


def image_understanding_text_is_meaningful(text: str, min_length: int = 16) -> bool:
    normalized = normalize_binary_snippet_text(text)
    if len(normalized) < min_length:
        return False
    signal_chars = [
        char for char in normalized
        if char.isalnum() or "\u4e00" <= char <= "\u9fff"
    ]
    return len(signal_chars) >= max(8, min_length // 2)


def sanitize_asset_filename(value: str, default_stem: str = "asset") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = cleaned.strip("._")
    if not cleaned:
        cleaned = default_stem
    return stabilize_filename_component(cleaned, max_bytes=120, separator="_")


def markdown_image_target_candidates(target_value: str) -> list[str]:
    normalized = target_value.strip()
    if not normalized:
        return []
    if normalized.startswith("<") and normalized.endswith(">"):
        normalized = normalized[1:-1].strip()
    normalized = normalized.replace("\\)", ")").replace("\\(", "(")
    return [normalized]


def markdown_image_asset_extension(
    *,
    parsed_url,
    source_name: str,
    content_type: str | None = None,
) -> str:
    candidate = Path(unquote(parsed_url.path or source_name)).suffix.lower()
    if candidate:
        return candidate
    if content_type:
        guessed = mimetypes.guess_extension(content_type.split(";", 1)[0].strip())
        if guessed:
            return guessed.lower()
    fallback = Path(source_name).suffix.lower()
    return fallback or ".bin"


def build_markdown_asset_path(assets_dir: Path, source_record: dict, image_index: int, source_name: str) -> Path:
    stem = sanitize_asset_filename(Path(source_name).stem or f"image_{image_index}", default_stem=f"image_{image_index}")
    suffix = Path(source_name).suffix.lower() or ".bin"
    relative_dir = Path(source_record["source_id"])
    relative_path = relative_dir / f"{image_index:03d}_{stem}{suffix}"
    return assets_dir / relative_path


def is_certificate_verification_error(exc: Exception) -> bool:
    if isinstance(exc, ssl.SSLCertVerificationError):
        return True
    if isinstance(exc, urllib.error.URLError):
        reason = exc.reason
        if isinstance(reason, ssl.SSLCertVerificationError):
            return True
        if isinstance(reason, ssl.SSLError):
            return "CERTIFICATE_VERIFY_FAILED" in str(reason)
    return "CERTIFICATE_VERIFY_FAILED" in str(exc)


def download_url_with_optional_insecure_retry(
    target_value: str,
    *,
    allow_insecure_downloads: bool,
) -> tuple[bytes, str | None, str, list[str]]:
    try:
        with urllib.request.urlopen(target_value, timeout=20) as response:
            return (
                response.read(),
                response.headers.get("Content-Type"),
                "verified",
                [],
            )
    except Exception as exc:
        if not allow_insecure_downloads or not is_certificate_verification_error(exc):
            raise

    insecure_context = ssl._create_unverified_context()
    with urllib.request.urlopen(target_value, timeout=20, context=insecure_context) as response:
        return (
            response.read(),
            response.headers.get("Content-Type"),
            "insecure_retry",
            ["markdown_remote_image_download_used_insecure_retry"],
        )


def download_markdown_image_to_assets(
    *,
    target: Path,
    source_record: dict,
    raw_dir: Path,
    image_index: int,
    target_value: str,
    allow_insecure_downloads: bool = True,
) -> dict:
    assets_dir = raw_assets_dir_for_workspace(target, raw_dir)
    parsed = urlparse(target_value)
    source_name = Path(unquote(parsed.path)).name or f"image_{image_index}"

    content, content_type, download_mode, warnings = download_url_with_optional_insecure_retry(
        target_value,
        allow_insecure_downloads=allow_insecure_downloads,
    )

    suffix = markdown_image_asset_extension(parsed_url=parsed, source_name=source_name, content_type=content_type)
    asset_path = build_markdown_asset_path(
        assets_dir,
        source_record,
        image_index,
        f"{Path(source_name).stem}{suffix}",
    )
    ensure_directory(asset_path.parent)
    asset_path.write_bytes(content)
    asset_hash = hashlib.sha256(content).hexdigest()
    return {
        "storage_kind": "downloaded",
        "asset_path": asset_path,
        "asset_hash": asset_hash,
        "content_type": content_type,
        "download_mode": download_mode,
        "warnings": warnings,
    }


def resolve_markdown_local_image_path(raw_path: Path, target_value: str, raw_dir: Path) -> Path:
    parsed = urlparse(target_value)
    if parsed.scheme and parsed.scheme != "file":
        raise ValueError(f"Unsupported local image reference scheme: {target_value}")
    if parsed.scheme == "file":
        candidate = Path(unquote(parsed.path)).expanduser()
    else:
        candidate = (raw_path.parent / unquote(target_value)).expanduser()
    return ensure_path_within_raw_root(candidate, raw_dir, purpose="Markdown image reference")


def convert_markdown_embedded_image_to_section(
    *,
    target: Path,
    asset_path: Path,
    asset_label: str,
    alt_text: str,
    image_context: dict | None = None,
) -> tuple[list[str], dict]:
    image_markdown, image_metadata = convert_image_to_markdown(
        asset_path,
        target=target,
        image_context=image_context,
    )
    body_lines = [line.rstrip() for line in image_markdown.splitlines()]
    if body_lines and body_lines[0].startswith("# "):
        body_lines = body_lines[1:]
        if body_lines and not body_lines[0]:
            body_lines = body_lines[1:]

    lines = [
        f"## 内嵌图片 {asset_label}",
        "",
        f"- alt: {alt_text or '(empty)'}",
        f"- asset_path: {asset_path}",
        f"- extraction_quality: {image_metadata.get('extraction_quality', 'partial')}",
    ]
    lines.extend(body_lines if body_lines else ["> 图片存在，但当前未生成附加文本内容。"])
    return lines, image_metadata


def enrich_markdown_with_embedded_images(
    *,
    target: Path,
    source_record: dict,
    raw_path: Path,
    raw_text: str,
    allow_insecure_downloads: bool = True,
) -> tuple[str, dict]:
    raw_dir = resolve_workspace_raw_dir(target)
    assets_dir = raw_assets_dir_for_workspace(target, raw_dir)
    ensure_directory(assets_dir)

    warnings: list[str] = []
    image_records: list[dict] = []
    section_lines: list[str] = []
    extraction_quality = "good"
    used_downloads = False
    used_ocr = False

    matches = list(MARKDOWN_IMAGE_PATTERN.finditer(raw_text))
    if not matches:
        return normalize_text_content("markdown", raw_text), {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good",
            "warnings": [],
            "location_map": {
                "type": "line_map",
                "source_path": source_record["source_path"],
            },
        }

    for image_index, match in enumerate(matches, start=1):
        alt_text = (match.group("alt") or "").strip()
        target_value = (match.group("target") or "").strip()
        candidate_values = markdown_image_target_candidates(target_value)
        if not candidate_values:
            warnings.append(f"markdown_image_invalid_target:{image_index}")
            extraction_quality = "partial"
            section_lines.extend([
                f"## 内嵌图片 image_{image_index}",
                "",
                f"- alt: {alt_text or '(empty)'}",
                "> 源文件包含图片，但图片地址为空或无法解析。",
                "",
            ])
            continue

        candidate = candidate_values[0]
        parsed = urlparse(candidate)
        try:
            if parsed.scheme in {"http", "https"}:
                asset_result = download_markdown_image_to_assets(
                    target=target,
                    source_record=source_record,
                    raw_dir=raw_dir,
                    image_index=image_index,
                    target_value=candidate,
                    allow_insecure_downloads=allow_insecure_downloads,
                )
                used_downloads = True
                asset_path = asset_result["asset_path"]
            else:
                asset_path = resolve_markdown_local_image_path(raw_path, candidate, raw_dir)
                asset_result = {
                    "storage_kind": "local_raw",
                    "asset_path": asset_path,
                    "asset_hash": file_sha256(asset_path),
                    "content_type": None,
                    "download_mode": "local_raw",
                    "warnings": [],
                }

            section_block, image_metadata = convert_markdown_embedded_image_to_section(
                target=target,
                asset_path=asset_path,
                asset_label=f"image_{image_index}",
                alt_text=alt_text,
                image_context={
                    "markdown_source_path": source_record["source_path"],
                    "image_index": image_index,
                    "image_alt": alt_text,
                    "image_target": candidate,
                },
            )
            used_ocr = used_ocr or bool(image_metadata.get("location_map", {}).get("ocr", {}).get("used"))
            image_quality = image_metadata.get("extraction_quality", "partial")
            if image_quality in {"failed", "poor", "partial"}:
                extraction_quality = "partial"
            warnings.extend(image_metadata.get("warnings", []))
            warnings.extend(asset_result.get("warnings", []))
            section_lines.extend(section_block)
            section_lines.append("")
            image_records.append({
                "index": image_index,
                "alt": alt_text,
                "target": candidate,
                "storage_kind": asset_result["storage_kind"],
                "asset_path": str(asset_path),
                "asset_hash": asset_result["asset_hash"],
                "content_type": asset_result.get("content_type"),
                "download_mode": asset_result.get("download_mode"),
                "image_metadata": image_metadata.get("location_map", {}),
            })
        except Exception as exc:
            warnings.append(f"markdown_image_conversion_failed:{image_index}:{type(exc).__name__}")
            extraction_quality = "partial"
            section_lines.extend([
                f"## 内嵌图片 image_{image_index}",
                "",
                f"- alt: {alt_text or '(empty)'}",
                f"- target: {candidate}",
                "> 源文件包含图片，但内容暂时无法转换为文本。",
                "",
            ])
            image_records.append({
                "index": image_index,
                "alt": alt_text,
                "target": candidate,
                "storage_kind": "failed",
                "error": f"{type(exc).__name__}: {exc}",
            })

    normalized_markdown = normalize_text_content("markdown", raw_text)
    if section_lines:
        normalized_markdown = (
            normalized_markdown.rstrip()
            + "\n\n## 内嵌图片内容 / Embedded Image Content\n\n"
            + "\n".join(section_lines).strip()
            + "\n"
        )

    extraction_method = "python_only"
    if used_downloads:
        extraction_method += "+remote_assets"
    if used_ocr:
        extraction_method += "+tesseract"

    return normalized_markdown, {
        "content_format": "markdown",
        "extraction_method": extraction_method,
        "extraction_quality": extraction_quality,
        "warnings": sorted(set(warnings)),
        "location_map": {
            "type": "markdown_with_embedded_images",
            "source_path": source_record["source_path"],
            "image_count": len(image_records),
            "images": image_records,
        },
    }


def normalize_agent_assisted_image_result(hook_result: dict | None, min_confidence: float) -> tuple[str, dict]:
    if not isinstance(hook_result, dict):
        return "", {
            "used": False,
            "ok": False,
            "quality": "unavailable",
            "warnings": [],
            "confidence": 0.0,
            "reason": "image_to_text_unavailable",
            "summary": "",
        }

    confidence = coerce_float(hook_result.get("confidence", 0.0), 0.0)
    reason = str(hook_result.get("reason", "")).strip() or "image_to_text_agent_result"
    warnings = [str(item).strip() for item in hook_result.get("warnings", []) if str(item).strip()]
    if confidence < min_confidence:
        return "", {
            "used": True,
            "ok": False,
            "quality": "low_confidence",
            "warnings": sorted(set([*warnings, "image_to_text_low_confidence"])),
            "confidence": confidence,
            "reason": reason,
            "summary": normalize_binary_snippet_text(str(hook_result.get("summary", ""))),
        }

    extracted_text = normalize_binary_snippet_text(
        str(hook_result.get("extracted_text") or hook_result.get("text") or "")
    )
    summary = normalize_binary_snippet_text(str(hook_result.get("summary", "")))
    combined_text = extracted_text
    if summary and summary not in combined_text:
        combined_text = (
            f"摘要: {summary}\n\n{extracted_text}".strip()
            if extracted_text else summary
        )

    if not combined_text:
        return "", {
            "used": True,
            "ok": False,
            "quality": "empty",
            "warnings": sorted(set([*warnings, "image_to_text_no_text"])),
            "confidence": confidence,
            "reason": reason,
            "summary": summary,
        }

    quality = "good" if image_understanding_text_is_meaningful(combined_text) else "partial"
    if quality == "partial":
        warnings.append("image_to_text_low_signal")
    return combined_text, {
        "used": True,
        "ok": True,
        "quality": quality,
        "warnings": sorted(set(warnings)),
        "confidence": confidence,
        "reason": reason,
        "summary": summary,
    }


def run_agent_assisted_image_to_text(
    target: Path | None,
    raw_path: Path,
    image_context: dict | None = None,
) -> tuple[str, dict]:
    if target is None:
        return "", {
            "used": False,
            "ok": False,
            "quality": "disabled",
            "warnings": [],
            "confidence": 0.0,
            "reason": "image_to_text_target_missing",
            "summary": "",
        }

    automation_config = load_automation_target_config(load_workspace_config(target), "image_to_text")
    if not automation_config.get("enabled"):
        return "", {
            "used": False,
            "ok": False,
            "quality": "disabled",
            "warnings": [],
            "confidence": 0.0,
            "reason": "image_to_text_agent_disabled",
            "summary": "",
        }

    payload = {
        "task": "describe_image",
        "image_path": str(raw_path.resolve()),
        "image_name": raw_path.name,
        "image_context": image_context or {},
    }
    hook_result = run_json_automation_command(
        target=target,
        command=automation_config.get("command", []),
        payload=payload,
        timeout_seconds=automation_config.get("timeout_seconds", 45),
    )
    return normalize_agent_assisted_image_result(
        hook_result,
        automation_config.get("min_confidence", 0.8),
    )


def run_tesseract_ocr(raw_path: Path) -> tuple[str, dict]:
    # 图片 OCR 作为增强路径存在：
    # - 有 tesseract：尽量提取正文
    # - 无 tesseract：走元数据占位
    # 这里统一返回结构化结果，方便 normalized / 日志 / 测试共用。
    if not command_exists("tesseract"):
        return "", {
            "used": False,
            "ok": False,
            "quality": "missing",
            "warnings": ["tesseract_missing"],
            "details": "tesseract command not found in PATH",
        }

    completed = subprocess.run(
        ["tesseract", str(raw_path), "stdout", "--psm", "3"],
        check=False,
        capture_output=True,
        text=True,
    )
    raw_text = (completed.stdout or "").strip()
    normalized_text = normalize_binary_snippet_text(raw_text)

    if completed.returncode != 0:
        stderr_text = (completed.stderr or "").strip() or "unknown_tesseract_error"
        return "", {
            "used": True,
            "ok": False,
            "quality": "failed",
            "warnings": ["tesseract_ocr_failed"],
            "details": stderr_text,
        }

    if not normalized_text:
        return "", {
            "used": True,
            "ok": True,
            "quality": "empty",
            "warnings": ["tesseract_ocr_no_text"],
            "details": (completed.stderr or "").strip(),
        }

    quality = "good" if ocr_text_is_meaningful(normalized_text) else "partial"
    warnings: list[str] = []
    if quality == "partial":
        warnings.append("tesseract_ocr_low_signal")
    return normalized_text, {
        "used": True,
        "ok": True,
        "quality": quality,
        "warnings": warnings,
        "details": (completed.stderr or "").strip(),
    }


def convert_image_to_markdown(
    raw_path: Path,
    *,
    target: Path | None = None,
    image_context: dict | None = None,
) -> tuple[str, dict]:
    # 图片标准化采用“元数据始终保底，OCR 视环境增强”的策略。
    warnings: list[str] = []
    stat = raw_path.stat()
    metadata_lines = [
        f"# {raw_path.stem}",
        "",
        f"- 文件名: {raw_path.name}",
        f"- 文件大小: {stat.st_size} bytes",
    ]
    location_map = {
        "type": "image_metadata",
        "source_path": str(raw_path),
    }

    width = None
    height = None
    image_format = None
    image_mode = "unknown"
    exif = {}

    if Image is None:
        warnings.append("pillow_missing")
        width, height, image_format = image_size_from_binary(raw_path)
    else:
        try:
            with Image.open(raw_path) as image:
                width = image.width
                height = image.height
                image_mode = image.mode
                image_format = image.format or "unknown"
                if hasattr(image, "getexif"):
                    exif_data = image.getexif()
                    if exif_data:
                        exif = {str(key): str(value) for key, value in exif_data.items()}
        except Exception as exc:
            warnings.append("pillow_image_open_failed")
            width, height, image_format = image_size_from_binary(raw_path)
            image_mode = "unknown"

    metadata_lines.extend([
        f"- 尺寸: {width}x{height}" if width and height else "- 尺寸: unknown",
        f"- 模式: {image_mode}",
        f"- 格式: {image_format}" if image_format else "- 格式: unknown",
    ])

    if exif:
        metadata_lines.append("")
        metadata_lines.append("## EXIF")
        for key, value in sorted(exif.items()):
            metadata_lines.append(f"- {key}: {value}")

    ocr_text, ocr_result = run_tesseract_ocr(raw_path)
    warnings.extend(ocr_result.get("warnings", []))
    llm_text = ""
    llm_result = {
        "used": False,
        "ok": False,
        "quality": "disabled",
        "warnings": [],
        "confidence": 0.0,
        "reason": "image_to_text_not_attempted",
        "summary": "",
    }

    should_try_llm = (
        not ocr_text
        or ocr_result.get("quality") in {"missing", "failed", "partial", "empty"}
    )
    if should_try_llm:
        llm_text, llm_result = run_agent_assisted_image_to_text(
            target=target,
            raw_path=raw_path,
            image_context=image_context,
        )
        warnings.extend(llm_result.get("warnings", []))

    extraction_quality = "partial"
    combined_sections: list[str] = []
    if ocr_text:
        combined_sections.extend([
            "## OCR 文本 / OCR Text",
            "",
            ocr_text,
        ])
        extraction_quality = "good" if ocr_result.get("quality") == "good" else "partial"
    if llm_text:
        combined_sections.extend([
            "## LLM 图片理解 / LLM Image Understanding",
            "",
            llm_text,
        ])
        if llm_result.get("quality") == "good":
            extraction_quality = "good"

    if combined_sections:
        metadata_lines.extend([
            "",
            *combined_sections,
        ])
    else:
        if ocr_result.get("quality") == "failed":
            metadata_lines.extend([
                "",
                "> tesseract 已安装，但本次 OCR 执行失败，当前仅保留图片元数据。",
            ])
        elif ocr_result.get("quality") == "missing":
            metadata_lines.extend([
                "",
                "> 当前环境未检测到 tesseract，图片仅生成元数据级 normalized 文档。",
            ])
        else:
            metadata_lines.extend([
                "",
                "> 本次 OCR 未提取到稳定正文，当前保留图片元数据供后续人工或 Agent 处理。",
            ])
        if llm_result.get("used") and not llm_text:
            metadata_lines.extend([
                "",
                "> 已尝试使用 LLM 识别图片内容，但当前未得到可稳定落盘的文本结果。",
            ])

    markdown = "\n".join(metadata_lines).strip() + "\n"
    location_map["image"] = {
        "has_exif": bool(exif),
        "width": width,
        "height": height,
        "mode": image_mode,
        "format": image_format,
    }
    location_map["ocr"] = {
        "used": ocr_result.get("used", False),
        "ok": ocr_result.get("ok", False),
        "quality": ocr_result.get("quality"),
        "char_count": len(ocr_text),
    }
    location_map["llm_image_understanding"] = {
        "used": llm_result.get("used", False),
        "ok": llm_result.get("ok", False),
        "quality": llm_result.get("quality"),
        "char_count": len(llm_text),
        "confidence": llm_result.get("confidence", 0.0),
        "reason": llm_result.get("reason"),
        "summary": llm_result.get("summary", ""),
    }
    extraction_method = "python_only"
    if ocr_result.get("used"):
        extraction_method += "+tesseract"
    if llm_result.get("used"):
        extraction_method += "+agent_assisted"
    return markdown, {
        "content_format": "markdown",
        "extraction_method": extraction_method,
        "extraction_quality": extraction_quality,
        "warnings": warnings if warnings else ([] if (ocr_text or llm_text) else ["image_metadata_only"]),
        "location_map": location_map,
    }


def convert_unknown_source_to_placeholder(raw_path: Path, source_type: str) -> tuple[str, dict]:
    # 对暂不支持的格式，至少生成一个可追踪占位文档，而不是静默丢掉。
    markdown = (
        f"# {raw_path.stem}\n\n"
        f"- source_type: {source_type}\n"
        f"- source_path: {raw_path}\n\n"
        "> 当前版本暂不支持该格式的自动标准化，请等待后续转换器或使用 Agent 辅助处理。\n"
    )
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "poor",
        "warnings": [f"unsupported_source_type:{source_type}"],
        "location_map": {
            "type": "placeholder",
            "source_path": str(raw_path),
        },
    }


def convert_source_to_normalized_markdown(raw_path: Path, source_type: str) -> tuple[str, dict]:
    # 多格式标准化统一从这里分派，后续新增格式时只需要扩展这一层。
    if source_type in {"markdown", "plain_text"}:
        raw_text = raw_path.read_text(encoding="utf-8")
        return normalize_text_content(source_type, raw_text), {
            "content_format": "markdown",
            "extraction_method": "python_only",
            "extraction_quality": "good",
            "warnings": [],
        }
    if source_type == "pdf":
        return convert_pdf_to_markdown(raw_path)
    if source_type == "docx":
        return convert_docx_to_markdown(raw_path)
    if source_type == "doc":
        return convert_legacy_doc_to_markdown(raw_path)
    if source_type == "spreadsheet":
        return convert_spreadsheet_to_markdown(raw_path)
    if source_type == "image":
        return convert_image_to_markdown(raw_path)
    return convert_unknown_source_to_placeholder(raw_path, source_type)


def build_failed_conversion_placeholder(raw_path: Path, source_type: str, exc: Exception) -> tuple[str, dict]:
    # 某个转换器失败时，不让整次 ingest 中断，而是生成一个可追踪的失败占位文档。
    markdown = (
        f"# {raw_path.stem}\n\n"
        f"- source_type: {source_type}\n"
        f"- source_path: {raw_path}\n"
        f"- converter_error: {type(exc).__name__}: {exc}\n\n"
        "> 当前版本在标准化该文件时失败，已生成占位文档等待后续修复或 Agent 辅助处理。\n"
    )
    return markdown, {
        "content_format": "markdown",
        "extraction_method": "python_only",
        "extraction_quality": "failed",
        "warnings": [f"converter_error:{type(exc).__name__}", str(exc)],
        "location_map": {
            "type": "conversion_error",
            "source_path": str(raw_path),
        },
    }


def normalize_source_record(
    target: Path,
    source_record: dict,
    *,
    allow_insecure_downloads: bool = True,
) -> dict | None:
    # 这一层负责把“来源登记记录”转成“标准化记录”。
    # 这里是 normalized 层的统一入口：不同类型都尽量产出 Markdown 形态的标准文本。
    source_type = source_record["source_type"]
    if source_type in {"markdown", "plain_text"}:
        return normalize_markdown_or_text_record(
            target,
            source_record,
            allow_insecure_downloads=allow_insecure_downloads,
        )

    raw_path = resolve_source_record_path(target, source_record["source_path"])
    ensure_path_within_raw_root(raw_path, resolve_workspace_raw_dir(target), purpose="Source record")
    try:
        normalized_text, metadata = convert_source_to_normalized_markdown(raw_path, source_type)
    except Exception as exc:
        # 多格式转换器允许单文件失败，但不应该拖垮整个 ingest 批次。
        normalized_text, metadata = build_failed_conversion_placeholder(raw_path, source_type, exc)
    normalized_text = normalize_text_content("markdown", normalized_text)
    normalized_rel_path = Path("normalized") / f"{source_record['source_id']}.md"
    normalized_abs_path = target / normalized_rel_path
    normalized_abs_path.write_text(normalized_text, encoding="utf-8")

    # 标准化后的 hash 要单独记录，因为后面 chunk/claim 更关心“规范化后的稳定文本”。
    normalized_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    title = raw_path.stem
    location_map = dict(metadata.get("location_map", {}))
    location_map["source_path"] = source_record["source_path"]

    if location_map.get("type") == "line_map" and "normalized_line_range" not in location_map:
        line_count = len(normalized_text.splitlines()) or 1
        location_map["normalized_line_range"] = f"1-{line_count}"

    return {
        "source_id": source_record["source_id"],
        "source_type": source_type,
        "source_path": source_record["source_path"],
        "normalized_path": str(normalized_rel_path),
        "title": title,
        "language": "unknown",
        "content_format": metadata.get("content_format", "markdown"),
        "raw_hash": source_record["source_hash"],
        "normalized_hash": normalized_hash,
        "normalizer_version": "normalize_v1",
        "document_kind": "note",
        "structure_quality": "unknown",
        "chunk_strategy_hint": "heading_first",
        "extraction_method": metadata.get("extraction_method", "python_only"),
        "extraction_quality": metadata.get("extraction_quality", "partial"),
        "warnings": metadata.get("warnings", []),
        "location_map": location_map,
        "updated_at": utc_now_iso(),
    }


def estimate_token_count(text: str) -> int:
    # 这里先用一个很保守的近似估算：中文/英文混排时，按字符数粗估 token 数量。
    # 真实 tokenizer 以后可以替换这里，但 V1 先保证逻辑可跑、阈值可控。
    compact = text.strip()
    if not compact:
        return 0
    return max(1, len(compact) // 4)


def summarize_chunk_text(text: str, max_chars: int = 120) -> str:
    # 给 chunk 生成一个非常短的摘要，优先用于调试、人工检查和后续 review 界面。
    compact = " ".join(line.strip() for line in text.splitlines() if line.strip())
    if len(compact) <= max_chars:
        return compact
    return compact[: max_chars - 3].rstrip() + "..."


def sanitize_section_label(value: str) -> str:
    # section_path 需要既可读又稳定，这里做一层轻量清洗。
    compact = re.sub(r"\s+", " ", value.strip())
    return compact or "未命名章节"


def build_section_hierarchy(section_parts: list[str]) -> dict:
    cleaned_parts = [sanitize_section_label(part) for part in section_parts if sanitize_section_label(part)]
    current_label = cleaned_parts[-1] if cleaned_parts else "未命名章节"
    parent_parts = cleaned_parts[:-1]
    return {
        "section_path_parts": cleaned_parts,
        "section_title": current_label,
        "parent_section_path": " > ".join(parent_parts),
        "heading_level": len(cleaned_parts),
    }


def parse_section_path(section_path: str) -> dict:
    return build_section_hierarchy([part.strip() for part in section_path.split(">") if part.strip()])


def markdown_heading_match(line: str):
    return re.match(r"^(#{1,6})\s+(.+?)\s*$", line)


def strip_markdown_inline_formatting(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    cleaned = cleaned.replace("**", "").replace("__", "").replace("~~", "")
    cleaned = cleaned.strip("`*_ ")
    return cleaned.strip()


def detect_structure_block_type(text: str) -> tuple[str, dict]:
    stripped = text.strip()
    if not stripped:
        return "blank", {}
    heading = markdown_heading_match(stripped)
    if heading:
        return "heading", {
            "heading_level": len(heading.group(1)),
            "heading_text": strip_markdown_inline_formatting(heading.group(2)),
        }
    if stripped.startswith("```"):
        return "code_block", {"fence": "```"}
    if re.match(r"^\s{0,3}>\s+", stripped):
        return "blockquote", {}
    if re.match(r"^\s*([-*+])\s+", stripped):
        marker = re.match(r"^\s*([-*+])\s+", stripped).group(1)
        indent = len(text) - len(text.lstrip(" "))
        return "list_item", {"list_marker": marker, "list_indent": indent}
    if re.match(r"^\s*\d+[.)]\s+", stripped):
        indent = len(text) - len(text.lstrip(" "))
        return "list_item", {"list_marker": "ordered", "list_indent": indent}
    if "|" in stripped and re.match(r"^\s*\|?.+\|.+\|?\s*$", stripped):
        if re.fullmatch(r"\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*", stripped):
            return "table_separator", {}
        return "table_row", {}
    return "paragraph", {}


def build_structure_block_id(source_id: str, start_line: int, end_line: int, text: str) -> str:
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()[:10]
    return f"sb_{source_id}_{start_line}_{end_line}_{digest}"


def build_markdown_structure_blocks(normalized_record: dict, normalized_text: str) -> list[dict]:
    source_id = normalized_record["source_id"]
    normalized_path = normalized_record["normalized_path"]
    source_path = normalized_record["source_path"]
    lines = normalized_text.splitlines()
    blocks: list[dict] = []
    heading_stack: list[tuple[int, str, str]] = []
    in_code_fence = False
    current_paragraph: list[tuple[int, str]] = []
    previous_block_id: str | None = None

    def current_heading_parts() -> list[str]:
        return [item[1] for item in heading_stack]

    def append_block(
        block_type: str,
        block_lines: list[tuple[int, str]],
        attributes: dict | None = None,
        *,
        heading_parts: list[str] | None = None,
        parent_block_id: str | None = None,
    ) -> dict | None:
        nonlocal previous_block_id
        if not block_lines:
            return None
        raw_markdown = "\n".join(line for _, line in block_lines).rstrip()
        text = raw_markdown.strip()
        if not text:
            return None
        start_line = block_lines[0][0]
        end_line = block_lines[-1][0]
        block_id = build_structure_block_id(source_id, start_line, end_line, raw_markdown)
        resolved_parent_block_id = parent_block_id
        if resolved_parent_block_id is None and block_type != "heading" and heading_stack:
            resolved_parent_block_id = heading_stack[-1][2] or None
        block = {
            "structure_block_id": block_id,
            "source_id": source_id,
            "source_path": source_path,
            "normalized_path": normalized_path,
            "block_type": block_type,
            "text": text,
            "raw_markdown": raw_markdown,
            "heading_path_parts": list(heading_parts if heading_parts is not None else current_heading_parts()),
            "parent_block_id": resolved_parent_block_id,
            "previous_block_id": previous_block_id,
            "next_block_id": None,
            "children_block_ids": [],
            "start_line": start_line,
            "end_line": end_line,
            "attributes": attributes or {},
            "hash": hashlib.sha256(raw_markdown.encode("utf-8")).hexdigest(),
            "created_at": utc_now_iso(),
        }
        if previous_block_id and blocks:
            blocks[-1]["next_block_id"] = block_id
        blocks.append(block)
        previous_block_id = block_id
        return block

    def flush_paragraph() -> None:
        nonlocal current_paragraph
        if current_paragraph:
            append_block("paragraph", current_paragraph)
            current_paragraph = []

    for line_no, line in enumerate(lines, start=1):
        stripped = line.strip()

        if in_code_fence:
            current_paragraph.append((line_no, line))
            if stripped.startswith("```"):
                append_block("code_block", current_paragraph, {"fence": "```"})
                current_paragraph = []
                in_code_fence = False
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            current_paragraph = [(line_no, line)]
            in_code_fence = True
            continue

        if not stripped:
            flush_paragraph()
            continue

        block_type, attributes = detect_structure_block_type(line)
        if block_type == "heading":
            flush_paragraph()
            level = attributes["heading_level"]
            title = sanitize_section_label(attributes["heading_text"])
            heading_stack = [item for item in heading_stack if item[0] < level]
            heading_stack.append((level, title, ""))
            heading_parts = [item[1] for item in heading_stack]
            parent_heading_id = next((item[2] for item in reversed(heading_stack[:-1]) if item[2]), None)
            block = append_block(
                "heading",
                [(line_no, line)],
                attributes,
                heading_parts=heading_parts,
                parent_block_id=parent_heading_id,
            )
            if block is not None:
                heading_stack[-1] = (level, title, block["structure_block_id"])
                if parent_heading_id:
                    for existing in blocks:
                        if existing["structure_block_id"] == parent_heading_id:
                            existing["children_block_ids"].append(block["structure_block_id"])
                            break
            continue

        if block_type in {"list_item", "blockquote", "table_row", "table_separator"}:
            flush_paragraph()
            append_block(block_type, [(line_no, line)], attributes)
            continue

        current_paragraph.append((line_no, line))

    flush_paragraph()
    return blocks


def parse_markdown_table_cells(raw_markdown: str) -> list[str]:
    stripped = raw_markdown.strip().strip("|")
    return [strip_markdown_inline_formatting(cell) for cell in stripped.split("|")]


def extract_metadata_from_text(text: str) -> dict:
    cleaned = strip_markdown_inline_formatting(clean_claim_candidate_text(text))
    match = re.match(r"^([^：:]{1,24})[：:]\s*(.+)$", cleaned)
    if not match:
        return {}
    key = match.group(1).strip()
    value = match.group(2).strip()
    if not key or not value:
        return {}
    if any(token in key.lower() for token in ("http", "https", "file")):
        return {}
    return {key: value}


def build_evidence_block_id(source_id: str, structure_block_ids: list[str], start_line: int, text: str) -> str:
    raw = "|".join([source_id, *structure_block_ids, str(start_line), text])
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:12]
    return f"ev_{source_id}_{start_line}_{digest}"


def normalize_list_item_text(text: str) -> str:
    cleaned = re.sub(r"^\s*[-*+]\s+", "", text.strip())
    cleaned = re.sub(r"^\s*\d+[.)]\s+", "", cleaned)
    return strip_markdown_inline_formatting(cleaned)


def append_semantic_feature(
    features: list[dict],
    tag: str,
    category: str,
    strength: str,
    evidence: str,
) -> None:
    feature = {
        "tag": tag,
        "category": category,
        "strength": strength,
        "evidence": evidence,
    }
    if feature not in features:
        features.append(feature)


def semantic_features_for_evidence(
    text: str,
    block_kind: str,
    metadata: dict,
    section_path_parts: list[str],
    local_heading: str | None,
) -> list[dict]:
    features: list[dict] = []
    combined_context = " ".join([
        *[str(item) for item in section_path_parts],
        str(local_heading or ""),
        text,
    ])

    if block_kind in {"table_row", "metadata_line"}:
        append_semantic_feature(features, "rules", "structure", "strong", block_kind)
        append_semantic_feature(features, "reference_structure", "structure", "strong", block_kind)
    if block_kind == "code_example":
        append_semantic_feature(features, "cases", "structure", "strong", block_kind)
        append_semantic_feature(features, "example_structure", "structure", "strong", block_kind)
    if block_kind == "list_item_with_body":
        append_semantic_feature(features, "local_heading_body", "structure", "medium", block_kind)

    cells = metadata.get("cells") if isinstance(metadata, dict) else None
    if isinstance(cells, list) and len([cell for cell in cells if str(cell).strip()]) >= 2:
        append_semantic_feature(features, "reference_structure", "structure", "strong", "table_cells")

    if isinstance(metadata, dict):
        metadata_keys = [str(key).strip() for key in metadata if str(key).strip() and key != "cells"]
        if metadata_keys:
            append_semantic_feature(features, "metadata_fact", "structure", "strong", "metadata_keys")

    if any(marker in combined_context for marker in ("案例：", "示例：", "场景：")):
        append_semantic_feature(features, "cases", "text_pattern", "strong", "explicit_example_label")
    elif "案例" in combined_context:
        append_semantic_feature(features, "cases", "text_pattern", "weak", "case_marker")

    if any(marker in combined_context for marker in ("规则清单", "参数列表", "字段表", "配置项", "FAQ")):
        append_semantic_feature(features, "rules", "text_pattern", "strong", "reference_label")
    elif "规则" in combined_context:
        append_semantic_feature(features, "rules", "text_pattern", "weak", "rule_marker")

    if any(marker in combined_context for marker in ("步骤", "首先", "然后", "最后", "第一步", "第二步", "第三步", "流程", "方法")):
        append_semantic_feature(features, "procedural_language", "text_pattern", "medium", "procedure_marker")
    if any(marker in combined_context for marker in ("时间线", "起初", "随后", "后来", "历程", "演变")):
        append_semantic_feature(features, "temporal_language", "text_pattern", "medium", "temporal_marker")
    if "培训" in combined_context:
        append_semantic_feature(features, "training", "text_pattern", "weak", "training_marker")
    if "指标" in combined_context or "数据" in combined_context:
        append_semantic_feature(features, "metrics", "text_pattern", "weak", "metrics_marker")

    return features


def content_tags_from_semantic_features(features: list[dict]) -> list[str]:
    structure_only_tags = {"local_heading_body", "metadata_fact", "reference_structure", "example_structure"}
    tags: list[str] = []
    for feature in features:
        if not isinstance(feature, dict):
            continue
        tag = str(feature.get("tag", "")).strip()
        category = str(feature.get("category", "")).strip()
        if category == "structure" or tag in structure_only_tags:
            continue
        if tag and tag not in tags:
            tags.append(tag)
    return sorted(tags)


def build_evidence_blocks_from_structure(structure_blocks: list[dict]) -> list[dict]:
    evidence_blocks: list[dict] = []

    index = 0
    while index < len(structure_blocks):
        block = structure_blocks[index]
        block_type = block.get("block_type")
        if block_type in {"table_separator"}:
            index += 1
            continue

        structure_group = [block]
        block_kind = block_type
        text = str(block.get("text", "")).strip()
        local_heading = None
        metadata = dict(block.get("attributes", {}))
        extraction_hint = "single_structure_block"

        if block_type == "heading":
            block_kind = "section_heading"
            local_heading = strip_markdown_inline_formatting(text.lstrip("#").strip())
            metadata["heading_path_parts"] = block.get("heading_path_parts", [])
            extraction_hint = "heading_as_structure_context"

        elif block_type == "list_item":
            local_heading = normalize_list_item_text(text)
            next_block = structure_blocks[index + 1] if index + 1 < len(structure_blocks) else None
            if (
                next_block is not None
                and next_block.get("block_type") == "paragraph"
                and next_block.get("heading_path_parts") == block.get("heading_path_parts")
            ):
                structure_group.append(next_block)
                block_kind = "list_item_with_body"
                text = f"{local_heading}\n{next_block.get('text', '').strip()}"
                extraction_hint = "local_heading_attached_to_body"
                index += 1
            else:
                block_kind = "list_item"
                text = local_heading
                extraction_hint = "list_item_as_evidence"

        elif block_type == "table_row":
            block_kind = "table_row"
            metadata["cells"] = parse_markdown_table_cells(block.get("raw_markdown", text))
            extraction_hint = "table_row_as_evidence"

        elif block_type == "code_block":
            block_kind = "code_example"
            extraction_hint = "code_block_preserved_as_evidence"

        elif block_type == "paragraph":
            metadata.update(extract_metadata_from_text(text))
            block_kind = "metadata_line" if metadata and len(metadata) == 1 else "paragraph"
            extraction_hint = "metadata_extracted_from_paragraph" if block_kind == "metadata_line" else "paragraph_as_evidence"

        start_line = min(item["start_line"] for item in structure_group)
        end_line = max(item["end_line"] for item in structure_group)
        structure_block_ids = [item["structure_block_id"] for item in structure_group]
        evidence_id = build_evidence_block_id(block["source_id"], structure_block_ids, start_line, text)
        heading_path_parts = block.get("heading_path_parts", [])
        semantic_features = semantic_features_for_evidence(
            text=text,
            block_kind=block_kind,
            metadata=metadata,
            section_path_parts=heading_path_parts,
            local_heading=local_heading,
        )
        evidence_blocks.append({
            "evidence_block_id": evidence_id,
            "source_id": block["source_id"],
            "source_path": block["source_path"],
            "normalized_path": block["normalized_path"],
            "structure_block_ids": structure_block_ids,
            "block_kind": block_kind,
            "text": text,
            "local_heading": local_heading,
            "context_before": None,
            "context_after": None,
            "section_path_parts": heading_path_parts,
            "start_line": start_line,
            "end_line": end_line,
            "metadata": metadata,
            "semantic_features": semantic_features,
            "content_tags": content_tags_from_semantic_features(semantic_features),
            "extraction_hint": extraction_hint,
            "hash": hashlib.sha256(text.encode("utf-8")).hexdigest(),
            "created_at": utc_now_iso(),
        })
        index += 1

    return evidence_blocks


def build_knowledge_unit_id(source_id: str, evidence_block_ids: list[str], text: str) -> str:
    raw = "|".join([source_id, *evidence_block_ids, normalize_claim_text(text)])
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:12]
    return f"ku_{source_id}_{digest}"


def knowledge_unit_kind_for_evidence(evidence_block: dict) -> str:
    block_kind = evidence_block.get("block_kind")
    if block_kind == "metadata_line":
        return "metadata_fact"
    if block_kind == "table_row":
        return "table_fact"
    if block_kind in {"section_heading"}:
        return "structural_shell"
    if block_kind == "code_example":
        return "code_example"
    return "statement"


def build_knowledge_units_from_evidence(evidence_blocks: list[dict]) -> list[dict]:
    knowledge_units: list[dict] = []
    for evidence_block in evidence_blocks:
        text = str(evidence_block.get("text", "")).strip()
        if not text:
            continue
        unit_kind = knowledge_unit_kind_for_evidence(evidence_block)
        evidence_block_ids = [evidence_block["evidence_block_id"]]
        knowledge_units.append({
            "knowledge_unit_id": build_knowledge_unit_id(evidence_block["source_id"], evidence_block_ids, text),
            "source_id": evidence_block["source_id"],
            "source_path": evidence_block["source_path"],
            "normalized_path": evidence_block["normalized_path"],
            "text": text,
            "normalized_text": normalize_claim_text(text),
            "unit_kind": unit_kind,
            "local_heading": evidence_block.get("local_heading"),
            "metadata": {
                **evidence_block.get("metadata", {}),
                "section_path_parts": evidence_block.get("section_path_parts", []),
            },
            "evidence_block_ids": evidence_block_ids,
            "source_refs": [
                {
                    "source_id": evidence_block["source_id"],
                    "normalized_path": evidence_block["normalized_path"],
                    "start_line": evidence_block.get("start_line"),
                    "end_line": evidence_block.get("end_line"),
                }
            ],
            "extraction_reason": evidence_block.get("extraction_hint", "evidence_block_compiled"),
            "quality_label": "structural_shell" if unit_kind == "structural_shell" else "standalone",
            "status": "draft",
            "lifecycle_status": "active",
            "semantic_decision_ids": [],
            "semantic_projection": {
                "content_tags": evidence_block.get("content_tags", []),
                "semantic_features": evidence_block.get("semantic_features", []),
            },
            "created_at": utc_now_iso(),
            "updated_at": utc_now_iso(),
        })
    return knowledge_units


def compile_structure_knowledge_records(normalized_record: dict, normalized_text: str) -> dict:
    structure_blocks = build_markdown_structure_blocks(normalized_record, normalized_text)
    evidence_blocks = build_evidence_blocks_from_structure(structure_blocks)
    knowledge_units = build_knowledge_units_from_evidence(evidence_blocks)
    return {
        "source_id": normalized_record["source_id"],
        "structure_blocks": structure_blocks,
        "evidence_blocks": evidence_blocks,
        "knowledge_units": knowledge_units,
        "updated_at": utc_now_iso(),
    }


def split_markdown_blocks(section_lines: list[tuple[int, str]]) -> list[dict]:
    # 这里把章节文本拆成“块”，优先按空行断开，但尽量不切开 fenced code block。
    blocks: list[dict] = []
    current_lines: list[tuple[int, str]] = []
    in_code_fence = False

    for line_no, line in section_lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_fence = not in_code_fence

        if not in_code_fence and stripped == "":
            if current_lines:
                block_text = "\n".join(item[1] for item in current_lines).strip("\n")
                blocks.append({
                    "text": block_text,
                    "start_line": current_lines[0][0],
                    "end_line": current_lines[-1][0],
                })
                current_lines = []
            continue

        current_lines.append((line_no, line))

    if current_lines:
        block_text = "\n".join(item[1] for item in current_lines).strip("\n")
        blocks.append({
            "text": block_text,
            "start_line": current_lines[0][0],
            "end_line": current_lines[-1][0],
        })

    return [block for block in blocks if block["text"].strip()]


def split_normalized_into_sections(normalized_text: str) -> list[dict]:
    # 第一版章节切分只识别 Markdown 标题。
    # 没有标题的文档会落到一个默认章节里，保证任何文本都能继续往下游走。
    sections: list[dict] = []
    lines = normalized_text.splitlines()
    heading_stack: list[str] = []
    current_section = {
        "section_path": ["文档开始"],
        "lines": [],
    }

    for line_no, line in enumerate(lines, start=1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            if current_section["lines"]:
                sections.append(current_section)

            level = len(match.group(1))
            title = sanitize_section_label(match.group(2))
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(title)
            current_section = {
                "section_path": heading_stack.copy(),
                "lines": [(line_no, line)],
            }
            continue

        current_section["lines"].append((line_no, line))

    if current_section["lines"]:
        sections.append(current_section)

    return sections


def split_normalized_into_paragraph_sections(normalized_text: str) -> list[dict]:
    # paragraph_first 用于结构较弱或纯文本型文档：
    # 不依赖标题层级，而是把连续段落作为 section 候选。
    sections: list[dict] = []
    lines = normalized_text.splitlines()
    current_lines: list[tuple[int, str]] = []

    def flush_section() -> None:
        nonlocal current_lines
        if not current_lines:
            return
        preview = next((line.strip() for _, line in current_lines if line.strip()), "文档段落")
        sections.append({
            "section_path": [sanitize_section_label(preview[:24]) or "文档段落"],
            "lines": current_lines,
        })
        current_lines = []

    for line_no, line in enumerate(lines, start=1):
        if re.match(r"^(#{1,6})\s+(.+?)\s*$", line):
            flush_section()
            sections.append({
                "section_path": [sanitize_section_label(re.sub(r"^(#{1,6})\s+", "", line).strip())],
                "lines": [(line_no, line)],
            })
            continue
        if not line.strip():
            flush_section()
            continue
        current_lines.append((line_no, line))

    flush_section()
    return sections


def choose_sections_for_chunking(normalized_record: dict, normalized_text: str) -> list[dict]:
    chunk_strategy_hint = str(normalized_record.get("chunk_strategy_hint", "heading_first")).strip() or "heading_first"
    if chunk_strategy_hint in {"paragraph_first", "chat_turn"}:
        sections = split_normalized_into_paragraph_sections(normalized_text)
        if sections:
            return sections
    return split_normalized_into_sections(normalized_text)


def build_chunk_records_for_section(
    source_id: str,
    source_path: str,
    normalized_rel_path: str,
    section: dict,
    chunk_offset: int,
) -> list[dict]:
    # 这一层把单个 section 继续切成 chunk。
    # 策略先求稳定和可解释：优先按块累积，超过上限就落一个 chunk。
    blocks = split_markdown_blocks(section["lines"])
    if not blocks:
        return []

    max_chars = DEFAULT_CHUNK_MAX_TOKENS * 4
    min_chars = DEFAULT_CHUNK_MIN_TOKENS * 4
    target_chars = DEFAULT_CHUNK_TARGET_TOKENS * 4
    section_path = " > ".join(section["section_path"])

    grouped_blocks: list[list[dict]] = []
    current_group: list[dict] = []
    current_chars = 0

    for block in blocks:
        block_chars = len(block["text"])
        # 如果单块本身就很长，允许它独立成块，避免为了凑阈值把结构切得更碎。
        if current_group and current_chars + block_chars > max_chars:
            grouped_blocks.append(current_group)
            current_group = [block]
            current_chars = block_chars
            continue

        current_group.append(block)
        current_chars += block_chars

        if current_chars >= target_chars:
            grouped_blocks.append(current_group)
            current_group = []
            current_chars = 0

    if current_group:
        if grouped_blocks:
            current_text = "\n\n".join(item["text"] for item in current_group)
            if len(current_text) < min_chars:
                grouped_blocks[-1].extend(current_group)
            else:
                grouped_blocks.append(current_group)
        else:
            grouped_blocks.append(current_group)

    chunk_records: list[dict] = []
    for index, block_group in enumerate(grouped_blocks, start=chunk_offset):
        chunk_text = "\n\n".join(block["text"] for block in block_group).strip() + "\n"
        start_line = block_group[0]["start_line"]
        end_line = block_group[-1]["end_line"]
        chunk_hash = hashlib.sha256(chunk_text.encode("utf-8")).hexdigest()
        section_key = sanitize_source_key(section_path)
        chunk_id = f"chk_{source_id}_{section_key}_{start_line}_{chunk_hash[:10]}"
        section_hierarchy = build_section_hierarchy(section["section_path"])

        chunk_records.append({
            "chunk_id": chunk_id,
            "source_id": source_id,
            "source_path": source_path,
            "normalized_path": normalized_rel_path,
            "section_path": section_path,
            "section_path_parts": section_hierarchy["section_path_parts"],
            "section_title": section_hierarchy["section_title"],
            "parent_section_path": section_hierarchy["parent_section_path"],
            "heading_level": section_hierarchy["heading_level"],
            "chunk_index": index,
            "start_line": start_line,
            "end_line": end_line,
            "page_range": None,
            "char_count": len(chunk_text),
            "token_estimate": estimate_token_count(chunk_text),
            "summary": summarize_chunk_text(chunk_text),
            "text": chunk_text,
            "previous_chunk": None,
            "next_chunk": None,
            "overlap_from_previous": 0,
            "hash": chunk_hash,
            "chunker_version": "chunk_v2",
            "updated_at": utc_now_iso(),
        })

    return chunk_records


def build_chunk_records(normalized_record: dict, normalized_text: str) -> list[dict]:
    # 这里负责文档级切块：先按 section 拆，再给每段 section 分配 chunk 序号。
    sections = choose_sections_for_chunking(normalized_record, normalized_text)
    chunk_records: list[dict] = []
    chunk_index = 0

    for section in sections:
        section_chunks = build_chunk_records_for_section(
            source_id=normalized_record["source_id"],
            source_path=normalized_record["source_path"],
            normalized_rel_path=normalized_record["normalized_path"],
            section=section,
            chunk_offset=chunk_index,
        )
        chunk_records.extend(section_chunks)
        chunk_index += len(section_chunks)

    # previous / next 引用最后再统一回填，避免在切块阶段一边生成一边回看。
        for index, record in enumerate(chunk_records):
            previous_chunk = chunk_records[index - 1]["chunk_id"] if index > 0 else None
            next_chunk = chunk_records[index + 1]["chunk_id"] if index + 1 < len(chunk_records) else None
            record["previous_chunk"] = previous_chunk
            record["next_chunk"] = next_chunk
            record["chunk_kind"] = normalized_record.get("chunk_strategy_hint", "heading_first")
            record["topicworthiness_hint"] = normalized_record.get("document_kind", "note")

    return chunk_records


def write_source_chunks(target: Path, source_id: str, chunk_records: list[dict]) -> str:
    # chunks/ 目录里按 source_id 保存一份局部 JSONL，方便人工单独查看某个来源的切块结果。
    chunk_rel_path = Path("chunks") / f"{source_id}.jsonl"
    chunk_abs_path = target / chunk_rel_path
    write_jsonl(chunk_abs_path, chunk_records)
    return str(chunk_rel_path)


def format_chunk_reference(from_page: Path, source_id: str, chunk_ref: dict) -> str:
    # chunk 目前按 source_id 聚合存成 JSONL，这里把 chunk_id 链到对应文件，并附上段落定位信息。
    chunk_file_link = markdown_link_between_pages(from_page, Path("chunks") / f"{source_id}.jsonl")
    section_path = chunk_ref.get("section_path") or "unknown section"
    start_line = chunk_ref.get("start_line")
    end_line = chunk_ref.get("end_line")
    location = (
        f"{section_path} (lines {start_line}-{end_line})"
        if start_line is not None and end_line is not None
        else section_path
    )
    return f"[`{chunk_ref['chunk_id']}`]({chunk_file_link}) {location}"


def chunk_normalized_record(target: Path, normalized_record: dict) -> dict | None:
    # poor / failed 的文档暂时不进入稳定 chunk 流程，避免把低质量文本继续放大。
    if normalized_record["extraction_quality"] not in {"good", "partial"}:
        return None

    normalized_path = target / normalized_record["normalized_path"]
    normalized_text = normalized_path.read_text(encoding="utf-8")
    chunk_records = build_chunk_records(normalized_record, normalized_text)
    if not chunk_records:
        return None

    chunk_file_path = write_source_chunks(target, normalized_record["source_id"], chunk_records)
    for record in chunk_records:
        record["chunk_file_path"] = chunk_file_path

    return {
        "source_id": normalized_record["source_id"],
        "chunk_file_path": chunk_file_path,
        "chunk_count": len(chunk_records),
        "chunks": chunk_records,
        "updated_at": utc_now_iso(),
    }


def write_json(path: Path, payload: dict) -> None:
    # 单个 claim / review 文件用普通 JSON 保存，人工查看会比 JSONL 更舒服。
    atomic_write_text(path, json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def markdown_to_plain_text(text: str) -> str:
    # Claim 草稿抽取先基于“较干净的正文文本”进行。
    # 这里只做保守清理，不追求完美去 markdown。
    cleaned = text
    cleaned = re.sub(r"<!--.*?-->", " ", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r"^#{1,6}\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^\s*[-*+]\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^\s*\d+\.\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"`{1,3}", "", cleaned)
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    cleaned = cleaned.replace("**", "").replace("__", "").replace("*", "").replace("~~", "")
    normalized_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in cleaned.splitlines()]
    cleaned = "\n".join(normalized_lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def strip_fenced_code_blocks(text: str) -> str:
    # 示例代码块里的 YAML/JSON/命令通常不是正文知识陈述，不应直接进入 claim 抽取。
    lines: list[str] = []
    in_code_fence = False

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_fence = not in_code_fence
            continue
        if in_code_fence:
            continue
        lines.append(line)

    return "\n".join(lines)


def normalize_heading_plus_body_claim_candidate(text: str) -> str:
    # Markdown 标题和正文在压平成单行后，常变成“Claim 是什么 Claim 是...”这种重复前缀。
    cleaned = clean_concept_title_text(text)
    if not cleaned:
        return ""

    suffix_match = re.match(r"^(.{1,32}?)\s*(?:是|指)?什么\s+(.+)$", cleaned, flags=re.IGNORECASE)
    if suffix_match:
        label = clean_concept_title_text(suffix_match.group(1))
        remainder = clean_concept_title_text(suffix_match.group(2))
        if label and remainder.startswith(label):
            return remainder

    prefix_match = re.match(r"^什么是\s+(.{1,32}?)\s+(.+)$", cleaned, flags=re.IGNORECASE)
    if prefix_match:
        label = clean_concept_title_text(prefix_match.group(1))
        remainder = clean_concept_title_text(prefix_match.group(2))
        if label and remainder.startswith(label):
            return remainder

    return text


def normalize_claim_text(text: str) -> str:
    # Claim 的规范文本用于去重、冲突判断和稳定生成 claim_id。
    cleaned = markdown_to_plain_text(text).lower()
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = cleaned.strip(" -:;,.!?。！？；：()[]{}\"'")
    return cleaned


def clean_claim_candidate_text(text: str) -> str:
    # 候选 claim 在进入规则判断前先做一轮轻量清洗：
    # - 去掉 Markdown 标题/引用/列表符号
    # - 去掉常见编号前缀
    # - 压缩多余空白
    cleaned = text.strip()
    cleaned = re.sub(r"<!--.*?-->", " ", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r"^\s{0,3}#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^\s*>\s*", "", cleaned)
    cleaned = re.sub(r"^\s*[-*+]\s*", "", cleaned)
    cleaned = re.sub(r"^\s*\d+[.)、:：]\s*", "", cleaned)
    cleaned = re.sub(r"^\s*[（(]?\d+[）)]\s*", "", cleaned)
    cleaned = re.sub(r"^\s*(因此|所以|同时|此外|另外|不过|但是|而且|并且|而是)\s*", "", cleaned)
    cleaned = normalize_heading_plus_body_claim_candidate(cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip(" -:;,.!?。！？；：，、()[]{}\"'")


def claim_candidate_has_short_gray_zone(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return False
    natural_chars = [
        char for char in cleaned
        if char.isalnum() or "\u4e00" <= char <= "\u9fff"
    ]
    if len(natural_chars) >= 10:
        return False
    if claim_candidate_is_noise(cleaned):
        return False
    return True


def text_is_iso_date_label(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    return bool(re.fullmatch(r"\d{4}-\d{2}-\d{2}", cleaned))


def text_is_question_like(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return False
    if cleaned.endswith(("？", "?")):
        return True
    return any(cleaned.startswith(prefix) for prefix in ("问题", "为什么", "如何", "怎么", "是否", "什么是"))


def claim_starts_with_dependent_prefix(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return False
    return any(cleaned.startswith(prefix) for prefix in CLAIM_DEPENDENT_PREFIXES)


def claim_starts_with_meta_prefix(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return False
    return any(cleaned.startswith(prefix) for prefix in CLAIM_META_PREFIXES)


def claim_has_standalone_predicate(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return False
    return any(marker in cleaned for marker in CLAIM_STANDALONE_PREDICATE_MARKERS)


def claim_can_stand_alone(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if claim_candidate_is_noise(cleaned):
        return False
    if claim_starts_with_dependent_prefix(cleaned):
        return False
    return claim_has_standalone_predicate(cleaned) or cleaned[:1].isalnum() or "\u4e00" <= cleaned[:1] <= "\u9fff"


def claim_is_definition_like_phrase(text: str) -> bool:
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return False
    if claim_starts_with_dependent_prefix(cleaned) or claim_starts_with_meta_prefix(cleaned):
        return False
    if len(cleaned) > 48:
        return False
    return bool(re.match(r"^(一种|一类|一个|一套|一组|一条|一项).+", cleaned))


def claim_candidate_is_noise(text: str) -> bool:
    # 这里过滤几类高噪声片段：
    # - 纯链接 / 文件路径味太重
    # - 表格分隔线
    # - 几乎没有自然语言内容的标题或目录碎片
    # 不再单纯因为“长度小于 12”就直接判噪声，避免误杀短而完整的中文陈述。
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return True
    if cleaned.startswith(("http://", "https://", "file://")):
        return True
    if any(marker in cleaned.lower() for marker in ("turn_id", "speaker:", "time:")):
        return True
    if re.match(r"^[A-Za-z][A-Za-z0-9_ -]{0,24}\s*:\s*", cleaned):
        return True
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", cleaned):
        return True
    if re.search(r"\b[A-Za-z][A-Za-z0-9_ -]{0,24}\s*:\s*", cleaned):
        return True
    if re.fullmatch(r"[-|: ]{3,}", cleaned):
        return True
    if cleaned.count("/") >= 3 and len(cleaned) < 48:
        return True
    if cleaned.lower().startswith(("raw/", "../raw/", "wiki/", "claims/", "chunks/", "normalized/")):
        return True

    natural_chars = [
        char for char in cleaned
        if char.isalnum() or "\u4e00" <= char <= "\u9fff"
    ]
    if len(natural_chars) < 4:
        return True
    if (
        len(natural_chars) < 8
        and not claim_has_standalone_predicate(cleaned)
        and not claim_is_definition_like_phrase(cleaned)
    ):
        return True
    return False


def split_long_claim_candidate(text: str, max_chars: int = 140) -> list[str]:
    # 很长的整段经常会把多个结论糊在一起。
    # 这里优先按中文逗号、顿号、分句连接词再切一次，但只取足够像独立陈述的片段。
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return []

    has_multi_clause_signal = any(
        marker in cleaned
        for marker in ("因此", "所以", "同时", "此外", "另外", "但是", "不过", "而且", "并且")
    ) and any(marker in cleaned for marker in ("，", ",", "；", ";"))
    if len(cleaned) <= max_chars and not has_multi_clause_signal:
        return [cleaned]

    secondary_parts = re.split(r"(?<=[，,；;。！？!?])\s*|(?=(?:但是|不过|因此|所以|同时|此外|另外|而且|并且|而是))", cleaned)
    refined_parts: list[str] = []
    for raw_part in secondary_parts:
        part = clean_claim_candidate_text(raw_part)
        if claim_candidate_is_noise(part):
            continue
        refined_parts.append(part)

    if len(refined_parts) >= 2:
        return refined_parts
    return [cleaned]


def split_claim_candidates_from_text(text: str) -> list[str]:
    # 这里优先抽取“像一句完整陈述”的片段：
    # 先按句号/分号/换行切，再做一轮候选清洗、去噪和长句拆分。
    # 中文资料里常见“句号后不加空格”的连写，这里也要能正常断句。
    candidates: list[str] = []
    raw_pieces = re.split(r"(?<=[。！？!?；;])\s*|\n{1,}|(?<=\.)\s{2,}", text)

    for raw_piece in raw_pieces:
        piece = clean_claim_candidate_text(raw_piece)
        if claim_candidate_is_noise(piece):
            continue
        sentence_candidates = [piece]
        refined_parts = split_long_claim_candidate(piece)
        if len(refined_parts) >= 2:
            sentence_candidates.extend(
                refined_piece
                for refined_piece in refined_parts
                if refined_piece != piece and claim_can_stand_alone(refined_piece)
            )

        for candidate_text in sentence_candidates:
            normalized_piece = clean_claim_candidate_text(candidate_text)
            if claim_candidate_is_noise(normalized_piece):
                continue
            if normalized_piece in candidates:
                continue
            candidates.append(normalized_piece)

    # 如果按句切之后一个都没留下，至少保留整段，避免 chunk 完全失去 claim 草稿。
    if not candidates and text.strip():
        fallback_piece = clean_claim_candidate_text(text.strip())
        if not claim_candidate_is_noise(fallback_piece):
            return [fallback_piece]
        return []
    return candidates


def classify_claim_type(text: str) -> str:
    # 先给 Claim 一个启发式类型，后面接入 Agent 时可以被重写或提升。
    cleaned = clean_claim_candidate_text(text)
    lowered = cleaned.lower()
    if any(keyword in text for keyword in ("注意", "警告", "风险", "不要", "禁止")):
        return "warning"
    if any(keyword in text for keyword in ("步骤", "做法", "如何", "怎么", "先", "然后")):
        return "procedure"
    if any(keyword in text for keyword in ("因为", "因此", "导致", "使得", "原因")):
        return "causal"
    if any(keyword in text for keyword in ("相比", "对比", "区别", "优于", "弱于")):
        return "comparison"
    if claim_is_definition_like_phrase(cleaned):
        return "definition"
    if any(keyword in text for keyword in ("是", "是指", "定义", "叫做")):
        return "definition"
    if any(keyword in lowered for keyword in ("better", "worse", "useful", "important", "effective")):
        return "evaluation"
    return "fact"


def format_claim_type_label(claim_type: str | None) -> str:
    # 用代码样式展示 claim 类型，避免 Markdown 方括号在部分查看器里被误解成可点击引用。
    return f"`{claim_type or 'unknown'}`"


def format_claim_reference(from_page: Path, claim_record: dict) -> str:
    # 优先把 claim_id 渲染成可跳转到 claims/*.json 的相对链接，方便沿证据链继续下钻。
    claim_id = claim_record["claim_id"]
    claim_file = claim_record.get("claim_file_path")
    if not claim_file:
        return f"`{claim_id}`"
    link = markdown_link_between_pages(from_page, Path(claim_file))
    return f"[`{claim_id}`]({link})"


def format_workspace_file_reference(from_page: Path, path_str: str) -> str:
    # 原始来源、标准化文件、chunk 文件统一渲染成工作区内相对链接，便于直接点开查看。
    link = markdown_link_between_pages(from_page, Path(path_str))
    return f"[`{path_str}`]({link})"


def format_page_label(from_page: Path, page_record: dict) -> str:
    link = markdown_link_between_pages(from_page, Path(page_record["page_path"]))
    return f"[{page_record['title']}]({link})"


def format_source_page_label(from_page: Path, source_page: dict) -> str:
    return format_page_label(from_page, source_page)


def format_source_page_meta(source_page: dict | None, source_ref: dict) -> str:
    # 内部 ID 仍然保留，但放到次级信息里，避免压过真正对人有用的标题和来源路径。
    parts = []
    if source_page is not None:
        parts.append(f"page=`{source_page['page_id']}`")
    parts.append(f"source=`{source_ref['source_id']}`")
    return ", ".join(parts)


def claim_lifecycle_status_for_record(claim_record: dict) -> str:
    # lifecycle_status 负责表达“这条 claim 现在是否仍然活跃可用”，
    # 不与 query 里的 draft / needs_review 混在一起。
    if claim_record.get("lifecycle_status") == "superseded":
        return "superseded"
    if claim_record.get("archived_at"):
        return "archived"
    if not claim_record.get("source_ids") or not claim_record.get("source_refs"):
        return "superseded"
    return "active"


def review_lifecycle_status_for_record(review_record: dict) -> str:
    if review_record.get("lifecycle_status") == "superseded":
        return "superseded"
    if review_record.get("archived_at"):
        return "archived"
    if not review_record.get("candidate_claim_ids") and not review_record.get("candidate_page_ids"):
        return "superseded"
    return "active"


def page_lifecycle_status_for_record(page_record: dict) -> str:
    if page_record.get("removed"):
        return "removed"
    if page_record.get("archived_at"):
        return "archived"
    return "active"


def ensure_claim_lifecycle_defaults(claim_record: dict) -> dict:
    claim_record.setdefault("superseded_by", [])
    claim_record.setdefault("archived_at", None)
    claim_record.setdefault("quality_label", None)
    claim_record.setdefault("quality_reason", None)
    claim_record.setdefault("quality_confidence", None)
    claim_record.setdefault("quality_review_required", False)
    claim_record.setdefault("quality_safe_auto_ready", None)
    claim_record.setdefault("quality_decision_source", None)
    claim_record["lifecycle_status"] = claim_lifecycle_status_for_record(claim_record)
    return sync_claim_semantic_projection(claim_record)


def ensure_review_lifecycle_defaults(review_record: dict) -> dict:
    review_record.setdefault("archived_at", None)
    review_record["lifecycle_status"] = review_lifecycle_status_for_record(review_record)
    return review_record


def ensure_page_lifecycle_defaults(page_record: dict) -> dict:
    # page 这一层和 claim/review 不同：
    # 我们希望在 state/pages.jsonl 里保留“曾经存在过但后来被移除”的自动页面痕迹，
    # 这样后续做页面历史、反向追踪和人工恢复时有抓手。
    page_record.setdefault("removed", False)
    page_record.setdefault("archived_at", None)
    page_record.setdefault("semantic_decision_ids", [])
    page_record.setdefault("page_route", {})
    page_record["lifecycle_status"] = page_lifecycle_status_for_record(page_record)
    return page_record


def is_live_page_record(page_record: dict) -> bool:
    # live page 指“当前应参与索引、检索、目录展示”的在线页面。
    # removed / archived 页面仍可保留在 state/pages.jsonl 里，但不进入在线视图。
    lifecycle_status = page_record.get("lifecycle_status")
    if lifecycle_status in {"removed", "archived"}:
        return False
    return not page_record.get("removed", False)


def filter_live_page_records(page_records: list[dict]) -> list[dict]:
    # 统一从完整页面账本中过滤出在线页面，避免 query / index / wiki index
    # 各自手写一遍过滤条件，后续语义更容易保持一致。
    return [record for record in page_records if is_live_page_record(record)]


def is_live_claim_record(claim_record: dict) -> bool:
    # claim 的在线态比 page 更严格：
    # 既要 lifecycle 是 active，也要仍然保有可用的 source/source_ref 追踪链。
    return (
        claim_record.get("lifecycle_status") == "active"
        and bool(claim_record.get("source_ids"))
        and bool(claim_record.get("source_refs"))
    )


def filter_live_claim_records(claim_records: list[dict]) -> list[dict]:
    return [record for record in claim_records if is_live_claim_record(record)]


def filter_live_stable_claim_records(claim_records: list[dict]) -> list[dict]:
    # 可读概念页只消费当前仍活跃、且已经被提升为 stable 的 claim。
    return [
        record for record in claim_records
        if is_live_claim_record(record) and record.get("status") == "stable"
    ]


def is_live_review_record(review_record: dict) -> bool:
    # review 只有在仍然挂着候选 claim、且 lifecycle 为 active 时，
    # 才应继续进入概念页和后续人工处理视图。
    if review_record.get("lifecycle_status") != "active":
        return False
    return bool(review_record.get("candidate_claim_ids") or review_record.get("candidate_page_ids"))


def filter_live_review_records(review_records: list[dict]) -> list[dict]:
    return [record for record in review_records if is_live_review_record(record)]


def build_ordered_claim_state_records(
    live_claims_by_id: dict[str, dict],
    historical_claims_by_id: dict[str, dict],
) -> list[dict]:
    # state/claims.jsonl 既保存在线 claim，也保存历史态 claim。
    # 这里统一做一次稳定排序，方便 diff 和排查。
    records = [*live_claims_by_id.values(), *historical_claims_by_id.values()]
    return sorted(
        records,
        key=lambda item: (
            item.get("created_at", ""),
            item.get("updated_at", ""),
            item.get("claim_id", ""),
        ),
    )


def build_ordered_review_state_records(
    live_reviews_by_id: dict[str, dict],
    historical_reviews_by_id: dict[str, dict],
) -> list[dict]:
    # review 账本里同一个 review_id 不应同时出现 live + historical 两份。
    # 这里优先保留 live 记录，并吞掉重复历史态。
    deduped_records_by_id = dict(historical_reviews_by_id)
    deduped_records_by_id.update(live_reviews_by_id)
    records = list(deduped_records_by_id.values())
    return sorted(
        records,
        key=lambda item: (
            item.get("created_at", ""),
            item.get("resolved_at", "") or "",
            item.get("review_id", ""),
        ),
    )


def load_claim_state_maps(target: Path) -> tuple[dict[str, dict], dict[str, dict], list[dict]]:
    # review 命令需要同时看到在线 claim 和历史 claim。
    # 这里统一加载，避免 list/apply 两个命令重复写一遍状态拆分逻辑。
    claim_records = [
        ensure_claim_lifecycle_defaults(record)
        for record in load_jsonl(target / "state" / "claims.jsonl")
    ]
    live_claims = {record["claim_id"]: record for record in filter_live_claim_records(claim_records)}
    historical_claims = {
        record["claim_id"]: record
        for record in claim_records
        if not is_live_claim_record(record)
    }
    return live_claims, historical_claims, claim_records


def load_review_state_maps(target: Path) -> tuple[dict[str, dict], dict[str, dict], list[dict]]:
    review_records = [
        ensure_review_lifecycle_defaults(record)
        for record in load_jsonl(target / "state" / "reviews.jsonl")
    ]
    live_reviews = {record["review_id"]: record for record in filter_live_review_records(review_records)}
    historical_reviews = {
        record["review_id"]: record
        for record in review_records
        if not is_live_review_record(record)
    }
    return live_reviews, historical_reviews, review_records


def build_claim_lookup_by_any_id(
    live_claims_by_id: dict[str, dict],
    historical_claims_by_id: dict[str, dict],
) -> dict[str, dict]:
    # review 历史里经常还保留“原始 claim_id”，
    # 因此这里同时支持按当前主键和 original_claim_id 反查。
    lookup = {}
    for record in [*live_claims_by_id.values(), *historical_claims_by_id.values()]:
        lookup[record["claim_id"]] = record
        original_claim_id = record.get("original_claim_id")
        if original_claim_id:
            lookup.setdefault(original_claim_id, record)
    return lookup


def claim_display_id(claim_record: dict) -> str:
    # 历史态 claim 的主键会加 __hist_ 后缀，展示给用户时优先露出原始 claim_id，
    # 这样 review 操作时不容易看花眼。
    return claim_record.get("original_claim_id") or claim_record["claim_id"]


def review_display_id(review_record: dict) -> str:
    return review_record.get("original_review_id") or review_record["review_id"]


def is_actionable_review_record(review_record: dict) -> bool:
    # 只有仍然 open 的 review 才应该继续影响概念页状态和后续人工待办。
    return is_live_review_record(review_record) and review_record.get("status") == "open"


def build_historical_claim_id(claim_record: dict) -> str:
    # 同一路径 source 原位更新后，新旧 claim 可能共享原始 claim_id。
    # 历史态记录需要搬到单独的命名空间里，避免和新一轮活跃 claim 撞 ID。
    archived_at = claim_record.get("archived_at") or utc_now_iso()
    archived_suffix = re.sub(r"[^0-9]", "", archived_at)[:20] or datetime.now().strftime("%Y%m%d%H%M%S")
    original_claim_id = claim_record.get("original_claim_id") or claim_record["claim_id"]
    return f"{original_claim_id}__hist_{archived_suffix}"


def build_historical_review_id(review_record: dict) -> str:
    archived_at = review_record.get("archived_at") or utc_now_iso()
    archived_suffix = re.sub(r"[^0-9]", "", archived_at)[:20] or datetime.now().strftime("%Y%m%d%H%M%S")
    original_review_id = review_record.get("original_review_id") or review_record["review_id"]
    return f"{original_review_id}__hist_{archived_suffix}"


def convert_claim_record_to_historical(claim_record: dict) -> dict:
    # 历史态 claim 仍保留原始 claim_id 供追踪，但 state/file 主键切换为历史态 ID。
    archived_record = dict(claim_record)
    archived_record["original_claim_id"] = archived_record.get("original_claim_id") or archived_record["claim_id"]
    archived_record["claim_id"] = build_historical_claim_id(archived_record)
    archived_record["claim_file_path"] = str(Path("claims") / f"{archived_record['claim_id']}.json")
    return archived_record


def convert_review_record_to_historical(review_record: dict) -> dict:
    archived_record = dict(review_record)
    archived_record["original_review_id"] = archived_record.get("original_review_id") or archived_record["review_id"]
    archived_record["review_id"] = build_historical_review_id(archived_record)
    archived_record["review_file_path"] = str(Path("reviews") / f"{archived_record['review_id']}.json")
    return archived_record


def has_negation(text: str) -> bool:
    lowered = text.lower()
    return any(marker in lowered for marker in NEGATION_MARKERS)


def normalize_claim_base_for_conflict(text: str) -> str:
    # 冲突判断先用一个粗糙但稳定的“去否定标记”版本做基线。
    normalized = normalize_claim_text(text)
    for marker in NEGATION_MARKERS:
        normalized = normalized.replace(marker.strip(), "")
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def build_similarity_bucket(text: str) -> str:
    # 这个 bucket 现在主要服务于概念页聚合与缺页补齐。
    # 它仍然保持“稳定、粗粒度、低成本”的特点，不承担最终相似判定责任。
    normalized = normalize_claim_text(text)
    return normalized[:24]


def build_claim_similarity_tokens(text: str) -> list[str]:
    # claim review 的相似检测要比概念页 bucket 更细一点：
    # - 先去掉否定词，让“需要/不需要”这类冲突句仍能进入同一候选池
    # - 再复用 query 的中英混合切词逻辑，保证检索和审核尽量共享一套词法直觉
    base_text = normalize_claim_base_for_conflict(text)
    seen: set[str] = set()
    tokens: list[str] = []

    for token in tokenize_for_search(base_text):
        cleaned_token = token.strip()
        if len(cleaned_token) < 2:
            continue
        if cleaned_token in seen:
            continue
        seen.add(cleaned_token)
        tokens.append(cleaned_token)
    return tokens


def claim_similarity_token_weight(token: str) -> float:
    # 更长的 token 往往语义更具体，给它更高一点的权重。
    # 这里故意保持简单，避免引入太多难以解释的启发式参数。
    latin_or_number = bool(re.fullmatch(r"[a-z0-9_]+", token))
    if latin_or_number:
        return max(1.0, min(len(token), 8) / 2.0)
    return float(min(len(token), 4))


def compute_weighted_token_overlap(left_tokens: list[str], right_tokens: list[str]) -> tuple[float, float, int]:
    # overlap_ratio 看“共享语义片段占较短句子的比例”；
    # jaccard_ratio 看“双方整体重合度”；
    # 两者一起用，能减少“只共享开头几个泛词”导致的误报。
    left_set = set(left_tokens)
    right_set = set(right_tokens)
    if not left_set or not right_set:
        return 0.0, 0.0, 0

    shared_tokens = left_set & right_set
    shared_weight = sum(claim_similarity_token_weight(token) for token in shared_tokens)
    left_weight = sum(claim_similarity_token_weight(token) for token in left_set)
    right_weight = sum(claim_similarity_token_weight(token) for token in right_set)

    overlap_ratio = shared_weight / max(1.0, min(left_weight, right_weight))
    jaccard_ratio = shared_weight / max(1.0, left_weight + right_weight - shared_weight)
    return overlap_ratio, jaccard_ratio, len(shared_tokens)


def measure_claim_text_similarity(left_text: str, right_text: str) -> dict:
    # 这里把“候选召回”和“最终是否足够相似”拆开：
    # 候选召回尽量宽一点，最终判定则结合字符序列和 token 重合度做保守收敛。
    left_base = normalize_claim_base_for_conflict(left_text)
    right_base = normalize_claim_base_for_conflict(right_text)
    left_tokens = build_claim_similarity_tokens(left_text)
    right_tokens = build_claim_similarity_tokens(right_text)
    overlap_ratio, jaccard_ratio, shared_token_count = compute_weighted_token_overlap(left_tokens, right_tokens)
    if left_base and right_base:
        matcher = difflib.SequenceMatcher(None, left_base, right_base)
        sequence_ratio = matcher.ratio()
        longest_common_span = max((block.size for block in matcher.get_matching_blocks()), default=0)
    else:
        sequence_ratio = 0.0
        longest_common_span = 0
    longest_common_span_ratio = longest_common_span / max(1, min(len(left_base), len(right_base)))

    shorter_base, longer_base = sorted([left_base, right_base], key=len)
    containment = bool(shorter_base) and len(shorter_base) >= 12 and shorter_base in longer_base

    return {
        "left_base": left_base,
        "right_base": right_base,
        "left_tokens": left_tokens,
        "right_tokens": right_tokens,
        "overlap_ratio": overlap_ratio,
        "jaccard_ratio": jaccard_ratio,
        "shared_token_count": shared_token_count,
        "sequence_ratio": sequence_ratio,
        "longest_common_span": longest_common_span,
        "longest_common_span_ratio": longest_common_span_ratio,
        "containment": containment,
    }


def claims_are_similar_for_review(left_text: str, right_text: str) -> bool:
    # 这一层不追求“语义理解”，只做 V1 足够稳定的近重复/冲突前置筛选：
    # - 完全同 base：直接视为同主题
    # - 一方基本包含另一方：通常是“扩写版/带前缀版”
    # - 其余情况需要同时满足字符近似 + token 重合，避免误把同领域句子都撞进 review
    metrics = measure_claim_text_similarity(left_text, right_text)
    left_base = metrics["left_base"]
    right_base = metrics["right_base"]
    if not left_base or not right_base:
        return False
    if left_base == right_base:
        return True
    if metrics["containment"]:
        return True
    if metrics["sequence_ratio"] >= 0.90:
        return True
    if (
        metrics["longest_common_span_ratio"] >= 0.62
        and metrics["shared_token_count"] >= 6
    ):
        return True
    if (
        metrics["sequence_ratio"] >= 0.72
        and metrics["overlap_ratio"] >= 0.60
        and metrics["shared_token_count"] >= 4
    ):
        return True
    if (
        metrics["overlap_ratio"] >= 0.82
        and metrics["jaccard_ratio"] >= 0.55
        and metrics["shared_token_count"] >= 5
    ):
        return True
    return False


def index_claim_similarity_tokens(
    similarity_index: dict[str, set[str]],
    claim_record: dict,
) -> None:
    # 这里维护一个很轻量的 token -> claim_id 倒排索引，
    # 让“前缀不同但核心短语相同”的 claim 也能被召回进入后续精判。
    for token in build_claim_similarity_tokens(claim_record.get("text", "")):
        similarity_index.setdefault(token, set()).add(claim_record["claim_id"])


def rebuild_claim_similarity_index(claim_records: list[dict]) -> dict[str, set[str]]:
    similarity_index: dict[str, set[str]] = {}
    for claim_record in claim_records:
        index_claim_similarity_tokens(similarity_index, claim_record)
    return similarity_index


def collect_claim_review_candidate_ids(
    claim_record: dict,
    claims_by_similarity_bucket: dict[str, list[dict]],
    claim_similarity_index: dict[str, set[str]],
) -> set[str]:
    # 候选召回分两路：
    # 1) 老的 bucket，成本低、兼容现有概念页聚合
    # 2) 新的 token 倒排，补足“句首不同但主体相同”的情况
    candidate_claim_ids: set[str] = set()
    similarity_bucket = build_similarity_bucket(claim_record["text"])
    candidate_claim_ids.update(
        item["claim_id"]
        for item in claims_by_similarity_bucket.get(similarity_bucket, [])
    )
    for token in build_claim_similarity_tokens(claim_record.get("text", "")):
        candidate_claim_ids.update(claim_similarity_index.get(token, set()))
    candidate_claim_ids.discard(claim_record["claim_id"])
    return candidate_claim_ids


def append_unique(items: list[str], value: str) -> None:
    if value not in items:
        items.append(value)


def sync_claim_semantic_projection(claim_record: dict) -> dict:
    updated = dict(claim_record)
    projection = dict(updated.get("semantic_projection") or {})
    projection["knowledge_role"] = updated.get("knowledge_role")
    projection["page_intent_hints"] = list(updated.get("page_intent_hints", []) or [])
    projection["concept_candidate_score"] = coerce_float(updated.get("concept_candidate_score", 0.0), 0.0)
    if updated.get("quality_label") is not None:
        projection["quality_label"] = updated.get("quality_label")
    if updated.get("quality_reason") is not None:
        projection["quality_reason"] = updated.get("quality_reason")
    if updated.get("quality_confidence") is not None:
        projection["quality_confidence"] = updated.get("quality_confidence")
    if updated.get("quality_safe_auto_ready") is not None:
        projection["quality_safe_auto_ready"] = updated.get("quality_safe_auto_ready")
    if updated.get("quality_review_required") is not None:
        projection["quality_review_required"] = updated.get("quality_review_required")
    updated["semantic_projection"] = projection
    updated.setdefault("semantic_decision_ids", [])
    return updated


def claim_semantic_projection(claim_record: dict) -> dict:
    projection = dict(claim_record.get("semantic_projection") or {})
    if "knowledge_role" not in projection:
        projection["knowledge_role"] = claim_record.get("knowledge_role")
    if "page_intent_hints" not in projection:
        projection["page_intent_hints"] = list(claim_record.get("page_intent_hints", []) or [])
    if "concept_candidate_score" not in projection:
        projection["concept_candidate_score"] = coerce_float(claim_record.get("concept_candidate_score", 0.0), 0.0)
    return projection


def claim_knowledge_role(claim_record: dict) -> str:
    return str(claim_semantic_projection(claim_record).get("knowledge_role") or "").strip().lower()


def claim_page_intent_hints(claim_record: dict) -> list[str]:
    return [
        str(item).strip().lower()
        for item in claim_semantic_projection(claim_record).get("page_intent_hints", []) or []
        if str(item).strip()
    ]


def claim_concept_candidate_score(claim_record: dict) -> float:
    return coerce_float(claim_semantic_projection(claim_record).get("concept_candidate_score", 0.0), 0.0)


def build_claim_record_from_chunk(chunk_record: dict, claim_text: str) -> dict:
    # 单条 Claim 草稿要把溯源线索一开始就带全，后面 page / review 都直接复用。
    normalized_text = normalize_claim_text(claim_text)
    claim_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    claim_id = f"clm_{chunk_record['source_id']}_{claim_hash[:12]}"
    now = utc_now_iso()
    section_path_parts = chunk_record.get("section_path_parts")
    if not section_path_parts:
        section_path_parts = parse_section_path(chunk_record.get("section_path", "")).get("section_path_parts", [])

    return sync_claim_semantic_projection({
        "claim_id": claim_id,
        "text": claim_text.strip(),
        "normalized_text": normalized_text,
        "claim_type": classify_claim_type(claim_text),
        "knowledge_role": None,
        "page_intent_hints": [],
        "concept_candidate_score": 0.0,
        "quality_label": None,
        "quality_reason": None,
        "quality_confidence": None,
        "quality_review_required": False,
        "quality_safe_auto_ready": None,
        "quality_decision_source": None,
        "status": "draft",
        "lifecycle_status": "active",
        "source_ids": [chunk_record["source_id"]],
        "knowledge_unit_ids": [],
        "evidence_block_ids": [],
        "chunk_ids": [chunk_record["chunk_id"]],
        "page_ids": [],
        "conflict_group": None,
        "duplicate_candidates": [],
        "review_reason": None,
        "superseded_by": [],
        "archived_at": None,
        "source_refs": [
            {
                "source_id": chunk_record["source_id"],
                "source_path": chunk_record["source_path"],
                "normalized_path": chunk_record["normalized_path"],
                "chunk_id": chunk_record["chunk_id"],
                "section_path": chunk_record["section_path"],
                "section_path_parts": section_path_parts,
                "section_title": chunk_record.get("section_title") or (section_path_parts[-1] if section_path_parts else ""),
                "parent_section_path": chunk_record.get("parent_section_path") or " > ".join(section_path_parts[:-1]),
                "heading_level": chunk_record.get("heading_level") or len(section_path_parts),
                "start_line": chunk_record["start_line"],
                "end_line": chunk_record["end_line"],
            }
        ],
        "extraction_method": "rule_based_chunk_v2",
        "created_at": now,
        "updated_at": now,
    })


def find_covering_chunk_for_knowledge_unit(knowledge_unit: dict, chunk_records: list[dict]) -> dict | None:
    source_refs = knowledge_unit.get("source_refs", [])
    first_ref = source_refs[0] if source_refs else {}
    start_line = first_ref.get("start_line")
    end_line = first_ref.get("end_line")
    normalized_path = knowledge_unit.get("normalized_path")
    if start_line is None or end_line is None:
        return None

    overlapping_chunks = []
    for chunk_record in chunk_records:
        if chunk_record.get("normalized_path") != normalized_path:
            continue
        chunk_start = chunk_record.get("start_line")
        chunk_end = chunk_record.get("end_line")
        if chunk_start is None or chunk_end is None:
            continue
        if chunk_start <= start_line and chunk_end >= end_line:
            return chunk_record
        if chunk_start <= end_line and chunk_end >= start_line:
            overlapping_chunks.append(chunk_record)
    return overlapping_chunks[0] if overlapping_chunks else None


def build_claim_record_from_knowledge_unit(
    knowledge_unit: dict,
    claim_text: str,
    chunk_record: dict | None = None,
) -> dict:
    normalized_text = normalize_claim_text(claim_text)
    claim_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    claim_id = f"clm_{knowledge_unit['source_id']}_{claim_hash[:12]}"
    now = utc_now_iso()
    chunk_id = chunk_record.get("chunk_id") if chunk_record else None
    chunk_section_parts = chunk_record.get("section_path_parts", []) if chunk_record else []
    if not isinstance(chunk_section_parts, list):
        chunk_section_parts = []
    source_refs = []
    for source_ref in knowledge_unit.get("source_refs", []):
        section_path_parts = knowledge_unit.get("metadata", {}).get("section_path_parts")
        if not section_path_parts:
            section_path_parts = chunk_section_parts
        source_refs.append({
            "source_id": knowledge_unit["source_id"],
            "source_path": knowledge_unit.get("source_path"),
            "normalized_path": knowledge_unit.get("normalized_path"),
            "chunk_id": chunk_id,
            "knowledge_unit_id": knowledge_unit["knowledge_unit_id"],
            "evidence_block_ids": knowledge_unit.get("evidence_block_ids", []),
            "section_path": chunk_record.get("section_path") if chunk_record else " > ".join(section_path_parts),
            "section_path_parts": section_path_parts,
            "section_title": (
                chunk_record.get("section_title")
                if chunk_record
                else section_path_parts[-1] if section_path_parts else ""
            ),
            "parent_section_path": (
                chunk_record.get("parent_section_path")
                if chunk_record
                else " > ".join(section_path_parts[:-1])
            ),
            "heading_level": chunk_record.get("heading_level") if chunk_record else len(section_path_parts),
            "start_line": source_ref.get("start_line"),
            "end_line": source_ref.get("end_line"),
        })
    if not source_refs:
        source_refs.append({
            "source_id": knowledge_unit["source_id"],
            "source_path": knowledge_unit.get("source_path"),
            "normalized_path": knowledge_unit.get("normalized_path"),
            "chunk_id": chunk_id,
            "knowledge_unit_id": knowledge_unit["knowledge_unit_id"],
            "evidence_block_ids": knowledge_unit.get("evidence_block_ids", []),
            "section_path": chunk_record.get("section_path") if chunk_record else "",
            "section_path_parts": chunk_section_parts,
            "section_title": chunk_record.get("section_title") if chunk_record else "",
            "parent_section_path": chunk_record.get("parent_section_path") if chunk_record else "",
            "heading_level": chunk_record.get("heading_level") if chunk_record else 0,
            "start_line": None,
            "end_line": None,
        })

    return sync_claim_semantic_projection({
        "claim_id": claim_id,
        "text": claim_text.strip(),
        "normalized_text": normalized_text,
        "claim_type": classify_claim_type(claim_text),
        "knowledge_role": None,
        "page_intent_hints": [],
        "concept_candidate_score": 0.0,
        "quality_label": None,
        "quality_reason": None,
        "quality_confidence": None,
        "quality_review_required": False,
        "quality_safe_auto_ready": None,
        "quality_decision_source": None,
        "status": "draft",
        "lifecycle_status": "active",
        "source_ids": [knowledge_unit["source_id"]],
        "knowledge_unit_ids": [knowledge_unit["knowledge_unit_id"]],
        "evidence_block_ids": list(knowledge_unit.get("evidence_block_ids", [])),
        "chunk_ids": [chunk_id] if chunk_id else [],
        "page_ids": [],
        "conflict_group": None,
        "duplicate_candidates": [],
        "review_reason": None,
        "superseded_by": [],
        "archived_at": None,
        "source_refs": source_refs,
        "extraction_method": "rule_based_knowledge_unit_v1",
        "created_at": now,
        "updated_at": now,
    })


def merge_claim_records(existing_record: dict, incoming_record: dict) -> dict:
    # 如果规范文本完全一致，就把它们视为同一 claim，并合并溯源关系。
    merged = dict(existing_record)
    for source_id in incoming_record["source_ids"]:
        append_unique(merged["source_ids"], source_id)
    for knowledge_unit_id in incoming_record.get("knowledge_unit_ids", []):
        append_unique(merged.setdefault("knowledge_unit_ids", []), knowledge_unit_id)
    for evidence_block_id in incoming_record.get("evidence_block_ids", []):
        append_unique(merged.setdefault("evidence_block_ids", []), evidence_block_id)
    for decision_id in incoming_record.get("semantic_decision_ids", []):
        append_unique(merged.setdefault("semantic_decision_ids", []), decision_id)
    for chunk_id in incoming_record["chunk_ids"]:
        append_unique(merged["chunk_ids"], chunk_id)
    for page_id in incoming_record.get("page_ids", []):
        append_unique(merged["page_ids"], page_id)

    existing_ref_keys = {
        (
            item.get("source_id"),
            item.get("chunk_id"),
            item.get("knowledge_unit_id"),
            tuple(item.get("evidence_block_ids", [])),
            item.get("start_line"),
            item.get("end_line"),
        )
        for item in merged.get("source_refs", [])
    }
    for source_ref in incoming_record.get("source_refs", []):
        ref_key = (
            source_ref.get("source_id"),
            source_ref.get("chunk_id"),
            source_ref.get("knowledge_unit_id"),
            tuple(source_ref.get("evidence_block_ids", [])),
            source_ref.get("start_line"),
            source_ref.get("end_line"),
        )
        if ref_key not in existing_ref_keys:
            merged.setdefault("source_refs", []).append(source_ref)
            existing_ref_keys.add(ref_key)

    if not merged.get("knowledge_role") and incoming_record.get("knowledge_role"):
        merged["knowledge_role"] = incoming_record.get("knowledge_role")
    incoming_intents = incoming_record.get("page_intent_hints", [])
    existing_intents = list(merged.get("page_intent_hints", []))
    for hint in incoming_intents:
        append_unique(existing_intents, hint)
    merged["page_intent_hints"] = existing_intents
    merged["concept_candidate_score"] = max(
        coerce_float(merged.get("concept_candidate_score", 0.0), 0.0),
        coerce_float(incoming_record.get("concept_candidate_score", 0.0), 0.0),
    )
    for field in (
        "quality_label",
        "quality_reason",
        "quality_confidence",
        "quality_review_required",
        "quality_safe_auto_ready",
        "quality_decision_source",
    ):
        incoming_value = incoming_record.get(field)
        if incoming_value is not None:
            merged[field] = incoming_value
    merged["updated_at"] = utc_now_iso()
    merged["lifecycle_status"] = claim_lifecycle_status_for_record(merged)
    return sync_claim_semantic_projection(merged)


def claim_file_path(target: Path, claim_id: str) -> Path:
    return target / "claims" / f"{claim_id}.json"


def write_claim_file(target: Path, claim_record: dict) -> str:
    # claim 文件是权威源；state/claims.jsonl 是便于扫描和索引的派生索引。
    claim_path = claim_file_path(target, claim_record["claim_id"])
    write_json(claim_path, claim_record)
    return str(Path("claims") / claim_path.name)


def review_file_path(target: Path, review_id: str) -> Path:
    return target / "reviews" / f"{review_id}.json"


def build_review_record(
    kind: str,
    candidate_claim_ids: list[str],
    reason: str,
    evidence: list[dict],
    recommended_action: str,
    signature_parts: list[str] | None = None,
) -> dict:
    # review item 尽量自解释：为什么进入审核、建议动作是什么、证据链在哪里。
    if signature_parts:
        signature = "|".join(signature_parts) + "|" + kind
    else:
        signature = "|".join(sorted(candidate_claim_ids)) + "|" + kind
    review_hash = hashlib.sha256(signature.encode("utf-8")).hexdigest()
    review_id = f"rev_{review_hash[:12]}"
    now = utc_now_iso()
    return {
        "review_id": review_id,
        "kind": kind,
        "status": "open",
        "lifecycle_status": "active",
        "candidate_claim_ids": sorted(candidate_claim_ids),
        "candidate_page_ids": [],
        "reason": reason,
        "recommended_action": recommended_action,
        "allowed_actions": ["merge", "keep_both", "archive_one", "edit_then_resume"],
        "resume_from": "claim_review",
        "evidence": evidence,
        "created_at": now,
        "resolved_at": None,
        "archived_at": None,
    }


def write_review_file(target: Path, review_record: dict) -> str:
    review_path = review_file_path(target, review_record["review_id"])
    write_json(review_path, review_record)
    return str(Path("reviews") / review_path.name)


def append_error_record(
    error_log_path: Path,
    task_id: str,
    source_id: str,
    stage: str,
    level: str,
    message: str,
    details: dict | None = None,
) -> dict:
    # 错误日志和 warning 日志统一写这里，后续 review / 报表 / 追查都能复用。
    record = {
        "task_id": task_id,
        "source_id": source_id,
        "stage": stage,
        "level": level,
        "message": message,
        "details": details or {},
        "created_at": utc_now_iso(),
    }
    append_jsonl(error_log_path, record)
    return record


def build_claims_from_chunk(chunk_record: dict) -> list[dict]:
    # 一个 chunk 里可能有多个可提取陈述；这里保留通过规则筛出的全部候选，
    # 避免长段落后半段的独立结论被前几条候选提前截断。
    plain_text = markdown_to_plain_text(strip_fenced_code_blocks(chunk_record["text"]))
    claim_candidates = split_claim_candidates_from_text(plain_text)
    claim_records: list[dict] = []

    for candidate_text in claim_candidates:
        claim_records.append(build_claim_record_from_chunk(chunk_record, candidate_text))

    if claim_records:
        return claim_records

    # 某些时间线型文档正文可能几乎全是对话/元数据噪声，但章节标题本身仍值得沉淀成概念入口。
    section_path = chunk_record.get("section_path", "")
    section_parts = [part.strip() for part in section_path.split(">") if part.strip()]
    if section_parts:
        fallback_label = clean_concept_title_text(section_parts[-1])
        if text_is_iso_date_label(fallback_label):
            return [build_claim_record_from_chunk(chunk_record, fallback_label)]

    return claim_records


def build_claims_from_knowledge_unit(knowledge_unit: dict, chunk_records: list[dict] | None = None) -> list[dict]:
    unit_kind = str(knowledge_unit.get("unit_kind", "")).strip()
    if unit_kind in {"structural_shell", "code_example"}:
        return []

    text = str(knowledge_unit.get("text", "")).strip()
    if not text:
        return []

    local_heading = str(knowledge_unit.get("local_heading") or "").strip()
    if local_heading and text.startswith(local_heading):
        body_text = text[len(local_heading):].strip()
        if body_text:
            text = f"{local_heading}：{body_text}"

    covering_chunk = find_covering_chunk_for_knowledge_unit(knowledge_unit, chunk_records or [])
    plain_text = markdown_to_plain_text(strip_fenced_code_blocks(text))
    claim_candidates = split_claim_candidates_from_text(plain_text)
    claim_records = [
        build_claim_record_from_knowledge_unit(knowledge_unit, candidate_text, covering_chunk)
        for candidate_text in claim_candidates
    ]
    if claim_records:
        return claim_records

    if unit_kind in {"metadata_fact", "table_fact"}:
        cleaned_text = clean_claim_candidate_text(text)
        if cleaned_text and not claim_candidate_is_noise(cleaned_text):
            return [build_claim_record_from_knowledge_unit(knowledge_unit, cleaned_text, covering_chunk)]

    return []


def build_claim_candidates_for_source(
    source_id: str,
    knowledge_units_by_source_id: dict[str, list[dict]],
    chunks_by_source_id: dict[str, list[dict]],
) -> list[dict]:
    knowledge_units = knowledge_units_by_source_id.get(source_id, [])
    if knowledge_units:
        claim_records: list[dict] = []
        source_chunks = chunks_by_source_id.get(source_id, [])
        for knowledge_unit in knowledge_units:
            claim_records.extend(build_claims_from_knowledge_unit(knowledge_unit, source_chunks))
        if claim_records:
            return claim_records

    claim_records = []
    for chunk_record in chunks_by_source_id.get(source_id, []):
        claim_records.extend(build_claims_from_chunk(chunk_record))
    return claim_records


def source_claim_stage_completed(source_record: dict) -> bool:
    # 只要来源已经进入 claimed / review_required / generated，
    # 就说明这个 source 的 chunk -> claim 抽取已经跑过一轮了。
    # 在当前“同内容文件不会重复导入为同一 source”的模型下，
    # 后续重复 ingest 不需要再把这个 source 的所有 chunk 全量重抽一遍 claim。
    return source_record.get("status") in {"claimed", "review_required", "generated"}


def workspace_can_skip_page_regeneration(
    sources_by_id: dict[str, dict],
    created_sources: list[dict],
    normalized_sources: list[dict],
    chunked_sources: list[dict],
    claims_created_by_source: dict[str, int],
    review_items: list[dict],
    semantic_claim_updates_applied: bool = False,
) -> bool:
    # 这是一个“无上游变化”的保守短路条件：
    # - 没有新 source
    # - 没有新 normalized/chunk/claim/review
    # - 没有仅由语义账本带来的 claim 角色/页面意图提示变化
    # - 所有来源都已经走到 generated 或 failed
    # 满足这些条件时，source-summary / concept / search index 在语义上都不该变化。
    if (
        created_sources
        or normalized_sources
        or chunked_sources
        or claims_created_by_source
        or review_items
        or semantic_claim_updates_applied
    ):
        return False
    return all(record.get("status") in {"generated", "failed"} for record in sources_by_id.values())


def choose_active_source_ids(sources_by_id: dict[str, dict]) -> set[str]:
    # 同一路径的原始文件在多次 ingest 后，可能会形成多个 source 版本。
    # 页面生成优先围绕“最新且未失败的版本”展开；如果某一路径最新版本失败了，
    # 就暂时回退到该路径最后一个未失败版本，避免因为一次转换失败把现有 wiki 整页抹掉。
    grouped: dict[str, list[dict]] = {}
    for record in sources_by_id.values():
        grouped.setdefault(record["source_path"], []).append(record)

    active_source_ids: set[str] = set()
    for source_path, records in grouped.items():
        sorted_records = sorted(records, key=lambda item: item.get("imported_at", ""))
        non_failed_records = [
            record for record in sorted_records
            if record.get("status") != "failed"
        ]
        chosen_record = non_failed_records[-1] if non_failed_records else sorted_records[-1]
        active_source_ids.add(chosen_record["source_id"])
    return active_source_ids


def expected_source_summary_page_id(source_id: str) -> str:
    return f"page_src_{source_id}"


def build_workspace_overview_page_id() -> str:
    return "page_ovw_workspace"


def workspace_overview_page_path() -> Path:
    return Path("wiki") / "overview" / "index.md"


def collect_missing_source_page_source_ids(
    active_source_ids: set[str],
    sources_by_id: dict[str, dict],
    page_records_by_id: dict[str, dict],
    claims_by_source_id: dict[str, list[dict]],
    chunks_by_source_id: dict[str, list[dict]],
) -> set[str]:
    # 如果某个来源已经有证据层产物，但缺少来源摘要页，就把它视为待补齐页面。
    missing_source_ids: set[str] = set()
    for source_id, source_record in sources_by_id.items():
        if source_id not in active_source_ids:
            continue
        if source_record.get("status") == "failed":
            continue
        if not claims_by_source_id.get(source_id) and not chunks_by_source_id.get(source_id):
            continue
        if expected_source_summary_page_id(source_id) not in page_records_by_id:
            missing_source_ids.add(source_id)
    return missing_source_ids


def collect_missing_concept_bucket_keys(
    claims_by_similarity_bucket: dict[str, list[dict]],
    page_records_by_id: dict[str, dict],
) -> set[str]:
    # 概念页的 page_id 可以由 bucket 稳定推导，因此可以快速补齐缺页。
    missing_bucket_keys: set[str] = set()
    for bucket_key, grouped_claims in claims_by_similarity_bucket.items():
        if not should_generate_concept_page(grouped_claims):
            continue
        if build_concept_page_id(bucket_key) not in page_records_by_id:
            missing_bucket_keys.add(bucket_key)
    return missing_bucket_keys


def expected_workspace_overview_concept_page_ids(
    claims_by_similarity_bucket: dict[str, list[dict]],
) -> set[str]:
    return {
        build_concept_page_id(bucket_key)
        for bucket_key, grouped_claims in claims_by_similarity_bucket.items()
        if should_generate_concept_page(grouped_claims)
    }


def collect_workspace_overview_concept_pages(
    claims_by_similarity_bucket: dict[str, list[dict]],
    page_records_by_id: dict[str, dict],
) -> list[dict]:
    concept_pages: list[dict] = []
    for page_id in sorted(expected_workspace_overview_concept_page_ids(claims_by_similarity_bucket)):
        page_record = page_records_by_id.get(page_id)
        if page_record is None or not is_live_page_record(page_record):
            continue
        if page_record.get("type") != "concept":
            continue
        concept_pages.append(page_record)
    return concept_pages


def should_generate_workspace_overview_page(concept_page_records: list[dict]) -> bool:
    # 综述页先保持保守：至少要有两个可读概念页，才值得给出工作区级总览。
    return len(concept_page_records) >= 2


def workspace_overview_page_missing(
    claims_by_similarity_bucket: dict[str, list[dict]],
    page_records_by_id: dict[str, dict],
) -> bool:
    concept_pages = collect_workspace_overview_concept_pages(
        claims_by_similarity_bucket=claims_by_similarity_bucket,
        page_records_by_id=page_records_by_id,
    )
    return (
        should_generate_workspace_overview_page(concept_pages)
        and build_workspace_overview_page_id() not in page_records_by_id
    )


def regroup_concept_claims_by_canonical_topic(
    concept_claim_groups: dict[str, list[dict]],
) -> dict[str, list[dict]]:
    # 第一层 bucket 主要解决“相似 claim 初步召回”。
    # 但真实页面层还需要再做一次“主题收口”：
    # 如果不同 bucket 最终推导出相同标题/规范键，就应该落到同一概念页，
    # 否则会出现多个 live page 共用同一 canonical_id，进而污染 alias index 和 lint。
    regrouped: dict[str, list[dict]] = {}

    for bucket_key, grouped_claims in concept_claim_groups.items():
        if not grouped_claims:
            continue
        group_topic_label = choose_group_topic_label(grouped_claims)
        canonical_claim = choose_canonical_claim(grouped_claims, group_topic_label)
        title = build_concept_title(canonical_claim, preferred_section_label=group_topic_label)
        canonical_key = build_concept_canonical_key(title)
        # 这里最终只按 canonical_key 收口。
        # 这样同一主题下的多条陈述即使来自不同 bucket，也会落到同一概念页，
        # 避免多个 live page 共享同一 canonical_id。
        regroup_key = canonical_key
        regrouped.setdefault(regroup_key, []).extend(grouped_claims)

    for regroup_key, grouped_claims in list(regrouped.items()):
        deduped_by_claim_id = {claim_record["claim_id"]: claim_record for claim_record in grouped_claims}
        regrouped[regroup_key] = list(deduped_by_claim_id.values())

    return regrouped


def remove_page_id_from_claim_records(claims_by_id: dict[str, dict], page_id: str) -> set[str]:
    dirty_claim_ids: set[str] = set()
    for claim_record in claims_by_id.values():
        page_ids = claim_record.get("page_ids", [])
        if page_id not in page_ids:
            continue
        claim_record["page_ids"] = [item for item in page_ids if item != page_id]
        claim_record["updated_at"] = utc_now_iso()
        dirty_claim_ids.add(claim_record["claim_id"])
    return dirty_claim_ids


def remove_page_id_from_review_records(reviews_by_id: dict[str, dict], page_id: str) -> set[str]:
    dirty_review_ids: set[str] = set()
    for review_record in reviews_by_id.values():
        candidate_page_ids = review_record.get("candidate_page_ids", [])
        if page_id not in candidate_page_ids:
            continue
        review_record["candidate_page_ids"] = [item for item in candidate_page_ids if item != page_id]
        dirty_review_ids.add(review_record["review_id"])
    return dirty_review_ids


def remove_source_refs_from_claim_record(claim_record: dict, source_id: str) -> bool:
    # 同一路径来源更新时，要把旧 source 对应的证据引用从 claim 中剥掉，
    # 这样后续页面聚合和阅读包才会逐步只围绕最新活动版本展开。
    original_source_ids = list(claim_record.get("source_ids", []))
    original_chunk_ids = list(claim_record.get("chunk_ids", []))
    original_source_refs = list(claim_record.get("source_refs", []))

    claim_record["source_ids"] = [
        item for item in claim_record.get("source_ids", [])
        if item != source_id
    ]
    claim_record["source_refs"] = [
        item for item in claim_record.get("source_refs", [])
        if item.get("source_id") != source_id
    ]
    active_chunk_ids = {
        item.get("chunk_id")
        for item in claim_record.get("source_refs", [])
        if item.get("chunk_id")
    }
    claim_record["chunk_ids"] = [
        chunk_id for chunk_id in claim_record.get("chunk_ids", [])
        if chunk_id in active_chunk_ids
    ]
    claim_record["lifecycle_status"] = claim_lifecycle_status_for_record(claim_record)

    return (
        claim_record["source_ids"] != original_source_ids
        or claim_record["chunk_ids"] != original_chunk_ids
        or claim_record["source_refs"] != original_source_refs
    )


def purge_source_from_claims(
    target: Path,
    claims_by_id: dict[str, dict],
    historical_claims_by_id: dict[str, dict],
    source_id: str,
) -> tuple[set[str], set[str]]:
    # 返回：
    # 1. 被修改但仍保留的 claim_id
    # 2. 因失去全部 source_ref 而被转入历史态的 claim_id
    dirty_claim_ids: set[str] = set()
    deleted_claim_ids: set[str] = set()

    for claim_id, claim_record in list(claims_by_id.items()):
        if source_id not in claim_record.get("source_ids", []):
            continue
        changed = remove_source_refs_from_claim_record(claim_record, source_id)
        if not changed:
            continue
        if not claim_record.get("source_ids") or not claim_record.get("source_refs"):
            claim_record["lifecycle_status"] = "superseded"
            claim_record["archived_at"] = utc_now_iso()
            claim_record["updated_at"] = utc_now_iso()
            deleted_claim_ids.add(claim_id)
            claims_by_id.pop(claim_id, None)
            historical_claim_record = convert_claim_record_to_historical(claim_record)
            historical_claims_by_id[historical_claim_record["claim_id"]] = historical_claim_record
            write_claim_file(target, historical_claim_record)
            continue
        claim_record["updated_at"] = utc_now_iso()
        dirty_claim_ids.add(claim_id)

    return dirty_claim_ids, deleted_claim_ids


def purge_deleted_claims_from_reviews(
    reviews_by_id: dict[str, dict],
    historical_reviews_by_id: dict[str, dict],
    deleted_claim_ids: set[str],
) -> tuple[set[str], set[str]]:
    # 被删掉的 claim 不该继续留在 review 候选里。
    # 返回：
    # 1. 仍保留但被修改的 review_id
    # 2. 因失去全部 candidate_claim_ids 而转入历史态的 review_id
    dirty_review_ids: set[str] = set()
    deleted_review_ids: set[str] = set()
    deleted_claim_id_set = set(deleted_claim_ids)

    for review_id, review_record in list(reviews_by_id.items()):
        original_claim_ids = list(review_record.get("candidate_claim_ids", []))
        remaining_claim_ids = [
            claim_id for claim_id in original_claim_ids
            if claim_id not in deleted_claim_id_set
        ]
        if remaining_claim_ids == original_claim_ids:
            continue
        if not remaining_claim_ids:
            review_record["lifecycle_status"] = "superseded"
            review_record["archived_at"] = utc_now_iso()
            reviews_by_id.pop(review_id, None)
            historical_review_record = convert_review_record_to_historical(review_record)
            historical_reviews_by_id[historical_review_record["review_id"]] = historical_review_record
            deleted_review_ids.add(review_id)
            continue
        review_record["candidate_claim_ids"] = remaining_claim_ids
        review_record["lifecycle_status"] = review_lifecycle_status_for_record(review_record)
        dirty_review_ids.add(review_id)

    return dirty_review_ids, deleted_review_ids


def prune_stale_auto_pages(
    target: Path,
    page_records_by_id: dict[str, dict],
    desired_auto_page_ids: set[str],
    claims_by_id: dict[str, dict],
    reviews_by_id: dict[str, dict],
    forced_stale_page_ids: set[str] | None = None,
) -> tuple[list[dict], set[str], set[str]]:
    # 这里负责清理“这轮模型下已经不该存在”的自动页面。
    # 典型场景是：同一路径文档被更新后，过期 source-summary 和概念页应退出主视图。
    removed_pages: list[dict] = []
    dirty_claim_ids: set[str] = set()
    dirty_review_ids: set[str] = set()
    forced_stale_page_ids = forced_stale_page_ids or set()
    auto_page_types = {
        "source-summary",
        "concept",
        "overview",
        "guide",
        "example",
        "topic",
        "reference",
        "timeline",
    }

    stale_page_ids = [
        page_id
        for page_id, page_record in page_records_by_id.items()
        if page_record.get("type") in auto_page_types
        and (
            page_id in forced_stale_page_ids
            or page_id not in desired_auto_page_ids
        )
    ]

    for page_id in stale_page_ids:
        page_record = dict(page_records_by_id[page_id])
        page_record["removed"] = True
        page_record["lifecycle_status"] = "removed"
        page_record["archived_at"] = utc_now_iso()
        page_record["updated"] = utc_now_iso()
        page_records_by_id[page_id] = page_record
        removed_pages.append(page_record)

        page_path = target / page_record["page_path"]
        if page_path.exists():
            page_path.unlink()

        dirty_claim_ids.update(remove_page_id_from_claim_records(claims_by_id, page_id))
        dirty_review_ids.update(remove_page_id_from_review_records(reviews_by_id, page_id))

    return removed_pages, dirty_claim_ids, dirty_review_ids


def sanitize_page_slug(value: str) -> str:
    # wiki 页文件名尽量稳定、可读、跨平台安全。
    slug = sanitize_source_key(value)
    stabilized = stabilize_filename_component(slug, separator="_")
    return stabilized or "page"


def sanitize_page_filename(value: str) -> str:
    # 面向最终导出的页面文件名尽量保留可读性，避免把标题压成一串下划线。
    cleaned = clean_concept_title_text(value)
    cleaned = re.sub(r"[\\/:*?\"<>|#]+", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    hash_source = re.sub(r"[\\/:*?\"<>|#]+", " ", str(value))
    hash_source = re.sub(r"\s+", " ", hash_source).strip(" .")
    if hash_source and hash_source != cleaned:
        digest = hashlib.sha256(hash_source.encode("utf-8")).hexdigest()[:FILENAME_HASH_LENGTH]
        cleaned = f"{cleaned}__{digest}" if cleaned else digest
    stabilized = stabilize_filename_component(cleaned)
    return stabilized or "page"


def summarize_claims_for_page(claim_records: list[dict], limit: int = 3) -> list[str]:
    # 来源摘要页先挑几条 claim 做“核心观点”。
    ranked = sorted(
        claim_records,
        key=claim_record_rank_key,
        reverse=True,
    )
    return [item["text"] for item in ranked[:limit]]


def source_summary_page_path(source_id: str, title: str) -> Path:
    slug = sanitize_page_slug(title)
    source_key = stabilize_filename_component(sanitize_source_key(source_id), separator="_") or "source"
    filename = stabilize_filename_component(
        f"{slug}__{source_key}",
        max_bytes=MAX_FILENAME_COMPONENT_BYTES - len(".md".encode("utf-8")),
    )
    return Path("wiki") / "sources" / f"{filename or 'page'}.md"


def shorten_title_text(value: str, limit: int = 32) -> str:
    # 页面标题不能无限长，否则文件名、索引页和终端输出都会变得很难读。
    compact = re.sub(r"\s+", " ", value).strip()
    if len(compact) <= limit:
        return compact
    return compact[: max(limit - 3, 1)].rstrip() + "..."


def build_concept_page_id(bucket_key: str) -> str:
    # 概念页 ID 基于最终分组键生成。
    # 在 V1 当前实现里，这个键已经包含“可读主题 + 规范概念键”的二次收口结果，
    # 可以避免不同 bucket 最终生成相同 canonical_id 却 page_id 不同的问题。
    bucket_hash = hashlib.sha256(bucket_key.encode("utf-8")).hexdigest()
    return f"page_cpt_{bucket_hash[:12]}"


def build_concept_group_key(claim_record: dict) -> str:
    # 概念页聚合键尽量与 review 检测使用同一套“主题归一化”直觉。
    # 这样 query、review、concept page 三者更容易围绕同一组 claim 收敛。
    section_label = extract_primary_section_label(claim_record)
    if section_label and not is_generic_concept_label(section_label):
        return build_concept_canonical_key(section_label)
    base_text = normalize_claim_base_for_conflict(claim_record.get("text", ""))
    similarity_tokens = build_claim_similarity_tokens(claim_record.get("text", ""))
    token_fingerprint = " ".join(similarity_tokens[:8])
    seed = base_text or claim_record.get("normalized_text", "") or claim_record.get("text", "")
    seed_hash = hashlib.sha256(f"{seed}|{token_fingerprint}".encode("utf-8")).hexdigest()[:12]
    readable_prefix = build_similarity_bucket(claim_record.get("text", ""))
    return f"{readable_prefix}|{seed_hash}"


def claim_role_blocks_concept_path(claim_record: dict) -> bool:
    role = claim_knowledge_role(claim_record)
    if role in {"procedure", "example", "meta", "structural_shell", "opinion"}:
        return True
    page_intent_hints = set(claim_page_intent_hints(claim_record))
    if "reject" in page_intent_hints:
        return True
    return False


def filter_claim_records_for_concept_path(claim_records: list[dict]) -> list[dict]:
    return [record for record in claim_records if not claim_role_blocks_concept_path(record)]


def concept_summary_page_path(page_id: str, title: str) -> Path:
    # 概念页文件名尽量贴近最终展示标题，避免导出到外部工具时把内部 page_id 暴露成主标题。
    filename = sanitize_page_filename(title)
    return Path("wiki") / "concepts" / page_id / f"{filename}.md"


def clean_concept_title_text(value: str) -> str:
    # 概念页标题要尽量像“页面名”，而不是原始 claim 文本残片。
    cleaned = value.replace("|", " ").replace("_", " ")
    # 仅清理“1. 标题”“2) 标题”这类前导编号，不要把 2026-05-24 这类日期误裁成 05-24。
    cleaned = re.sub(r"^\s*\d+\s*[.)、:：]\s*", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -:;,.!?。！？；：|")
    return cleaned


def normalize_question_style_concept_label(label: str) -> str:
    # FAQ/目录型标题常写成“1. Claim 是什么”或“什么是 Claim”，这里尽量还原成概念名本身。
    cleaned = clean_concept_title_text(label)
    if not cleaned:
        return ""

    suffix_match = re.fullmatch(r"(.+?)\s*(?:是|指)?什么[？?]?", cleaned, flags=re.IGNORECASE)
    if suffix_match:
        candidate = clean_concept_title_text(suffix_match.group(1))
        if candidate:
            return candidate

    prefix_match = re.fullmatch(r"(?:什么是|什么叫|何谓)\s*(.+?)[？?]?", cleaned, flags=re.IGNORECASE)
    if prefix_match:
        candidate = clean_concept_title_text(prefix_match.group(1))
        if candidate:
            return candidate

    return cleaned


def section_path_parts_from_claim_record(claim_record: dict) -> list[str]:
    for source_ref in claim_record.get("source_refs", []):
        parts = source_ref.get("section_path_parts")
        if isinstance(parts, list):
            cleaned_parts = [normalize_question_style_concept_label(str(part)) for part in parts if str(part).strip()]
            cleaned_parts = [part for part in cleaned_parts if part]
            if cleaned_parts:
                return cleaned_parts
        section_path = source_ref.get("section_path", "")
        if not section_path:
            continue
        parsed = parse_section_path(section_path)
        parts = [
            normalize_question_style_concept_label(part)
            for part in parsed.get("section_path_parts", [])
            if part
        ]
        parts = [part for part in parts if part]
        if parts:
            return parts
    return []


def section_label_is_meaningful_context(label: str) -> bool:
    cleaned = normalize_question_style_concept_label(label)
    if not cleaned or is_generic_concept_label(cleaned):
        return False
    if text_is_question_like(cleaned):
        return False
    if len(cleaned) >= 18 and claim_has_standalone_predicate(cleaned):
        return False
    if any(marker in cleaned for marker in ("：", ":", "。", "？", "?")) and len(cleaned) >= 10:
        return False
    return True


def build_hierarchical_section_label(section_parts: list[str], max_parts: int = 3) -> str:
    if not section_parts:
        return ""
    meaningful_parts = [part for part in section_parts if section_label_is_meaningful_context(part)]
    if not meaningful_parts:
        meaningful_parts = [normalize_question_style_concept_label(part) for part in section_parts if part]
        meaningful_parts = [part for part in meaningful_parts if part]
    selected_parts = meaningful_parts[-max_parts:]
    if len(selected_parts) <= 1:
        return selected_parts[0] if selected_parts else ""
    return " / ".join(selected_parts)


def extract_primary_section_label(claim_record: dict) -> str:
    # 对概念页命名来说，section_path 往往比整句 claim 更接近“主题名”。
    parts = section_path_parts_from_claim_record(claim_record)
    if not parts:
        return ""
    return build_hierarchical_section_label(parts)


def choose_group_topic_label(claim_records: list[dict]) -> str:
    # 概念页首先要回答“这一组 claim 到底在讲什么主题”。
    # 这里优先采用来源 section label 的共识，而不是直接相信某一条 claim 的可读性。
    label_scores: dict[str, float] = {}

    for claim_record in claim_records:
        label = extract_primary_section_label(claim_record)
        if not label or is_generic_concept_label(label):
            continue
        label_scores[label] = label_scores.get(label, 0.0) + max(1, len(claim_record.get("source_ids", [])))

    if not label_scores:
        return ""

    ranked = sorted(
        label_scores.items(),
        key=lambda item: (item[1], len(build_claim_similarity_tokens(item[0])), len(item[0])),
        reverse=True,
    )
    return ranked[0][0]


def collect_section_label_aliases(claim_records: list[dict]) -> list[str]:
    aliases: list[str] = []
    for claim_record in claim_records:
        parts = section_path_parts_from_claim_record(claim_record)
        if not parts:
            continue
        candidates = [
            build_hierarchical_section_label(parts),
            normalize_question_style_concept_label(parts[-1]),
            " > ".join(parts),
        ]
        for candidate in candidates:
            cleaned = clean_concept_title_text(candidate)
            if cleaned and cleaned not in aliases:
                aliases.append(cleaned)
    return aliases


def is_generic_concept_label(label: str) -> bool:
    # 有些 section label 太泛，比如“文档开始”“sample”“表格 1”，单独拿来做页面名会很弱。
    normalized = clean_concept_title_text(label).lower()
    if normalized in {"", "文档开始", "sample"}:
        return True
    generic_exact_values = {
        "示例", "总结", "小结", "说明", "原因", "背景", "方法", "流程", "步骤",
        "注意", "补充", "附录", "表格", "代码", "引用", "问题", "状态", "初始化",
        "别名", "为什么", "如何", "怎么做", "做法", "概述", "介绍",
    }
    if normalized in generic_exact_values:
        return True
    if len(normalized) <= 1 and re.search(r"[\u4e00-\u9fff]", normalized):
        return True
    if re.fullmatch(r"(?:问题|示例|总结|步骤|方法)\s*\d*", normalized):
        return True
    if re.fullmatch(r"[一二三四五六七八九十]+[、.]?.{0,2}", normalized):
        return True
    if re.fullmatch(r"表格\s*\d+", normalized):
        return True
    return False


def concept_title_is_whitelisted_short_label(label: str) -> bool:
    normalized = clean_concept_title_text(label)
    if not normalized:
        return False
    if normalized in {"AI", "RAG", "MCP", "CLI", "SDK", "API", "Tauri", "React", "Rust", "BM25"}:
        return True
    return bool(re.fullmatch(r"[A-Za-z0-9._+-]{2,8}", normalized))


def concept_title_quality_details(
    title: str,
    canonical_claim: dict,
    claim_records: list[dict],
    preferred_section_label: str = "",
) -> dict:
    normalized_title = clean_concept_title_text(title)
    normalized_lower = normalized_title.lower()
    section_label = preferred_section_label or extract_primary_section_label(canonical_claim)
    section_label_normalized = clean_concept_title_text(section_label)
    claim_text = clean_claim_candidate_text(canonical_claim.get("text", ""))
    source_ids = {
        source_id
        for claim_record in claim_records
        for source_id in claim_record.get("source_ids", [])
    }
    reasons: list[str] = []
    score = 0

    if not normalized_title:
        reasons.append("empty_title")
        score -= 10

    if is_generic_concept_label(normalized_title):
        reasons.append("generic_title")
        score -= 8
    else:
        score += 3

    if len(normalized_title) <= 1 and not concept_title_is_whitelisted_short_label(normalized_title):
        reasons.append("too_short")
        score -= 10
    elif len(normalized_title) <= 3 and not concept_title_is_whitelisted_short_label(normalized_title):
        reasons.append("very_short")
        score -= 4
    elif len(normalized_title) >= 4:
        score += 1

    if concept_title_is_whitelisted_short_label(normalized_title):
        reasons.append("short_whitelisted")
        score += 4

    if claim_has_standalone_predicate(claim_text):
        score += 2
    else:
        reasons.append("claim_without_predicate")
        score -= 2

    if claim_starts_with_dependent_prefix(claim_text):
        reasons.append("dependent_prefix_claim")
        score -= 3

    if text_is_question_like(claim_text):
        reasons.append("question_like_claim")
        score -= 4

    if claim_is_topic_shell_text(canonical_claim, normalized_title):
        reasons.append("topic_shell_claim")
        score -= 6

    topic_alignment = claim_topic_alignment_score(canonical_claim, normalized_title)
    if topic_alignment >= 8:
        score += 4
    elif topic_alignment >= 4:
        score += 2
    else:
        reasons.append("low_topic_alignment")
        score -= 3

    if section_label_normalized and normalized_lower == section_label_normalized.lower():
        score += 1
    if section_label_normalized and is_generic_concept_label(section_label_normalized):
        reasons.append("generic_section_label")
        score -= 4

    if len(claim_records) >= 2:
        score += 2
    else:
        reasons.append("single_claim_only")
        score -= 1

    if len(source_ids) >= 2:
        reasons.append("cross_source_support")
        score += 4
    else:
        reasons.append("single_source_only")
        score -= 2

    if canonical_claim.get("claim_type") == "definition":
        score += 2

    hard_reject_reasons = {"empty_title", "generic_title", "too_short", "question_like_claim"}
    classification = "strong"
    if any(reason in hard_reject_reasons for reason in reasons):
        classification = "reject"
    elif score < 6:
        classification = "gray"

    return {
        "title": normalized_title,
        "score": score,
        "classification": classification,
        "reasons": reasons,
        "topic_alignment": topic_alignment,
        "section_label": section_label_normalized,
        "source_count": len(source_ids),
        "claim_count": len(claim_records),
    }


def load_concept_quality_review_config(config: dict) -> dict:
    render_config = load_page_render_config(config, "concept_update")
    return {
        "mode": render_config.get("mode"),
        "command": render_config.get("command", []),
        "timeout_seconds": render_config.get("timeout_seconds", 20),
    }


def run_llm_assisted_concept_title_review(
    target: Path,
    review_config: dict,
    title: str,
    canonical_claim: dict,
    claim_records: list[dict],
    preferred_section_label: str = "",
) -> dict | None:
    if review_config.get("mode") != "llm_assisted":
        return None
    command = review_config.get("command", [])
    if not command:
        return None

    payload = {
        "task": "review_concept_candidate",
        "candidate_title": title,
        "preferred_section_label": preferred_section_label,
        "canonical_claim": {
            "claim_id": canonical_claim.get("claim_id"),
            "text": canonical_claim.get("text"),
            "claim_type": canonical_claim.get("claim_type"),
        },
        "supporting_claims": [
            {
                "claim_id": claim_record.get("claim_id"),
                "text": claim_record.get("text"),
                "claim_type": claim_record.get("claim_type"),
                "section_label": extract_primary_section_label(claim_record),
                "source_count": len(claim_record.get("source_ids", [])),
            }
            for claim_record in claim_records[:6]
        ],
        "instructions": (
            "Judge whether this title is a valid reusable concept title or just a structural heading. "
            "If invalid, suggest a better concept title when the evidence clearly supports one. "
            "Return strict JSON only."
        ),
    }

    try:
        completed = subprocess.run(
            command,
            cwd=target,
            input=json.dumps(payload, ensure_ascii=False),
            capture_output=True,
            text=True,
            timeout=review_config.get("timeout_seconds", 20),
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None

    if completed.returncode != 0:
        return None

    stdout = (completed.stdout or "").strip()
    if not stdout:
        return None

    try:
        result = json.loads(stdout)
    except json.JSONDecodeError:
        return None
    if not isinstance(result, dict):
        return None

    suggested_title = clean_concept_title_text(result.get("suggested_title", ""))
    decision = str(result.get("decision", "") or "").strip().lower()
    reason = str(result.get("reason", "") or "").strip()
    confidence = result.get("confidence", 0.0)
    if decision not in {"accept", "reject", "rename"}:
        return None
    if suggested_title and is_generic_concept_label(suggested_title):
        suggested_title = ""
    return {
        "decision": decision,
        "suggested_title": suggested_title,
        "reason": reason,
        "confidence": float(confidence) if isinstance(confidence, (int, float)) else 0.0,
    }


def extract_markdown_table_rows(text: str) -> list[list[str]]:
    # 当 claim 来自表格时，原始文本里常带 Markdown 表格；这里抽出单元格以便后续命名。
    rows: list[list[str]] = []
    for line in text.splitlines():
        stripped = line.strip()
        if "|" not in stripped:
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if not cells:
            continue
        if all(re.fullmatch(r"-+", cell or "") for cell in cells):
            continue
        rows.append([cell for cell in cells if cell])
    return rows


def extract_concept_phrase_from_table_rows(rows: list[list[str]]) -> str:
    # 表格里的“最后一行有效数据”通常最像当前 chunk 的具体陈述。
    if not rows:
        return ""
    data_rows = rows[1:] if len(rows) >= 2 else rows
    if not data_rows:
        data_rows = rows
    last_row = data_rows[-1]
    generic_headers = {"字段", "键", "值", "说明", "项目", "claim"}
    if len(last_row) >= 2:
        head, tail = last_row[0], last_row[1]
        if head in generic_headers:
            return clean_concept_title_text(tail)
        if head.lower().startswith("工作表"):
            return clean_concept_title_text(tail)
        return clean_concept_title_text(f"{head} {tail}")
    return clean_concept_title_text(" ".join(last_row))


def extract_concept_phrase_from_claim(claim_text: str, section_label: str) -> str:
    # 如果 section label 不够具体，就再从 claim 本体里提一段更适合当标题的短语。
    table_rows = extract_markdown_table_rows(claim_text)
    if table_rows:
        phrase = extract_concept_phrase_from_table_rows(table_rows)
        if phrase:
            return phrase

    plain_text = clean_concept_title_text(markdown_to_plain_text(claim_text))
    if section_label and plain_text.lower().startswith(section_label.lower()):
        plain_text = clean_concept_title_text(plain_text[len(section_label):])

    # 句子型 claim 先截到第一个停顿点，避免整句都变成标题。
    pieces = re.split(r"[。！？!?；;:：]", plain_text, maxsplit=1)
    candidate = clean_concept_title_text(pieces[0] if pieces else plain_text)
    return candidate


def build_concept_title(canonical_claim: dict, preferred_section_label: str = "") -> str:
    # 概念页标题优先用 section label，必要时再拼一个来自 claim 的补充短语。
    section_label = preferred_section_label or extract_primary_section_label(canonical_claim)
    claim_phrase = extract_concept_phrase_from_claim(canonical_claim.get("text", ""), section_label)

    if section_label and not is_generic_concept_label(section_label):
        if section_label.startswith("工作表") and claim_phrase and claim_phrase not in section_label:
            return shorten_title_text(clean_concept_title_text(f"{section_label} - {claim_phrase}"), limit=28)
        return shorten_title_text(section_label, limit=28)

    if section_label.startswith("表格") and claim_phrase:
        if claim_phrase.lower().startswith(section_label.lower()):
            claim_phrase = clean_concept_title_text(claim_phrase[len(section_label):])
        if claim_phrase:
            return shorten_title_text(clean_concept_title_text(f"{section_label} - {claim_phrase}"), limit=28)

    if section_label and claim_phrase:
        return shorten_title_text(clean_concept_title_text(f"{section_label} - {claim_phrase}"), limit=28)
    if claim_phrase:
        return shorten_title_text(claim_phrase, limit=28)
    return shorten_title_text(clean_concept_title_text(canonical_claim.get("text", "")), limit=28)


def build_concept_canonical_key(title: str) -> str:
    # canonical key 给后续 alias、redirect、检索归一化用，尽量保持短而稳定。
    cleaned = clean_concept_title_text(title)
    if not cleaned:
        return "concept"
    compact = sanitize_source_key(cleaned)
    return compact or "concept"


def claim_topic_alignment_score(claim_record: dict, group_topic_label: str = "") -> float:
    # 代表陈述首先应和页面主题对齐，其次才是“读起来像一句完整的话”。
    if not group_topic_label:
        return 0.0

    normalized_topic = clean_concept_title_text(group_topic_label).lower()
    if not normalized_topic:
        return 0.0

    claim_text = clean_claim_candidate_text(claim_record.get("text", ""))
    claim_text_lower = claim_text.lower()
    section_label = extract_primary_section_label(claim_record)
    section_label_lower = clean_concept_title_text(section_label).lower()
    topic_tokens = build_claim_similarity_tokens(group_topic_label)
    claim_tokens = build_claim_similarity_tokens(claim_text)
    overlap_ratio, _, shared_token_count = compute_weighted_token_overlap(topic_tokens, claim_tokens)

    score = overlap_ratio * 10.0 + float(shared_token_count)
    if section_label_lower == normalized_topic:
        score += 8.0
    elif section_label_lower and normalized_topic in section_label_lower:
        score += 4.0
    if claim_text_lower.startswith(normalized_topic):
        score += 6.0
    elif normalized_topic in claim_text_lower:
        score += 3.0
    return score


def claim_is_topic_shell_text(claim_record: dict, group_topic_label: str = "") -> bool:
    # 有些 claim 实际上只是 section 标题本身，被抽出来后并没有承载独立结论。
    # 这类“壳句”适合留作 supporting context，但不应抢占代表陈述。
    cleaned = clean_claim_candidate_text(claim_record.get("text", ""))
    if not cleaned:
        return False

    normalized_cleaned = cleaned.lower()
    candidate_labels = []
    section_label = extract_primary_section_label(claim_record)
    if section_label:
        candidate_labels.append(clean_concept_title_text(section_label).lower())
    if group_topic_label:
        candidate_labels.append(clean_concept_title_text(group_topic_label).lower())

    candidate_labels = [label for label in candidate_labels if label]
    if not candidate_labels:
        return False
    if normalized_cleaned not in candidate_labels:
        return False
    return not claim_has_standalone_predicate(cleaned)


def claim_record_readability_score(claim_record: dict, group_topic_label: str = "") -> int:
    text = claim_record.get("text", "")
    cleaned = clean_claim_candidate_text(text)
    if not cleaned:
        return -10

    score = 0
    topic_alignment = claim_topic_alignment_score(claim_record, group_topic_label)
    if topic_alignment >= 12:
        score += 5
    elif topic_alignment >= 8:
        score += 3
    elif topic_alignment >= 4:
        score += 1
    if claim_record.get("claim_type") == "definition":
        score += 2
    if not claim_starts_with_dependent_prefix(cleaned):
        score += 2
    if claim_has_standalone_predicate(cleaned):
        score += 1
    if claim_is_definition_like_phrase(cleaned):
        score += 3

    section_label = extract_primary_section_label(claim_record)
    if section_label and cleaned.lower().startswith(section_label.lower()):
        score += 2

    if 16 <= len(cleaned) <= 72:
        score += 1
    elif len(cleaned) > 120:
        score -= 1

    if claim_starts_with_dependent_prefix(cleaned):
        score -= 3
    if claim_starts_with_meta_prefix(cleaned):
        score -= 3
    if claim_is_topic_shell_text(claim_record, group_topic_label):
        score -= 6
    return score


def claim_record_rank_key(claim_record: dict, group_topic_label: str = "") -> tuple:
    text = claim_record.get("text", "")
    is_topic_shell = claim_is_topic_shell_text(claim_record, group_topic_label)
    return (
        0 if is_topic_shell else 1,
        claim_topic_alignment_score(claim_record, group_topic_label),
        len(claim_record.get("source_ids", [])),
        len(claim_record.get("source_refs", [])),
        claim_record_readability_score(claim_record, group_topic_label),
        -abs(len(text) - 42),
        len(text),
    )


def build_display_claim_text(claim_record: dict, concept_title: str = "") -> str:
    # 概念页里的代表陈述优先展示成“概念名 + 定义短语”，
    # 这样比孤立的“一种……”更像一句可读的说明。
    raw_text = markdown_to_plain_text(claim_record.get("text", ""))
    cleaned = clean_claim_candidate_text(raw_text)
    if not cleaned:
        return raw_text.strip()

    title = clean_concept_title_text(concept_title)
    if title and claim_is_definition_like_phrase(cleaned) and not cleaned.lower().startswith(title.lower()):
        return clean_concept_title_text(f"{title} {cleaned}")
    return cleaned


def choose_canonical_claim(claim_records: list[dict], group_topic_label: str = "") -> dict:
    # 一组 claim 需要选一个“代表陈述”，后面会用它来命名页面和生成摘要。
    ranked = sorted(
        claim_records,
        key=lambda item: claim_record_rank_key(item, group_topic_label),
        reverse=True,
    )
    return ranked[0]


def should_generate_concept_page(claim_records: list[dict]) -> bool:
    # 概念页不必给每条 claim 都生成一份。
    # V1 先优先保留三类更有价值的候选：
    # 1. 多条相似 claim 汇聚到一起；
    # 2. 单条 claim 但有多个来源支撑；
    # 3. 单条 claim 但表达完整、主题明确，值得先沉淀成主题入口。
    concept_claim_records = filter_claim_records_for_concept_path(claim_records)
    if not concept_claim_records:
        return False

    source_ids = {
        source_id
        for claim_record in concept_claim_records
        for source_id in claim_record.get("source_ids", [])
    }
    if len(concept_claim_records) >= 2:
        return True
    if len(source_ids) >= 2:
        return True
    canonical_claim = choose_canonical_claim(concept_claim_records, choose_group_topic_label(concept_claim_records))
    section_label = choose_group_topic_label(concept_claim_records)
    concept_title = build_concept_title(canonical_claim, preferred_section_label=choose_group_topic_label(concept_claim_records))
    quality = concept_title_quality_details(
        title=concept_title,
        canonical_claim=canonical_claim,
        claim_records=concept_claim_records,
        preferred_section_label=section_label,
    )
    claim_text = canonical_claim.get("text", "")
    cleaned_claim_text = clean_claim_candidate_text(claim_text)
    # 一些明显是“转换占位提示”的文本先不提升成概念页，避免 Wiki 被环境提示刷屏。
    if any(marker in claim_text for marker in ("当前环境缺少", "当前环境未启用", "仅生成占位", "估计页数:")):
        return False
    if text_is_iso_date_label(section_label) or text_is_iso_date_label(cleaned_claim_text):
        return True
    if quality["classification"] == "reject":
        return False
    if text_is_iso_date_label(cleaned_claim_text):
        return True
    if canonical_claim.get("claim_type") == "definition" and len(cleaned_claim_text) >= 14:
        return True
    if section_label and not is_generic_concept_label(section_label) and len(section_label) >= 4:
        return True
    concept_candidate_score = claim_concept_candidate_score(canonical_claim)
    if concept_candidate_score >= 0.75:
        return True
    return len(cleaned_claim_text) >= 18 and concept_candidate_score >= 0.3 and claim_can_stand_alone(cleaned_claim_text)


def resolve_concept_title_candidate(
    target: Path,
    config: dict,
    canonical_claim: dict,
    claim_records: list[dict],
    preferred_section_label: str = "",
) -> tuple[str, dict]:
    title = build_concept_title(canonical_claim, preferred_section_label=preferred_section_label)
    quality = concept_title_quality_details(
        title=title,
        canonical_claim=canonical_claim,
        claim_records=claim_records,
        preferred_section_label=preferred_section_label,
    )

    llm_review: dict | None = None
    if quality["classification"] == "gray":
        llm_review = run_llm_assisted_concept_title_review(
            target=target,
            review_config=load_concept_quality_review_config(config),
            title=title,
            canonical_claim=canonical_claim,
            claim_records=claim_records,
            preferred_section_label=preferred_section_label,
        )
        if llm_review and llm_review.get("decision") == "rename" and llm_review.get("suggested_title"):
            title = llm_review["suggested_title"]
            quality = concept_title_quality_details(
                title=title,
                canonical_claim=canonical_claim,
                claim_records=claim_records,
                preferred_section_label=preferred_section_label,
            )
        elif llm_review and llm_review.get("decision") == "reject":
            quality = dict(quality)
            quality["classification"] = "reject"
            quality["reasons"] = list(quality.get("reasons", [])) + ["llm_rejected_gray_candidate"]
        elif llm_review and llm_review.get("decision") == "accept":
            quality = dict(quality)
            quality["classification"] = "strong"
            quality["reasons"] = list(quality.get("reasons", [])) + ["llm_accepted_gray_candidate"]

    quality = dict(quality)
    quality["llm_review"] = llm_review
    if quality["classification"] == "reject":
        fallback_candidates = [
            extract_concept_phrase_from_claim(canonical_claim.get("text", ""), ""),
            preferred_section_label if text_is_iso_date_label(preferred_section_label) else "",
            clean_concept_title_text(shorten_title_text(markdown_to_plain_text(canonical_claim.get("text", "")), limit=28)),
        ]
        for fallback_title in fallback_candidates:
            fallback_title = clean_concept_title_text(fallback_title)
            if not fallback_title:
                continue
            if is_generic_concept_label(fallback_title) and not text_is_iso_date_label(fallback_title):
                continue
            title = shorten_title_text(fallback_title, limit=28)
            quality = concept_title_quality_details(
                title=title,
                canonical_claim=canonical_claim,
                claim_records=claim_records,
                preferred_section_label=preferred_section_label,
            )
            quality = dict(quality)
            quality["llm_review"] = llm_review
            quality["fallback_title_applied"] = True
            break
    return title, quality


def aggregate_source_refs_for_page(claim_records: list[dict]) -> list[dict]:
    # 页面层只保留“按来源聚合后的证据索引”，正文再去展开更细的 claim / chunk 关系。
    aggregated: dict[str, dict] = {}
    for claim_record in claim_records:
        for source_ref in claim_record.get("source_refs", []):
            source_id = source_ref["source_id"]
            if source_id not in aggregated:
                aggregated[source_id] = {
                    "source_id": source_id,
                    "source_path": source_ref["source_path"],
                    "chunk_ids": [],
                    "claim_ids": [],
                    "chunks": [],
                }
            if source_ref["chunk_id"] not in aggregated[source_id]["chunk_ids"]:
                aggregated[source_id]["chunk_ids"].append(source_ref["chunk_id"])
                aggregated[source_id]["chunks"].append({
                    "chunk_id": source_ref["chunk_id"],
                    "section_path": source_ref.get("section_path"),
                    "start_line": source_ref.get("start_line"),
                    "end_line": source_ref.get("end_line"),
                })
            append_unique(aggregated[source_id]["claim_ids"], claim_record["claim_id"])
    for record in aggregated.values():
        record["chunks"].sort(
            key=lambda item: (
                item.get("start_line") if item.get("start_line") is not None else math.inf,
                item.get("chunk_id", ""),
            )
        )
    return sorted(aggregated.values(), key=lambda item: item["source_id"])


def aggregate_source_refs_for_pages(page_records: list[dict]) -> list[dict]:
    # 更高层页面（如 overview）复用下层页面已经整理过的 source_refs，
    # 避免再次按 claim 全量回扫。
    aggregated: dict[str, dict] = {}
    for page_record in page_records:
        for source_ref in page_record.get("source_refs", []):
            source_id = source_ref["source_id"]
            if source_id not in aggregated:
                aggregated[source_id] = {
                    "source_id": source_id,
                    "source_path": source_ref["source_path"],
                    "chunk_ids": [],
                    "claim_ids": [],
                    "chunks": [],
                }
            for chunk_id in source_ref.get("chunk_ids", []):
                if chunk_id not in aggregated[source_id]["chunk_ids"]:
                    aggregated[source_id]["chunk_ids"].append(chunk_id)
            for claim_id in source_ref.get("claim_ids", []):
                append_unique(aggregated[source_id]["claim_ids"], claim_id)
            for chunk_ref in source_ref.get("chunks", []):
                if any(item.get("chunk_id") == chunk_ref.get("chunk_id") for item in aggregated[source_id]["chunks"]):
                    continue
                aggregated[source_id]["chunks"].append({
                    "chunk_id": chunk_ref.get("chunk_id"),
                    "section_path": chunk_ref.get("section_path"),
                    "start_line": chunk_ref.get("start_line"),
                    "end_line": chunk_ref.get("end_line"),
                })
    for record in aggregated.values():
        record["chunks"].sort(
            key=lambda item: (
                item.get("start_line") if item.get("start_line") is not None else math.inf,
                item.get("chunk_id", ""),
            )
        )
    return sorted(aggregated.values(), key=lambda item: item["source_id"])


def markdown_link_between_pages(from_page: Path, to_page: Path) -> str:
    # 页面正文里尽量使用相对链接，这样整个工作区换目录后链接仍然有效。
    relative = os.path.relpath(to_page, start=from_page.parent)
    return quote(relative.replace(os.sep, "/"), safe="/._-~")


def markdown_link_target(path: str) -> str:
    # 目录页链接也需要对空格等字符做转义，避免某些 Markdown 查看器截断路径。
    return quote(path.replace(os.sep, "/"), safe="/._-~")


def collect_source_summary_pages_for_claims(claim_records: list[dict], page_records_by_id: dict[str, dict]) -> list[dict]:
    # 概念页里最需要引用的是“来源摘要页”，因此这里把相关来源页单独筛出来。
    source_pages: list[dict] = []
    seen_page_ids: set[str] = set()
    for claim_record in claim_records:
        for page_id in claim_record.get("page_ids", []):
            page_record = page_records_by_id.get(page_id)
            if page_record is None or page_record.get("type") != "source-summary":
                continue
            if page_id in seen_page_ids:
                continue
            seen_page_ids.add(page_id)
            source_pages.append(page_record)
    return sorted(source_pages, key=lambda item: item.get("title", "").lower())


def collect_review_ids_for_claims(claim_ids: list[str], review_records: list[dict]) -> list[str]:
    # review 记录本身是“人工裁决入口”，页面记录里也要能反查到它们。
    claim_id_set = set(claim_ids)
    matched: list[str] = []
    for review_record in review_records:
        if not is_actionable_review_record(review_record):
            continue
        candidate_claim_ids = set(review_record.get("candidate_claim_ids", []))
        if claim_id_set & candidate_claim_ids:
            matched.append(review_record["review_id"])
    return sorted(set(matched))


def find_live_page_by_canonical_id_and_type(
    page_records_by_id: dict[str, dict],
    canonical_id: str,
    page_type: str,
) -> dict | None:
    for page_record in page_records_by_id.values():
        if not is_live_page_record(page_record):
            continue
        if page_record.get("canonical_id") != canonical_id:
            continue
        if page_record.get("type") != page_type:
            continue
        return page_record
    return None


def render_claim_as_sentence(claim_record: dict, concept_title: str = "") -> str:
    sentence = build_display_claim_text(claim_record, concept_title).strip()
    if not sentence:
        return ""
    if sentence.endswith(("。", "！", "？", ".", "!", "?")):
        return sentence
    return f"{sentence}。"


def build_readable_concept_summary_text(
    title: str,
    canonical_claim: dict,
    stable_claim_records: list[dict],
    source_refs: list[dict],
) -> str:
    intro = render_claim_as_sentence(canonical_claim, title)
    if not intro:
        intro = f"{title} 目前已经沉淀出可复用的稳定结论。"
    coverage = f"当前版本基于 {len(stable_claim_records)} 条稳定 Claim、{len(source_refs)} 个来源整理。"
    return f"{intro} {coverage}".strip()


def extract_first_sentence(text: str) -> str:
    cleaned = markdown_to_plain_text(text).strip()
    if not cleaned:
        return ""
    match = re.split(r"(?<=[。！？.!?])\s+", cleaned, maxsplit=1)
    return match[0].strip() if match else cleaned


def split_text_into_sentences(text: str) -> list[str]:
    cleaned = markdown_to_plain_text(text).strip()
    if not cleaned:
        return []
    parts = re.split(r"(?<=[。！？.!?])\s+", cleaned)
    return [part.strip() for part in parts if part.strip()]


def collect_page_claim_records(page_record: dict, claim_records_by_id: dict[str, dict]) -> list[dict]:
    return [
        claim_records_by_id[claim_id]
        for claim_id in page_record.get("claim_ids", [])
        if claim_id in claim_records_by_id
    ]


def count_claim_types(claim_records: list[dict]) -> Counter:
    return Counter(
        claim_record.get("claim_type", "unknown")
        for claim_record in claim_records
    )


def concept_page_overview_rank_key(page_record: dict, claim_records_by_id: dict[str, dict]) -> tuple[int, int, int, str]:
    claim_records = collect_page_claim_records(page_record, claim_records_by_id)
    claim_type_counts = count_claim_types(claim_records)
    operational_signal = sum(
        claim_type_counts.get(claim_type, 0)
        for claim_type in {"procedure", "warning", "comparison", "causal", "evaluation"}
    )
    return (
        len(page_record.get("source_refs", [])),
        len(page_record.get("claim_ids", [])),
        operational_signal,
        page_record.get("title", "").lower(),
    )


def summarize_concept_page_for_overview(page_record: dict) -> str:
    sentence = extract_first_sentence(page_record.get("summary", ""))
    if sentence:
        return sentence
    return f"{page_record.get('title', '该主题')} 已经沉淀出稳定结论。"


def build_workspace_overview_key_theme_rows(
    concept_pages: list[dict],
    claim_records_by_id: dict[str, dict],
    limit: int = 6,
) -> list[dict]:
    rows: list[dict] = []
    ranked_pages = sorted(
        concept_pages,
        key=lambda item: concept_page_overview_rank_key(item, claim_records_by_id),
        reverse=True,
    )
    for page_record in ranked_pages[:limit]:
        claim_records = collect_page_claim_records(page_record, claim_records_by_id)
        claim_type_counts = count_claim_types(claim_records)
        foundational_signal = sum(
            claim_type_counts.get(claim_type, 0)
            for claim_type in {"definition", "fact"}
        )
        operational_signal = sum(
            claim_type_counts.get(claim_type, 0)
            for claim_type in {"procedure", "warning", "comparison", "causal", "evaluation"}
        )
        if operational_signal > foundational_signal:
            theme_kind = "operational"
        elif foundational_signal > 0:
            theme_kind = "foundation"
        else:
            theme_kind = "mixed"
        rows.append({
            "page_record": page_record,
            "summary": summarize_concept_page_for_overview(page_record),
            "theme_kind": theme_kind,
            "source_count": len(page_record.get("source_refs", [])),
            "claim_count": len(page_record.get("claim_ids", [])),
            "review_count": len(page_record.get("review_ids", [])),
        })
    return rows


def build_workspace_source_coverage_rows(
    concept_pages: list[dict],
    source_pages_by_id: dict[str, dict],
    limit: int = 8,
) -> list[dict]:
    grouped: dict[str, dict] = {}
    for concept_page in concept_pages:
        for source_ref in concept_page.get("source_refs", []):
            source_id = source_ref["source_id"]
            if source_id not in grouped:
                grouped[source_id] = {
                    "source_ref": source_ref,
                    "source_page": source_pages_by_id.get(expected_source_summary_page_id(source_id)),
                    "concept_titles": [],
                    "claim_ids": [],
                    "chunk_ids": [],
                }
            append_unique(grouped[source_id]["concept_titles"], concept_page.get("title", ""))
            for claim_id in source_ref.get("claim_ids", []):
                append_unique(grouped[source_id]["claim_ids"], claim_id)
            for chunk_id in source_ref.get("chunk_ids", []):
                append_unique(grouped[source_id]["chunk_ids"], chunk_id)
    rows = sorted(
        grouped.values(),
        key=lambda item: (
            len(item["concept_titles"]),
            len(item["claim_ids"]),
            item["source_ref"].get("source_id", ""),
        ),
        reverse=True,
    )
    return rows[:limit]


def build_workspace_overview_summary_text(
    concept_pages: list[dict],
    source_refs: list[dict],
    claim_records_by_id: dict[str, dict],
) -> str:
    claim_ids = {
        claim_id
        for page_record in concept_pages
        for claim_id in page_record.get("claim_ids", [])
    }
    key_theme_rows = build_workspace_overview_key_theme_rows(
        concept_pages=concept_pages,
        claim_records_by_id=claim_records_by_id,
        limit=3,
    )
    key_theme_titles = [item["page_record"].get("title", "") for item in key_theme_rows if item["page_record"].get("title")]
    key_theme_text = "、".join(key_theme_titles[:3]) if key_theme_titles else "若干稳定主题"
    operational_theme_count = sum(1 for item in key_theme_rows if item["theme_kind"] == "operational")
    summary_parts = [
        f"{key_theme_text} 是当前工作区里已经沉淀出的稳定主题。",
        f"这些主题当前覆盖 {len(claim_ids)} 条稳定 Claim 和 {len(source_refs)} 个来源。",
    ]
    if operational_theme_count:
        summary_parts.append(f"其中有 {operational_theme_count} 个主题带有更强的操作或判断信号。")
    else:
        summary_parts.append("这些主题目前以基础概念和事实定义为主。")
    return " ".join(summary_parts)


def format_overview_title_phrase(titles: list[str]) -> str:
    cleaned_titles = [str(title).strip() for title in titles if str(title).strip()]
    if not cleaned_titles:
        return ""
    if len(cleaned_titles) == 1:
        return cleaned_titles[0]
    if len(cleaned_titles) == 2:
        return f"{cleaned_titles[0]} 和 {cleaned_titles[1]}"
    return "、".join(cleaned_titles[:-1]) + f" 和 {cleaned_titles[-1]}"


def format_overview_theme_count_phrase(count: int) -> str:
    small_count_labels = {
        1: "一个",
        2: "两个",
        3: "三个",
        4: "四个",
        5: "五个",
        6: "六个",
    }
    return small_count_labels.get(count, f"{count} 个")


def build_workspace_overview_summary_grounding_references(
    concept_pages: list[dict],
    source_refs: list[dict],
    claim_records_by_id: dict[str, dict],
) -> list[str]:
    claim_ids = {
        claim_id
        for page_record in concept_pages
        for claim_id in page_record.get("claim_ids", [])
    }
    key_theme_rows = build_workspace_overview_key_theme_rows(
        concept_pages=concept_pages,
        claim_records_by_id=claim_records_by_id,
        limit=3,
    )
    key_theme_titles = [
        item["page_record"].get("title", "")
        for item in key_theme_rows
        if item["page_record"].get("title")
    ]
    operational_theme_count = sum(
        1 for item in key_theme_rows if item["theme_kind"] == "operational"
    )
    references = [
        build_workspace_overview_summary_text(
            concept_pages=concept_pages,
            source_refs=source_refs,
            claim_records_by_id=claim_records_by_id,
        ),
        f"这些主题当前覆盖 {len(claim_ids)} 条稳定 Claim 和 {len(source_refs)} 个来源。",
    ]
    if key_theme_titles:
        key_theme_text = "、".join(key_theme_titles[:3])
        references.append(f"{key_theme_text} 是当前工作区里已经沉淀出的稳定主题。")
        readable_title_text = format_overview_title_phrase(key_theme_titles[:3])
        if len(key_theme_titles) == 1:
            references.append(f"这个工作区主要围绕 {readable_title_text} 这个稳定主题展开。")
        else:
            references.append(
                f"这个工作区主要围绕 {readable_title_text} "
                f"{format_overview_theme_count_phrase(len(key_theme_titles))}稳定主题展开。"
            )
            references.append(f"工作区当前沉淀出的稳定主题包括 {readable_title_text}。")
    if operational_theme_count:
        references.append(f"其中有 {operational_theme_count} 个主题带有更强的操作或判断信号。")
    else:
        references.append("这些主题目前以基础概念和事实定义为主。")
    for page_record in concept_pages:
        references.append(page_record.get("title", ""))
        references.append(page_record.get("summary", ""))
    return [reference for reference in references if str(reference).strip()]


def text_is_grounded_in_reference(
    text: str,
    reference_text: str,
    *,
    min_overlap: int = 2,
    min_ratio: float = 0.35,
) -> bool:
    cleaned = markdown_to_plain_text(str(text)).strip()
    reference_cleaned = markdown_to_plain_text(str(reference_text)).strip()
    if not cleaned or not reference_cleaned:
        return False
    normalized = normalize_claim_text(cleaned)
    reference_normalized = normalize_claim_text(reference_cleaned)
    if reference_normalized and (reference_normalized in normalized or normalized in reference_normalized):
        return True

    grounded_tokens = set(tokenize_for_search(reference_cleaned))
    candidate_tokens = set(tokenize_for_search(cleaned))
    if not candidate_tokens or not grounded_tokens:
        return False

    overlap = grounded_tokens.intersection(candidate_tokens)
    if len(overlap) < min_overlap:
        return False
    return (len(overlap) / len(candidate_tokens)) >= min_ratio


def strip_overview_rewrite_framing(text: str) -> str:
    cleaned = markdown_to_plain_text(text).strip()
    if not cleaned:
        return ""
    cleaned = re.sub(r"\s*\(claims=\d+\s*,\s*sources=\d+\)\s*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.replace("|", " ")
    cleaned = re.sub(
        r"^如果你[^，。；:：]{0,32}[，,:：]\s*",
        "",
        cleaned,
    )
    cleaned = re.sub(r"[。.!?]+$", "", cleaned)
    cleaned = re.sub(
        r"^这个主题(?:会|将)?(?:解释|说明|介绍|展示|聚焦|围绕|讨论|主要讲)(?:了)?\s*",
        "",
        cleaned,
    )
    cleaned = re.sub(
        r"^如果你(?:更)?想[^，。；:：]{0,24}?(?:先读|先看|再看|接着读|优先从)\s*",
        "",
        cleaned,
    )
    cleaned = re.sub(r"^(?:建议|可以|可先|优先)?\s*(?:先读|先看|再看|接着读|优先从|从)\s*", "", cleaned)
    cleaned = re.sub(
        r"(?:\s*(?:主题|页面|页))?(?:\s*(?:往下钻|继续往下钻|开始|入手|了解))?\s*$",
        "",
        cleaned,
    )
    return cleaned.strip(" -:;,.!?。！？；：，、()[]{}\"'")


def llm_assisted_rewrite_text_is_grounded(text: str, claim_record: dict, title: str) -> bool:
    cleaned = markdown_to_plain_text(str(text)).strip()
    if not cleaned:
        return False
    normalized = normalize_claim_text(cleaned)
    claim_text = claim_record.get("text", "")
    claim_normalized = normalize_claim_text(claim_text)
    if not claim_normalized:
        return False
    if claim_normalized in normalized or normalized in claim_normalized:
        return True

    grounded_tokens = set(tokenize_for_search(claim_text))
    grounded_tokens.update(tokenize_for_search(title))
    candidate_tokens = set(tokenize_for_search(cleaned))
    if not candidate_tokens or not grounded_tokens:
        return False

    overlap = grounded_tokens.intersection(candidate_tokens)
    if len(overlap) < 2:
        return False
    return (len(overlap) / len(candidate_tokens)) >= 0.35


def llm_assisted_rewrite_text_is_grounded_in_page(text: str, page_record: dict) -> bool:
    cleaned = markdown_to_plain_text(str(text)).strip()
    if not cleaned:
        return False
    page_summary = markdown_to_plain_text(page_record.get("summary", "")).strip()
    title = page_record.get("title", "")
    allowed_text = " ".join(part for part in [title, page_summary] if part).strip()
    if text_is_grounded_in_reference(cleaned, allowed_text):
        return True

    stripped = strip_overview_rewrite_framing(cleaned)
    normalized_title = normalize_claim_text(title)
    stripped_normalized = normalize_claim_text(stripped)
    if normalized_title and stripped_normalized in {
        normalized_title,
        normalize_claim_text(f"{title} 主题"),
        normalize_claim_text(f"{title} 页面"),
    }:
        return True
    if stripped and stripped != cleaned and text_is_grounded_in_reference(stripped, allowed_text):
        return True
    return False


def llm_assisted_rewrite_text_is_grounded_in_pages(text: str, page_records: list[dict]) -> bool:
    return any(
        llm_assisted_rewrite_text_is_grounded_in_page(text, page_record)
        for page_record in page_records
    )


def llm_assisted_overview_summary_is_grounded(
    text: str,
    page_records: list[dict],
    claim_records_by_id: dict[str, dict],
) -> bool:
    sentences = split_text_into_sentences(text)
    if not sentences:
        return False
    source_refs = aggregate_source_refs_for_pages(page_records)
    reference_texts = build_workspace_overview_summary_grounding_references(
        concept_pages=page_records,
        source_refs=source_refs,
        claim_records_by_id=claim_records_by_id,
    )
    if not reference_texts:
        return False
    for sentence in sentences:
        if any(
            text_is_grounded_in_reference(sentence, reference_text, min_overlap=2, min_ratio=0.25)
            for reference_text in reference_texts
        ):
            continue
        if not llm_assisted_rewrite_text_is_grounded_in_pages(sentence, page_records):
            return False
    return True


def normalize_llm_assisted_rewrite_items(
    raw_items,
    allowed_claims_by_id: dict[str, dict],
    title: str,
    limit: int,
) -> list[dict]:
    normalized_items: list[dict] = []
    if not isinstance(raw_items, list):
        return normalized_items

    for raw_item in raw_items:
        if len(normalized_items) >= limit:
            break
        if not isinstance(raw_item, dict):
            continue
        claim_id = str(raw_item.get("claim_id", "")).strip()
        text = str(raw_item.get("text", "")).strip()
        claim_record = allowed_claims_by_id.get(claim_id)
        if claim_record is None:
            continue
        if not llm_assisted_rewrite_text_is_grounded(text, claim_record, title):
            continue
        normalized_items.append({
            "claim_record": claim_record,
            "text": markdown_to_plain_text(text).strip(),
        })
    return normalized_items


def normalize_llm_assisted_page_items(
    raw_items,
    allowed_pages_by_id: dict[str, dict],
    limit: int,
) -> list[dict]:
    normalized_items: list[dict] = []
    if not isinstance(raw_items, list):
        return normalized_items

    for raw_item in raw_items:
        if len(normalized_items) >= limit:
            break
        if not isinstance(raw_item, dict):
            continue
        page_id = str(raw_item.get("page_id", "")).strip()
        text = str(raw_item.get("text", "")).strip()
        page_record = allowed_pages_by_id.get(page_id)
        if page_record is None:
            continue
        if not llm_assisted_rewrite_text_is_grounded_in_page(text, page_record):
            continue
        normalized_items.append({
            "page_record": page_record,
            "text": markdown_to_plain_text(text).strip(),
        })
    return normalized_items


def extract_markdown_section_text(page_text: str, heading: str) -> str:
    lines = page_text.splitlines()
    collecting = False
    collected: list[str] = []

    for line in lines:
        stripped = line.strip()
        if stripped == heading:
            collecting = True
            continue
        if collecting and stripped.startswith("## "):
            break
        if collecting:
            collected.append(line)
    return "\n".join(collected).strip()


def extract_markdown_bullet_lines(section_text: str) -> list[str]:
    bullet_lines: list[str] = []
    for line in section_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("- "):
            bullet_lines.append(stripped)
    return bullet_lines


def parse_claim_id_from_markdown_reference(text: str) -> str | None:
    match = re.search(r"\[`([^`]+)`\]\([^)]+\)", text)
    if match:
        return match.group(1)
    return None


def parse_page_id_from_markdown_page_link(text: str) -> str | None:
    match = re.search(r"\]\([^)]+/([^/]+)/[^/)]+\.md\)", text)
    if match:
        return match.group(1)
    return None


def strip_claim_reference_from_bullet(text: str) -> str:
    cleaned = re.sub(r"\s*\(\[`[^`]+`\]\([^)]+\)\)\s*$", "", text).strip()
    return cleaned


def readable_concept_page_grounding_issues(
    target: Path,
    page_record: dict,
    claim_records_by_id: dict[str, dict],
) -> list[str]:
    if page_record.get("type") != "concept":
        return []

    page_path = page_record.get("page_path")
    if not page_path:
        return ["missing_page_path"]
    page_file = target / page_path
    if not page_file.exists():
        return ["missing_page_file"]

    stable_claim_records = [
        claim_records_by_id[claim_id]
        for claim_id in page_record.get("claim_ids", [])
        if claim_id in claim_records_by_id
    ]
    if not stable_claim_records:
        return ["missing_live_claims"]

    title = page_record.get("title", "")
    canonical_claim = choose_canonical_claim(stable_claim_records, title)
    page_text = page_file.read_text(encoding="utf-8")

    summary_text = extract_markdown_section_text(page_text, "## 摘要 / Summary")
    if summary_text:
        summary_ok = any(
            llm_assisted_rewrite_text_is_grounded(summary_text, claim_record, title)
            for claim_record in stable_claim_records
        )
        if not summary_ok:
            return ["summary_not_grounded"]

    core_definition_text = extract_markdown_section_text(page_text, "## 核心定义 / Core Definition")
    if core_definition_text and not llm_assisted_rewrite_text_is_grounded(core_definition_text, canonical_claim, title):
        return ["core_definition_not_grounded"]

    key_points_text = extract_markdown_section_text(page_text, "## 关键要点 / Key Points")
    for bullet in extract_markdown_bullet_lines(key_points_text):
        claim_id = parse_claim_id_from_markdown_reference(bullet)
        if not claim_id:
            return [f"key_point_missing_claim_ref:{bullet}"]
        claim_record = claim_records_by_id.get(claim_id)
        if claim_record is None:
            return [f"key_point_unknown_claim_ref:{claim_id}"]
        bullet_text = strip_claim_reference_from_bullet(bullet[2:])
        if not llm_assisted_rewrite_text_is_grounded(bullet_text, claim_record, title):
            return [f"key_point_not_grounded:{claim_id}"]

    practical_text = extract_markdown_section_text(page_text, "## 使用提示 / Practical Notes")
    for bullet in extract_markdown_bullet_lines(practical_text):
        fallback_text = "- 当前稳定结论以概念定义和基础事实为主，尚未整理出更多操作性提示。"
        if bullet == fallback_text:
            continue
        claim_id = parse_claim_id_from_markdown_reference(bullet)
        if not claim_id:
            return [f"practical_note_missing_claim_ref:{bullet}"]
        claim_record = claim_records_by_id.get(claim_id)
        if claim_record is None:
            return [f"practical_note_unknown_claim_ref:{claim_id}"]
        bullet_text = strip_claim_reference_from_bullet(bullet[2:])
        if not llm_assisted_rewrite_text_is_grounded(bullet_text, claim_record, title):
            return [f"practical_note_not_grounded:{claim_id}"]

    return []


def overview_page_grounding_issues(
    target: Path,
    page_record: dict,
    page_records_by_id: dict[str, dict],
    claim_records_by_id: dict[str, dict],
) -> list[str]:
    if page_record.get("type") != "overview":
        return []

    page_path = page_record.get("page_path")
    if not page_path:
        return ["missing_page_path"]
    page_file = target / page_path
    if not page_file.exists():
        return ["missing_page_file"]

    concept_pages = {
        record["page_id"]: record
        for record in page_records_by_id.values()
        if is_live_page_record(record) and record.get("type") == "concept"
    }
    if not concept_pages:
        return ["missing_concept_pages"]

    page_text = page_file.read_text(encoding="utf-8")
    summary_text = extract_markdown_section_text(page_text, "## 工作区综述 / Workspace Overview")
    if summary_text:
        summary_ok = llm_assisted_overview_summary_is_grounded(
            summary_text,
            list(concept_pages.values()),
            claim_records_by_id=claim_records_by_id,
        )
        if not summary_ok:
            return ["overview_summary_not_grounded"]

    theme_map_text = extract_markdown_section_text(page_text, "## 主题导览 / Theme Map")
    theme_map_bullets = extract_markdown_bullet_lines(theme_map_text)
    for bullet in theme_map_bullets:
        if bullet == "- 先读这些基础主题：" or bullet == "- 再看这些更偏操作或判断的主题：" or bullet == "- 当前还没有足够的稳定主题可用于生成综述。":
            continue
        page_id = parse_page_id_from_markdown_page_link(bullet)
        if not page_id:
            return [f"theme_map_missing_page_ref:{bullet}"]
        concept_page = concept_pages.get(page_id)
        if concept_page is None:
            return [f"theme_map_unknown_page_ref:{page_id}"]
        bullet_text = strip_claim_reference_from_bullet(bullet[2:])
        if not llm_assisted_rewrite_text_is_grounded_in_page(bullet_text, concept_page):
            return [f"theme_map_not_grounded:{page_id}"]

    reading_path_text = extract_markdown_section_text(page_text, "## 推荐阅读路径 / Suggested Reading Path")
    for bullet in extract_markdown_bullet_lines(reading_path_text):
        page_id = parse_page_id_from_markdown_page_link(bullet)
        if not page_id:
            return [f"reading_path_missing_page_ref:{bullet}"]
        concept_page = concept_pages.get(page_id)
        if concept_page is None:
            return [f"reading_path_unknown_page_ref:{page_id}"]
        bullet_text = strip_claim_reference_from_bullet(bullet[2:])
        if not llm_assisted_rewrite_text_is_grounded_in_page(bullet_text, concept_page):
            return [f"reading_path_not_grounded:{page_id}"]

    return []


def rendered_page_grounding_issues(
    target: Path,
    page_record: dict,
    claim_records_by_id: dict[str, dict],
    page_records_by_id: dict[str, dict] | None = None,
) -> list[str]:
    render_target = page_record_render_target(page_record)
    if render_target == "readable_concept":
        return readable_concept_page_grounding_issues(
            target=target,
            page_record=page_record,
            claim_records_by_id=claim_records_by_id,
        )
    if render_target == "overview" and page_records_by_id is not None:
        return overview_page_grounding_issues(
            target=target,
            page_record=page_record,
            page_records_by_id=page_records_by_id,
            claim_records_by_id=claim_records_by_id,
        )
    return []


def concept_page_quality_issues(page_record: dict, claim_records_by_id: dict[str, dict]) -> list[str]:
    if page_record.get("type") != "concept":
        return []
    title = page_record.get("title", "") or ""
    claim_ids = page_record.get("claim_ids", []) or []
    claim_records = [
        claim_records_by_id[claim_id]
        for claim_id in claim_ids
        if claim_id in claim_records_by_id
    ]
    if not claim_records:
        return ["missing_claim_records"]
    canonical_claim = claim_records[0]
    quality = concept_title_quality_details(
        title=title,
        canonical_claim=canonical_claim,
        claim_records=claim_records,
    )
    issues: list[str] = []
    if quality["classification"] == "reject":
        issues.append(f"rejected_title:{title}")
    if "generic_title" in quality["reasons"]:
        issues.append("generic_title")
    if "too_short" in quality["reasons"] or "very_short" in quality["reasons"]:
        issues.append("too_short")
    if "question_like_claim" in quality["reasons"]:
        issues.append("question_like_claim")
    return issues


def page_semantic_consistency_issues(page_record: dict, claim_records_by_id: dict[str, dict]) -> list[str]:
    page_type = str(page_record.get("type", "")).strip().lower()
    if page_type not in {"concept", "guide", "example", "topic", "reference", "timeline"}:
        return []

    page_intent = str(page_record.get("page_intent", "")).strip().lower()
    page_route = page_record.get("page_route", {}) if isinstance(page_record.get("page_route"), dict) else {}
    route_target = str(page_route.get("route_target", "")).strip().lower()
    route_decision_id = str(page_route.get("semantic_decision_id", "") or "").strip()
    semantic_decision_ids = set(page_record.get("semantic_decision_ids", []) or [])
    claim_ids = page_record.get("claim_ids", []) or []
    claim_records = [
        claim_records_by_id[claim_id]
        for claim_id in claim_ids
        if claim_id in claim_records_by_id
    ]
    if not claim_records:
        return []

    roles = {
        claim_knowledge_role(record)
        for record in claim_records
        if claim_knowledge_role(record)
    }
    intent_hints = {
        hint
        for record in claim_records
        for hint in claim_page_intent_hints(record)
    }
    issues: list[str] = []

    expected_intent_by_type = {
        "concept": "concept",
        "guide": "guide",
        "example": "example",
        "topic": "topic",
        "reference": "reference",
        "timeline": "timeline",
    }
    expected_intent = expected_intent_by_type.get(page_type)
    if page_intent and expected_intent and page_intent != expected_intent:
        issues.append(f"page_type_intent_mismatch:{page_type}!={page_intent}")
    if route_target and expected_intent and route_target != expected_intent:
        issues.append(f"page_route_target_mismatch:{page_type}!={route_target}")
    if not route_decision_id:
        issues.append("page_route_decision_missing")
    elif route_decision_id not in semantic_decision_ids:
        issues.append("page_route_decision_not_linked")

    if page_type == "concept":
        blocked_roles = sorted(role for role in roles if role in {"procedure", "example", "meta", "structural_shell", "opinion"})
        if blocked_roles:
            issues.append(f"concept_page_blocked_roles:{','.join(blocked_roles)}")
        if "reject" in intent_hints:
            issues.append("concept_page_reject_intent_hint")
    elif page_type == "guide":
        if roles and "procedure" not in roles:
            issues.append(f"guide_page_missing_procedure_role:{','.join(sorted(roles))}")
    elif page_type == "example":
        if roles and "example" not in roles:
            issues.append(f"example_page_missing_example_role:{','.join(sorted(roles))}")
    elif page_type == "topic":
        if roles and roles.issubset({"procedure", "example", "meta", "structural_shell"}):
            issues.append(f"topic_page_semantically_thin:{','.join(sorted(roles))}")
    elif page_type == "reference":
        if roles and roles.issubset({"procedure", "example", "meta"}):
            issues.append(f"reference_page_semantically_thin:{','.join(sorted(roles))}")
    elif page_type == "timeline":
        if roles and roles.issubset({"meta", "structural_shell"}):
            issues.append(f"timeline_page_semantically_thin:{','.join(sorted(roles))}")

    return issues


def page_intent_brake_issues(page_record: dict) -> list[str]:
    page_route = page_record.get("page_route", {}) if isinstance(page_record.get("page_route"), dict) else {}
    route_reason = str(page_route.get("route_reason", "")).strip()
    if route_reason.startswith("page_intent_validation_downgraded_"):
        return [route_reason]
    return []


def claim_semantic_risk_issues(
    claim_record: dict,
    semantic_decisions_by_id: dict[str, dict],
) -> list[str]:
    issues: list[str] = []
    if not is_live_claim_record(claim_record):
        return issues
    if claim_record.get("status") == "needs_review" or claim_record.get("review_reason"):
        return issues
    for decision_id in claim_record.get("semantic_decision_ids", []) or []:
        decision_record = semantic_decisions_by_id.get(str(decision_id))
        if not decision_record or decision_record.get("task_type") != "claim_role":
            continue
        risk_flags = normalize_string_list(decision_record.get("risk_flags"))
        ambiguous_flags = sorted(flag for flag in risk_flags if "ambiguous" in flag)
        if ambiguous_flags:
            issues.append(f"{claim_record.get('claim_id')}:{','.join(ambiguous_flags)}")
            break
    return issues


def run_llm_assisted_readable_concept_render(
    target: Path,
    render_config: dict,
    title: str,
    canonical_claim: dict,
    stable_claim_records: list[dict],
    key_point_claims: list[dict],
    practical_claims: list[dict],
    summary_text: str,
) -> dict | None:
    if render_config.get("mode") != "llm_assisted":
        return None
    command = render_config.get("command", [])
    if not command:
        return None

    payload = {
        "task": "render_readable_concept_page",
        "title": title,
        "canonical_claim": {
            "claim_id": canonical_claim["claim_id"],
            "text": canonical_claim["text"],
            "claim_type": canonical_claim.get("claim_type"),
        },
        "stable_claims": [
            {
                "claim_id": claim_record["claim_id"],
                "text": claim_record["text"],
                "claim_type": claim_record.get("claim_type"),
                "status": claim_record.get("status"),
            }
            for claim_record in stable_claim_records
        ],
        "default_summary": summary_text,
        "default_key_points": [
            {
                "claim_id": claim_record["claim_id"],
                "text": render_claim_as_sentence(claim_record, title),
            }
            for claim_record in key_point_claims
        ],
        "default_practical_notes": [
            {
                "claim_id": claim_record["claim_id"],
                "text": render_claim_as_sentence(claim_record, title),
            }
            for claim_record in practical_claims
        ],
        "instructions": (
            "Only rewrite for readability. Do not add new facts. "
            "Every rewritten bullet must stay grounded in the referenced claim."
        ),
    }

    try:
        completed = subprocess.run(
            command,
            cwd=target,
            input=json.dumps(payload, ensure_ascii=False),
            capture_output=True,
            text=True,
            timeout=render_config.get("timeout_seconds", 20),
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None

    if completed.returncode != 0:
        return None

    stdout = (completed.stdout or "").strip()
    if not stdout:
        return None

    try:
        raw_result = json.loads(stdout)
    except json.JSONDecodeError:
        return None
    if not isinstance(raw_result, dict):
        return None

    allowed_claims_by_id = {
        claim_record["claim_id"]: claim_record
        for claim_record in stable_claim_records
    }
    assisted_summary = str(raw_result.get("summary", "")).strip()
    if not llm_assisted_rewrite_text_is_grounded(assisted_summary, canonical_claim, title):
        assisted_summary = ""

    key_points = normalize_llm_assisted_rewrite_items(
        raw_result.get("key_points", []),
        allowed_claims_by_id=allowed_claims_by_id,
        title=title,
        limit=max(len(key_point_claims), 1),
    )
    practical_notes = normalize_llm_assisted_rewrite_items(
        raw_result.get("practical_notes", []),
        allowed_claims_by_id=allowed_claims_by_id,
        title=title,
        limit=max(len(practical_claims), 1),
    )

    if not assisted_summary and not key_points and not practical_notes:
        return None
    return {
        "summary": assisted_summary,
        "key_points": key_points,
        "practical_notes": practical_notes,
    }


def run_llm_assisted_overview_render(
    target: Path,
    render_config: dict,
    title: str,
    summary_text: str,
    theme_rows: list[dict],
    reading_path_rows: list[dict],
    claim_records_by_id: dict[str, dict],
) -> dict | None:
    if render_config.get("mode") != "llm_assisted":
        return None
    command = render_config.get("command", [])
    if not command:
        return None

    payload = {
        "task": "render_workspace_overview_page",
        "title": title,
        "default_summary": summary_text,
        "theme_rows": [
            {
                "page_id": item["page_record"]["page_id"],
                "title": item["page_record"].get("title", ""),
                "summary": item["summary"],
                "theme_kind": item["theme_kind"],
                "claim_count": item["claim_count"],
                "source_count": item["source_count"],
                "review_count": item["review_count"],
            }
            for item in theme_rows
        ],
        "reading_path_rows": [
            {
                "page_id": item["page_record"]["page_id"],
                "title": item["page_record"].get("title", ""),
                "summary": summarize_concept_page_for_overview(item["page_record"]),
            }
            for item in reading_path_rows
        ],
        "instructions": (
            "Only rewrite for readability. Do not add new facts. "
            "Every rewritten theme summary and reading-path bullet must stay grounded in the referenced concept page."
        ),
    }

    try:
        completed = subprocess.run(
            command,
            cwd=target,
            input=json.dumps(payload, ensure_ascii=False),
            capture_output=True,
            text=True,
            timeout=render_config.get("timeout_seconds", 20),
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None

    if completed.returncode != 0:
        return None

    stdout = (completed.stdout or "").strip()
    if not stdout:
        return None

    try:
        raw_result = json.loads(stdout)
    except json.JSONDecodeError:
        return None
    if not isinstance(raw_result, dict):
        return None

    allowed_pages_by_id = {
        item["page_record"]["page_id"]: item["page_record"]
        for item in theme_rows
    }
    allowed_pages_by_id.update({
        item["page_record"]["page_id"]: item["page_record"]
        for item in reading_path_rows
    })

    assisted_summary = str(raw_result.get("summary", "")).strip()
    if not llm_assisted_overview_summary_is_grounded(
        assisted_summary,
        list(allowed_pages_by_id.values()),
        claim_records_by_id=claim_records_by_id,
    ):
        assisted_summary = ""

    theme_rows_rewrite = normalize_llm_assisted_page_items(
        raw_result.get("theme_rows", []),
        allowed_pages_by_id=allowed_pages_by_id,
        limit=max(len(theme_rows), 1),
    )
    reading_path_rewrite = normalize_llm_assisted_page_items(
        raw_result.get("reading_path", []),
        allowed_pages_by_id=allowed_pages_by_id,
        limit=max(len(reading_path_rows), 1),
    )

    if not assisted_summary and not theme_rows_rewrite and not reading_path_rewrite:
        return None
    return {
        "summary": assisted_summary,
        "theme_rows": theme_rows_rewrite,
        "reading_path": reading_path_rewrite,
    }


def build_source_summary_page(
    source_record: dict,
    page_rel_path: Path,
    normalized_record: dict | None,
    claim_records: list[dict],
    chunk_records: list[dict],
) -> tuple[str, dict]:
    # 第一版 wiki 页面先从“来源摘要页”做起，稳定承载这一份资料的处理结果。
    page_id = f"page_src_{source_record['source_id']}"
    title = normalized_record["title"] if normalized_record else Path(source_record["source_path"]).stem
    summary_claims = summarize_claims_for_page(claim_records)
    summary_text = summary_claims[0] if summary_claims else f"Source summary for {title}"

    lines = [
        "---",
        f'page_id: "{page_id}"',
        f'title: "{title}"',
        'type: "source-summary"',
        f'canonical_id: "{page_id}"',
        'status: "draft"',
        'automation_level: "auto_with_log"',
        f'source_id: "{source_record["source_id"]}"',
        f'claim_count: {len(claim_records)}',
        f'chunk_count: {len(chunk_records)}',
        "---",
        "",
        f"# {title}",
        "",
        "## 原文概览 / Source Overview",
        "",
        f"- 来源路径: `{source_record['source_path']}`",
        f"- 来源类型: `{source_record['source_type']}`",
        f"- 当前状态: `{source_record.get('status', 'unknown')}`",
    ]

    if normalized_record is not None:
        lines.extend([
            f"- 标准化文件: `{normalized_record['normalized_path']}`",
            f"- 提取方式: `{normalized_record['extraction_method']}`",
            f"- 提取质量: `{normalized_record['extraction_quality']}`",
        ])
        if normalized_record.get("warnings"):
            lines.append(f"- 标准化警告: `{', '.join(normalized_record['warnings'])}`")

    lines.extend([
        "",
        "## 核心观点 / Key Points",
        "",
    ])
    if summary_claims:
        for claim_text in summary_claims:
            lines.append(f"- {claim_text}")
    else:
        lines.append("- 当前尚未生成可用 claim。")

    lines.extend([
        "",
        "## 知识声明 / Claims",
        "",
    ])
    if claim_records:
        for claim_record in claim_records:
            lines.append(
                f"- {format_claim_reference(page_rel_path, claim_record)} {format_claim_type_label(claim_record.get('claim_type'))} "
                f"{claim_record['text']}"
            )
    else:
        lines.append("- 暂无 claims。")

    lines.extend([
        "",
        "## 证据切块 / Chunks",
        "",
    ])
    if chunk_records:
        for chunk_record in chunk_records[:10]:
            lines.append(
                f"- {format_chunk_reference(page_rel_path, source_record['source_id'], chunk_record)}"
            )
        if len(chunk_records) > 10:
            lines.append(f"- ... 其余 {len(chunk_records) - 10} 个 chunk 省略")
    else:
        lines.append("- 暂无 chunks。")

    lines.extend([
        "",
        "## 后续建议 / Next Steps",
        "",
        "- 检查是否需要将高价值 claim 提升为概念页或综述页。",
        "- 若存在 partial / failed 提取，优先补强转换路径或人工审核。",
    ])

    page_text = "\n".join(lines).strip() + "\n"
    page_record = {
        "page_id": page_id,
        "title": title,
        "type": "source-summary",
        "canonical_id": page_id,
        "status": "draft",
        "lifecycle_status": "active",
        "automation_level": "auto_with_log",
        "review_reason": None,
        "summary": summary_text,
        "aliases": [],
        "redirect_to": None,
        "claim_ids": [item["claim_id"] for item in claim_records],
        "source_refs": [
            {
                "source_id": source_record["source_id"],
                "source_path": source_record["source_path"],
                "chunk_ids": [item["chunk_id"] for item in chunk_records],
            }
        ],
        "created": utc_now_iso(),
        "updated": utc_now_iso(),
        "archived_at": None,
    }
    return page_text, page_record


def build_concept_page(
    target: Path,
    bucket_key: str,
    page_rel_path: Path,
    claim_records: list[dict],
    page_records_by_id: dict[str, dict],
    review_records: list[dict],
    render_config: dict | None = None,
) -> tuple[str, dict]:
    render_target = page_record_render_target({"type": "concept"}) or "readable_concept"
    config = load_workspace_config(target)
    concept_claim_records = filter_claim_records_for_concept_path(claim_records)
    if concept_claim_records:
        primary_claim_records = concept_claim_records
    else:
        primary_claim_records = claim_records
    stable_claim_records = filter_live_stable_claim_records(primary_claim_records)
    if stable_claim_records:
        render_claim_records = stable_claim_records
    else:
        render_claim_records = primary_claim_records

    group_topic_label = choose_group_topic_label(render_claim_records)
    canonical_claim = choose_canonical_claim(render_claim_records, group_topic_label)
    page_id = build_concept_page_id(bucket_key)
    title, title_quality = resolve_concept_title_candidate(
        target=target,
        config=config,
        canonical_claim=canonical_claim,
        claim_records=render_claim_records,
        preferred_section_label=group_topic_label,
    )
    canonical_id = f"concept:{build_concept_canonical_key(title)}"
    canonical_display_text = render_claim_as_sentence(canonical_claim, title)
    review_ids = collect_review_ids_for_claims(
        [claim_record["claim_id"] for claim_record in render_claim_records],
        review_records,
    )
    source_pages = collect_source_summary_pages_for_claims(render_claim_records, page_records_by_id)
    source_refs = aggregate_source_refs_for_page(render_claim_records)

    sorted_claims = sorted(
        render_claim_records,
        key=lambda item: claim_record_rank_key(item, group_topic_label),
        reverse=True,
    )
    supporting_claims = [
        claim_record
        for claim_record in sorted_claims
        if claim_record["claim_id"] != canonical_claim["claim_id"]
        and not claim_is_topic_shell_text(claim_record, group_topic_label)
    ]
    key_point_claims = supporting_claims[:4]
    practical_claims = [
        claim_record
        for claim_record in supporting_claims
        if claim_record.get("claim_type") in {"comparison", "causal", "procedure", "warning", "evaluation"}
    ][:3]
    aliases = [
        alias
        for alias in {
            clean_concept_title_text(shorten_title_text(claim_record["text"], limit=36))
            for claim_record in render_claim_records
            if claim_record["text"] != canonical_claim["text"]
        }
        if alias
    ]
    section_alias = extract_primary_section_label(canonical_claim)
    if section_alias and section_alias != title:
        aliases.append(section_alias)
    aliases.extend(collect_section_label_aliases(render_claim_records))

    summary_text = build_readable_concept_summary_text(
        title=title,
        canonical_claim=canonical_claim,
        stable_claim_records=render_claim_records,
        source_refs=source_refs,
    )
    assisted_render = run_llm_assisted_readable_concept_render(
        target=target,
        render_config=render_config or {"mode": "deterministic", "command": [], "timeout_seconds": 20},
        title=title,
        canonical_claim=canonical_claim,
        stable_claim_records=render_claim_records,
        key_point_claims=key_point_claims,
        practical_claims=practical_claims,
        summary_text=summary_text,
    )
    rendered_summary_text = assisted_render.get("summary") if assisted_render else ""
    if not rendered_summary_text:
        rendered_summary_text = summary_text
    rendered_key_points = assisted_render.get("key_points", []) if assisted_render else []
    rendered_practical_notes = assisted_render.get("practical_notes", []) if assisted_render else []
    requested_render_mode = (render_config or {}).get("mode", "deterministic")
    render_status = (
        "llm_assisted"
        if assisted_render
        else "deterministic_fallback"
        if requested_render_mode == "llm_assisted"
        else "deterministic"
    )

    lines = [
        "---",
        f'page_id: "{page_id}"',
        f'title: "{title}"',
        'type: "concept"',
        f'render_target: "{render_target}"',
        f'canonical_id: "{canonical_id}"',
        f'status: "{"needs_review" if review_ids else "stable"}"',
        'automation_level: "auto_with_log"',
        f'render_mode: "{requested_render_mode}"',
        f'render_status: "{render_status}"',
        f'claim_count: {len(stable_claim_records)}',
        f'source_count: {len(source_refs)}',
        "---",
        "",
        f"# {title}",
        "",
        "## 概念摘要 / Concept Summary",
        "",
        f"- 规范概念键: `{build_concept_canonical_key(title)}`",
        f"- 聚类键: `{bucket_key}`",
        f"- 代表陈述: {canonical_display_text}",
        f"- 关联 Claim 数量: `{len(render_claim_records)}`",
        f"- 关联来源数量: `{len(source_refs)}`",
        f"- 关联审核项数量: `{len(review_ids)}`",
        "",
        "## 摘要 / Summary",
        "",
        rendered_summary_text,
        "",
        "## 核心定义 / Core Definition",
        "",
        canonical_display_text,
        "",
        f"支撑 Claim: {format_claim_reference(page_rel_path, canonical_claim)} {format_claim_type_label(canonical_claim.get('claim_type'))}",
        "",
        "## 核心陈述 / Canonical Claim",
        "",
        f"- {format_claim_reference(page_rel_path, canonical_claim)} {format_claim_type_label(canonical_claim.get('claim_type'))} {canonical_display_text}",
        "",
        "## 关键要点 / Key Points",
        "",
    ]

    if rendered_key_points:
        for item in rendered_key_points:
            lines.append(
                f"- {item['text']} "
                f"({format_claim_reference(page_rel_path, item['claim_record'])})"
            )
    elif key_point_claims:
        for claim_record in key_point_claims:
            lines.append(
                f"- {render_claim_as_sentence(claim_record, title)} "
                f"({format_claim_reference(page_rel_path, claim_record)})"
            )
    else:
        lines.append(f"- {canonical_display_text} ({format_claim_reference(page_rel_path, canonical_claim)})")

    lines.extend([
        "",
        "## 使用提示 / Practical Notes",
        "",
    ])
    if rendered_practical_notes:
        for item in rendered_practical_notes:
            lines.append(
                f"- {item['text']} "
                f"({format_claim_reference(page_rel_path, item['claim_record'])})"
            )
    elif practical_claims:
        for claim_record in practical_claims:
            lines.append(
                f"- {render_claim_as_sentence(claim_record, title)} "
                f"({format_claim_reference(page_rel_path, claim_record)})"
            )
    else:
        lines.append("- 当前稳定结论以概念定义和基础事实为主，尚未整理出更多操作性提示。")

    lines.extend([
        "",
        "## 支撑声明 / Supporting Claims",
        "",
    ])
    for claim_record in sorted_claims:
        lines.append(
            f"- {format_claim_reference(page_rel_path, claim_record)} {format_claim_type_label(claim_record.get('claim_type'))} {claim_record['text']} "
            f"(sources={len(claim_record.get('source_ids', []))}, "
            f"chunks={len(claim_record.get('chunk_ids', []))})"
        )

    lines.extend([
        "",
        "## 来源页面 / Source Pages",
        "",
    ])
    if source_pages:
        source_pages_by_id = {source_page["page_id"]: source_page for source_page in source_pages}
        for source_ref in source_refs:
            source_page = source_pages_by_id.get(f"page_src_{source_ref['source_id']}")
            if source_page is None:
                continue
            lines.append(f"- 来源摘要页: {format_source_page_label(page_rel_path, source_page)}")
            lines.append(f"  原始文件: {format_workspace_file_reference(page_rel_path, source_ref['source_path'])}")
            lines.append(f"  标识: {format_source_page_meta(source_page, source_ref)}")
    else:
        lines.append("- 当前还没有可链接的来源摘要页。")

    lines.extend([
        "",
        "## 证据入口 / Evidence Trail",
        "",
    ])
    for source_ref in source_refs:
        source_page = next(
            (item for item in source_pages if item["page_id"] == f"page_src_{source_ref['source_id']}"),
            None,
        )
        source_label = (
            format_source_page_label(page_rel_path, source_page)
            if source_page is not None
            else "`未生成来源摘要页`"
        )
        lines.append(
            f"- {source_label} | 原始文件: {format_workspace_file_reference(page_rel_path, source_ref['source_path'])} | "
            f"claims={len(source_ref['claim_ids'])}, chunks={len(source_ref['chunk_ids'])}"
        )
        lines.append(f"  标识: {format_source_page_meta(source_page, source_ref)}")
        if source_ref.get("chunks"):
            lines.append("  证据切块:")
            for chunk_ref in source_ref["chunks"][:6]:
                lines.append(f"  - {format_chunk_reference(page_rel_path, source_ref['source_id'], chunk_ref)}")
            if len(source_ref["chunks"]) > 6:
                lines.append(f"  - ... 其余 {len(source_ref['chunks']) - 6} 个 chunk")

    lines.extend([
        "",
        "## 维护状态 / Maintenance",
        "",
        f"- 页面状态: `{'needs_review' if review_ids else 'stable'}`",
        f"- 聚合 Claim 数量: `{len(render_claim_records)}`",
        f"- 稳定 Claim 数量: `{len(stable_claim_records)}`",
        f"- 覆盖来源数量: `{len(source_refs)}`",
        f"- 关联审核项数量: `{len(review_ids)}`",
    ])
    if review_ids:
        lines.append("- 当前已有稳定结论，但仍有未关闭的审核项，阅读时请结合证据页一起查看。")
    else:
        lines.append("- 当前页面由稳定 Claim 自动编译，适合作为优先阅读入口。")

    page_text = "\n".join(lines).strip() + "\n"
    page_record = {
        "page_id": page_id,
        "title": title,
        "type": "concept",
        "render_target": render_target,
        "canonical_id": canonical_id,
        "status": "needs_review" if review_ids else "stable",
        "lifecycle_status": "active",
        "automation_level": "auto_with_log",
        "render_mode": requested_render_mode,
        "render_status": render_status,
        "concept_title_quality": title_quality,
        "review_reason": "claim_reviews_attached" if review_ids else None,
        "summary": rendered_summary_text,
        "aliases": sorted(set(alias for alias in aliases if alias and alias != title))[:8],
        "redirect_to": None,
        "claim_ids": [claim_record["claim_id"] for claim_record in render_claim_records],
        "review_ids": review_ids,
        "source_refs": source_refs,
        "created": utc_now_iso(),
        "updated": utc_now_iso(),
        "archived_at": None,
    }
    return page_text, page_record


def build_workspace_overview_page(
    target: Path,
    page_rel_path: Path,
    concept_pages: list[dict],
    page_records_by_id: dict[str, dict],
    claim_records_by_id: dict[str, dict],
    render_config: dict | None = None,
) -> tuple[str, dict]:
    render_target = "overview"
    config = load_workspace_config(target)
    project_name = (
        str(config.get("project", {}).get("name", "")).strip()
        if isinstance(config.get("project"), dict)
        else ""
    ) or target.name
    page_id = build_workspace_overview_page_id()
    title = f"{project_name} 综述"
    canonical_id = "overview:workspace"
    source_refs = aggregate_source_refs_for_pages(concept_pages)
    source_pages_by_id = {
        page_record["page_id"]: page_record
        for page_record in page_records_by_id.values()
        if is_live_page_record(page_record) and page_record.get("type") == "source-summary"
    }
    claim_ids = sorted({
        claim_id
        for page_record in concept_pages
        for claim_id in page_record.get("claim_ids", [])
    })
    review_ids = sorted({
        review_id
        for page_record in concept_pages
        for review_id in page_record.get("review_ids", [])
    })
    summary_text = build_workspace_overview_summary_text(
        concept_pages=concept_pages,
        source_refs=source_refs,
        claim_records_by_id=claim_records_by_id,
    )
    key_theme_rows = build_workspace_overview_key_theme_rows(
        concept_pages=concept_pages,
        claim_records_by_id=claim_records_by_id,
        limit=6,
    )
    source_coverage_rows = build_workspace_source_coverage_rows(
        concept_pages=concept_pages,
        source_pages_by_id=source_pages_by_id,
        limit=8,
    )
    operational_rows = [item for item in key_theme_rows if item["theme_kind"] == "operational"]
    foundational_rows = [item for item in key_theme_rows if item["theme_kind"] != "operational"]
    requested_render_mode = (render_config or {}).get("mode", "deterministic")
    deterministic_reading_path_rows: list[dict] = []
    if foundational_rows:
        deterministic_reading_path_rows.append({"page_record": foundational_rows[0]["page_record"]})
    if operational_rows:
        deterministic_reading_path_rows.append({"page_record": operational_rows[0]["page_record"]})
    if key_theme_rows:
        densest_page = max(
            key_theme_rows,
            key=lambda item: (item["source_count"], item["claim_count"], item["review_count"]),
        )["page_record"]
        if all(item["page_record"]["page_id"] != densest_page["page_id"] for item in deterministic_reading_path_rows):
            deterministic_reading_path_rows.append({"page_record": densest_page})

    assisted_render = run_llm_assisted_overview_render(
        target=target,
        render_config=render_config or {"mode": "deterministic", "command": [], "timeout_seconds": 20},
        title=title,
        summary_text=summary_text,
        theme_rows=key_theme_rows,
        reading_path_rows=deterministic_reading_path_rows,
        claim_records_by_id=claim_records_by_id,
    )
    rendered_summary_text = assisted_render.get("summary") if assisted_render else ""
    if not rendered_summary_text:
        rendered_summary_text = summary_text
    rendered_theme_rows = assisted_render.get("theme_rows", []) if assisted_render else []
    rendered_reading_path = assisted_render.get("reading_path", []) if assisted_render else []
    render_status = (
        "llm_assisted"
        if assisted_render
        else "deterministic_fallback"
        if requested_render_mode == "llm_assisted"
        else "deterministic"
    )

    lines = [
        "---",
        f'page_id: "{page_id}"',
        f'title: "{title}"',
        'type: "overview"',
        f'render_target: "{render_target}"',
        f'canonical_id: "{canonical_id}"',
        f'status: "{"needs_review" if review_ids else "stable"}"',
        'automation_level: "auto_with_log"',
        f'render_mode: "{requested_render_mode}"',
        f'render_status: "{render_status}"',
        f'claim_count: {len(claim_ids)}',
        f'source_count: {len(source_refs)}',
        "---",
        "",
        f"# {title}",
        "",
        "## 工作区综述 / Workspace Overview",
        "",
        rendered_summary_text,
        "",
        "## 主题导览 / Theme Map",
        "",
    ]

    if rendered_theme_rows:
        for item in rendered_theme_rows:
            concept_page = item["page_record"]
            lines.append(
                f"- {format_page_label(page_rel_path, concept_page)} | {item['text']}"
            )
    elif foundational_rows:
        lines.append("- 先读这些基础主题：")
        for item in foundational_rows[:3]:
            concept_page = item["page_record"]
            lines.append(
                f"  - {format_page_label(page_rel_path, concept_page)} | {item['summary']} "
                f"(claims={item['claim_count']}, sources={item['source_count']})"
            )
    if operational_rows:
        lines.append("- 再看这些更偏操作或判断的主题：")
        for item in operational_rows[:3]:
            concept_page = item["page_record"]
            lines.append(
                f"  - {format_page_label(page_rel_path, concept_page)} | {item['summary']} "
                f"(claims={item['claim_count']}, sources={item['source_count']})"
            )
    if not key_theme_rows:
        lines.append("- 当前还没有足够的稳定主题可用于生成综述。")

    lines.extend([
        "",
        "## 推荐阅读路径 / Suggested Reading Path",
        "",
    ])
    if rendered_reading_path:
        for item in rendered_reading_path:
            concept_page = item["page_record"]
            lines.append(f"- {item['text']} ({format_page_label(page_rel_path, concept_page)})")
    else:
        if foundational_rows:
            first_page = foundational_rows[0]["page_record"]
            lines.append(f"- 如果你想先建立全局认识，建议先读 {format_page_label(page_rel_path, first_page)}。")
        if operational_rows:
            first_operational_page = operational_rows[0]["page_record"]
            lines.append(f"- 如果你更关心做法、风险或取舍，接着读 {format_page_label(page_rel_path, first_operational_page)}。")
        if key_theme_rows:
            densest_page = max(
                key_theme_rows,
                key=lambda item: (item["source_count"], item["claim_count"], item["review_count"]),
            )["page_record"]
            lines.append(f"- 如果你想追证据覆盖面，优先从 {format_page_label(page_rel_path, densest_page)} 往下钻。")

    if render_status == "llm_assisted":
        lines.extend([
            "",
            "## 改写回绑 / Rewrite Traceability",
            "",
            "<details>",
            "<summary>查看 overview 改写句与其回绑页面</summary>",
            "",
        ])
        if rendered_summary_text:
            lines.append(
                f"- 工作区综述摘要 -> 基于这些主题页聚合改写: "
                f"{', '.join(format_page_label(page_rel_path, item['page_record']) for item in key_theme_rows[:3]) or '`无可用主题页`'}"
            )
        for item in rendered_theme_rows:
            lines.append(
                f"- 主题导览句: `{item['text']}` -> {format_page_label(page_rel_path, item['page_record'])}"
            )
        for item in rendered_reading_path:
            lines.append(
                f"- 推荐阅读句: `{item['text']}` -> {format_page_label(page_rel_path, item['page_record'])}"
            )
        if not rendered_summary_text and not rendered_theme_rows and not rendered_reading_path:
            lines.append("- 当前没有可展示的 overview 改写回绑项。")
        lines.extend([
            "",
            "</details>",
        ])

    lines.extend([
        "",
        "## 来源覆盖 / Source Coverage",
        "",
    ])
    for row in source_coverage_rows:
        source_ref = row["source_ref"]
        source_page = row["source_page"]
        source_label = (
            format_source_page_label(page_rel_path, source_page)
            if source_page is not None
            else "`未生成来源摘要页`"
        )
        lines.append(
            f"- 来源页: {source_label} | 原始文件: {format_workspace_file_reference(page_rel_path, source_ref['source_path'])} | "
            f"concepts={len(row['concept_titles'])}, claims={len(row['claim_ids'])}, chunks={len(row['chunk_ids'])}"
        )
        if row["concept_titles"]:
            lines.append(f"  关联主题: {', '.join(row['concept_titles'][:4])}")

    lines.extend([
        "",
        "## 维护状态 / Maintenance",
        "",
        f"- 可读概念页数量: `{len(concept_pages)}`",
        f"- 稳定 Claim 数量: `{len(claim_ids)}`",
        f"- 覆盖来源数量: `{len(source_refs)}`",
        f"- 关联审核项数量: `{len(review_ids)}`",
    ])
    if review_ids:
        lines.append("- 当前综述仍关联未关闭审核项，阅读时请优先回到对应概念页与证据页确认边界。")
    else:
        lines.append("- 当前综述由稳定概念页自动汇总，适合作为工作区级入口。")

    page_text = "\n".join(lines).strip() + "\n"
    page_record = {
        "page_id": page_id,
        "title": title,
        "type": "overview",
        "render_target": render_target,
        "canonical_id": canonical_id,
        "status": "needs_review" if review_ids else "stable",
        "lifecycle_status": "active",
        "automation_level": "auto_with_log",
        "render_mode": requested_render_mode,
        "render_status": render_status,
        "review_reason": "claim_reviews_attached" if review_ids else None,
        "summary": rendered_summary_text,
        "aliases": [],
        "redirect_to": None,
        "claim_ids": claim_ids,
        "review_ids": review_ids,
        "source_refs": source_refs,
        "created": utc_now_iso(),
        "updated": utc_now_iso(),
        "archived_at": None,
    }
    return page_text, page_record


def page_intent_page_id(bucket_key: str, page_intent: str) -> str:
    bucket_hash = hashlib.sha256(f"{page_intent}|{bucket_key}".encode("utf-8")).hexdigest()
    return f"page_{page_intent[:3]}_{bucket_hash[:12]}"


def page_intent_page_path(page_intent: str, page_id: str, title: str) -> Path:
    filename = sanitize_page_filename(title)
    folder = {
        "guide": "guides",
        "example": "examples",
        "topic": "topics",
        "reference": "references",
        "timeline": "timelines",
    }.get(page_intent, "topics")
    return Path("wiki") / folder / page_id / f"{filename}.md"


def build_intent_routed_page(
    target: Path,
    config: dict,
    bucket_key: str,
    page_intent: str,
    page_rel_path: Path,
    claim_records: list[dict],
    page_records_by_id: dict[str, dict],
    review_records: list[dict],
) -> tuple[str, dict]:
    group_topic_label = choose_group_topic_label(claim_records)
    canonical_claim = choose_canonical_claim(claim_records, group_topic_label)
    if page_intent == "guide":
        title = clean_concept_title_text(group_topic_label or canonical_claim.get("text", "")) or "指南"
        canonical_id = f"guide:{build_concept_canonical_key(title)}"
        summary = f"{title} 的操作步骤与执行提示。"
        section_title = "步骤摘要 / Steps"
    elif page_intent == "example":
        title = clean_concept_title_text(group_topic_label or canonical_claim.get("text", "")) or "示例"
        canonical_id = f"example:{build_concept_canonical_key(title)}"
        summary = f"{title} 的样例与案例说明。"
        section_title = "示例内容 / Examples"
    elif page_intent == "reference":
        title = clean_concept_title_text(group_topic_label or canonical_claim.get("text", "")) or "参考"
        canonical_id = f"reference:{build_concept_canonical_key(title)}"
        summary = f"{title} 的参考信息、规则条目与检索入口。"
        section_title = "参考条目 / Reference Notes"
    elif page_intent == "timeline":
        title = clean_concept_title_text(group_topic_label or canonical_claim.get("text", "")) or "时间线"
        canonical_id = f"timeline:{build_concept_canonical_key(title)}"
        summary = f"{title} 的时间顺序事实与演变节点。"
        section_title = "时间节点 / Timeline Notes"
    else:
        title = clean_concept_title_text(group_topic_label or canonical_claim.get("text", "")) or "主题"
        canonical_id = f"topic:{build_concept_canonical_key(title)}"
        summary = f"{title} 的主题概览与相关证据入口。"
        section_title = "主题要点 / Topic Notes"

    page_id = page_intent_page_id(bucket_key, page_intent)
    review_ids = collect_review_ids_for_claims(
        [claim_record["claim_id"] for claim_record in claim_records],
        review_records,
    )
    source_refs = aggregate_source_refs_for_page(claim_records)
    source_pages = collect_source_summary_pages_for_claims(claim_records, page_records_by_id)

    lines = [
        "---",
        f'page_id: "{page_id}"',
        f'title: "{title}"',
        f'type: "{page_intent}"',
        f'canonical_id: "{canonical_id}"',
        f'status: "{"needs_review" if review_ids else "stable"}"',
        'automation_level: "auto_with_log"',
        f'claim_count: {len(claim_records)}',
        f'source_count: {len(source_refs)}',
        "---",
        "",
        f"# {title}",
        "",
        "## 摘要 / Summary",
        "",
        summary,
        "",
        f"## {section_title}",
        "",
    ]
    for claim_record in claim_records[:6]:
        lines.append(
            f"- {render_claim_as_sentence(claim_record, title)} "
            f"({format_claim_reference(page_rel_path, claim_record)})"
        )
    if page_intent == "timeline":
        lines.extend([
            "",
            "## 时间线来源 / Timeline Sources",
            "",
        ])
        for source_ref in source_refs[:8]:
            lines.append(
                f"- {source_ref.get('section_path') or '未标注章节'} | "
                f"{format_workspace_file_reference(page_rel_path, source_ref['source_path'])}"
            )
    lines.extend([
        "",
        "## 证据入口 / Evidence Trail",
        "",
    ])
    for source_ref in source_refs:
        source_page = next(
            (item for item in source_pages if item["page_id"] == f"page_src_{source_ref['source_id']}"),
            None,
        )
        source_label = (
            format_source_page_label(page_rel_path, source_page)
            if source_page is not None
            else "`未生成来源摘要页`"
        )
        lines.append(
            f"- {source_label} | 原始文件: {format_workspace_file_reference(page_rel_path, source_ref['source_path'])}"
        )

    page_text = "\n".join(lines).strip() + "\n"
    page_record = {
        "page_id": page_id,
        "title": title,
        "type": page_intent,
        "canonical_id": canonical_id,
        "status": "needs_review" if review_ids else "stable",
        "lifecycle_status": "active",
        "automation_level": "auto_with_log",
        "review_reason": "claim_reviews_attached" if review_ids else None,
        "page_intent": page_intent,
        "summary": summary,
        "aliases": [],
        "redirect_to": None,
        "claim_ids": [claim_record["claim_id"] for claim_record in claim_records],
        "review_ids": review_ids,
        "source_refs": source_refs,
        "created": utc_now_iso(),
        "updated": utc_now_iso(),
        "archived_at": None,
    }
    return page_text, page_record


def write_wiki_page(target: Path, relative_path: Path, page_text: str) -> None:
    # Wiki 页面属于最终产物，改写时也尽量走原子写，避免意外中断留下半截 Markdown。
    page_path = target / relative_path
    page_path.parent.mkdir(parents=True, exist_ok=True)
    atomic_write_text(page_path, page_text, encoding="utf-8")


def remove_stale_page_file(target: Path, previous_path: str, current_path: str) -> None:
    # 页面改名后，把旧文件清掉，避免 wiki 目录里残留同一 page_id 的历史壳文件。
    if not previous_path or previous_path == current_path:
        return
    previous_page_path = target / previous_path
    if previous_page_path.exists():
        previous_page_path.unlink()

    stop_dirs = {
        (target / "wiki").resolve(),
        (target / "wiki" / "concepts").resolve(),
        (target / "wiki" / "sources").resolve(),
    }
    parent = previous_page_path.parent
    while parent.exists() and parent.resolve() not in stop_dirs:
        try:
            parent.rmdir()
        except OSError:
            break
        parent = parent.parent


def build_page_signature(page_record: dict, page_text: str) -> str:
    # 页面签名用于判断“这页内容是否真的变了”。
    # 如果签名不变，就没必要重写页面文件、日志和页面索引记录。
    signature_payload = {
        "page_path": page_record.get("page_path", ""),
        "title": page_record.get("title", ""),
        "type": page_record.get("type", ""),
        "status": page_record.get("status", ""),
        "summary": page_record.get("summary", ""),
        "aliases": page_record.get("aliases", []),
        "canonical_id": page_record.get("canonical_id"),
        "claim_ids": page_record.get("claim_ids", []),
        "review_ids": page_record.get("review_ids", []),
        "source_refs": page_record.get("source_refs", []),
        "page_text": page_text,
    }
    return hashlib.sha256(
        json.dumps(signature_payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()


def upsert_wiki_page(
    target: Path,
    page_records_by_id: dict[str, dict],
    page_record: dict,
    page_text: str,
) -> tuple[dict, bool]:
    # 统一处理页面落盘、页面索引更新和“是否真的发生变化”的判断。
    page_record = dict(page_record)
    page_record["lifecycle_status"] = page_lifecycle_status_for_record(page_record)
    page_record["page_signature"] = build_page_signature(page_record, page_text)

    previous_record = page_records_by_id.get(page_record["page_id"])
    if previous_record is not None and previous_record.get("page_signature") == page_record["page_signature"]:
        # 内容没变时，保留旧 created / updated / signature，避免制造无意义噪声。
        return previous_record, False

    if previous_record is not None:
        remove_stale_page_file(
            target=target,
            previous_path=previous_record.get("page_path", ""),
            current_path=page_record.get("page_path", ""),
        )
        page_record["created"] = previous_record.get("created", page_record.get("created"))
    page_record["updated"] = utc_now_iso()

    write_wiki_page(target, Path(page_record["page_path"]), page_text)
    page_records_by_id[page_record["page_id"]] = page_record
    return page_record, True


def rebuild_wiki_index(target: Path, page_records: list[dict]) -> None:
    # index 先作为“wiki 总目录”存在，按页面类型分组后会比单一长列表更容易读。
    page_records = filter_live_page_records(page_records)
    lines = [
        "# Wiki 索引 / Wiki Index",
        "",
        "## 阅读页 / Readable Pages",
        "",
    ]

    concept_pages = [
        record for record in page_records
        if record.get("type") == "concept"
    ]
    overview_pages = [
        record for record in page_records
        if record.get("type") == "overview"
    ]
    source_pages = [
        record for record in page_records
        if record.get("type") == "source-summary"
    ]
    other_pages = [
        record for record in page_records
        if record.get("type") not in {"overview", "concept", "source-summary"}
    ]

    if overview_pages or concept_pages:
        for record in sorted(overview_pages, key=lambda item: item["title"].lower()):
            page_path = markdown_link_target(record.get("page_path", ""))
            lines.append(
                f"- [{record['title']}]({page_path}) "
                f"({record['type']}, claims={len(record.get('claim_ids', []))}, reviews={len(record.get('review_ids', []))}) "
                f"- {record['summary']}"
            )
        for record in sorted(concept_pages, key=lambda item: item["title"].lower()):
            page_path = markdown_link_target(record.get("page_path", ""))
            lines.append(
                f"- [{record['title']}]({page_path}) "
                f"({record['type']}, claims={len(record.get('claim_ids', []))}, reviews={len(record.get('review_ids', []))}) "
                f"- {record['summary']}"
            )
    else:
        lines.append("- 暂无可读概念页。")

    lines.extend([
        "",
        "## 来源页 / Source Pages",
        "",
    ])
    if source_pages:
        for record in sorted(source_pages, key=lambda item: item["title"].lower()):
            page_path = markdown_link_target(record.get("page_path", ""))
            lines.append(
                f"- [{record['title']}]({page_path}) "
                f"({record['type']}, claims={len(record.get('claim_ids', []))}) - {record['summary']}"
            )
    else:
        lines.append("- 暂无来源页。")

    if other_pages:
        lines.extend([
            "",
            "## 其他页面 / Other Pages",
            "",
        ])
        for record in sorted(other_pages, key=lambda item: item["title"].lower()):
            page_path = markdown_link_target(record.get("page_path", ""))
            lines.append(
                f"- [{record['title']}]({page_path}) "
                f"({record['type']}, claims={len(record.get('claim_ids', []))}) - {record['summary']}"
            )

    # index.md 每次 ingest 都会整体重建，因此也适合直接走原子覆盖写。
    atomic_write_text(target / "wiki" / "index.md", "\n".join(lines).strip() + "\n", encoding="utf-8")


def append_wiki_log(target: Path, task_id: str, changed_pages: list[dict]) -> None:
    # log 先走 append-only，记录每次 ingest 真实写入、更新或清理了哪些页面。
    log_path = target / "wiki" / "log.md"
    lines = [f"## [{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] ingest | {task_id}"]
    if changed_pages:
        for page in changed_pages:
            action = "removed" if page.get("removed") else "generated"
            lines.append(f"- {action} {page['type']} page: `{page['page_id']}` -> `{page['page_path']}`")
    else:
        lines.append("- no wiki pages generated in this run")
    lines.append("")
    with log_path.open("a", encoding="utf-8") as fh:
        fh.write("\n".join(lines))


def link_claims_to_page_in_memory(
    claim_records: list[dict],
    page_id: str,
    claims_by_id: dict[str, dict],
) -> set[str]:
    # 页面生成后，把 claim -> page 的反向关系先写进内存。
    # 这样一轮 ingest 内如果有很多页面命中同一批 claims，就不会每条 claim 都去重写一遍 claims.jsonl。
    dirty_claim_ids: set[str] = set()
    for claim_record in claim_records:
        if page_id in claim_record.get("page_ids", []):
            continue
        claim_record.setdefault("page_ids", []).append(page_id)
        claim_record["updated_at"] = utc_now_iso()
        claim_record["lifecycle_status"] = claim_lifecycle_status_for_record(claim_record)
        claims_by_id[claim_record["claim_id"]] = claim_record
        dirty_claim_ids.add(claim_record["claim_id"])
    return dirty_claim_ids


def link_reviews_to_page_in_memory(
    review_records: list[dict],
    page_id: str,
    claim_ids: list[str],
    reviews_by_id: dict[str, dict],
) -> set[str]:
    # review 记录也先在内存里补 page 反链，最后统一写回 reviews.jsonl 和 review 文件。
    claim_id_set = set(claim_ids)
    dirty_review_ids: set[str] = set()
    for review_record in review_records:
        if not claim_id_set.intersection(review_record.get("candidate_claim_ids", [])):
            continue
        if page_id in review_record.get("candidate_page_ids", []):
            continue
        review_record.setdefault("candidate_page_ids", []).append(page_id)
        review_record["lifecycle_status"] = review_lifecycle_status_for_record(review_record)
        reviews_by_id[review_record["review_id"]] = review_record
        dirty_review_ids.add(review_record["review_id"])
    return dirty_review_ids


def extract_markdown_headings(text: str) -> list[str]:
    # query 阶段会单独给 headings 打分，因此这里把 Markdown 标题抽出来。
    headings: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped.startswith("#"):
            continue
        heading_text = re.sub(r"^#{1,6}\s*", "", stripped).strip()
        if heading_text and heading_text.lower() not in QUERY_HEADING_BLACKLIST:
            headings.append(heading_text)
    return headings


def strip_frontmatter(text: str) -> str:
    # wiki 页面大多带 frontmatter，query 做正文检索时应当把这层元数据先剥掉。
    lines = text.splitlines()
    if len(lines) >= 3 and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                return "\n".join(lines[index + 1:])
    return text


def build_searchable_body_text(page_text: str) -> str:
    # body 字段不应该把“每页都重复的模板栏目标题”当成正文。
    # 这里先做一个保守清洗：去 frontmatter、去公共标题行，保留其余正文和列表内容。
    cleaned_lines: list[str] = []
    for line in strip_frontmatter(page_text).splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            heading_text = re.sub(r"^#{1,6}\s*", "", stripped).strip().lower()
            if heading_text in QUERY_HEADING_BLACKLIST:
                continue
        cleaned_lines.append(line)
    return markdown_to_plain_text("\n".join(cleaned_lines))


def tokenize_for_search(text: str) -> list[str]:
    # V1 先用一个非常保守、纯 Python 的中英混合切词器：
    # - 英文/数字连续串作为一个 token
    # - 中文按双字/三字滑窗补一层召回
    # 这样虽然不如专业中文分词细，但不引入额外依赖也能先把检索跑起来。
    normalized = normalize_claim_text(text)
    latin_tokens = re.findall(r"[a-z0-9_]+", normalized)
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", normalized)

    tokens = [token for token in latin_tokens if len(token) >= 2]
    if chinese_chars:
        joined = "".join(chinese_chars)
        # 单个中文字符召回虽然高，但噪声也很大；V1 默认从双字开始。
        for width in (2, 3):
            if len(joined) >= width:
                for index in range(len(joined) - width + 1):
                    tokens.append(joined[index:index + width])

    # 最后再补一层原始空格切分，给中英混排句子多一些命中机会。
    for part in normalized.split():
        if len(part) >= 2 and part not in tokens:
            tokens.append(part)
    return tokens


def normalize_query_text(text: str) -> str:
    # query normalization 先保持轻量：
    # 统一空白、大小写和常见标点，得到一个稳定查询串。
    return normalize_claim_text(text)


def detect_query_intent(query_text: str, normalized_query: str) -> str:
    # V1 的意图识别先保持可解释：
    # 只看一些高频提示词，不做复杂分类模型。
    combined_text = f"{query_text.lower()} {normalized_query}"
    # 某些意图比 definition 更具体，例如“如何建立来源追踪”“来源证据是什么”。
    # 这里先把 how_to / evidence 提到前面，避免被“是什么”或“来源”误吞。
    if any(marker in combined_text for marker in QUERY_INTENT_MARKERS["how_to"]):
        return "how_to"
    if any(marker in combined_text for marker in QUERY_INTENT_MARKERS["evidence"]):
        return "evidence"
    if any(marker in combined_text for marker in QUERY_INTENT_MARKERS["reference"]):
        return "reference"
    if any(marker in combined_text for marker in QUERY_INTENT_MARKERS["overview"]):
        return "overview"
    for intent, markers in QUERY_INTENT_MARKERS.items():
        if intent in {"how_to", "evidence", "reference", "overview"}:
            continue
        if any(marker in combined_text for marker in markers):
            return intent
    return "lookup"


def alias_match_boost(page_record: dict, normalized_query: str, alias_hits: list[dict]) -> tuple[float, list[str]]:
    # alias/canonical 命中如果只回传不参与排序，用户感受会很奇怪。
    # 这里给精确命中一个温和加权，不会压过真正强相关的正文命中。
    boost_reasons: list[str] = []
    title_norm = normalize_query_text(page_record.get("title", ""))
    aliases_norm = [normalize_query_text(alias) for alias in page_record.get("aliases", [])]
    canonical_norm = normalize_query_text(page_record.get("canonical_id", "") or "")

    boost = 1.0
    if normalized_query and title_norm == normalized_query:
        boost = max(boost, QUERY_EXACT_MATCH_MAX_BOOST)
        boost_reasons.append("title_exact")
    if normalized_query and canonical_norm == normalized_query:
        boost = max(boost, 1.30)
        boost_reasons.append("canonical_exact")
    if normalized_query and normalized_query in aliases_norm:
        boost = max(boost, 1.25)
        boost_reasons.append("alias_exact")

    alias_hit_page_ids = {item.get("page_id") for item in alias_hits}
    if page_record.get("page_id") in alias_hit_page_ids:
        boost = max(boost, 1.20)
        boost_reasons.append("alias_registry_hit")

    return boost, boost_reasons


def query_intent_page_type_boost(intent: str, page_record: dict) -> tuple[float, str | None]:
    # query intent 只做一层轻微调权，用来把“看定义”和“看证据”这类问题往更合适的页型推一点。
    page_type = page_record.get("type", "")
    page_status = page_record.get("status", "")

    if intent == "overview" and page_type == "overview":
        return 1.85, "intent_overview_prefers_overview_page"
    if intent == "overview" and page_type == "concept":
        return 1.05, "intent_overview_falls_back_to_readable_concept"
    if intent == "overview" and page_type == "source-summary":
        return 0.70, "intent_overview_deprioritizes_source_summary"
    if intent == "definition" and page_type == "concept":
        return 1.7, "intent_definition_prefers_concept"
    if intent == "compare" and page_type == "concept":
        return 1.12, "intent_compare_prefers_concept"
    if intent == "reference" and page_type == "reference":
        return 2.40, "intent_reference_prefers_reference_page"
    if intent == "reference" and page_type == "source-summary":
        return 1.45, "intent_reference_prefers_source"
    if intent == "reference" and page_type == "concept":
        return 0.35, "intent_reference_deprioritizes_concept_views"
    if intent == "how_to" and page_type == "source-summary":
        return 1.05, "intent_how_to_prefers_source"
    if intent == "evidence" and page_type == "source-summary":
        return 2.6, "intent_evidence_prefers_source"
    if intent == "evidence" and page_type == "topic":
        return 0.55, "intent_evidence_deprioritizes_topic_page"
    if intent == "evidence" and page_type == "guide":
        return 0.50, "intent_evidence_deprioritizes_guide_page"
    if intent == "evidence" and page_type == "example":
        return 0.50, "intent_evidence_deprioritizes_example_page"
    if intent == "evidence" and page_type == "concept":
        return 0.40, "intent_evidence_deprioritizes_concept"
    if intent == "timeline" and page_status == "stable":
        return 1.05, "intent_timeline_prefers_stable"
    return 1.0, None


def query_intent_field_multiplier(intent: str, field_name: str) -> float:
    # 字段乘子只做轻量调权，尽量不破坏 BM25 的主体排序逻辑。
    return QUERY_INTENT_FIELD_MULTIPLIERS.get(intent, {}).get(field_name, 1.0)


def expand_query_with_alias_registry(
    query_text: str,
    alias_index: dict,
) -> dict:
    # 这一步负责把用户输入变成“可检索对象”：
    # - 原始查询
    # - 规范化查询
    # - alias 命中的扩展词
    # - 如果 alias 唯一映射到某个 canonical_id，就一并记录规范目标
    normalized_query = normalize_query_text(query_text)
    alias_map = alias_index.get("alias_map", {})
    canonical_map = alias_index.get("canonical_map", {})
    matched_alias_entries = alias_map.get(normalized_query, [])

    alias_expansions: list[str] = []
    canonical_targets: list[dict] = []
    for entry in matched_alias_entries:
        canonical_id = entry.get("canonical_id")
        canonical_record = canonical_map.get(canonical_id, {})
        for candidate in [
            entry.get("title", ""),
            canonical_id or "",
            *canonical_record.get("aliases", []),
        ]:
            normalized_candidate = normalize_query_text(candidate)
            if normalized_candidate and normalized_candidate not in alias_expansions:
                alias_expansions.append(normalized_candidate)
        if canonical_record and canonical_record not in canonical_targets:
            canonical_targets.append(canonical_record)

    expanded_parts = [normalized_query, *alias_expansions]
    expanded_query = " ".join(part for part in expanded_parts if part).strip()
    expanded_tokens = tokenize_for_search(expanded_query)
    detected_intent = detect_query_intent(query_text, normalized_query)

    return {
        "raw_query": query_text,
        "normalized_query": normalized_query,
        "expanded_query": expanded_query or normalized_query,
        "query_tokens": expanded_tokens,
        "alias_hits": matched_alias_entries,
        "canonical_targets": canonical_targets,
        "intent": detected_intent,
    }


def build_page_field_texts(page_record: dict, page_text: str, claim_records_by_id: dict[str, dict]) -> dict[str, str]:
    # query 读的是“页面视角”，但 claim/source 相关内容也要折叠进字段里参与排序。
    claim_texts = []
    for claim_id in page_record.get("claim_ids", []):
        claim_record = claim_records_by_id.get(claim_id)
        if claim_record is None:
            continue
        claim_texts.append(claim_record.get("text", ""))

    source_ref_parts = []
    hierarchy_parts = []
    for source_ref in page_record.get("source_refs", []):
        source_ref_parts.append(source_ref.get("source_id", ""))
        source_ref_parts.append(source_ref.get("source_path", ""))
        hierarchy_parts.append(source_ref.get("section_path", ""))
        hierarchy_parts.append(source_ref.get("section_title", ""))
        hierarchy_parts.append(source_ref.get("parent_section_path", ""))
        section_path_parts = source_ref.get("section_path_parts", [])
        if isinstance(section_path_parts, list):
            hierarchy_parts.extend(str(part) for part in section_path_parts if str(part).strip())
        for chunk_ref in source_ref.get("chunks", []):
            chunk_section_path = chunk_ref.get("section_path", "")
            hierarchy_parts.append(chunk_section_path)
            parsed = parse_section_path(chunk_section_path)
            hierarchy_parts.extend(parsed.get("section_path_parts", []))
    hierarchy_parts.append(page_record.get("title", ""))
    hierarchy_parts.extend(page_record.get("aliases", []))

    return {
        "title": page_record.get("title", ""),
        "aliases": "\n".join(page_record.get("aliases", [])),
        "hierarchy": "\n".join(part for part in hierarchy_parts if part),
        "summary": page_record.get("summary", ""),
        "headings": "\n".join(extract_markdown_headings(page_text)),
        "body": build_searchable_body_text(page_text),
        "claim_text": "\n".join(claim_texts),
        "source_refs": "\n".join(source_ref_parts),
    }


def compute_document_frequency(documents: list[dict[str, list[str]]], field_name: str) -> dict[str, int]:
    # BM25 需要知道“某个 token 出现在多少文档里”，这里按字段分别统计。
    frequency: dict[str, int] = {}
    for document in documents:
        unique_tokens = set(document.get(field_name, []))
        for token in unique_tokens:
            frequency[token] = frequency.get(token, 0) + 1
    return frequency


def bm25_score(
    query_tokens: list[str],
    document_tokens: list[str],
    document_frequency: dict[str, int],
    total_documents: int,
    average_length: float,
    k1: float = QUERY_BM25_K1,
    b: float = QUERY_BM25_B,
) -> float:
    # 这里实现标准 BM25 基线公式，不做额外花活，方便后面替换成索引引擎时对齐行为。
    if not query_tokens or not document_tokens or total_documents == 0:
        return 0.0

    token_counts = Counter(document_tokens)
    document_length = len(document_tokens)
    score = 0.0

    for token in query_tokens:
        term_frequency = token_counts.get(token, 0)
        if term_frequency == 0:
            continue

        doc_freq = document_frequency.get(token, 0)
        inverse_document_frequency = math.log(1 + (total_documents - doc_freq + 0.5) / (doc_freq + 0.5))
        numerator = term_frequency * (k1 + 1)
        denominator = term_frequency + k1 * (1 - b + b * (document_length / max(average_length, 1e-9)))
        score += inverse_document_frequency * (numerator / denominator)

    return score


def query_page_type_weight(page_record: dict) -> float:
    # 设计文档里的类型名和当前实现的页面类型不完全一样，这里做一层兼容映射。
    page_type = page_record.get("type", "source-summary")
    return QUERY_PAGE_TYPE_WEIGHTS.get(page_type, QUERY_PAGE_TYPE_WEIGHTS.get("draft", 0.70))


def query_page_status_weight(page_record: dict) -> float:
    status = page_record.get("status", "draft")
    return QUERY_PAGE_STATUS_WEIGHTS.get(status, QUERY_PAGE_STATUS_WEIGHTS["draft"])


def page_type_profile(page_type: str) -> str:
    normalized = str(page_type or "").strip().lower()
    if normalized == "guide":
        return "guide"
    if normalized == "example":
        return "example"
    if normalized == "topic":
        return "topic"
    if normalized == "reference":
        return "reference"
    if normalized == "timeline":
        return "timeline"
    if normalized == "concept":
        return "concept"
    if normalized == "source-summary":
        return "source"
    if normalized == "overview":
        return "overview"
    return "generic"


def build_query_documents(
    target: Path,
    page_records: list[dict],
    claim_records_by_id: dict[str, dict],
) -> list[dict]:
    # 页面记录是 query 的主入口；正文和 claim 文本会在这里汇总成待检索文档。
    documents: list[dict] = []
    for page_record in filter_live_page_records(page_records):
        page_path = target / page_record["page_path"]
        if not page_path.exists():
            continue
        page_text = page_path.read_text(encoding="utf-8")
        field_texts = build_page_field_texts(page_record, page_text, claim_records_by_id)
        field_tokens = {
            field_name: tokenize_for_search(field_text)
            for field_name, field_text in field_texts.items()
        }
        documents.append({
            "page_record": page_record,
            "page_text": page_text,
            "field_texts": field_texts,
            "field_tokens": field_tokens,
        })
    return documents


def build_search_index_record(document: dict) -> dict:
    # 持久化索引只保存 query 真正需要的字段，避免把整页正文重复存得过重。
    page_record = document["page_record"]
    field_texts = document["field_texts"]
    field_tokens = document["field_tokens"]
    signature_payload = {
        "page_path": page_record.get("page_path", ""),
        "title": page_record.get("title", ""),
        "type": page_record.get("type", ""),
        "status": page_record.get("status", ""),
        "summary": page_record.get("summary", ""),
        "aliases": page_record.get("aliases", []),
        "canonical_id": page_record.get("canonical_id"),
        "claim_ids": page_record.get("claim_ids", []),
        "review_ids": page_record.get("review_ids", []),
        "source_refs": page_record.get("source_refs", []),
        "field_texts": field_texts,
    }
    document_signature = hashlib.sha256(
        json.dumps(signature_payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()
    return {
        "page_id": page_record["page_id"],
        "title": page_record.get("title", ""),
        "page_path": page_record.get("page_path", ""),
        "type": page_record.get("type", ""),
        "status": page_record.get("status", ""),
        "summary": page_record.get("summary", ""),
        "aliases": page_record.get("aliases", []),
        "canonical_id": page_record.get("canonical_id"),
        "claim_ids": page_record.get("claim_ids", []),
        "review_ids": page_record.get("review_ids", []),
        "source_refs": page_record.get("source_refs", []),
        "field_texts": field_texts,
        "field_tokens": field_tokens,
        # 把页面签名一并写进索引，后续 ingest 就可以先看页面有没有变，
        # 没变时直接复用整条索引记录，连页面文件都不用再读。
        "page_signature": page_record.get("page_signature"),
        "document_signature": document_signature,
        "indexed_at": utc_now_iso(),
        "index_version": SEARCH_PAGES_INDEX_VERSION,
    }


def write_search_pages_index(
    target: Path,
    page_records: list[dict],
    claim_records_by_id: dict[str, dict],
    previous_records: list[dict] | None = None,
) -> dict:
    # search_pages.jsonl 是 query 的派生索引：
    # 页面权威内容仍然在 wiki/*.md 和 state/pages.jsonl，索引只负责加速读取。
    index_path = target / SEARCH_PAGES_INDEX_REL_PATH
    index_path.parent.mkdir(parents=True, exist_ok=True)
    previous_by_page_id = {
        record["page_id"]: record for record in (previous_records or [])
        if record.get("page_id")
    }
    records: list[dict] = []
    reused_count = 0
    rebuilt_count = 0

    for page_record in filter_live_page_records(page_records):
        previous_record = previous_by_page_id.get(page_record["page_id"])
        if (
            previous_record is not None
            and previous_record.get("page_signature") == page_record.get("page_signature")
            and previous_record.get("index_version") == SEARCH_PAGES_INDEX_VERSION
        ):
            reused_record = dict(previous_record)
            # 索引命中完全相同内容时，保留旧 indexed_at，便于观察哪些记录真的被重建过。
            records.append(reused_record)
            reused_count += 1
            continue

        page_path = target / page_record["page_path"]
        if not page_path.exists():
            continue

        page_text = page_path.read_text(encoding="utf-8")
        field_texts = build_page_field_texts(page_record, page_text, claim_records_by_id)
        field_tokens = {
            field_name: tokenize_for_search(field_text)
            for field_name, field_text in field_texts.items()
        }
        record = build_search_index_record({
            "page_record": page_record,
            "page_text": page_text,
            "field_texts": field_texts,
            "field_tokens": field_tokens,
        })

        records.append(record)
        rebuilt_count += 1

    write_jsonl(index_path, records)
    return {
        "index_path": str(SEARCH_PAGES_INDEX_REL_PATH),
        "record_count": len(records),
        "rebuilt_count": rebuilt_count,
        "reused_count": reused_count,
        "index_version": SEARCH_PAGES_INDEX_VERSION,
        "updated_at": utc_now_iso(),
    }


def load_search_pages_index(target: Path) -> list[dict]:
    index_path = target / SEARCH_PAGES_INDEX_REL_PATH
    if not index_path.exists():
        return []
    return load_jsonl(index_path)


def index_records_to_query_documents(index_records: list[dict]) -> list[dict]:
    # query 从索引恢复时，拼回与“现算文档”一致的形状，这样评分逻辑就可以复用。
    documents: list[dict] = []
    for record in index_records:
        page_record = {
            "page_id": record["page_id"],
            "title": record.get("title", ""),
            "page_path": record.get("page_path", ""),
            "type": record.get("type", ""),
            "status": record.get("status", ""),
            "summary": record.get("summary", ""),
            "aliases": record.get("aliases", []),
            "canonical_id": record.get("canonical_id"),
            "claim_ids": record.get("claim_ids", []),
            "review_ids": record.get("review_ids", []),
            "source_refs": record.get("source_refs", []),
        }
        documents.append({
            "page_record": page_record,
            "page_text": None,
            "field_texts": record.get("field_texts", {}),
            "field_tokens": record.get("field_tokens", {}),
        })
    return documents


def ensure_query_documents(
    target: Path,
    page_records: list[dict],
    claim_records_by_id: dict[str, dict],
) -> tuple[list[dict], str]:
    # query 优先读持久化索引；索引缺失时再降级现场构建，保证功能永远可用。
    index_records = load_search_pages_index(target)
    if index_records:
        return index_records_to_query_documents(index_records), "search_pages_index"
    return build_query_documents(target, page_records, claim_records_by_id), "live_scan"


def load_chunks_by_id(target: Path) -> dict[str, dict]:
    # chunk 阅读包要频繁按 chunk_id 回查，因此这里先建一个内存映射。
    chunks_path = target / "state" / "chunks.jsonl"
    if not chunks_path.exists():
        return {}
    chunk_records = load_jsonl(chunks_path)
    return {record["chunk_id"]: record for record in chunk_records}


def score_claim_for_query(query_tokens: list[str], claim_record: dict) -> tuple[float, list[str]]:
    # claim 排序先用一个轻量相关度：命中 token 越多、claim 自身越短小直接，越优先展示。
    claim_tokens = tokenize_for_search(claim_record.get("text", ""))
    matched_tokens = select_top_matches(query_tokens, claim_tokens, limit=8)
    if not matched_tokens:
        return 0.0, []
    score = float(len(matched_tokens))
    return score, matched_tokens


def score_chunk_for_query(query_tokens: list[str], chunk_record: dict) -> tuple[float, list[str]]:
    # chunk 相关度同样保持简单：看 chunk 文本与摘要里实际命中的 token。
    chunk_text = "\n".join([chunk_record.get("summary", ""), chunk_record.get("text", "")])
    chunk_tokens = tokenize_for_search(chunk_text)
    matched_tokens = select_top_matches(query_tokens, chunk_tokens, limit=8)
    section_tokens = tokenize_for_search(
        "\n".join(
            [
                chunk_record.get("section_title", ""),
                chunk_record.get("parent_section_path", ""),
                chunk_record.get("section_path", ""),
                "\n".join(chunk_record.get("section_path_parts", []) if isinstance(chunk_record.get("section_path_parts", []), list) else []),
            ]
        )
    )
    section_matches = select_top_matches(query_tokens, section_tokens, limit=8)
    if not matched_tokens and not section_matches:
        return 0.0, []
    score = float(len(matched_tokens))
    score += float(len(section_matches)) * 0.45
    if chunk_record.get("section_title") and section_matches:
        score += 0.2
    # 更短的 chunk 往往更聚焦，给一个很轻的偏好。
    score += 0.25 if chunk_record.get("char_count", 0) <= 600 else 0.0
    combined_matches = []
    for token in [*matched_tokens, *section_matches]:
        if token not in combined_matches:
            combined_matches.append(token)
    return score, combined_matches[:8]


def build_source_brief(source_ref: dict) -> dict:
    # 阅读包里不需要把 source 全量展开，先给一个短摘要，保持结果紧凑。
    return {
        "source_id": source_ref.get("source_id"),
        "source_path": source_ref.get("source_path"),
        "normalized_path": source_ref.get("normalized_path"),
        "start_line": source_ref.get("start_line"),
        "end_line": source_ref.get("end_line"),
        "section_path": source_ref.get("section_path"),
        "chunk_id": source_ref.get("chunk_id"),
    }


def build_chunk_reading_brief(chunk_record: dict) -> dict:
    # 设计文档要求命中 chunk 时附带 section_path / previous / next，这里固定打包出来。
    return {
        "chunk_id": chunk_record.get("chunk_id"),
        "section_path": chunk_record.get("section_path"),
        "start_line": chunk_record.get("start_line"),
        "end_line": chunk_record.get("end_line"),
        "summary": chunk_record.get("summary"),
        "text": chunk_record.get("text"),
        "previous_chunk": chunk_record.get("previous_chunk"),
        "next_chunk": chunk_record.get("next_chunk"),
        "source_id": chunk_record.get("source_id"),
        "source_path": chunk_record.get("source_path"),
        "normalized_path": chunk_record.get("normalized_path"),
    }


def build_timeline_sources(chunk_matches: list[dict]) -> list[dict]:
    # timeline query 需要的不是“更多 chunk”，而是“按来源组织的时序证据入口”。
    grouped: dict[str, dict] = {}
    for chunk in chunk_matches:
        source_id = chunk.get("source_id")
        if not source_id:
            continue
        if source_id not in grouped:
            grouped[source_id] = {
                "source_id": source_id,
                "source_path": chunk.get("source_path"),
                "normalized_path": chunk.get("normalized_path"),
                "chunk_ids": [],
                "section_paths": [],
            }
        if chunk.get("chunk_id") and chunk["chunk_id"] not in grouped[source_id]["chunk_ids"]:
            grouped[source_id]["chunk_ids"].append(chunk["chunk_id"])
        if chunk.get("section_path") and chunk["section_path"] not in grouped[source_id]["section_paths"]:
            grouped[source_id]["section_paths"].append(chunk["section_path"])
    return sorted(grouped.values(), key=lambda item: item["source_id"])


def build_source_trail(claim_matches: list[dict], chunk_matches: list[dict]) -> list[dict]:
    grouped: dict[str, dict] = {}
    for claim in claim_matches:
        for source_ref in claim.get("source_refs", []):
            source_id = source_ref.get("source_id")
            if not source_id:
                continue
            if source_id not in grouped:
                grouped[source_id] = {
                    "source_id": source_id,
                    "source_path": source_ref.get("source_path"),
                    "normalized_path": source_ref.get("normalized_path"),
                    "claim_ids": [],
                    "chunk_ids": [],
                    "section_paths": [],
                }
            if claim.get("claim_id") and claim["claim_id"] not in grouped[source_id]["claim_ids"]:
                grouped[source_id]["claim_ids"].append(claim["claim_id"])
            chunk_id = source_ref.get("chunk_id")
            if chunk_id and chunk_id not in grouped[source_id]["chunk_ids"]:
                grouped[source_id]["chunk_ids"].append(chunk_id)
            section_path = source_ref.get("section_path")
            if section_path and section_path not in grouped[source_id]["section_paths"]:
                grouped[source_id]["section_paths"].append(section_path)

    for chunk in chunk_matches:
        source_id = chunk.get("source_id")
        if not source_id:
            continue
        if source_id not in grouped:
            grouped[source_id] = {
                "source_id": source_id,
                "source_path": chunk.get("source_path"),
                "normalized_path": chunk.get("normalized_path"),
                "claim_ids": [],
                "chunk_ids": [],
                "section_paths": [],
            }
        if chunk.get("chunk_id") and chunk["chunk_id"] not in grouped[source_id]["chunk_ids"]:
            grouped[source_id]["chunk_ids"].append(chunk["chunk_id"])
        if chunk.get("section_path") and chunk["section_path"] not in grouped[source_id]["section_paths"]:
            grouped[source_id]["section_paths"].append(chunk["section_path"])

    return sorted(
        grouped.values(),
        key=lambda item: (len(item["claim_ids"]), len(item["chunk_ids"]), item["source_id"]),
        reverse=True,
    )


def build_hierarchy_match_explanation(result: dict, matched_chunks: list[dict]) -> dict:
    hierarchy_tokens = result.get("field_hits", {}).get("hierarchy", []) or []
    hierarchy_paths: list[str] = []
    matched_parent = False
    matched_leaf = False

    for chunk in matched_chunks:
        section_path = chunk.get("section_path")
        if section_path and section_path not in hierarchy_paths:
            hierarchy_paths.append(section_path)

    for source_ref in result.get("source_refs", []) or []:
        section_path = source_ref.get("section_path")
        if section_path and section_path not in hierarchy_paths:
            hierarchy_paths.append(section_path)
        for chunk_ref in source_ref.get("chunks", []) or []:
            chunk_section_path = chunk_ref.get("section_path")
            if chunk_section_path and chunk_section_path not in hierarchy_paths:
                hierarchy_paths.append(chunk_section_path)

    for section_path in hierarchy_paths:
        parsed = parse_section_path(section_path)
        section_parts = parsed.get("section_path_parts", [])
        if not section_parts:
            continue
        leaf_tokens = set(tokenize_for_search(section_parts[-1]))
        parent_tokens = set(tokenize_for_search(" ".join(section_parts[:-1])))
        if leaf_tokens.intersection(hierarchy_tokens):
            matched_leaf = True
        if parent_tokens.intersection(hierarchy_tokens):
            matched_parent = True

    if matched_parent and matched_leaf:
        anchor_reason = "matched_parent_and_leaf"
    elif matched_parent:
        anchor_reason = "matched_parent_only"
    elif matched_leaf:
        anchor_reason = "matched_leaf_only"
    else:
        anchor_reason = "matched_hierarchy_context"

    anchor_reason_text = {
        "matched_parent_and_leaf": "同时命中了父级路径和叶子标题，因此更偏向这个层级分支。",
        "matched_parent_only": "主要命中了父级路径，因此结果更偏向这个上层分类。",
        "matched_leaf_only": "主要命中了叶子标题，因此结果更偏向这个具体节点。",
        "matched_hierarchy_context": "命中了层级相关上下文，因此结果参考了章节路径信息。",
    }.get(anchor_reason, "命中了层级路径信息。")

    return {
        "matched_tokens": hierarchy_tokens,
        "matched_paths": hierarchy_paths[:5],
        "anchor_reason": anchor_reason,
        "anchor_reason_text": anchor_reason_text,
    }


def query_reading_focus(query_intent: str, page_type: str = "") -> str:
    if query_intent == "overview":
        return "workspace_overview"
    if query_intent == "compare":
        return "compare_claims"
    if query_intent == "timeline":
        return "timeline_evidence"
    if query_intent == "how_to":
        return "guide_steps" if page_type == "guide" else "procedural_chunks"
    if query_intent == "evidence":
        return "source_evidence"
    if page_type == "example":
        return "worked_examples"
    if page_type == "topic":
        return "topic_orientation"
    return "general_lookup"


def build_answer_guardrails(
    query_intent: str,
    page_status: str,
    page_type: str,
    review_ids: list[str],
    matched_claims: list[dict],
    matched_chunks: list[dict],
    timeline_sources: list[dict],
    source_trail: list[dict],
) -> dict:
    risk_flags: list[str] = []
    if review_ids:
        risk_flags.append("has_open_reviews")
    if page_status == "needs_review":
        risk_flags.append("page_needs_review")
    elif page_status == "disputed":
        risk_flags.append("page_disputed")
    elif page_status == "outdated":
        risk_flags.append("page_outdated")
    elif page_status == "draft":
        risk_flags.append("page_draft")
    if not matched_claims:
        risk_flags.append("no_matched_claims")
    if query_intent in {"how_to", "timeline", "evidence"} and not matched_chunks:
        risk_flags.append("no_matched_chunks")
    if query_intent == "timeline" and not timeline_sources:
        risk_flags.append("no_timeline_sources")
    if query_intent == "evidence" and not source_trail and not any(
        claim.get("source_refs") for claim in matched_claims
    ):
        risk_flags.append("weak_source_trace")

    can_answer_from_summary_only = (
        query_intent in {"lookup", "definition", "overview"}
        and page_type not in {"guide", "example"}
        and page_status not in {"needs_review", "disputed", "outdated"}
        and not review_ids
    )

    return {
        "can_answer_from_summary_only": can_answer_from_summary_only,
        "must_read_claims": query_intent in {"compare", "timeline", "evidence"},
        "must_read_chunks": query_intent in {"how_to", "timeline", "evidence"},
        "must_read_sources": query_intent in {"timeline", "evidence"},
        "cite_expectation": (
            "strong" if query_intent in {"timeline", "evidence"}
            else "light" if query_intent in {"compare", "how_to"}
            else "none"
        ),
        "risk_flags": risk_flags,
    }


def build_answer_handoff(query_intent: str, answer_guardrails: dict, page_type: str) -> dict:
    if query_intent in {"timeline", "evidence"}:
        answer_mode = "sources_first"
        recommended_read_order = [
            "retrieval_context.focus",
            "evidence_context.matched_claims",
            "evidence_context.matched_chunks",
            "evidence_context.timeline_sources" if query_intent == "timeline" else "evidence_context.source_trail",
            "page_context.summary",
        ]
    elif page_type == "guide":
        answer_mode = "chunks_first"
        recommended_read_order = [
            "retrieval_context.focus",
            "page_context.summary",
            "evidence_context.matched_chunks",
            "evidence_context.matched_claims",
        ]
    elif page_type == "example":
        answer_mode = "claims_first"
        recommended_read_order = [
            "retrieval_context.focus",
            "page_context.summary",
            "evidence_context.matched_claims",
            "evidence_context.matched_chunks",
        ]
    elif page_type == "topic" and query_intent in {"lookup", "definition", "overview"}:
        answer_mode = "summary_first"
        recommended_read_order = [
            "retrieval_context.focus",
            "page_context.summary",
            "evidence_context.matched_claims",
            "evidence_context.matched_chunks",
        ]
    elif query_intent == "how_to":
        answer_mode = "chunks_first"
        recommended_read_order = [
            "retrieval_context.focus",
            "evidence_context.matched_chunks",
            "evidence_context.matched_claims",
            "page_context.summary",
        ]
    elif query_intent == "compare":
        answer_mode = "claims_first"
        recommended_read_order = [
            "retrieval_context.focus",
            "evidence_context.matched_claims",
            "evidence_context.matched_chunks",
            "page_context.summary",
        ]
    else:
        answer_mode = "summary_first"
        recommended_read_order = [
            "page_context.summary",
            "evidence_context.matched_claims",
            "retrieval_context.focus",
        ]

    required_evidence_paths = []
    if answer_guardrails.get("must_read_claims"):
        required_evidence_paths.append("evidence_context.matched_claims")
    if answer_guardrails.get("must_read_chunks"):
        required_evidence_paths.append("evidence_context.matched_chunks")
    if answer_guardrails.get("must_read_sources"):
        required_evidence_paths.append(
            "evidence_context.timeline_sources" if query_intent == "timeline" else "evidence_context.source_trail"
        )

    risk_flags = answer_guardrails.get("risk_flags", [])
    if risk_flags:
        fallback_action = "answer_with_uncertainty"
    elif answer_guardrails.get("can_answer_from_summary_only"):
        fallback_action = "answer_from_summary_and_claims"
    else:
        fallback_action = "read_required_evidence_before_answering"

    return {
        "answer_mode": answer_mode,
        "recommended_read_order": recommended_read_order,
        "required_evidence_paths": required_evidence_paths,
        "should_cite_sources": answer_guardrails.get("cite_expectation") in {"light", "strong"},
        "should_surface_uncertainty": bool(risk_flags),
        "fallback_action": fallback_action,
    }


def build_result_reading_pack(
    result: dict,
    query_text: str,
    normalized_query: str,
    query_tokens: list[str],
    claim_records_by_id: dict[str, dict],
    chunk_records_by_id: dict[str, dict],
    claim_limit: int,
    chunk_limit: int,
    query_intent: str,
) -> dict:
    # 阅读包是 query V1 的关键补强：
    # 先返回“为什么命中这页”，再给最相关的 claims/chunks/sources，方便 Agent 继续读。
    claim_matches: list[dict] = []
    chunk_matches: list[dict] = []
    seen_chunk_ids: set[str] = set()

    for claim_id in result.get("claim_ids", []):
        claim_record = claim_records_by_id.get(claim_id)
        if claim_record is None:
            continue
        claim_score, claim_hits = score_claim_for_query(query_tokens, claim_record)
        if claim_score <= 0:
            continue

        claim_matches.append({
            "claim_id": claim_record["claim_id"],
            "text": claim_record.get("text", ""),
            "claim_type": claim_record.get("claim_type"),
            "status": claim_record.get("status"),
            "matched_tokens": claim_hits,
            "source_refs": [build_source_brief(item) for item in claim_record.get("source_refs", [])],
            "_score": (
                claim_score + (len(claim_record.get("source_refs", [])) * 0.25)
                if query_intent == "evidence"
                else claim_score + 0.4
                if query_intent == "compare" and claim_record.get("claim_type") in {"comparison", "evaluation", "causal"}
                else claim_score + 0.2
                if query_intent == "timeline" and claim_record.get("source_refs")
                else claim_score
            ),
        })

        for chunk_id in claim_record.get("chunk_ids", []):
            if chunk_id in seen_chunk_ids:
                continue
            chunk_record = chunk_records_by_id.get(chunk_id)
            if chunk_record is None:
                continue
            chunk_score, chunk_hits = score_chunk_for_query(query_tokens, chunk_record)
            if chunk_score <= 0:
                continue
            seen_chunk_ids.add(chunk_id)
            chunk_matches.append({
                "matched_tokens": chunk_hits,
                "_score": (
                    chunk_score + 0.5
                    if query_intent == "evidence" and chunk_record.get("source_id")
                    else chunk_score + 0.45
                    if query_intent == "how_to" and any(
                        marker in (chunk_record.get("text", "") + chunk_record.get("summary", ""))
                        for marker in ("步骤", "首先", "然后", "最后", "如何", "怎么")
                    )
                    else chunk_score + 0.25
                    if query_intent == "timeline" and chunk_record.get("source_id")
                    else chunk_score
                ),
                **build_chunk_reading_brief(chunk_record),
            })

    claim_matches.sort(key=lambda item: item["_score"], reverse=True)
    chunk_matches.sort(key=lambda item: item["_score"], reverse=True)

    trimmed_claims = []
    for item in claim_matches[:claim_limit]:
        cleaned = dict(item)
        cleaned.pop("_score", None)
        trimmed_claims.append(cleaned)

    trimmed_chunks = []
    for item in chunk_matches[:chunk_limit]:
        cleaned = dict(item)
        cleaned.pop("_score", None)
        trimmed_chunks.append(cleaned)

    reading_depth = result.get("reading_depth", "standard")
    source_trail = build_source_trail(trimmed_claims, trimmed_chunks) if reading_depth == "deep" else []
    timeline_sources = build_timeline_sources(trimmed_chunks) if query_intent == "timeline" else []
    page_type = result.get("type", "")
    focus = query_reading_focus(query_intent, page_type=page_type)
    hierarchy_explanation = build_hierarchy_match_explanation(result, trimmed_chunks)
    ranking_reasons = []
    if result.get("exact_match_reasons"):
        ranking_reasons.extend(result["exact_match_reasons"])
    if result.get("intent_boost_reason"):
        ranking_reasons.append(result["intent_boost_reason"])
    if hierarchy_explanation["matched_tokens"] or hierarchy_explanation["matched_paths"]:
        ranking_reasons.append(f"hierarchy_{hierarchy_explanation['anchor_reason']}")
    ranking_reasons.extend(sorted(result.get("field_scores", {}).keys()))
    answer_guardrails = build_answer_guardrails(
        query_intent=query_intent,
        page_status=result.get("status", ""),
        page_type=page_type,
        review_ids=result.get("review_ids", []),
        matched_claims=trimmed_claims,
        matched_chunks=trimmed_chunks,
        timeline_sources=timeline_sources,
        source_trail=source_trail,
    )
    answer_handoff = build_answer_handoff(
        query_intent=query_intent,
        answer_guardrails=answer_guardrails,
        page_type=page_type,
    )
    return {
        "contract_version": QUERY_ANSWER_HANDOFF_CONTRACT_VERSION,
        "handoff_kind": "reading_pack",
        "page_summary": result.get("summary", ""),
        "query_intent": query_intent,
        "reading_depth": reading_depth,
        "matched_claims": trimmed_claims,
        "matched_chunks": trimmed_chunks,
        "source_trail": source_trail,
        "timeline_sources": timeline_sources,
        "review_ids": result.get("review_ids", []),
        "focus": focus,
        "query": {
            "text": query_text,
            "normalized_text": normalized_query,
            "intent": query_intent,
            "reading_depth": reading_depth,
        },
        "page_context": {
            "page_id": result.get("page_id"),
            "title": result.get("title", ""),
            "page_path": result.get("page_path", ""),
            "type": result.get("type", ""),
            "page_type_profile": page_type_profile(result.get("type", "")),
            "status": result.get("status", ""),
            "summary": result.get("summary", ""),
            "canonical_id": result.get("canonical_id"),
            "aliases": result.get("aliases", []),
        },
        "retrieval_context": {
            "focus": focus,
            "matched_fields": sorted(result.get("field_hits", {}).keys()),
            "ranking_reasons": ranking_reasons,
            "review_ids": result.get("review_ids", []),
            "hierarchy_hits": hierarchy_explanation["matched_tokens"],
            "hierarchy_paths": hierarchy_explanation["matched_paths"],
            "hierarchy_anchor_reason": hierarchy_explanation["anchor_reason"],
            "hierarchy_anchor_reason_text": hierarchy_explanation["anchor_reason_text"],
        },
        "evidence_context": {
            "matched_claims": trimmed_claims,
            "matched_chunks": trimmed_chunks,
            "timeline_sources": timeline_sources,
            "source_trail": source_trail,
        },
        "answer_guardrails": answer_guardrails,
        "answer_handoff": answer_handoff,
    }


def build_answer_ready_payload(query_payload: dict) -> dict:
    base_payload = {
        "contract_version": ANSWER_READY_OUTPUT_VERSION,
        "query_contract_version": query_payload.get("contract_version"),
        "workspace": query_payload.get("workspace"),
        "workspace_summary": query_payload.get("workspace_summary"),
        "query": query_payload.get("query"),
        "normalized_query": query_payload.get("normalized_query"),
        "expanded_query": query_payload.get("expanded_query"),
        "intent": query_payload.get("intent"),
        "reading_depth": query_payload.get("reading_depth"),
        "selected_result": None,
        "alternatives": [],
        "agent_brief": {
            "answer_mode": "no_match",
            "page_type_profile": "unknown",
            "recommended_read_order": [],
            "required_evidence_paths": [],
            "should_cite_sources": False,
            "should_surface_uncertainty": True,
            "fallback_action": "broaden_or_rephrase_query",
            "risk_flags": ["no_query_results"],
        },
        "answer_context": {
            "page_summary": "",
            "answer_shape": "unknown",
            "key_claims": [],
            "key_chunks": [],
            "key_sources": [],
        },
        "agent_summary": "No matched page was found. Broaden or rephrase the query before attempting an answer.",
    }

    results = query_payload.get("results", [])
    if not results:
        return base_payload

    top_result = results[0]
    reading_pack = top_result.get("reading_pack", {})
    answer_guardrails = reading_pack.get("answer_guardrails", {})
    answer_handoff = reading_pack.get("answer_handoff", {})
    evidence_context = reading_pack.get("evidence_context", {})
    retrieval_context = reading_pack.get("retrieval_context", {})
    matched_fields = retrieval_context.get("matched_fields", [])
    weak_match = (
        not matched_fields
        or (
            top_result.get("score", 0.0) < 1.0
            and not evidence_context.get("matched_claims")
            and not evidence_context.get("matched_chunks")
        )
    )
    if weak_match:
        base_payload["agent_brief"] = {
            "answer_mode": "no_match",
            "page_type_profile": "unknown",
            "recommended_read_order": [],
            "required_evidence_paths": [],
            "should_cite_sources": False,
            "should_surface_uncertainty": True,
            "fallback_action": "broaden_or_rephrase_query",
            "risk_flags": ["weak_top_match"],
        }
        base_payload["selected_result"] = {
            "rank": 1,
            "page_id": top_result.get("page_id"),
            "title": top_result.get("title", ""),
            "page_path": top_result.get("page_path", ""),
            "type": top_result.get("type", ""),
            "status": top_result.get("status", ""),
            "summary": top_result.get("summary", ""),
            "score": top_result.get("score"),
            "focus": reading_pack.get("focus"),
            "page_type_profile": "unknown",
            "ready_state": "answer_with_uncertainty",
        }
        base_payload["alternatives"] = [
            {
                "rank": index,
                "page_id": result.get("page_id"),
                "title": result.get("title", ""),
                "page_path": result.get("page_path", ""),
                "type": result.get("type", ""),
                "status": result.get("status", ""),
                "score": result.get("score"),
            }
            for index, result in enumerate(results[1:4], start=2)
        ]
        base_payload["agent_summary"] = (
            "Top query result is too weak to serve as a safe answer anchor. "
            "Broaden or rephrase the query before attempting an answer."
        )
        base_payload["answer_context"]["answer_shape"] = "unknown"
        return base_payload

    key_claims = [
        {
            "claim_id": claim.get("claim_id"),
            "text": claim.get("text", ""),
            "claim_type": claim.get("claim_type"),
            "status": claim.get("status"),
            "source_ref_count": len(claim.get("source_refs", [])),
        }
        for claim in evidence_context.get("matched_claims", [])[:3]
    ]
    key_chunks = [
        {
            "chunk_id": chunk.get("chunk_id"),
            "section_path": chunk.get("section_path"),
            "summary": chunk.get("summary") or chunk.get("text", "")[:180],
            "start_line": chunk.get("start_line"),
            "end_line": chunk.get("end_line"),
            "source_path": chunk.get("source_path"),
        }
        for chunk in evidence_context.get("matched_chunks", [])[:3]
    ]
    key_sources = (
        evidence_context.get("timeline_sources")
        or evidence_context.get("source_trail")
        or [
            {
                "source_id": source_ref.get("source_id"),
                "source_path": source_ref.get("source_path"),
                "section_path": source_ref.get("section_path"),
                "chunk_id": source_ref.get("chunk_id"),
            }
            for claim in evidence_context.get("matched_claims", [])
            for source_ref in claim.get("source_refs", [])
        ][:5]
    )
    hierarchy_hits = retrieval_context.get("hierarchy_hits", [])
    hierarchy_paths = retrieval_context.get("hierarchy_paths", [])
    hierarchy_anchor_reason = retrieval_context.get("hierarchy_anchor_reason")
    hierarchy_anchor_reason_text = retrieval_context.get("hierarchy_anchor_reason_text")
    page_type = str(top_result.get("type", "")).strip().lower()
    page_profile = page_type_profile(page_type)
    answer_shape = (
        "step_by_step" if answer_handoff.get("answer_mode") == "chunks_first"
        else "worked_example" if page_profile == "example"
        else "topic_orientation" if page_profile == "topic"
        else "reference_sheet" if page_profile == "reference"
        else "timeline_evidence" if page_profile == "timeline"
        else "evidence_trace" if answer_handoff.get("answer_mode") == "sources_first"
        else "concept_summary" if page_profile == "concept"
        else "generic_summary"
    )

    if answer_handoff.get("fallback_action") == "answer_with_uncertainty":
        ready_state = "answer_with_uncertainty"
    elif answer_guardrails.get("can_answer_from_summary_only"):
        ready_state = "summary_ready"
    else:
        ready_state = "evidence_required"

    risk_flags = answer_guardrails.get("risk_flags", [])
    risk_text = ", ".join(risk_flags) if risk_flags else "none"
    selected_title = top_result.get("title", "")
    agent_summary_lines = [
        f"Use page '{selected_title}' as the answer anchor.",
        f"Answer mode: {answer_handoff.get('answer_mode', 'summary_first')}.",
        f"Read order: {' -> '.join(answer_handoff.get('recommended_read_order', [])) or 'page_context.summary'}.",
        f"Fallback action: {answer_handoff.get('fallback_action', 'read_required_evidence_before_answering')}.",
        f"Risk flags: {risk_text}.",
    ]
    if hierarchy_paths:
        agent_summary_lines.append(
            f"Hierarchy anchor: {' | '.join(hierarchy_paths[:3])}."
        )
    if hierarchy_anchor_reason_text:
        agent_summary_lines.append(f"Hierarchy reason: {hierarchy_anchor_reason_text}")
    elif hierarchy_anchor_reason:
        agent_summary_lines.append(f"Hierarchy reason: {hierarchy_anchor_reason}.")

    base_payload.update({
        "selected_result": {
            "rank": 1,
            "page_id": top_result.get("page_id"),
            "title": top_result.get("title", ""),
            "page_path": top_result.get("page_path", ""),
            "type": top_result.get("type", ""),
            "status": top_result.get("status", ""),
            "summary": top_result.get("summary", ""),
            "score": top_result.get("score"),
            "focus": reading_pack.get("focus"),
            "page_type_profile": page_profile,
            "ready_state": ready_state,
            "hierarchy_hits": hierarchy_hits,
            "hierarchy_paths": hierarchy_paths,
            "hierarchy_anchor_reason": hierarchy_anchor_reason,
            "hierarchy_anchor_reason_text": hierarchy_anchor_reason_text,
        },
        "alternatives": [
            {
                "rank": index,
                "page_id": result.get("page_id"),
                "title": result.get("title", ""),
                "page_path": result.get("page_path", ""),
                "type": result.get("type", ""),
                "status": result.get("status", ""),
                "score": result.get("score"),
            }
            for index, result in enumerate(results[1:4], start=2)
        ],
        "agent_brief": {
            "answer_mode": answer_handoff.get("answer_mode"),
            "page_type_profile": page_profile,
            "recommended_read_order": answer_handoff.get("recommended_read_order", []),
            "required_evidence_paths": answer_handoff.get("required_evidence_paths", []),
            "should_cite_sources": answer_handoff.get("should_cite_sources", False),
            "should_surface_uncertainty": answer_handoff.get("should_surface_uncertainty", False),
            "fallback_action": answer_handoff.get("fallback_action"),
            "risk_flags": risk_flags,
        },
        "answer_context": {
            "page_summary": reading_pack.get("page_summary", top_result.get("summary", "")),
            "answer_shape": answer_shape,
            "key_claims": key_claims,
            "key_chunks": key_chunks,
            "key_sources": key_sources,
            "hierarchy_hits": hierarchy_hits,
            "hierarchy_paths": hierarchy_paths,
            "hierarchy_anchor_reason": hierarchy_anchor_reason,
            "hierarchy_anchor_reason_text": hierarchy_anchor_reason_text,
        },
        "agent_summary": "\n".join(agent_summary_lines),
    })
    return base_payload


def render_answer_ready_message(answer_ready_payload: dict) -> str:
    lines = [
        f"Query: {answer_ready_payload['query']}",
        f"Intent: {answer_ready_payload['intent']}",
    ]
    selected_result = answer_ready_payload.get("selected_result")
    if selected_result is None:
        lines.append("")
        lines.append("Answer-Ready Summary:")
        lines.append("  No matched page was found.")
        lines.append("  Action: broaden or rephrase the query before answering.")
        return "\n".join(lines)

    agent_brief = answer_ready_payload.get("agent_brief", {})
    answer_context = answer_ready_payload.get("answer_context", {})
    lines.extend([
        "",
        "Answer-Ready Summary:",
        f"  anchor_page: {selected_result.get('title')} [{selected_result.get('type')}, status={selected_result.get('status')}]",
        f"  path: {selected_result.get('page_path')}",
        f"  ready_state: {selected_result.get('ready_state')}",
        f"  answer_mode: {agent_brief.get('answer_mode')}",
        f"  page_type_profile: {agent_brief.get('page_type_profile')}",
        f"  answer_shape: {answer_context.get('answer_shape')}",
        f"  fallback_action: {agent_brief.get('fallback_action')}",
        f"  read_order: {' -> '.join(agent_brief.get('recommended_read_order', []))}",
    ])
    if agent_brief.get("required_evidence_paths"):
        lines.append(f"  required_evidence: {', '.join(agent_brief['required_evidence_paths'])}")
    if agent_brief.get("risk_flags"):
        lines.append(f"  risk_flags: {', '.join(agent_brief['risk_flags'])}")
    if answer_context.get("hierarchy_paths"):
        lines.append(f"  hierarchy: {' | '.join(answer_context['hierarchy_paths'])}")
    if answer_context.get("hierarchy_anchor_reason_text"):
        lines.append(f"  hierarchy_reason: {answer_context.get('hierarchy_anchor_reason_text')}")
    elif answer_context.get("hierarchy_anchor_reason"):
        lines.append(f"  hierarchy_reason: {answer_context.get('hierarchy_anchor_reason')}")
    lines.append(f"  summary: {answer_context.get('page_summary', '')}")

    key_claims = answer_context.get("key_claims", [])
    if key_claims:
        lines.append("  key_claims:")
        for claim in key_claims:
            lines.append(f"    - {claim['claim_id']} {claim['text']}")

    key_chunks = answer_context.get("key_chunks", [])
    if key_chunks:
        lines.append("  key_chunks:")
        for chunk in key_chunks:
            lines.append(
                f"    - {chunk['chunk_id']} {chunk.get('section_path')} "
                f"(lines {chunk.get('start_line')}-{chunk.get('end_line')})"
            )

    key_sources = answer_context.get("key_sources", [])
    if key_sources:
        lines.append("  key_sources:")
        for source in key_sources[:5]:
            lines.append(
                f"    - {source.get('source_id')} {source.get('source_path')} "
                f"{source.get('section_path') or ''}".rstrip()
            )

    alternatives = answer_ready_payload.get("alternatives", [])
    if alternatives:
        lines.append("  alternatives:")
        for alternative in alternatives:
            lines.append(
                f"    - #{alternative['rank']} {alternative['title']} "
                f"[{alternative['type']}, status={alternative['status']}]"
            )
    return "\n".join(lines)


def render_answer_ready_prompt(answer_ready_payload: dict) -> str:
    selected_result = answer_ready_payload.get("selected_result")
    agent_brief = answer_ready_payload.get("agent_brief", {})
    answer_context = answer_ready_payload.get("answer_context", {})

    lines = [
        "You are the answer layer for a MyAgentWiki query handoff.",
        "Use only the provided handoff context to answer.",
        "If evidence is weak or risk flags are present, explicitly say what is uncertain.",
        "Do not invent citations or unsupported details.",
        "",
        "## Query",
        f"- user_query: {answer_ready_payload.get('query', '')}",
        f"- intent: {answer_ready_payload.get('intent', '')}",
        f"- reading_depth: {answer_ready_payload.get('reading_depth', '')}",
        "",
        "## Handoff",
        f"- answer_mode: {agent_brief.get('answer_mode', '')}",
        f"- page_type_profile: {agent_brief.get('page_type_profile', '')}",
        f"- answer_shape: {answer_context.get('answer_shape', '')}",
        f"- fallback_action: {agent_brief.get('fallback_action', '')}",
        f"- should_cite_sources: {agent_brief.get('should_cite_sources', False)}",
        f"- should_surface_uncertainty: {agent_brief.get('should_surface_uncertainty', False)}",
        f"- risk_flags: {', '.join(agent_brief.get('risk_flags', [])) or 'none'}",
        f"- recommended_read_order: {' -> '.join(agent_brief.get('recommended_read_order', [])) or 'none'}",
        f"- required_evidence_paths: {', '.join(agent_brief.get('required_evidence_paths', [])) or 'none'}",
        "",
        "## Selected Result",
    ]

    if selected_result is None:
        lines.extend([
            "- selected_result: none",
            "",
            "## Answer Instruction",
            "Explain that no reliable answer anchor was found and suggest broadening or rephrasing the query.",
        ])
        return "\n".join(lines)

    lines.extend([
        f"- title: {selected_result.get('title', '')}",
        f"- page_type: {selected_result.get('type', '')}",
        f"- page_status: {selected_result.get('status', '')}",
        f"- ready_state: {selected_result.get('ready_state', '')}",
        f"- page_path: {selected_result.get('page_path', '')}",
        f"- page_summary: {answer_context.get('page_summary', '')}",
    ])
    if answer_context.get("hierarchy_paths"):
        lines.append(f"- hierarchy_anchor: {' | '.join(answer_context.get('hierarchy_paths', [])[:3])}")
    if answer_context.get("hierarchy_hits"):
        lines.append(f"- hierarchy_hits: {'/'.join(answer_context.get('hierarchy_hits', []))}")
    if answer_context.get("hierarchy_anchor_reason_text"):
        lines.append(f"- hierarchy_reason: {answer_context.get('hierarchy_anchor_reason_text')}")
    elif answer_context.get("hierarchy_anchor_reason"):
        lines.append(f"- hierarchy_reason: {answer_context.get('hierarchy_anchor_reason')}")
    lines.extend([
        "",
        "## Key Claims",
    ])

    key_claims = answer_context.get("key_claims", [])
    if key_claims:
        for claim in key_claims:
            lines.append(
                f"- {claim.get('claim_id')}: {claim.get('text', '')} "
                f"(type={claim.get('claim_type')}, status={claim.get('status')})"
            )
    else:
        lines.append("- none")

    lines.extend([
        "",
        "## Key Chunks",
    ])
    key_chunks = answer_context.get("key_chunks", [])
    if key_chunks:
        for chunk in key_chunks:
            lines.append(
                f"- {chunk.get('chunk_id')}: {chunk.get('summary', '')} "
                f"[section={chunk.get('section_path')}, lines={chunk.get('start_line')}-{chunk.get('end_line')}, source={chunk.get('source_path')}]"
            )
    else:
        lines.append("- none")

    lines.extend([
        "",
        "## Key Sources",
    ])
    key_sources = answer_context.get("key_sources", [])
    if key_sources:
        for source in key_sources:
            lines.append(
                f"- source_id={source.get('source_id')} path={source.get('source_path')} "
                f"section={source.get('section_path') or ''} chunk_id={source.get('chunk_id') or ''}".rstrip()
            )
    else:
        lines.append("- none")

    lines.extend([
        "",
        "## Answer Instruction",
        "Write a concise answer for the user grounded only in the handoff above.",
        "If `page_type_profile` is `guide` or `answer_shape` is `step_by_step`, prefer a procedural step-by-step answer.",
        "If `page_type_profile` is `example` or `answer_shape` is `worked_example`, prefer explaining through a concrete example.",
        "If `page_type_profile` is `topic` or `answer_shape` is `topic_orientation`, first orient the user to the topic, then cite the most relevant claims.",
        "If `page_type_profile` is `reference` or `answer_shape` is `reference_sheet`, prefer a structured reference-style answer with concise keyed items.",
        "If `page_type_profile` is `timeline` or `answer_shape` is `timeline_evidence`, present events in chronological order and foreground dated evidence.",
        "If `answer_shape` is `evidence_trace`, foreground sources, chunks, and provenance instead of giving a bare summary.",
        "If `answer_shape` is `concept_summary`, prefer a concise definition-first explanation.",
        "If `should_cite_sources` is true, mention the supporting source paths or sections in the answer.",
        "If `should_surface_uncertainty` is true, include a short uncertainty note.",
        "If `fallback_action` is not `answer_from_summary_and_claims`, obey that fallback instead of overclaiming.",
    ])
    return "\n".join(lines)


def build_answer_ready_messages(answer_ready_payload: dict) -> list[dict]:
    prompt_text = render_answer_ready_prompt(answer_ready_payload)
    return [
        {
            "role": "system",
            "content": (
                "You are the answer layer for a MyAgentWiki handoff. "
                "Answer only from the provided context, surface uncertainty when required, "
                "and do not invent unsupported claims or citations."
            ),
        },
        {
            "role": "user",
            "content": prompt_text,
        },
    ]


def render_answer_ready_chatml(answer_ready_payload: dict) -> str:
    messages = build_answer_ready_messages(answer_ready_payload)
    blocks = []
    for message in messages:
        blocks.append(f"<|im_start|>{message['role']}\n{message['content']}\n<|im_end|>")
    return "\n".join(blocks)


def select_top_matches(query_tokens: list[str], field_tokens: list[str], limit: int = 5) -> list[str]:
    # 返回具体命中的 token，方便 CLI 输出“为什么这条结果排这么前面”。
    if not query_tokens or not field_tokens:
        return []
    field_counter = Counter(field_tokens)
    matches = []
    for token in query_tokens:
        if field_counter.get(token, 0) > 0 and token not in matches:
            matches.append(token)
    return matches[:limit]


def build_query_payload(
    target: Path,
    query_text: str,
    limit: int,
    claim_limit: int,
    chunk_limit: int,
    reading_depth: str = "standard",
) -> dict:
    # query 当前走“现算现查”策略：
    # 直接读取 state/pages.jsonl + claims + wiki 页面，避免先做一层复杂索引器。
    pages_path = target / "state" / "pages.jsonl"
    claims_path = target / "state" / "claims.jsonl"
    if not pages_path.exists():
        raise FileNotFoundError(f"Missing pages index: {pages_path}")

    page_records = [ensure_page_lifecycle_defaults(record) for record in load_jsonl(pages_path)]
    claim_records = [ensure_claim_lifecycle_defaults(record) for record in (load_jsonl(claims_path) if claims_path.exists() else [])]
    live_claim_records = filter_live_claim_records(claim_records)
    claim_records_by_id = {record["claim_id"]: record for record in live_claim_records}
    chunk_records_by_id = load_chunks_by_id(target)
    alias_index = load_alias_index(target)
    normalized_query_payload = expand_query_with_alias_registry(query_text, alias_index)

    normalized_query = normalized_query_payload["normalized_query"]
    query_tokens = normalized_query_payload["query_tokens"]
    query_intent = normalized_query_payload["intent"]
    documents, document_source = ensure_query_documents(target, page_records, claim_records_by_id)

    field_document_frequencies = {
        field_name: compute_document_frequency(
            [document["field_tokens"] for document in documents],
            field_name,
        )
        for field_name in QUERY_FIELD_WEIGHTS
    }
    field_average_lengths = {
        field_name: (
            sum(len(document["field_tokens"].get(field_name, [])) for document in documents) / len(documents)
            if documents else 0.0
        )
        for field_name in QUERY_FIELD_WEIGHTS
    }

    scored_results = []
    for document in documents:
        page_record = document["page_record"]
        field_scores: dict[str, float] = {}
        weighted_field_sum = 0.0
        field_hits: dict[str, list[str]] = {}

        for field_name, field_weight in QUERY_FIELD_WEIGHTS.items():
            document_tokens = document["field_tokens"].get(field_name, [])
            raw_score = bm25_score(
                query_tokens=query_tokens,
                document_tokens=document_tokens,
                document_frequency=field_document_frequencies[field_name],
                total_documents=len(documents),
                average_length=field_average_lengths[field_name],
            )
            if raw_score <= 0:
                continue
            intent_field_multiplier = query_intent_field_multiplier(query_intent, field_name)
            weighted_score = raw_score * field_weight * intent_field_multiplier
            field_scores[field_name] = round(weighted_score, 6)
            weighted_field_sum += weighted_score
            field_hits[field_name] = select_top_matches(query_tokens, document_tokens)

        if weighted_field_sum <= 0:
            continue

        page_type_weight = query_page_type_weight(page_record)
        page_status_weight = query_page_status_weight(page_record)
        exact_match_boost, exact_match_reasons = alias_match_boost(
            page_record=page_record,
            normalized_query=normalized_query,
            alias_hits=normalized_query_payload["alias_hits"],
        )
        intent_boost, intent_boost_reason = query_intent_page_type_boost(query_intent, page_record)
        final_score = weighted_field_sum * page_type_weight * page_status_weight * exact_match_boost * intent_boost
        result_record = {
            "page_id": page_record["page_id"],
            "title": page_record.get("title", ""),
            "page_path": page_record.get("page_path", ""),
            "type": page_record.get("type", ""),
            "canonical_id": page_record.get("canonical_id"),
            "status": page_record.get("status", ""),
            "summary": page_record.get("summary", ""),
            "aliases": page_record.get("aliases", []),
            "claim_ids": page_record.get("claim_ids", []),
            "review_ids": page_record.get("review_ids", []),
            "score": round(final_score, 6),
            "field_scores": field_scores,
            "field_hits": field_hits,
            "page_type_weight": page_type_weight,
            "page_status_weight": page_status_weight,
            "exact_match_boost": round(exact_match_boost, 4),
            "exact_match_reasons": exact_match_reasons,
            "intent": query_intent,
            "intent_boost": round(intent_boost, 4),
            "intent_boost_reason": intent_boost_reason,
            "reading_depth": reading_depth,
        }
        result_record["reading_pack"] = build_result_reading_pack(
            result=result_record,
            query_text=query_text,
            normalized_query=normalized_query,
            query_tokens=query_tokens,
            claim_records_by_id=claim_records_by_id,
            chunk_records_by_id=chunk_records_by_id,
            claim_limit=claim_limit,
            chunk_limit=chunk_limit,
            query_intent=query_intent,
        )
        scored_results.append(result_record)

    scored_results.sort(key=lambda item: (item["score"], item["title"]), reverse=True)
    return {
        "workspace": str(target),
        "workspace_summary": build_workspace_summary(target),
        "contract_version": QUERY_ANSWER_HANDOFF_CONTRACT_VERSION,
        "query": query_text,
        "normalized_query": normalized_query,
        "expanded_query": normalized_query_payload["expanded_query"],
        "query_tokens": query_tokens,
        "intent": query_intent,
        "reading_depth": reading_depth,
        "alias_hits": normalized_query_payload["alias_hits"],
        "canonical_targets": normalized_query_payload["canonical_targets"],
        "weights": {
            "fields": QUERY_FIELD_WEIGHTS,
            "intent_field_multipliers": QUERY_INTENT_FIELD_MULTIPLIERS,
            "page_types": QUERY_PAGE_TYPE_WEIGHTS,
            "page_status": QUERY_PAGE_STATUS_WEIGHTS,
            "exact_match_max_boost": QUERY_EXACT_MATCH_MAX_BOOST,
        },
        "reading_depth_limits": {
            "claim_limit": claim_limit,
            "chunk_limit": chunk_limit,
        },
        "document_source": document_source,
        "results": scored_results[:limit],
        "summary": {
            "candidate_page_count": len(documents),
            "matched_page_count": len(scored_results),
            "returned_page_count": min(limit, len(scored_results)),
        },
    }


def command_query(args: argparse.Namespace) -> CommandResult:
    # query 的 V1 目标是“先给出靠谱候选页 + 可解释排序原因”，而不是直接替用户写答案。
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    reading_depth = str(getattr(args, "reading_depth", "standard") or "standard").strip().lower()
    if reading_depth not in QUERY_READING_DEPTH_LIMITS:
        reading_depth = "standard"
    depth_limits = QUERY_READING_DEPTH_LIMITS[reading_depth]
    claim_limit = args.claim_limit if getattr(args, "claim_limit", None) is not None else depth_limits["claim_limit"]
    chunk_limit = args.chunk_limit if getattr(args, "chunk_limit", None) is not None else depth_limits["chunk_limit"]
    payload = build_query_payload(
        target=target,
        query_text=args.text,
        limit=args.limit,
        claim_limit=claim_limit,
        chunk_limit=chunk_limit,
        reading_depth=reading_depth,
    )
    if getattr(args, "answer_ready", False):
        answer_ready_payload = build_answer_ready_payload(payload)
        answer_ready_format = str(getattr(args, "format", "summary") or "summary").strip().lower()
        if answer_ready_format == "prompt":
            answer_ready_payload["prompt_text"] = render_answer_ready_prompt(answer_ready_payload)
        elif answer_ready_format == "messages":
            answer_ready_payload["messages"] = build_answer_ready_messages(answer_ready_payload)
        elif answer_ready_format == "chatml":
            answer_ready_payload["messages"] = build_answer_ready_messages(answer_ready_payload)
            answer_ready_payload["chatml_text"] = render_answer_ready_chatml(answer_ready_payload)
        if args.json:
            return CommandResult(
                payload=answer_ready_payload,
                message=render_workspace_summary_message(
                    "Answer-ready query completed.",
                    target_dir=target,
                    extra_lines=[
                        f"Query: {payload['query']}",
                        f"Intent: {payload['intent']}",
                        (
                            "Summary: "
                            f"candidates={payload['summary']['candidate_page_count']}, "
                            f"matched={payload['summary']['matched_page_count']}, "
                            f"returned={payload['summary']['returned_page_count']}"
                        ),
                    ],
                ),
            )
        if answer_ready_format == "prompt":
            return CommandResult(
                payload=answer_ready_payload,
                message=answer_ready_payload["prompt_text"],
            )
        if answer_ready_format == "messages":
            return CommandResult(
                payload=answer_ready_payload,
                message=json.dumps(answer_ready_payload["messages"], ensure_ascii=False, indent=2),
            )
        if answer_ready_format == "chatml":
            return CommandResult(
                payload=answer_ready_payload,
                message=answer_ready_payload["chatml_text"],
            )
        return CommandResult(
            payload=answer_ready_payload,
            message=render_workspace_summary_message(
                "Answer-ready query completed.",
                target_dir=target,
                extra_lines=[
                    f"Query: {payload['query']}",
                    f"Intent: {payload['intent']}",
                    "",
                    render_answer_ready_message(answer_ready_payload),
                ],
            ),
        )

    if args.json:
        return CommandResult(
            payload=payload,
            message=render_workspace_summary_message(
                "Query completed.",
                target_dir=target,
                extra_lines=[
                    f"Query: {payload['query']}",
                    f"Intent: {payload['intent']}",
                    (
                        "Summary: "
                        f"candidates={payload['summary']['candidate_page_count']}, "
                        f"matched={payload['summary']['matched_page_count']}, "
                        f"returned={payload['summary']['returned_page_count']}"
                    ),
                ],
            ),
        )

    if not payload["results"]:
        return CommandResult(
            payload=payload,
            message=render_workspace_summary_message(
                f"No wiki results matched query: {args.text}",
                target_dir=target,
                extra_lines=[
                    f"Intent: {payload['intent']}",
                    (
                        "Summary: "
                        f"candidates={payload['summary']['candidate_page_count']}, "
                        f"matched={payload['summary']['matched_page_count']}, "
                        f"returned={payload['summary']['returned_page_count']}"
                    ),
                ],
            ),
        )

    lines = [
        render_workspace_summary_message(
            "Query completed.",
            target_dir=target,
            extra_lines=[
                f'Query: {payload["query"]}',
                f'Normalized: {payload["normalized_query"]}',
                f'Expanded: {payload["expanded_query"]}',
                f'Intent: {payload["intent"]}',
                (
                    "Summary: "
                    f"candidates={payload['summary']['candidate_page_count']}, "
                    f"matched={payload['summary']['matched_page_count']}, "
                    f"returned={payload['summary']['returned_page_count']}"
                ),
                "",
                "Top Results:",
            ],
        )
    ]
    for index, result in enumerate(payload["results"], start=1):
        lines.append(
            f"{index}. {result['title']} [{result['type']}, status={result['status']}, score={result['score']:.4f}]"
        )
        lines.append(f"   path: {result['page_path']}")
        lines.append(f"   summary: {result['summary']}")
        if result["field_scores"]:
            explanation = ", ".join(
                f"{field}={score:.3f}"
                for field, score in sorted(result["field_scores"].items(), key=lambda item: item[1], reverse=True)
            )
            lines.append(f"   field_scores: {explanation}")
        if result.get("exact_match_reasons"):
            lines.append(
                f"   exact_match_boost: {result['exact_match_boost']:.3f} ({'/'.join(result['exact_match_reasons'])})"
            )
        if result.get("intent_boost_reason"):
            lines.append(
                f"   intent_boost: {result['intent_boost']:.3f} ({result['intent_boost_reason']})"
            )
        if result["field_hits"]:
            hit_explanation = ", ".join(
                f"{field}:{'/'.join(tokens)}"
                for field, tokens in result["field_hits"].items()
                if tokens
            )
            if hit_explanation:
                lines.append(f"   hits: {hit_explanation}")
        reading_pack = result.get("reading_pack", {})
        hierarchy_hits = reading_pack.get("retrieval_context", {}).get("hierarchy_hits", [])
        hierarchy_paths = reading_pack.get("retrieval_context", {}).get("hierarchy_paths", [])
        hierarchy_anchor_reason = reading_pack.get("retrieval_context", {}).get("hierarchy_anchor_reason")
        hierarchy_anchor_reason_text = reading_pack.get("retrieval_context", {}).get("hierarchy_anchor_reason_text")
        if hierarchy_hits or hierarchy_paths:
            hierarchy_explanation_parts = []
            if hierarchy_hits:
                hierarchy_explanation_parts.append(f"hits={'/'.join(hierarchy_hits)}")
            if hierarchy_paths:
                hierarchy_explanation_parts.append(f"paths={' | '.join(hierarchy_paths)}")
            if hierarchy_anchor_reason_text:
                hierarchy_explanation_parts.append(f"reason={hierarchy_anchor_reason_text}")
            elif hierarchy_anchor_reason:
                hierarchy_explanation_parts.append(f"reason={hierarchy_anchor_reason}")
            lines.append(f"   hierarchy: {'; '.join(hierarchy_explanation_parts)}")
        matched_claims = reading_pack.get("matched_claims", [])
        matched_chunks = reading_pack.get("matched_chunks", [])
        if matched_claims:
            lines.append("   matched_claims:")
            for claim in matched_claims:
                lines.append(
                    f"     - {claim['claim_id']} {format_claim_type_label(claim.get('claim_type'))} "
                    f"{claim['text']} (hits={'/'.join(claim.get('matched_tokens', []))})"
                )
        if matched_chunks:
            lines.append("   matched_chunks:")
            for chunk in matched_chunks:
                lines.append(
                    f"     - {chunk['chunk_id']} {chunk['section_path']} "
                    f"(lines {chunk['start_line']}-{chunk['end_line']}, hits={'/'.join(chunk.get('matched_tokens', []))})"
                )
                lines.append(
                    f"       prev={chunk.get('previous_chunk')} next={chunk.get('next_chunk')}"
                )
        source_trail = reading_pack.get("source_trail", [])
        if source_trail:
            lines.append("   source_trail:")
            for source in source_trail:
                lines.append(
                    f"     - {source['source_id']} "
                    f"(claims={len(source.get('claim_ids', []))}, chunks={len(source.get('chunk_ids', []))}) "
                    f"{source.get('source_path')}"
                )
    return CommandResult(payload=payload, message="\n".join(lines))


def command_answer_query(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    reading_depth = str(getattr(args, "reading_depth", "standard") or "standard").strip().lower()
    if reading_depth not in QUERY_READING_DEPTH_LIMITS:
        reading_depth = "standard"
    depth_limits = QUERY_READING_DEPTH_LIMITS[reading_depth]
    claim_limit = args.claim_limit if getattr(args, "claim_limit", None) is not None else depth_limits["claim_limit"]
    chunk_limit = args.chunk_limit if getattr(args, "chunk_limit", None) is not None else depth_limits["chunk_limit"]
    query_payload = build_query_payload(
        target=target,
        query_text=args.text,
        limit=args.limit,
        claim_limit=claim_limit,
        chunk_limit=chunk_limit,
        reading_depth=reading_depth,
    )
    answer_ready_payload = build_answer_ready_payload(query_payload)
    answer_ready_format = str(getattr(args, "format", "summary") or "summary").strip().lower()
    if answer_ready_format == "prompt":
        answer_ready_payload["prompt_text"] = render_answer_ready_prompt(answer_ready_payload)
    elif answer_ready_format == "messages":
        answer_ready_payload["messages"] = build_answer_ready_messages(answer_ready_payload)
    elif answer_ready_format == "chatml":
        answer_ready_payload["messages"] = build_answer_ready_messages(answer_ready_payload)
        answer_ready_payload["chatml_text"] = render_answer_ready_chatml(answer_ready_payload)
    if args.json:
        return CommandResult(
            payload=answer_ready_payload,
            message=render_workspace_summary_message(
                "Answer-ready query completed.",
                target_dir=target,
                extra_lines=[
                    f"Query: {query_payload['query']}",
                    f"Intent: {query_payload['intent']}",
                    (
                        "Summary: "
                        f"candidates={query_payload['summary']['candidate_page_count']}, "
                        f"matched={query_payload['summary']['matched_page_count']}, "
                        f"returned={query_payload['summary']['returned_page_count']}"
                    ),
                ],
            ),
        )
    if answer_ready_format == "prompt":
        return CommandResult(
            payload=answer_ready_payload,
            message=answer_ready_payload["prompt_text"],
        )
    if answer_ready_format == "messages":
        return CommandResult(
            payload=answer_ready_payload,
            message=json.dumps(answer_ready_payload["messages"], ensure_ascii=False, indent=2),
        )
    if answer_ready_format == "chatml":
        return CommandResult(
            payload=answer_ready_payload,
            message=answer_ready_payload["chatml_text"],
        )
    return CommandResult(
        payload=answer_ready_payload,
        message=render_workspace_summary_message(
            "Answer-ready query completed.",
            target_dir=target,
            extra_lines=[
                f"Query: {query_payload['query']}",
                f"Intent: {query_payload['intent']}",
                "",
                render_answer_ready_message(answer_ready_payload),
            ],
        ),
    )


def build_review_list_payload(target: Path, status_filter: str | None = None) -> dict:
    live_reviews_by_id, historical_reviews_by_id, all_review_records = load_review_state_maps(target)
    live_claims_by_id, historical_claims_by_id, _ = load_claim_state_maps(target)
    claim_lookup = build_claim_lookup_by_any_id(live_claims_by_id, historical_claims_by_id)
    refresh_alias_conflict_reviews(
        target=target,
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )
    all_review_records = build_ordered_review_state_records(
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )
    write_jsonl(target / "state" / "reviews.jsonl", all_review_records)
    for review_record in all_review_records:
        write_review_file(target, review_record)
    cleanup_superseded_record_files(
        target=target,
        historical_claims_by_id=historical_claims_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )

    review_records = all_review_records
    if status_filter:
        review_records = [
            record for record in review_records
            if record.get("status") == status_filter
        ]
    else:
        review_records = [
            record for record in review_records
            if record.get("lifecycle_status") == "active" and record.get("status") == "open"
        ]

    items = []
    for review_record in sorted(
        review_records,
        key=lambda item: (
            1 if item.get("lifecycle_status") == "active" else 0,
            1 if item.get("status") == "open" else 0,
            item.get("created_at", ""),
            review_display_id(item),
        ),
        reverse=True,
    ):
        candidate_claims = []
        for claim_id in review_record.get("candidate_claim_ids", []):
            claim_record = claim_lookup.get(claim_id)
            if claim_record is None:
                continue
            candidate_claims.append({
                "claim_id": claim_id,
                "display_id": claim_display_id(claim_record),
                "lifecycle_status": claim_record.get("lifecycle_status"),
                "status": claim_record.get("status"),
                "text": claim_record.get("text", ""),
                "source_count": len(claim_record.get("source_ids", [])),
                "page_count": len(claim_record.get("page_ids", [])),
            })

        items.append({
            "review_id": review_display_id(review_record),
            "state_review_id": review_record["review_id"],
            "display_id": review_display_id(review_record),
            "kind": review_record.get("kind"),
            "status": review_record.get("status"),
            "lifecycle_status": review_record.get("lifecycle_status"),
            "reason": review_record.get("reason"),
            "recommended_action": review_record.get("recommended_action"),
            "allowed_actions": review_record.get("allowed_actions", []),
            "candidate_page_ids": review_record.get("candidate_page_ids", []),
            "candidate_claims": candidate_claims,
            "created_at": review_record.get("created_at"),
            "resolved_at": review_record.get("resolved_at"),
            "archived_at": review_record.get("archived_at"),
        })

    return {
        "workspace": str(target),
        "workspace_summary": build_workspace_summary(target),
        "items": items,
        "summary": {
            "review_count": len(items),
            "live_review_count": len(live_reviews_by_id),
            "historical_review_count": len(historical_reviews_by_id),
        },
    }


def command_review_list(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    payload = build_review_list_payload(target, status_filter=args.status)
    if args.json:
        return CommandResult(
            payload=payload,
            message=render_workspace_summary_message(
                "Review list completed.",
                target_dir=target,
                extra_lines=[
                    f"Status filter: {args.status or 'all'}",
                    (
                        "Summary: "
                        f"reviews={payload['summary']['review_count']}, "
                        f"live={payload['summary']['live_review_count']}, "
                        f"historical={payload['summary']['historical_review_count']}"
                    ),
                ],
            ),
        )

    lines = [
        render_workspace_summary_message(
            "Review list completed.",
            target_dir=target,
            extra_lines=[
                f"Status filter: {args.status or 'all'}",
                (
                    "Summary: "
                    f"reviews={payload['summary']['review_count']}, "
                    f"live={payload['summary']['live_review_count']}, "
                    f"historical={payload['summary']['historical_review_count']}"
                ),
                "",
                "Review Items:",
            ],
        )
    ]
    for index, item in enumerate(payload["items"], start=1):
        lines.append(
            f"{index}. {item['display_id']} [{item['kind']}, status={item['status']}, lifecycle={item['lifecycle_status']}]"
        )
        lines.append(f"   reason: {item['reason']}")
        lines.append(f"   recommended: {item['recommended_action']}")
        lines.append(f"   allowed: {', '.join(item['allowed_actions'])}")
        for claim in item["candidate_claims"][:4]:
            lines.append(
                f"   - {claim['display_id']} [{claim['lifecycle_status']}/{claim['status']}] {claim['text']}"
            )
    if len(lines) == 1:
        lines.append("No review items found.")
    return CommandResult(payload=payload, message="\n".join(lines))


def choose_auto_merge_primary_claim_id(candidate_claim_ids: list[str], live_claims_by_id: dict[str, dict]) -> str | None:
    live_candidates = [live_claims_by_id[claim_id] for claim_id in candidate_claim_ids if claim_id in live_claims_by_id]
    if len(live_candidates) != len(candidate_claim_ids):
        return None
    if len(live_candidates) != 2:
        return None
    ranked = sorted(
        live_candidates,
        key=lambda item: (
            len(item.get("source_ids", [])),
            len(clean_claim_candidate_text(item.get("text", ""))),
            item.get("created_at", ""),
            item["claim_id"],
        ),
        reverse=True,
    )
    return ranked[0]["claim_id"] if ranked else None


def build_review_auto_decision_payload(
    review_record: dict,
    live_claims_by_id: dict[str, dict],
    target: Path,
) -> dict:
    candidate_claims = []
    for claim_id in review_record.get("candidate_claim_ids", []):
        claim_record = live_claims_by_id.get(claim_id)
        if claim_record is None:
            continue
        candidate_claims.append({
            "claim_id": claim_id,
            "text": claim_record.get("text", ""),
            "status": claim_record.get("status"),
            "claim_type": claim_record.get("claim_type"),
            "source_count": len(set(claim_record.get("source_ids", []))),
            "source_ref_count": len(claim_record.get("source_refs", [])),
            "duplicate_candidates": claim_record.get("duplicate_candidates", []),
            "conflict_group": claim_record.get("conflict_group"),
        })
    payload = {
        "task": "review_auto_decision",
        "review": {
            "review_id": review_record.get("review_id"),
            "kind": review_record.get("kind"),
            "reason": review_record.get("reason"),
            "recommended_action": review_record.get("recommended_action"),
            "allowed_actions": review_record.get("allowed_actions", []),
            "candidate_claim_ids": review_record.get("candidate_claim_ids", []),
            "candidate_page_ids": review_record.get("candidate_page_ids", []),
            "evidence": review_record.get("evidence", []),
        },
        "candidate_claims": candidate_claims,
        "candidate_pages": [],
    }
    candidate_page_ids = review_record.get("candidate_page_ids", [])
    if candidate_page_ids:
        pages_path = target / "state" / "pages.jsonl"
        if pages_path.exists():
            page_records = [
                ensure_page_lifecycle_defaults(record)
                for record in load_jsonl(pages_path)
            ]
            page_lookup = {record.get("page_id"): record for record in page_records}
            for page_id in candidate_page_ids:
                page_record = page_lookup.get(page_id)
                if page_record is None:
                    continue
                payload["candidate_pages"].append({
                    "page_id": page_id,
                    "canonical_id": page_record.get("canonical_id") or page_id,
                    "title": page_record.get("title", ""),
                    "aliases": page_record.get("aliases", []),
                    "type": page_record.get("type"),
                    "status": page_record.get("status"),
                })
    if review_record.get("kind") == "alias_conflict":
        evidence = review_record.get("evidence", [])
        alias_value = str(evidence[0].get("alias", "")).strip() if evidence else ""
        alias_index = load_alias_index(target)
        alias_matches = alias_index_matches_for_value(alias_index, alias_value) if alias_value else []
        payload["alias_conflict_context"] = {
            "alias_value": alias_value,
            "canonical_ids": evidence[0].get("canonical_ids", []) if evidence else [],
            "matched_pages": alias_matches,
        }
    return payload


def normalize_review_auto_hook_plan(
    hook_result: dict,
    review_record: dict,
    base_plan: dict,
    live_claims_by_id: dict[str, dict],
    min_confidence: float,
) -> dict | None:
    decision = str(hook_result.get("decision", "")).strip().lower()
    if decision != "auto_apply":
        return None

    confidence = coerce_float(hook_result.get("confidence", 0.0), 0.0)
    if confidence < min_confidence:
        return None

    action = str(hook_result.get("action", "")).strip()
    allowed_actions = set(review_record.get("allowed_actions", []))
    if action not in allowed_actions:
        return None

    candidate_claim_ids = list(review_record.get("candidate_claim_ids", []))
    candidate_page_ids = list(review_record.get("candidate_page_ids", []))
    plan = dict(base_plan)
    plan.update({
        "decision": "auto_apply",
        "action": action,
        "reason": str(hook_result.get("reason", "agent_hook_auto_apply")).strip() or "agent_hook_auto_apply",
        "confidence": confidence,
    })

    if action in {"merge", "archive_one"}:
        primary_claim_id = str(hook_result.get("primary_claim_id", "")).strip()
        if primary_claim_id not in candidate_claim_ids:
            return None
        plan["primary_claim_id"] = primary_claim_id
        if action == "merge":
            secondary_claim_id = str(hook_result.get("secondary_claim_id", "")).strip()
            if secondary_claim_id not in candidate_claim_ids or secondary_claim_id == primary_claim_id:
                ranked_primary = choose_auto_merge_primary_claim_id(candidate_claim_ids, live_claims_by_id)
                if ranked_primary is None:
                    return None
                primary_claim_id = ranked_primary
                secondary_claim_id = next(
                    claim_id for claim_id in candidate_claim_ids
                    if claim_id != primary_claim_id
                )
                plan["primary_claim_id"] = primary_claim_id
            plan["secondary_claim_id"] = secondary_claim_id
    elif action in {"assign_alias", "remove_alias"}:
        primary_page_id = str(hook_result.get("primary_page_id", "")).strip()
        if primary_page_id not in candidate_page_ids:
            return None
        alias_value = str(hook_result.get("alias_value", "")).strip()
        if not alias_value:
            evidence = review_record.get("evidence", [])
            alias_value = str(evidence[0].get("alias", "")).strip() if evidence else ""
        if not alias_value:
            return None
        plan["primary_page_id"] = primary_page_id
        plan["alias_value"] = alias_value
    return plan


def maybe_get_agent_assisted_review_plan(
    target: Path,
    review_record: dict,
    live_claims_by_id: dict[str, dict],
    automation_config: dict,
    base_plan: dict,
) -> dict | None:
    if not automation_config.get("enabled"):
        return None

    payload = build_review_auto_decision_payload(
        review_record=review_record,
        live_claims_by_id=live_claims_by_id,
        target=target,
    )
    hook_result = run_json_automation_command(
        target=target,
        command=automation_config.get("command", []),
        payload=payload,
        timeout_seconds=automation_config.get("timeout_seconds", 45),
    )
    if hook_result is None:
        return None

    return normalize_review_auto_hook_plan(
        hook_result=hook_result,
        review_record=review_record,
        base_plan=base_plan,
        live_claims_by_id=live_claims_by_id,
        min_confidence=automation_config.get("min_confidence", 0.8),
    )


def review_action_plain_label(action: str) -> str:
    labels = {
        "merge": "合并成一条更清楚的结论",
        "keep_both": "两条都保留",
        "archive_one": "保留一条并归档另一条",
        "edit_then_resume": "先手工修改再继续",
        "assign_alias": "把别名指定给某个页面",
        "remove_alias": "移除这个别名",
    }
    return labels.get(action, action)


def claim_record_is_safe_auto_stable_candidate(
    claim_record: dict,
    live_reviews_by_id: dict[str, dict],
) -> tuple[bool, str | None]:
    if claim_record.get("lifecycle_status") != "active":
        return False, "claim_not_active"
    if claim_record.get("status") != "draft":
        return False, f"claim_status_is_{claim_record.get('status')}"
    if not claim_record.get("source_ids") or not claim_record.get("source_refs"):
        return False, "claim_missing_source_traceability"
    active_review_ids = [
        review_record["review_id"]
        for review_record in live_reviews_by_id.values()
        if is_actionable_review_record(review_record)
        and claim_record["claim_id"] in review_record.get("candidate_claim_ids", [])
    ]
    if active_review_ids:
        return False, "claim_still_attached_to_open_review"
    if claim_record.get("duplicate_candidates"):
        return False, "claim_still_has_duplicate_candidates"
    if claim_record.get("conflict_group"):
        return False, "claim_still_has_conflict_group"
    cleaned_text = clean_claim_candidate_text(claim_record.get("text", ""))
    if claim_candidate_is_noise(cleaned_text):
        return False, "claim_text_is_noise"
    if claim_starts_with_dependent_prefix(cleaned_text):
        return False, "claim_text_is_dependent_fragment"
    if text_is_question_like(cleaned_text):
        return False, "claim_text_is_question_like"
    quality_label = str(claim_record.get("quality_label") or "").strip().lower()
    if quality_label in {"noise", "fragment", "title_shell"}:
        return False, f"claim_quality_label_is_{quality_label}"
    if claim_candidate_has_short_gray_zone(cleaned_text):
        if claim_record.get("quality_safe_auto_ready") is True:
            return True, "semantic_quality_marked_short_claim_safe_auto_ready"
        return False, "short_claim_requires_semantic_quality_clearance"
    if not claim_can_stand_alone(cleaned_text) and claim_record.get("claim_type") != "definition":
        return False, "claim_text_not_standalone_enough"
    return True, None


def build_stable_promotion_payload(claim_record: dict) -> dict:
    return {
        "task": "claim_stable_promotion",
        "claim": {
            "claim_id": claim_record.get("claim_id"),
            "text": claim_record.get("text", ""),
            "status": claim_record.get("status"),
            "claim_type": claim_record.get("claim_type"),
            "source_ids": claim_record.get("source_ids", []),
            "source_refs": claim_record.get("source_refs", []),
            "duplicate_candidates": claim_record.get("duplicate_candidates", []),
            "conflict_group": claim_record.get("conflict_group"),
            "page_ids": claim_record.get("page_ids", []),
        },
    }


def maybe_get_agent_assisted_stable_promotion(
    target: Path,
    claim_record: dict,
    automation_config: dict,
) -> tuple[bool, str | None]:
    if not automation_config.get("enabled"):
        return False, None

    hook_result = run_json_automation_command(
        target=target,
        command=automation_config.get("command", []),
        payload=build_stable_promotion_payload(claim_record),
        timeout_seconds=automation_config.get("timeout_seconds", 45),
    )
    if hook_result is None:
        return False, None

    decision = str(hook_result.get("decision", "")).strip().lower()
    confidence = coerce_float(hook_result.get("confidence", 0.0), 0.0)
    if decision != "promote" or confidence < automation_config.get("min_confidence", 0.8):
        return False, None
    reason = str(hook_result.get("reason", "agent_hook_promoted_to_stable")).strip() or "agent_hook_promoted_to_stable"
    return True, reason


def build_review_auto_escalation_entry(
    review_record: dict,
    plan: dict,
    live_claims_by_id: dict[str, dict],
) -> dict:
    candidate_claims = []
    for claim_id in review_record.get("candidate_claim_ids", [])[:3]:
        claim_record = live_claims_by_id.get(claim_id)
        if claim_record is None:
            continue
        candidate_claims.append({
            "claim_id": claim_id,
            "text": claim_record.get("text", ""),
            "status": claim_record.get("status"),
            "source_count": len(claim_record.get("source_ids", [])),
        })

    choice_options = [
        {
            "action": action,
            "label": review_action_plain_label(action),
            "is_recommended": action == review_record.get("recommended_action"),
        }
        for action in review_record.get("allowed_actions", [])
    ]

    if review_record.get("kind") == "alias_conflict":
        evidence = review_record.get("evidence", [])
        alias_value = evidence[0].get("alias") if evidence else None
        issue_summary = (
            f"别名 `{alias_value}` 同时指向多个页面，当前还不能安全地自动决定归属。"
            if alias_value
            else "有一个别名同时指向多个页面，当前还不能安全地自动决定归属。"
        )
    elif review_record.get("kind") == "claim_duplicate":
        issue_summary = "这组 claim 可能在说同一件事，但当前仍需要人确认是否真的该合并。"
    else:
        issue_summary = "这条 review 超出了当前保守自动规则，需要人进一步判断。"

    if plan.get("reason") == "alias_conflict_needs_human_owner_choice":
        why_human_needed = "多个页面都可能拥有这个 alias，自动分配存在误伤风险。"
    elif plan.get("reason") == "duplicate_review_needs_human_merge_choice":
        why_human_needed = "候选 claim 不止两条，或主次关系不够明确，自动合并风险偏高。"
    else:
        why_human_needed = "当前证据还不足以支持保守自动裁决。"

    entry = {
        "review_id": review_record["review_id"],
        "display_id": review_display_id(review_record),
        "kind": review_record.get("kind"),
        "issue_summary": issue_summary,
        "why_human_needed": why_human_needed,
        "recommended_action": review_record.get("recommended_action"),
        "choice_options": choice_options,
        "candidate_claims": candidate_claims,
        "candidate_page_ids": review_record.get("candidate_page_ids", []),
        "suggested_user_prompt": (
            f"请帮我判断审核单 {review_display_id(review_record)}。"
            "如果你已经知道怎么处理，可以直接告诉我保留、合并、归档，或先手工修改再继续。"
        ),
    }
    return entry


def build_review_auto_agent_handoff(
    auto_apply_plans: list[dict],
    escalated_entries: list[dict],
    promoted_claims: list[dict],
    review_automation_config: dict,
    stable_automation_config: dict,
) -> tuple[dict, str]:
    should_ask_user = bool(escalated_entries)
    next_action = (
        "ask_user_to_decide_escalated_reviews"
        if should_ask_user
        else "continue_with_normal_workflow"
    )
    agent_brief = {
        "mode": "needs_user_decision" if should_ask_user else "auto_resolved_all_safe_reviews",
        "should_ask_user": should_ask_user,
        "next_action": next_action,
        "auto_apply_review_count": len(auto_apply_plans),
        "escalated_review_count": len(escalated_entries),
        "promoted_claim_count": len(promoted_claims),
        "review_auto_strategy": review_automation_config.get("strategy"),
        "stable_promotion_strategy": stable_automation_config.get("strategy"),
    }
    if should_ask_user:
        first_item = escalated_entries[0]
        agent_summary = (
            f"已自动处理 {len(auto_apply_plans)} 条高把握 review"
            f"（review_auto={review_automation_config.get('strategy')}），"
            f"但还有 {len(escalated_entries)} 条需要用户判断。"
            f"优先向用户解释 `{first_item['display_id']}`：{first_item['issue_summary']}"
        )
    else:
        agent_summary = (
            f"本轮高把握 review 已全部自动收口，共处理 {len(auto_apply_plans)} 条，"
            f"并额外提升了 {len(promoted_claims)} 条 claim 为 stable"
            f"（stable_promotion={stable_automation_config.get('strategy')}）。"
        )
    return agent_brief, agent_summary


def render_review_auto_message(review_auto_payload: dict) -> str:
    lines = [
        "Review Auto Summary:",
        f"  dry_run: {review_auto_payload.get('dry_run')}",
        (
            "  counts: "
            f"planned={review_auto_payload['summary']['planned_review_count']}, "
            f"auto_apply={review_auto_payload['summary']['auto_apply_count']}, "
            f"escalated={review_auto_payload['summary']['escalated_count']}, "
            f"applied={review_auto_payload['summary']['applied_count']}, "
            f"promoted_claims={review_auto_payload['summary']['promoted_claim_count']}"
        ),
        f"  next_action: {review_auto_payload.get('agent_brief', {}).get('next_action')}",
        f"  agent_summary: {review_auto_payload.get('agent_summary', '')}",
    ]
    applied_actions = review_auto_payload.get("applied_actions", [])
    if applied_actions:
        lines.append("  applied_actions:")
        for item in applied_actions[:5]:
            lines.append(
                f"    - {item.get('display_id')} [{item.get('kind')}] action={item.get('action')} reason={item.get('reason')}"
            )
    escalations = review_auto_payload.get("escalation_handoff", [])
    if escalations:
        lines.append("  escalations:")
        for item in escalations[:5]:
            lines.append(
                f"    - {item.get('display_id')} [{item.get('kind')}] {item.get('issue_summary')}"
            )
            lines.append(f"      why_human_needed: {item.get('why_human_needed')}")
    return "\n".join(lines)


def render_review_auto_prompt(review_auto_payload: dict) -> str:
    agent_brief = review_auto_payload.get("agent_brief", {})
    lines = [
        "You are the workflow layer for a MyAgentWiki review-auto handoff.",
        "Use the handoff below to decide whether to continue automatically or ask the user for a focused review decision.",
        "Do not invent review outcomes that are not supported by the handoff.",
        "",
        "## Review Auto Run",
        f"- contract_version: {review_auto_payload.get('contract_version', '')}",
        f"- dry_run: {review_auto_payload.get('dry_run')}",
        f"- workspace: {review_auto_payload.get('workspace', '')}",
        f"- planned_review_count: {review_auto_payload.get('summary', {}).get('planned_review_count')}",
        f"- auto_apply_count: {review_auto_payload.get('summary', {}).get('auto_apply_count')}",
        f"- escalated_count: {review_auto_payload.get('summary', {}).get('escalated_count')}",
        f"- applied_count: {review_auto_payload.get('summary', {}).get('applied_count')}",
        f"- promoted_claim_count: {review_auto_payload.get('summary', {}).get('promoted_claim_count')}",
        "",
        "## Agent Brief",
        f"- mode: {agent_brief.get('mode')}",
        f"- should_ask_user: {agent_brief.get('should_ask_user')}",
        f"- next_action: {agent_brief.get('next_action')}",
        f"- agent_summary: {review_auto_payload.get('agent_summary', '')}",
        "",
        "## Applied Actions",
    ]
    applied_actions = review_auto_payload.get("applied_actions", [])
    if applied_actions:
        for item in applied_actions:
            lines.append(
                f"- {item.get('display_id')}: kind={item.get('kind')} action={item.get('action')} "
                f"reason={item.get('reason')} changed_claim_ids={','.join(item.get('changed_claim_ids', [])) or 'none'} "
                f"changed_page_ids={','.join(item.get('changed_page_ids', [])) or 'none'}"
            )
    else:
        lines.append("- none")
    lines.extend([
        "",
        "## Escalations",
    ])
    escalations = review_auto_payload.get("escalation_handoff", [])
    if escalations:
        for item in escalations:
            lines.append(
                f"- {item.get('display_id')}: kind={item.get('kind')} issue={item.get('issue_summary')} "
                f"why_human_needed={item.get('why_human_needed')} recommended_action={item.get('recommended_action')}"
            )
            if item.get("candidate_claims"):
                for claim in item["candidate_claims"]:
                    lines.append(
                        f"  candidate_claim: {claim.get('claim_id')} status={claim.get('status')} "
                        f"sources={claim.get('source_count')} text={claim.get('text')}"
                    )
            if item.get("candidate_page_ids"):
                lines.append(f"  candidate_page_ids: {', '.join(item['candidate_page_ids'])}")
            if item.get("choice_options"):
                for option in item["choice_options"]:
                    lines.append(
                        f"  option: action={option.get('action')} label={option.get('label')} "
                        f"recommended={option.get('is_recommended')}"
                    )
            lines.append(f"  suggested_user_prompt: {item.get('suggested_user_prompt')}")
    else:
        lines.append("- none")
    lines.extend([
        "",
        "## Instruction",
        "If `should_ask_user` is false, continue the broader workflow without asking the user to re-judge already resolved reviews.",
        "If `should_ask_user` is true, ask only about the escalated reviews and explain the options in plain language.",
        "Prefer the recommended action when presenting options, but make uncertainty explicit when the handoff says human judgment is needed.",
    ])
    return "\n".join(lines)


def build_review_auto_messages(review_auto_payload: dict) -> list[dict]:
    prompt_text = render_review_auto_prompt(review_auto_payload)
    return [
        {
            "role": "system",
            "content": (
                "You are the workflow layer for a MyAgentWiki review-auto handoff. "
                "Continue automatically when the handoff says it is safe, and ask the user only for escalated decisions."
            ),
        },
        {
            "role": "user",
            "content": prompt_text,
        },
    ]


def render_review_auto_chatml(review_auto_payload: dict) -> str:
    messages = build_review_auto_messages(review_auto_payload)
    blocks = []
    for message in messages:
        blocks.append(f"<|im_start|>{message['role']}\n{message['content']}\n<|im_end|>")
    return "\n".join(blocks)


def run_post_ingest_review_auto(target: Path) -> dict:
    review_auto_result = command_review_auto(argparse.Namespace(
        target_dir=str(target),
        dry_run=False,
        format="summary",
        json=True,
    ))
    return review_auto_result.payload


def render_post_ingest_review_auto_summary(review_auto_payload: dict | None) -> str | None:
    if not review_auto_payload:
        return None
    summary = review_auto_payload.get("summary", {})
    agent_brief = review_auto_payload.get("agent_brief", {})
    return (
        "Auto review: "
        f"applied={summary.get('applied_count', 0)}, "
        f"escalated={summary.get('escalated_count', 0)}, "
        f"promoted_claims={summary.get('promoted_claim_count', 0)}, "
        f"next_action={agent_brief.get('next_action', 'continue_with_normal_workflow')}"
    )


def propose_review_auto_action(
    target: Path,
    review_record: dict,
    live_claims_by_id: dict[str, dict],
    automation_config: dict,
) -> dict:
    review_id = review_record["review_id"]
    kind = review_record.get("kind")
    recommended_action = review_record.get("recommended_action")
    allowed_actions = set(review_record.get("allowed_actions", []))

    plan = {
        "review_id": review_id,
        "display_id": review_display_id(review_record),
        "kind": kind,
        "recommended_action": recommended_action,
        "decision": "escalate",
        "reason": "review_requires_human_judgment",
        "action": None,
        "primary_claim_id": None,
        "secondary_claim_id": None,
        "primary_page_id": None,
        "alias_value": None,
        "confidence": None,
    }

    if not is_actionable_review_record(review_record):
        plan["reason"] = "review_not_actionable"
        return plan

    if kind == "alias_conflict":
        evidence = review_record.get("evidence", [])
        alias_value = evidence[0].get("alias") if evidence else None
        candidate_page_ids = list(review_record.get("candidate_page_ids", []))
        if (
            recommended_action == "remove_alias"
            and "remove_alias" in allowed_actions
            and alias_value
            and len(candidate_page_ids) >= 1
        ):
            plan.update({
                "decision": "auto_apply",
                "reason": "alias_conflict_can_safely_remove_shared_alias",
                "action": "remove_alias",
                "primary_page_id": candidate_page_ids[0],
                "alias_value": alias_value,
            })
            return plan
        if (
            recommended_action == "assign_alias"
            and "assign_alias" in allowed_actions
            and alias_value
            and len(candidate_page_ids) == 1
        ):
            plan.update({
                "decision": "auto_apply",
                "reason": "alias_conflict_has_single_candidate_owner",
                "action": "assign_alias",
                "primary_page_id": candidate_page_ids[0],
                "alias_value": alias_value,
            })
            return plan
        plan["reason"] = "alias_conflict_needs_human_owner_choice"
        agent_plan = maybe_get_agent_assisted_review_plan(
            target=target,
            review_record=review_record,
            live_claims_by_id=live_claims_by_id,
            automation_config=automation_config,
            base_plan=plan,
        )
        return agent_plan or plan

    if kind == "claim_duplicate" and recommended_action == "merge" and "merge" in allowed_actions:
        candidate_claim_ids = list(review_record.get("candidate_claim_ids", []))
        primary_claim_id = choose_auto_merge_primary_claim_id(candidate_claim_ids, live_claims_by_id)
        if primary_claim_id is None:
            plan["reason"] = "duplicate_review_needs_human_merge_choice"
            agent_plan = maybe_get_agent_assisted_review_plan(
                target=target,
                review_record=review_record,
                live_claims_by_id=live_claims_by_id,
                automation_config=automation_config,
                base_plan=plan,
            )
            return agent_plan or plan
        secondary_claim_id = next(
            claim_id for claim_id in candidate_claim_ids
            if claim_id != primary_claim_id
        )
        plan.update({
            "decision": "auto_apply",
            "reason": "duplicate_review_has_safe_two_claim_merge",
            "action": "merge",
            "primary_claim_id": primary_claim_id,
            "secondary_claim_id": secondary_claim_id,
        })
        return plan

    if kind == "claim_conflict":
        agent_plan = maybe_get_agent_assisted_review_plan(
            target=target,
            review_record=review_record,
            live_claims_by_id=live_claims_by_id,
            automation_config=automation_config,
            base_plan=plan,
        )
        return agent_plan or plan

    agent_plan = maybe_get_agent_assisted_review_plan(
        target=target,
        review_record=review_record,
        live_claims_by_id=live_claims_by_id,
        automation_config=automation_config,
        base_plan=plan,
    )
    if agent_plan is not None:
        return agent_plan
    plan["reason"] = "review_kind_or_action_not_in_safe_auto_rules"
    return plan


def command_review_auto(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    claims_path = target / "state" / "claims.jsonl"
    reviews_path = target / "state" / "reviews.jsonl"
    config = load_workspace_config(target)
    review_automation_config = load_automation_target_config(config, "review_auto")
    stable_automation_config = load_automation_target_config(config, "stable_promotion")

    live_claims_by_id, historical_claims_by_id, _ = load_claim_state_maps(target)
    live_reviews_by_id, historical_reviews_by_id, _ = load_review_state_maps(target)
    _, _, archived_alias_review_ids = refresh_alias_conflict_reviews(
        target=target,
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )
    if archived_alias_review_ids:
        refreshed_review_state_records = build_ordered_review_state_records(
            live_reviews_by_id=live_reviews_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )
        write_jsonl(reviews_path, refreshed_review_state_records)
        for review_state_record in refreshed_review_state_records:
            write_review_file(target, review_state_record)
        cleanup_superseded_record_files(
            target=target,
            historical_claims_by_id=historical_claims_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )

    planned_actions = [
        propose_review_auto_action(
            target=target,
            review_record=review_record,
            live_claims_by_id=live_claims_by_id,
            automation_config=review_automation_config,
        )
        for review_record in sorted(
            live_reviews_by_id.values(),
            key=lambda item: (item.get("created_at", ""), review_display_id(item)),
        )
        if is_actionable_review_record(review_record)
    ]

    auto_apply_plans = [item for item in planned_actions if item["decision"] == "auto_apply"]
    escalated_plans = [item for item in planned_actions if item["decision"] != "auto_apply"]
    applied_actions: list[dict] = []
    promoted_claims: list[dict] = []

    auto_apply_failures: list[dict] = []

    if not args.dry_run:
        for plan in auto_apply_plans:
            review_record = live_reviews_by_id.get(plan["review_id"])
            if review_record is None:
                continue
            current_plan = propose_review_auto_action(
                target=target,
                review_record=review_record,
                live_claims_by_id=live_claims_by_id,
                automation_config=review_automation_config,
            )
            if current_plan.get("decision") != "auto_apply":
                continue
            try:
                result = apply_review_action(
                    target=target,
                    review_record=review_record,
                    action=current_plan["action"],
                    primary_claim_id=current_plan["primary_claim_id"],
                    secondary_claim_id=current_plan["secondary_claim_id"],
                    primary_page_id=current_plan["primary_page_id"],
                    alias_value=current_plan["alias_value"],
                    live_claims_by_id=live_claims_by_id,
                    historical_claims_by_id=historical_claims_by_id,
                    live_reviews_by_id=live_reviews_by_id,
                    historical_reviews_by_id=historical_reviews_by_id,
                )
            except ValueError as exc:
                auto_apply_failures.append({
                    "review_id": current_plan["review_id"],
                    "display_id": current_plan["display_id"],
                    "kind": current_plan["kind"],
                    "reason": "auto_apply_failed_validation",
                    "validation_error": str(exc),
                })
                continue
            live_reviews_by_id[review_record["review_id"]] = review_record
            applied_actions.append({
                "review_id": current_plan["review_id"],
                "display_id": current_plan["display_id"],
                "kind": current_plan["kind"],
                "reason": current_plan["reason"],
                "action": result.get("action"),
                "changed_claim_ids": result.get("changed_claim_ids", []),
                "changed_page_ids": result.get("changed_page_ids", []),
            })

        for claim_record in sorted(live_claims_by_id.values(), key=lambda item: item["claim_id"]):
            is_safe, reason = claim_record_is_safe_auto_stable_candidate(claim_record, live_reviews_by_id)
            if not is_safe:
                promoted_by_hook, hook_reason = maybe_get_agent_assisted_stable_promotion(
                    target=target,
                    claim_record=claim_record,
                    automation_config=stable_automation_config,
                )
                if not promoted_by_hook:
                    continue
                reason = hook_reason
            claim_record["status"] = "stable"
            claim_record["review_reason"] = None
            claim_record["updated_at"] = utc_now_iso()
            promoted_claims.append({
                "claim_id": claim_record["claim_id"],
                "reason": reason or "safe_auto_promoted_to_stable",
            })

        write_jsonl(
            claims_path,
            build_ordered_claim_state_records(
                live_claims_by_id=live_claims_by_id,
                historical_claims_by_id=historical_claims_by_id,
            ),
        )
        for claim_record in [*live_claims_by_id.values(), *historical_claims_by_id.values()]:
            write_claim_file(target, claim_record)

        write_jsonl(
            reviews_path,
            build_ordered_review_state_records(
                live_reviews_by_id=live_reviews_by_id,
                historical_reviews_by_id=historical_reviews_by_id,
            ),
        )
        for review_state_record in [*live_reviews_by_id.values(), *historical_reviews_by_id.values()]:
            write_review_file(target, review_state_record)
        cleanup_superseded_record_files(
            target=target,
            historical_claims_by_id=historical_claims_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )

        rebuild_review_affected_pages(
            target=target,
            live_claims_by_id=live_claims_by_id,
            live_reviews_by_id=live_reviews_by_id,
        )
        _, _, archived_alias_review_ids = refresh_alias_conflict_reviews(
            target=target,
            live_reviews_by_id=live_reviews_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )
        if archived_alias_review_ids:
            write_jsonl(
                reviews_path,
                build_ordered_review_state_records(
                    live_reviews_by_id=live_reviews_by_id,
                    historical_reviews_by_id=historical_reviews_by_id,
                ),
            )
            for review_state_record in [*live_reviews_by_id.values(), *historical_reviews_by_id.values()]:
                write_review_file(target, review_state_record)
            cleanup_superseded_record_files(
                target=target,
                historical_claims_by_id=historical_claims_by_id,
                historical_reviews_by_id=historical_reviews_by_id,
            )

    failure_plan_by_review_id = {
        item["review_id"]: item
        for item in auto_apply_failures
    }
    final_escalated_plans = list(escalated_plans) + [
        {
            "review_id": item["review_id"],
            "display_id": item["display_id"],
            "kind": item["kind"],
            "recommended_action": None,
            "decision": "escalate",
            "reason": item["reason"],
            "action": None,
            "primary_claim_id": None,
            "secondary_claim_id": None,
            "primary_page_id": None,
            "alias_value": None,
            "confidence": None,
        }
        for item in auto_apply_failures
        if item["review_id"] in live_reviews_by_id
    ]

    escalated_entries = [
        build_review_auto_escalation_entry(
            review_record=live_reviews_by_id[plan["review_id"]],
            plan=plan,
            live_claims_by_id=live_claims_by_id,
        )
        for plan in final_escalated_plans
        if plan["review_id"] in live_reviews_by_id
    ]
    for entry in escalated_entries:
        failure_meta = failure_plan_by_review_id.get(entry["review_id"])
        if failure_meta is not None:
            entry["why_human_needed"] = (
                "自动裁决在最终收敛校验时失败，需要人工确认别名归属或改为保留多义并存。"
            )
            entry["validation_error"] = failure_meta["validation_error"]
    agent_brief, agent_summary = build_review_auto_agent_handoff(
        auto_apply_plans=auto_apply_plans,
        escalated_entries=escalated_entries,
        promoted_claims=promoted_claims,
        review_automation_config=review_automation_config,
        stable_automation_config=stable_automation_config,
    )

    payload = {
        "contract_version": REVIEW_AUTO_HANDOFF_CONTRACT_VERSION,
        "workspace": str(target),
        "workspace_summary": build_workspace_summary(target),
        "dry_run": bool(args.dry_run),
        "planned_actions": planned_actions,
        "applied_actions": applied_actions,
        "escalated_reviews": final_escalated_plans,
        "escalation_handoff": escalated_entries,
        "auto_apply_failures": auto_apply_failures,
        "promoted_claims": promoted_claims,
        "agent_brief": agent_brief,
        "agent_summary": agent_summary,
        "summary": {
            "planned_review_count": len(planned_actions),
            "auto_apply_count": len(auto_apply_plans),
            "escalated_count": len(final_escalated_plans),
            "applied_count": len(applied_actions),
            "auto_apply_failure_count": len(auto_apply_failures),
            "promoted_claim_count": len(promoted_claims),
        },
        "automation": {
            "review_auto": {
                "strategy": review_automation_config.get("strategy"),
                "enabled": review_automation_config.get("enabled"),
                "min_confidence": review_automation_config.get("min_confidence"),
            },
            "stable_promotion": {
                "strategy": stable_automation_config.get("strategy"),
                "enabled": stable_automation_config.get("enabled"),
                "min_confidence": stable_automation_config.get("min_confidence"),
            },
        },
    }
    handoff_format = str(getattr(args, "format", "summary") or "summary").strip().lower()
    if handoff_format == "prompt":
        payload["prompt_text"] = render_review_auto_prompt(payload)
    elif handoff_format == "messages":
        payload["messages"] = build_review_auto_messages(payload)
    elif handoff_format == "chatml":
        payload["messages"] = build_review_auto_messages(payload)
        payload["chatml_text"] = render_review_auto_chatml(payload)

    if args.json:
        return CommandResult(
            payload=payload,
            message=render_workspace_summary_message(
                "Review auto pass completed." if not args.dry_run else "Review auto dry-run completed.",
                target_dir=target,
                extra_lines=[
                    f"Dry run: {bool(args.dry_run)}",
                    f"Format: {handoff_format}",
                    (
                        "Summary: "
                        f"planned={payload['summary']['planned_review_count']}, "
                        f"auto_apply={payload['summary']['auto_apply_count']}, "
                        f"escalated={payload['summary']['escalated_count']}, "
                        f"applied={payload['summary']['applied_count']}, "
                        f"promoted_claims={payload['summary']['promoted_claim_count']}"
                    ),
                ],
            ),
        )
    if handoff_format == "prompt":
        return CommandResult(payload=payload, message=payload["prompt_text"])
    if handoff_format == "messages":
        return CommandResult(
            payload=payload,
            message=json.dumps(payload["messages"], ensure_ascii=False, indent=2),
        )
    if handoff_format == "chatml":
        return CommandResult(payload=payload, message=payload["chatml_text"])

    return CommandResult(
        payload=payload,
        message=render_workspace_summary_message(
            "Review auto pass completed." if not args.dry_run else "Review auto dry-run completed.",
            target_dir=target,
            extra_lines=[
                f"Dry run: {bool(args.dry_run)}",
                (
                    "Summary: "
                    f"planned={payload['summary']['planned_review_count']}, "
                    f"auto_apply={payload['summary']['auto_apply_count']}, "
                    f"escalated={payload['summary']['escalated_count']}, "
                    f"applied={payload['summary']['applied_count']}, "
                    f"promoted_claims={payload['summary']['promoted_claim_count']}"
                ),
                "",
                render_review_auto_message(payload),
            ],
        ),
    )


def rebuild_review_affected_pages(
    target: Path,
    live_claims_by_id: dict[str, dict],
    live_reviews_by_id: dict[str, dict],
) -> None:
    # review 动作完成后，如果不刷新页面，wiki 页面与 query 索引会滞后。
    # 这里走一个“小范围账本重建”：
    # 1. 重新根据 live claims / live reviews 计算 source-summary / concept
    # 2. 移除不再需要的自动页
    # 3. 重建 pages.jsonl / wiki/index.md / search index
    config = load_workspace_config(target)
    readable_concept_render_config = load_readable_concept_render_config(config)
    overview_render_config = load_page_render_config(config, "overview")
    page_intent_config = load_semantic_task_config(config, "page_intent")
    sources_path = target / "state" / "sources.jsonl"
    normalized_path = target / "state" / "normalized.jsonl"
    pages_path = target / "state" / "pages.jsonl"
    chunks_path = target / "state" / "chunks.jsonl"

    sources_by_id = {
        record["source_id"]: record
        for record in load_jsonl(sources_path)
    }
    normalized_records_by_source = {
        record["source_id"]: record
        for record in load_jsonl(normalized_path)
    }
    chunk_records = load_jsonl(chunks_path)
    chunks_by_source_id: dict[str, list[dict]] = {}
    for chunk_record in chunk_records:
        chunks_by_source_id.setdefault(chunk_record["source_id"], []).append(chunk_record)

    page_records = [ensure_page_lifecycle_defaults(record) for record in load_jsonl(pages_path)]
    page_records_by_id = {record["page_id"]: record for record in page_records}
    active_source_ids = choose_active_source_ids(sources_by_id)
    live_claim_records = list(live_claims_by_id.values())
    live_review_records = [
        record for record in live_reviews_by_id.values()
        if is_actionable_review_record(record)
    ]

    # 先把旧的自动页反链从 live claim / review 中清掉，再按当前 live 集合重建。
    for page_record in list(page_records_by_id.values()):
        if page_record.get("type") not in {"source-summary", "concept", "overview"}:
            continue
        page_id = page_record["page_id"]
        for claim_record in live_claim_records:
            if page_id in claim_record.get("page_ids", []):
                claim_record["page_ids"] = [item for item in claim_record["page_ids"] if item != page_id]
                claim_record["updated_at"] = utc_now_iso()
        for review_record in live_review_records:
            if page_id in review_record.get("candidate_page_ids", []):
                review_record["candidate_page_ids"] = [
                    item for item in review_record["candidate_page_ids"] if item != page_id
                ]

    claims_by_source_id: dict[str, list[dict]] = {}
    for claim_record in live_claim_records:
        for source_id in claim_record.get("source_ids", []):
            if source_id in active_source_ids:
                claims_by_source_id.setdefault(source_id, []).append(claim_record)

    for source_id in sorted(active_source_ids):
        source_record = sources_by_id.get(source_id)
        if source_record is None or source_record.get("status") == "failed":
            continue
        source_claims = claims_by_source_id.get(source_id, [])
        source_chunks = chunks_by_source_id.get(source_id, [])
        if not source_claims and not source_chunks:
            continue
        source_record_for_page = dict(source_record)
        source_record_for_page["status"] = "generated"
        normalized_record = normalized_records_by_source.get(source_id)
        page_rel_path = source_summary_page_path(source_id, normalized_record["title"] if normalized_record else Path(source_record["source_path"]).stem)
        page_text, page_record = build_source_summary_page(
            source_record=source_record_for_page,
            page_rel_path=page_rel_path,
            normalized_record=normalized_record,
            claim_records=source_claims,
            chunk_records=source_chunks,
        )
        page_record = apply_page_alias_overrides(target, page_record)
        page_record["page_path"] = str(page_rel_path)
        stored_page_record, _ = upsert_wiki_page(
            target=target,
            page_records_by_id=page_records_by_id,
            page_record=page_record,
            page_text=page_text,
        )
        link_claims_to_page_in_memory(source_claims, stored_page_record["page_id"], live_claims_by_id)

    concept_claim_groups: dict[str, list[dict]] = {}
    for claim_record in live_claim_records:
        active_claim_source_ids = [
            source_id for source_id in claim_record.get("source_ids", [])
            if source_id in active_source_ids
        ]
        if not active_claim_source_ids:
            continue
        bucket_key = build_concept_group_key(claim_record)
        concept_claim_groups.setdefault(bucket_key, []).append(claim_record)
    concept_claim_groups = regroup_concept_claims_by_canonical_topic(concept_claim_groups)
    page_routes_by_bucket = apply_page_intent_decisions_to_claim_groups(
        target=target,
        concept_claim_groups=concept_claim_groups,
        task_config=page_intent_config,
    )

    for bucket_key, grouped_claims in sorted(concept_claim_groups.items()):
        page_route = page_route_for_bucket(page_routes_by_bucket, bucket_key)
        page_intent = preferred_page_intent_for_claim_group(
            grouped_claims,
            page_route.get("page_intent", "topic"),
        )
        page_route["page_intent"] = page_intent
        page_route["route_target"] = page_intent
        if page_intent == "reject":
            continue
        if page_intent == "concept" and should_generate_concept_page(grouped_claims):
            group_topic_label = choose_group_topic_label(grouped_claims)
            canonical_claim = choose_canonical_claim(grouped_claims, group_topic_label)
            concept_page_id = build_concept_page_id(bucket_key)
            concept_title, concept_title_quality = resolve_concept_title_candidate(
                target=target,
                config=config,
                canonical_claim=canonical_claim,
                claim_records=grouped_claims,
                preferred_section_label=group_topic_label,
            )
            if concept_title_quality["classification"] == "reject":
                stable_claims = filter_live_stable_claim_records(grouped_claims)
                if not stable_claims:
                    continue
            page_rel_path = concept_summary_page_path(
                concept_page_id,
                concept_title,
            )
            page_text, page_record = build_concept_page(
                target=target,
                bucket_key=bucket_key,
                page_rel_path=page_rel_path,
                claim_records=grouped_claims,
                page_records_by_id=page_records_by_id,
                review_records=live_review_records,
                render_config=readable_concept_render_config,
            )
            page_record = apply_page_route_to_page_record(page_record, page_route)
            page_record = apply_page_alias_overrides(target, page_record)
            page_record["page_path"] = str(page_rel_path)
            stored_page_record, _ = upsert_wiki_page(
                target=target,
                page_records_by_id=page_records_by_id,
                page_record=page_record,
                page_text=page_text,
            )
            link_claims_to_page_in_memory(grouped_claims, stored_page_record["page_id"], live_claims_by_id)
            link_reviews_to_page_in_memory(
                review_records=live_review_records,
                page_id=stored_page_record["page_id"],
                claim_ids=stored_page_record["claim_ids"],
                reviews_by_id=live_reviews_by_id,
            )
        elif page_intent in {"guide", "example", "topic", "reference", "timeline"}:
            page_title_source = choose_group_topic_label(grouped_claims) or choose_canonical_claim(grouped_claims).get("text", "")
            page_rel_path = page_intent_page_path(page_intent, page_intent_page_id(bucket_key, page_intent), page_title_source)
            page_text, page_record = build_intent_routed_page(
                target=target,
                config=config,
                bucket_key=bucket_key,
                page_intent=page_intent,
                page_rel_path=page_rel_path,
                claim_records=grouped_claims,
                page_records_by_id=page_records_by_id,
                review_records=live_review_records,
            )
            page_record = apply_page_route_to_page_record(page_record, page_route)
            page_record = apply_page_alias_overrides(target, page_record)
            page_record["page_path"] = str(page_rel_path)
            stored_page_record, _ = upsert_wiki_page(
                target=target,
                page_records_by_id=page_records_by_id,
                page_record=page_record,
                page_text=page_text,
            )
            link_claims_to_page_in_memory(grouped_claims, stored_page_record["page_id"], live_claims_by_id)
            link_reviews_to_page_in_memory(
                review_records=live_review_records,
                page_id=stored_page_record["page_id"],
                claim_ids=stored_page_record["claim_ids"],
                reviews_by_id=live_reviews_by_id,
            )

    overview_concept_pages = collect_workspace_overview_concept_pages(
        claims_by_similarity_bucket=concept_claim_groups,
        page_records_by_id=page_records_by_id,
    )
    if should_generate_workspace_overview_page(overview_concept_pages):
        overview_page_rel_path = workspace_overview_page_path()
        overview_page_text, overview_page_record = build_workspace_overview_page(
            target=target,
            page_rel_path=overview_page_rel_path,
            concept_pages=overview_concept_pages,
            page_records_by_id=page_records_by_id,
            claim_records_by_id=live_claims_by_id,
            render_config=overview_render_config,
        )
        overview_page_record = apply_page_alias_overrides(target, overview_page_record)
        overview_page_record["page_path"] = str(overview_page_rel_path)
        stored_overview_page, _ = upsert_wiki_page(
            target=target,
            page_records_by_id=page_records_by_id,
            page_record=overview_page_record,
            page_text=overview_page_text,
        )
        link_claims_to_page_in_memory(
            [
                live_claims_by_id[claim_id]
                for claim_id in stored_overview_page["claim_ids"]
                if claim_id in live_claims_by_id
            ],
            stored_overview_page["page_id"],
            live_claims_by_id,
        )
        link_reviews_to_page_in_memory(
            review_records=live_review_records,
            page_id=stored_overview_page["page_id"],
            claim_ids=stored_overview_page["claim_ids"],
            reviews_by_id=live_reviews_by_id,
        )

    desired_auto_page_ids = {
        expected_source_summary_page_id(source_id)
        for source_id in active_source_ids
        if claims_by_source_id.get(source_id) or chunks_by_source_id.get(source_id)
    }
    forced_stale_page_ids: set[str] = set()
    for bucket_key, grouped_claims in concept_claim_groups.items():
        page_route = page_route_for_bucket(page_routes_by_bucket, bucket_key)
        page_intent = preferred_page_intent_for_claim_group(
            grouped_claims,
            page_route.get("page_intent", "topic"),
        )
        forced_stale_page_ids.update(
            {
                build_concept_page_id(bucket_key),
                *{
                    page_intent_page_id(bucket_key, stale_intent)
                    for stale_intent in {"guide", "example", "topic", "reference", "timeline"}
                    if stale_intent != page_intent
                },
            }
        )
        if page_intent == "concept" and should_generate_concept_page(grouped_claims):
            desired_auto_page_ids.add(build_concept_page_id(bucket_key))
        elif page_intent in {"guide", "example", "topic", "reference", "timeline"}:
            desired_auto_page_ids.add(page_intent_page_id(bucket_key, page_intent))
    if should_generate_workspace_overview_page(overview_concept_pages):
        desired_auto_page_ids.add(build_workspace_overview_page_id())

    prune_stale_auto_pages(
        target=target,
        page_records_by_id=page_records_by_id,
        desired_auto_page_ids=desired_auto_page_ids,
        claims_by_id=live_claims_by_id,
        reviews_by_id=live_reviews_by_id,
        forced_stale_page_ids=forced_stale_page_ids - desired_auto_page_ids,
    )

    write_jsonl(
        target / "state" / "claims.jsonl",
        build_ordered_claim_state_records(
            live_claims_by_id=live_claims_by_id,
            historical_claims_by_id={
                record["claim_id"]: record
                for record in load_jsonl(target / "state" / "claims.jsonl")
                if ensure_claim_lifecycle_defaults(record).get("lifecycle_status") != "active"
            },
        ),
    )
    for claim_record in load_jsonl(target / "state" / "claims.jsonl"):
        write_claim_file(target, claim_record)

    existing_historical_reviews_by_id = {
        record["review_id"]: ensure_review_lifecycle_defaults(record)
        for record in load_jsonl(target / "state" / "reviews.jsonl")
        if not is_live_review_record(ensure_review_lifecycle_defaults(record))
    }
    review_state_records = build_ordered_review_state_records(
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=existing_historical_reviews_by_id,
    )
    write_jsonl(target / "state" / "reviews.jsonl", review_state_records)
    for review_record in review_state_records:
        write_review_file(target, review_record)

    write_jsonl(pages_path, list(page_records_by_id.values()))
    rebuild_wiki_index(target, list(page_records_by_id.values()))
    write_alias_index(target, list(page_records_by_id.values()))
    refresh_alias_conflict_reviews(
        target=target,
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=existing_historical_reviews_by_id,
        page_records=list(page_records_by_id.values()),
    )
    updated_review_state_records = build_ordered_review_state_records(
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=existing_historical_reviews_by_id,
    )
    write_jsonl(target / "state" / "reviews.jsonl", updated_review_state_records)
    for review_record in updated_review_state_records:
        write_review_file(target, review_record)
    cleanup_superseded_record_files(
        target=target,
        historical_claims_by_id={
            record["claim_id"]: ensure_claim_lifecycle_defaults(record)
            for record in load_jsonl(target / "state" / "claims.jsonl")
            if ensure_claim_lifecycle_defaults(record).get("lifecycle_status") != "active"
        },
        historical_reviews_by_id=existing_historical_reviews_by_id,
    )
    previous_search_index_records = load_search_pages_index(target)
    write_search_pages_index(
        target=target,
        page_records=list(page_records_by_id.values()),
        claim_records_by_id=live_claims_by_id,
        previous_records=previous_search_index_records,
    )

def resolve_claim_record_for_action(
    claim_id: str,
    live_claims_by_id: dict[str, dict],
    historical_claims_by_id: dict[str, dict],
) -> dict:
    claim_record = build_claim_lookup_by_any_id(live_claims_by_id, historical_claims_by_id).get(claim_id)
    if claim_record is None:
        raise KeyError(f"Unknown claim_id: {claim_id}")
    return claim_record


def archive_live_claim(
    claim_record: dict,
    live_claims_by_id: dict[str, dict],
    historical_claims_by_id: dict[str, dict],
    archived_by_claim_id: str | None = None,
) -> dict:
    # archive_one / merge 都会复用这段逻辑：
    # 把在线 claim 迁到历史态，并保留被谁替代的线索。
    live_claims_by_id.pop(claim_record["claim_id"], None)
    claim_record = dict(claim_record)
    claim_record["lifecycle_status"] = "superseded"
    claim_record["archived_at"] = utc_now_iso()
    claim_record["updated_at"] = utc_now_iso()
    if archived_by_claim_id:
        append_unique(claim_record.setdefault("superseded_by", []), archived_by_claim_id)
    historical_claim_record = convert_claim_record_to_historical(claim_record)
    historical_claims_by_id[historical_claim_record["claim_id"]] = historical_claim_record
    return historical_claim_record


def normalize_claim_review_flags(claim_record: dict) -> None:
    # claim 上的 review 标志位要和当前人工状态同步。
    # 如果重复候选、冲突组都已经不再存在，就把它还原成普通 draft。
    duplicate_candidates = [
        item for item in claim_record.get("duplicate_candidates", [])
        if item != claim_record["claim_id"]
    ]
    claim_record["duplicate_candidates"] = sorted(set(duplicate_candidates))
    if claim_record.get("status") == "needs_review":
        has_duplicate_signal = bool(claim_record.get("duplicate_candidates"))
        has_conflict_signal = bool(claim_record.get("conflict_group"))
        if not has_duplicate_signal and not has_conflict_signal:
            claim_record["status"] = "draft"
            claim_record["review_reason"] = None
    claim_record["updated_at"] = utc_now_iso()


def sync_claim_review_state_from_open_reviews(
    claim_ids: list[str],
    live_claims_by_id: dict[str, dict],
    live_reviews_by_id: dict[str, dict],
) -> set[str]:
    # 某条 review 被处理后，相关 claim 的“待审状态”不能只看旧标记，
    # 要重新基于当前仍然 open 的 review 集合计算一次。
    dirty_claim_ids: set[str] = set()
    target_claim_ids = {claim_id for claim_id in claim_ids if claim_id in live_claims_by_id}

    for claim_id in sorted(target_claim_ids):
        claim_record = live_claims_by_id.get(claim_id)
        if claim_record is None:
            continue

        duplicate_candidates: set[str] = set()
        conflict_group = claim_record.get("conflict_group")
        review_reason = None

        for review_record in live_reviews_by_id.values():
            if not is_actionable_review_record(review_record):
                continue
            candidate_claim_ids = set(review_record.get("candidate_claim_ids", []))
            if claim_id not in candidate_claim_ids:
                continue

            if review_record.get("kind") == "claim_duplicate":
                duplicate_candidates.update(candidate_claim_ids - {claim_id})
                review_reason = "possible_duplicate_claim"
            elif review_record.get("kind") == "claim_conflict":
                if not conflict_group:
                    conflict_group = claim_record.get("conflict_group")
                review_reason = "conflicting_claims_detected"

        claim_record["duplicate_candidates"] = sorted(duplicate_candidates)
        claim_record["conflict_group"] = conflict_group if review_reason == "conflicting_claims_detected" else None
        if review_reason:
            claim_record["status"] = "needs_review"
            claim_record["review_reason"] = review_reason
        else:
            claim_record["review_reason"] = None
        normalize_claim_review_flags(claim_record)
        dirty_claim_ids.add(claim_id)

    return dirty_claim_ids


def rewrite_open_reviews_for_claim_change(
    changed_review_id: str,
    removed_claim_id: str,
    replacement_claim_id: str | None,
    live_claims_by_id: dict[str, dict],
    live_reviews_by_id: dict[str, dict],
) -> tuple[set[str], set[str]]:
    # review 动作会改变候选 claim 集合：
    # - archive_one: 某个候选 claim 彻底退出
    # - merge: 某个候选 claim 被并入另一个仍存活的 claim
    # 这里把其他 open review 里对旧 claim 的引用一并修正，
    # 避免 review 队列继续挂着已经失效的候选项。
    dirty_review_ids: set[str] = set()
    touched_live_claim_ids: set[str] = set()

    for review_id, candidate_review in live_reviews_by_id.items():
        if review_id == changed_review_id:
            continue
        if candidate_review.get("status") != "open":
            continue

        original_claim_ids = list(candidate_review.get("candidate_claim_ids", []))
        if removed_claim_id not in original_claim_ids:
            continue

        rewritten_claim_ids: list[str] = []
        for claim_id in original_claim_ids:
            if claim_id != removed_claim_id:
                if claim_id not in rewritten_claim_ids:
                    rewritten_claim_ids.append(claim_id)
                continue
            if replacement_claim_id and replacement_claim_id not in rewritten_claim_ids:
                rewritten_claim_ids.append(replacement_claim_id)

        if rewritten_claim_ids == original_claim_ids:
            continue

        candidate_review["candidate_claim_ids"] = rewritten_claim_ids
        candidate_review["candidate_page_ids"] = []

        if len(rewritten_claim_ids) < 2:
            candidate_review["status"] = "resolved"
            candidate_review["resolved_at"] = utc_now_iso()
        candidate_review["lifecycle_status"] = review_lifecycle_status_for_record(candidate_review)
        dirty_review_ids.add(review_id)

        touched_live_claim_ids.update(set(original_claim_ids) | set(rewritten_claim_ids))

    sync_claim_review_state_from_open_reviews(
        claim_ids=sorted(touched_live_claim_ids),
        live_claims_by_id=live_claims_by_id,
        live_reviews_by_id=live_reviews_by_id,
    )

    return dirty_review_ids, touched_live_claim_ids


def cleanup_superseded_record_files(
    target: Path,
    historical_claims_by_id: dict[str, dict],
    historical_reviews_by_id: dict[str, dict],
) -> None:
    # state/*.jsonl 才是账本真相，但磁盘上的旧文件也要跟着收口。
    # 否则同一个 claim/review 在进入历史态后，会同时留下“旧活跃文件 + 新历史文件”两份。
    for claim_record in historical_claims_by_id.values():
        original_claim_id = claim_record.get("original_claim_id")
        if not original_claim_id or original_claim_id == claim_record["claim_id"]:
            continue
        stale_claim_path = claim_file_path(target, original_claim_id)
        if stale_claim_path.exists():
            stale_claim_path.unlink()

    for review_record in historical_reviews_by_id.values():
        original_review_id = review_record.get("original_review_id")
        if not original_review_id or original_review_id == review_record["review_id"]:
            continue
        stale_review_path = review_file_path(target, original_review_id)
        if stale_review_path.exists():
            stale_review_path.unlink()


def reload_claims_from_disk_for_review(
    target: Path,
    claim_ids: list[str],
    live_claims_by_id: dict[str, dict],
) -> set[str]:
    # edit_then_resume 的前提是“人已经先改了 claim 文件”。
    # 这里把指定 claim 的磁盘内容重新加载回内存账本。
    reloaded_claim_ids: set[str] = set()

    for claim_id in claim_ids:
        claim_path = claim_file_path(target, claim_id)
        if not claim_path.exists():
            raise FileNotFoundError(
                f"Claim file does not exist for edit_then_resume: {claim_path}"
            )
        disk_record = ensure_claim_lifecycle_defaults(load_json(claim_path))
        disk_record["claim_id"] = claim_id
        disk_record["claim_file_path"] = str(Path("claims") / claim_path.name)
        disk_record.setdefault("page_ids", live_claims_by_id.get(claim_id, {}).get("page_ids", []))
        live_claims_by_id[claim_id] = disk_record
        reloaded_claim_ids.add(claim_id)

    return reloaded_claim_ids


def apply_review_action(
    target: Path,
    review_record: dict,
    action: str,
    primary_claim_id: str | None,
    secondary_claim_id: str | None,
    primary_page_id: str | None,
    alias_value: str | None,
    live_claims_by_id: dict[str, dict],
    historical_claims_by_id: dict[str, dict],
    live_reviews_by_id: dict[str, dict],
    historical_reviews_by_id: dict[str, dict],
) -> dict:
    allowed_actions = set(review_record.get("allowed_actions", []))
    if action not in allowed_actions:
        raise ValueError(f"Action `{action}` is not allowed for review `{review_display_id(review_record)}`.")

    candidate_claim_ids = list(review_record.get("candidate_claim_ids", []))
    candidate_page_ids = list(review_record.get("candidate_page_ids", []))

    if review_record.get("kind") == "alias_conflict" and action in {"assign_alias", "remove_alias"}:
        if not primary_page_id:
            raise ValueError(f"{action} requires --primary-page-id.")
        if primary_page_id not in candidate_page_ids:
            raise ValueError("primary page must belong to candidate_page_ids.")

        evidence = review_record.get("evidence", [])
        alias_from_review = evidence[0].get("alias") if evidence else None
        alias_to_assign = alias_value or alias_from_review
        if not alias_to_assign:
            raise ValueError(f"{action} requires alias value from review evidence or --alias-value.")

        live_aliases_by_page_id = load_live_page_aliases_by_id(target)
        page_state_records = [
            ensure_page_lifecycle_defaults(record)
            for record in load_jsonl(target / "state" / "pages.jsonl")
        ]

        def update_and_validate(overrides: dict) -> dict:
            updated_overrides = apply_alias_override_action(
                overrides=overrides,
                live_aliases_by_page_id=live_aliases_by_page_id,
                candidate_page_ids=candidate_page_ids,
                primary_page_id=primary_page_id,
                alias_value=alias_to_assign,
                action=action,
            )
            page_records = [
                apply_page_alias_overrides_payload(record, updated_overrides)
                for record in page_state_records
            ]
            projected_alias_index = build_alias_index(page_records)
            projected_matches = alias_index_matches_for_value(projected_alias_index, alias_to_assign)
            projected_page_ids = sorted({match.get("page_id") for match in projected_matches if match.get("page_id")})
            projected_canonical_ids = sorted({
                match.get("canonical_id") or match.get("page_id")
                for match in projected_matches
                if match.get("canonical_id") or match.get("page_id")
            })
            primary_canonical_id = next(
                (
                    record.get("canonical_id") or record.get("page_id")
                    for record in page_state_records
                    if record.get("page_id") == primary_page_id
                ),
                primary_page_id,
            )

            if action == "assign_alias":
                if projected_canonical_ids != [primary_canonical_id]:
                    raise ValueError(
                        "assign_alias did not converge alias ownership. "
                        "Alias "
                        f"`{alias_to_assign}` would remain on canonical_ids={projected_canonical_ids} "
                        f"(page_ids={projected_page_ids})."
                    )
            elif projected_page_ids:
                raise ValueError(
                    "remove_alias did not fully clear alias ownership. "
                    f"Alias `{alias_to_assign}` would remain on page_ids={projected_page_ids}."
                )

            return clear_accepted_alias_conflict(
                updated_overrides,
                alias_to_assign,
                projected_canonical_ids,
            )

        update_page_alias_overrides_with_lock(target, update_and_validate)

        review_record["status"] = "resolved"
        review_record["resolved_at"] = utc_now_iso()
        review_record["lifecycle_status"] = "active"
        return {
            "action": action,
            "changed_page_ids": candidate_page_ids,
            "resolved_review_id": review_record["review_id"],
            "alias_value": alias_to_assign,
            "primary_page_id": primary_page_id,
        }

    if action == "keep_both":
        if review_record.get("kind") == "alias_conflict":
            evidence = review_record.get("evidence", [])
            alias_from_review = str(evidence[0].get("alias", "")).strip() if evidence else ""
            canonical_ids = [
                str(value).strip()
                for value in (evidence[0].get("canonical_ids", []) if evidence else [])
                if str(value).strip()
            ]
            if alias_from_review and canonical_ids:
                def accept_alias_conflict(overrides: dict) -> dict:
                    return persist_accepted_alias_conflict(
                        overrides=overrides,
                        alias_value=alias_from_review,
                        canonical_ids=canonical_ids,
                    )

                update_page_alias_overrides_with_lock(target, accept_alias_conflict)
        review_record["status"] = "resolved"
        review_record["resolved_at"] = utc_now_iso()
        review_record["lifecycle_status"] = "active"
        sync_claim_review_state_from_open_reviews(
            claim_ids=candidate_claim_ids,
            live_claims_by_id=live_claims_by_id,
            live_reviews_by_id=live_reviews_by_id,
        )
        return {
            "action": action,
            "changed_claim_ids": candidate_claim_ids,
            "resolved_review_id": review_record["review_id"],
        }

    if action == "edit_then_resume":
        reloaded_claim_ids = reload_claims_from_disk_for_review(
            target=target,
            claim_ids=candidate_claim_ids,
            live_claims_by_id=live_claims_by_id,
        )
        review_record["status"] = "resolved"
        review_record["resolved_at"] = utc_now_iso()
        review_record["lifecycle_status"] = "active"
        sync_claim_review_state_from_open_reviews(
            claim_ids=sorted(reloaded_claim_ids),
            live_claims_by_id=live_claims_by_id,
            live_reviews_by_id=live_reviews_by_id,
        )
        return {
            "action": action,
            "changed_claim_ids": sorted(reloaded_claim_ids),
            "resolved_review_id": review_record["review_id"],
        }

    if action == "archive_one":
        if not primary_claim_id:
            raise ValueError("archive_one requires --primary-claim-id.")
        if primary_claim_id not in candidate_claim_ids:
            raise ValueError("primary claim must belong to candidate_claim_ids.")
        claim_record = resolve_claim_record_for_action(primary_claim_id, live_claims_by_id, historical_claims_by_id)
        if claim_record["claim_id"] not in live_claims_by_id:
            raise ValueError("archive_one currently only supports active claims.")
        historical_claim_record = archive_live_claim(
            claim_record=claim_record,
            live_claims_by_id=live_claims_by_id,
            historical_claims_by_id=historical_claims_by_id,
        )
        review_record["status"] = "resolved"
        review_record["resolved_at"] = utc_now_iso()
        review_record["lifecycle_status"] = "active"
        rewrite_open_reviews_for_claim_change(
            changed_review_id=review_record["review_id"],
            removed_claim_id=primary_claim_id,
            replacement_claim_id=None,
            live_claims_by_id=live_claims_by_id,
            live_reviews_by_id=live_reviews_by_id,
        )
        return {
            "action": action,
            "changed_claim_ids": [historical_claim_record["claim_id"]],
            "resolved_review_id": review_record["review_id"],
        }

    if action == "merge":
        if not primary_claim_id or not secondary_claim_id:
            raise ValueError("merge requires --primary-claim-id and --secondary-claim-id.")
        if primary_claim_id == secondary_claim_id:
            raise ValueError("primary and secondary claim ids must be different.")
        if primary_claim_id not in candidate_claim_ids or secondary_claim_id not in candidate_claim_ids:
            raise ValueError("Both primary and secondary claims must belong to candidate_claim_ids.")

        primary_record = resolve_claim_record_for_action(primary_claim_id, live_claims_by_id, historical_claims_by_id)
        secondary_record = resolve_claim_record_for_action(secondary_claim_id, live_claims_by_id, historical_claims_by_id)
        if primary_record["claim_id"] not in live_claims_by_id or secondary_record["claim_id"] not in live_claims_by_id:
            raise ValueError("merge currently only supports active claims.")

        merged_record = merge_claim_records(primary_record, secondary_record)
        merged_record["duplicate_candidates"] = [
            item for item in merged_record.get("duplicate_candidates", [])
            if item not in {merged_record["claim_id"], secondary_record["claim_id"]}
        ]
        merged_record["conflict_group"] = None
        merged_record["review_reason"] = None
        merged_record["status"] = "draft"
        merged_record["updated_at"] = utc_now_iso()
        live_claims_by_id[merged_record["claim_id"]] = merged_record

        historical_secondary = archive_live_claim(
            claim_record=secondary_record,
            live_claims_by_id=live_claims_by_id,
            historical_claims_by_id=historical_claims_by_id,
            archived_by_claim_id=merged_record["claim_id"],
        )
        review_record["status"] = "resolved"
        review_record["resolved_at"] = utc_now_iso()
        review_record["lifecycle_status"] = "active"
        rewrite_open_reviews_for_claim_change(
            changed_review_id=review_record["review_id"],
            removed_claim_id=secondary_claim_id,
            replacement_claim_id=merged_record["claim_id"],
            live_claims_by_id=live_claims_by_id,
            live_reviews_by_id=live_reviews_by_id,
        )
        normalize_claim_review_flags(merged_record)
        return {
            "action": action,
            "changed_claim_ids": [merged_record["claim_id"], historical_secondary["claim_id"]],
            "resolved_review_id": review_record["review_id"],
        }

    raise NotImplementedError(f"Action `{action}` is not implemented yet.")


def command_review_apply(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    claims_path = target / "state" / "claims.jsonl"
    reviews_path = target / "state" / "reviews.jsonl"

    live_claims_by_id, historical_claims_by_id, _ = load_claim_state_maps(target)
    live_reviews_by_id, historical_reviews_by_id, _ = load_review_state_maps(target)
    _, _, archived_alias_review_ids = refresh_alias_conflict_reviews(
        target=target,
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )
    if archived_alias_review_ids:
        refreshed_review_state_records = build_ordered_review_state_records(
            live_reviews_by_id=live_reviews_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )
        write_jsonl(reviews_path, refreshed_review_state_records)
        for review_state_record in refreshed_review_state_records:
            write_review_file(target, review_state_record)
        cleanup_superseded_record_files(
            target=target,
            historical_claims_by_id=historical_claims_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )

    review_record = live_reviews_by_id.get(args.review_id) or historical_reviews_by_id.get(args.review_id)
    if review_record is None:
        matched_live_review = next(
            (
                record for record in live_reviews_by_id.values()
                if review_display_id(record) == args.review_id
            ),
            None,
        )
        if matched_live_review is not None:
            review_record = matched_live_review
    if review_record is None:
        raise KeyError(f"Unknown review_id: {args.review_id}")
    if review_record["review_id"] not in live_reviews_by_id:
        raise ValueError("review_apply currently only supports active review items.")

    result = apply_review_action(
        target=target,
        review_record=review_record,
        action=args.action,
        primary_claim_id=args.primary_claim_id,
        secondary_claim_id=args.secondary_claim_id,
        primary_page_id=args.primary_page_id,
        alias_value=args.alias_value,
        live_claims_by_id=live_claims_by_id,
        historical_claims_by_id=historical_claims_by_id,
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )
    live_reviews_by_id[review_record["review_id"]] = review_record

    write_jsonl(
        claims_path,
        build_ordered_claim_state_records(
            live_claims_by_id=live_claims_by_id,
            historical_claims_by_id=historical_claims_by_id,
        ),
    )
    for claim_record in [*live_claims_by_id.values(), *historical_claims_by_id.values()]:
        write_claim_file(target, claim_record)

    write_jsonl(
        reviews_path,
        build_ordered_review_state_records(
            live_reviews_by_id=live_reviews_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        ),
    )
    for review_state_record in [*live_reviews_by_id.values(), *historical_reviews_by_id.values()]:
        write_review_file(target, review_state_record)
    cleanup_superseded_record_files(
        target=target,
        historical_claims_by_id=historical_claims_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )

    rebuild_review_affected_pages(
        target=target,
        live_claims_by_id=live_claims_by_id,
        live_reviews_by_id=live_reviews_by_id,
    )
    _, _, archived_alias_review_ids = refresh_alias_conflict_reviews(
        target=target,
        live_reviews_by_id=live_reviews_by_id,
        historical_reviews_by_id=historical_reviews_by_id,
    )
    if archived_alias_review_ids:
        write_jsonl(
            reviews_path,
            build_ordered_review_state_records(
                live_reviews_by_id=live_reviews_by_id,
                historical_reviews_by_id=historical_reviews_by_id,
            ),
        )
        for review_state_record in [*live_reviews_by_id.values(), *historical_reviews_by_id.values()]:
            write_review_file(target, review_state_record)
        cleanup_superseded_record_files(
            target=target,
            historical_claims_by_id=historical_claims_by_id,
            historical_reviews_by_id=historical_reviews_by_id,
        )

    payload = {
        "workspace": str(target),
        "workspace_summary": build_workspace_summary(target),
        "review_id": review_record["review_id"],
        "display_id": review_display_id(review_record),
        **result,
    }
    return CommandResult(
        payload=payload,
        message=render_workspace_summary_message(
            "Review action applied.",
            target_dir=target,
            extra_lines=[
                f"Review: {review_display_id(review_record)}",
                f"Action: {result.get('action')}",
                f"Review id: {review_record['review_id']}",
            ],
        ),
    )


def command_init(args: argparse.Namespace) -> CommandResult:
    # init 负责把“普通资料目录”初始化成“可被 Agent 驱动的 MyAgentWiki 工作区”。
    root = find_project_root()
    raw_dir = Path(args.source_dir).expanduser().resolve() if args.source_dir else None
    target_dir = Path(args.target_dir).expanduser().resolve() if args.target_dir else None

    if raw_dir is None and target_dir is None:
        target_dir = (Path.cwd() / args.project_name).resolve()
        raw_dir = (target_dir.parent / "raw").resolve()
    elif raw_dir is not None and target_dir is None:
        target_dir = (raw_dir.parent / args.project_name).resolve()
    elif raw_dir is None and target_dir is not None:
        raw_dir = (target_dir.parent / "raw").resolve()

    assert raw_dir is not None
    assert target_dir is not None
    if raw_dir.name != "raw":
        raise ValueError(f"Raw directory must be named 'raw': {raw_dir}")
    if raw_dir.parent != target_dir.parent:
        raise ValueError(
            f"Raw directory must be a sibling of the workspace: raw={raw_dir} target={target_dir}"
        )
    if raw_dir.exists() and not raw_dir.is_dir():
        raise FileExistsError(f"Raw path exists but is not a directory: {raw_dir}")

    ensure_clean_target(target_dir)
    # 这里即使目录不存在也没关系，mkdir 会顺手把父目录一起建出来。
    ensure_directory(target_dir)
    raw_dir_preexisting = raw_dir.exists()
    assets_dir = (raw_dir.parent / "assets").resolve()
    assets_dir_preexisting = assets_dir.exists()
    ensure_directory(raw_dir)
    ensure_directory(assets_dir)
    raw_dir_relative_path = os.path.relpath(raw_dir, start=target_dir).replace(os.sep, "/")
    assets_dir_relative_path = os.path.relpath(assets_dir, start=target_dir).replace(os.sep, "/")

    context = {
        "project_name": args.project_name,
        "source_dir_name": raw_dir.name,
        "source_dir_path": str(raw_dir),
        "raw_dir_name": raw_dir.name,
        "raw_dir_path": str(raw_dir),
        "raw_dir_relative_path": raw_dir_relative_path,
        "assets_dir_name": assets_dir.name,
        "assets_dir_path": str(assets_dir),
        "assets_dir_relative_path": assets_dir_relative_path,
        "python_executable": sys.executable,
    }

    for directory in (
        "normalized",
        "chunks",
        "claims",
        "semantic",
        "semantic/batches",
        "wiki",
        "indexes",
        "state",
        "reviews",
        "logs",
        "outputs",
        "config",
        "reports/lint",
    ):
        # 工作区目录先一次性建齐，后面各阶段脚本就可以假定这些路径始终存在。
        ensure_directory(target_dir / directory)

    template_root = root / "templates" / "project"
    # 模板文件保存的是“用户工作区初始化时需要复制的骨架文件”。
    template_files = {
        "AGENTS.md.tmpl": target_dir / "AGENTS.md",
        "CLAUDE.md.tmpl": target_dir / "CLAUDE.md",
        "gitignore.tmpl": target_dir / ".gitignore",
        "wiki/index.md.tmpl": target_dir / "wiki" / "index.md",
        "wiki/log.md.tmpl": target_dir / "wiki" / "log.md",
        "config/project.yml.tmpl": target_dir / "config" / "project.yml",
        "config/runtime_manifest.yml.tmpl": target_dir / "config" / "runtime_manifest.yml",
    }

    for template_name, output_path in template_files.items():
        # 用同一份 context 渲染，能保证 README、配置、Agent 说明里的项目名一致。
        rendered = render_template(template_root / template_name, context)
        output_path.write_text(rendered, encoding="utf-8")

    metadata_files = {
        target_dir / "state" / "sources.jsonl": [],
        target_dir / "state" / "ingest_state.jsonl": [],
        target_dir / "state" / "error_log.jsonl": [],
        target_dir / "state" / "normalized.jsonl": [],
        target_dir / STRUCTURE_BLOCKS_REL_PATH: [],
        target_dir / EVIDENCE_BLOCKS_REL_PATH: [],
        target_dir / KNOWLEDGE_UNITS_REL_PATH: [],
        target_dir / "state" / "chunks.jsonl": [],
        target_dir / "state" / "claims.jsonl": [],
        target_dir / "state" / "reviews.jsonl": [],
        target_dir / "state" / "pages.jsonl": [],
        target_dir / SEMANTIC_DECISIONS_REL_PATH: [],
        target_dir / SEARCH_PAGES_INDEX_REL_PATH: [],
    }
    for path, records in metadata_files.items():
        # 这些状态文件先创建成空 JSONL，后面脚本就不用额外判“文件是否存在”。
        write_jsonl(path, records)

    write_alias_index(target_dir, [])

    git_steps: list[str] = []
    if not (target_dir / ".git").exists():
        git_steps = git_init_and_commit(target_dir)

    payload = {
        "project_name": args.project_name,
        "source_dir": str(raw_dir),
        "raw_dir": str(raw_dir),
        "assets_dir": str(assets_dir),
        "target_dir": str(target_dir),
        "workspace_summary": build_workspace_summary(target_dir, raw_dir),
        "created_directories": [
            str(raw_dir),
            str(assets_dir),
            *[
                str(target_dir / path)
                for path in (
                    "normalized",
                    "chunks",
                    "claims",
                    "semantic",
                    "semantic/batches",
                    "wiki",
                    "indexes",
                    "state",
                    "reviews",
                    "logs",
                    "outputs",
                    "config",
                    "reports/lint",
                )
            ],
        ],
        "raw_dir_relative_path": raw_dir_relative_path,
        "raw_dir_preexisting": raw_dir_preexisting,
        "assets_dir_relative_path": assets_dir_relative_path,
        "assets_dir_preexisting": assets_dir_preexisting,
        "metadata_files": [str(path) for path in metadata_files],
        "git_steps": git_steps,
    }
    return CommandResult(
        payload=payload,
        message=render_workspace_summary_message(
            "Workspace initialized.",
            target_dir=target_dir,
            raw_dir=raw_dir,
            extra_lines=[
                f"Project name: {args.project_name}",
                f"Raw directory existed before init: {'yes' if raw_dir_preexisting else 'no'}",
            ],
        ),
    )


def command_ingest(args: argparse.Namespace) -> CommandResult:
    # ingest 是第一条真正处理用户资料的流水线命令。
    # 它现在包含五步：来源登记、标准化、切块、Claim 草稿抽取、Wiki 页面生成。
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    config = load_workspace_config(target)
    post_ingest_config = load_post_ingest_review_auto_config(config)
    document_analysis_config = load_semantic_task_config(config, "document_analysis")
    claim_candidate_quality_config = load_semantic_task_config(config, "claim_candidate_quality")
    claim_role_config = load_semantic_task_config(config, "claim_role")
    page_intent_config = load_semantic_task_config(config, "page_intent")
    readable_concept_render_config = load_readable_concept_render_config(config)
    overview_render_config = load_page_render_config(config, "overview")
    raw_dir = resolve_workspace_raw_dir(target)
    if not raw_dir.exists():
        raise FileNotFoundError(f"Raw directory does not exist: {raw_dir}")

    sources_path = target / "state" / "sources.jsonl"
    ingest_state_path = target / "state" / "ingest_state.jsonl"
    normalized_path = target / "state" / "normalized.jsonl"
    structure_blocks_path = target / STRUCTURE_BLOCKS_REL_PATH
    evidence_blocks_path = target / EVIDENCE_BLOCKS_REL_PATH
    knowledge_units_path = target / KNOWLEDGE_UNITS_REL_PATH
    chunks_path = target / "state" / "chunks.jsonl"
    claims_path = target / "state" / "claims.jsonl"
    reviews_path = target / "state" / "reviews.jsonl"
    error_log_path = target / "state" / "error_log.jsonl"
    pages_path = target / "state" / "pages.jsonl"

    existing_sources = load_jsonl(sources_path)
    existing_by_hash = {record["source_hash"]: record for record in existing_sources}
    sources_by_id = {record["source_id"]: record for record in existing_sources}
    latest_source_by_path = build_latest_source_record_by_path(existing_sources)
    existing_normalized = {record["source_id"]: record for record in load_jsonl(normalized_path)}
    existing_structure_source_ids = {record["source_id"] for record in load_jsonl(structure_blocks_path)}
    existing_evidence_source_ids = {record["source_id"] for record in load_jsonl(evidence_blocks_path)}
    existing_knowledge_source_ids = {record["source_id"] for record in load_jsonl(knowledge_units_path)}
    existing_structured_source_ids = (
        existing_structure_source_ids
        & existing_evidence_source_ids
        & existing_knowledge_source_ids
    )
    existing_chunked = {record["source_id"]: record for record in load_jsonl(chunks_path)}
    existing_claim_records = [ensure_claim_lifecycle_defaults(record) for record in load_jsonl(claims_path)]
    live_existing_claims = filter_live_claim_records(existing_claim_records)
    historical_existing_claims = [
        record for record in existing_claim_records
        if not is_live_claim_record(record)
    ]
    claims_by_id = {record["claim_id"]: record for record in live_existing_claims}
    historical_claims_by_id = {record["claim_id"]: record for record in historical_existing_claims}
    claims_by_normalized_text = {record["normalized_text"]: record for record in live_existing_claims}
    claims_by_similarity_bucket: dict[str, list[dict]] = {}
    for record in live_existing_claims:
        claims_by_similarity_bucket.setdefault(build_similarity_bucket(record["text"]), []).append(record)
    claim_similarity_index = rebuild_claim_similarity_index(live_existing_claims)
    existing_review_records = [
        ensure_review_lifecycle_defaults(record)
        for record in (load_jsonl(reviews_path) if reviews_path.exists() else [])
    ]
    live_existing_reviews = filter_live_review_records(existing_review_records)
    historical_existing_reviews = [
        record for record in existing_review_records
        if not is_live_review_record(record)
    ]
    existing_reviews = {
        record["review_id"]: record for record in live_existing_reviews
    }
    historical_reviews_by_id = {
        record["review_id"]: ensure_review_lifecycle_defaults(record)
        for record in historical_existing_reviews
    }
    # 这里按 hash 做首轮去重，意味着内容完全一样的文件只会登记一次。
    # 后面如果要支持“同内容不同来源”的更细建模，可以再加来源关系表。

    created_sources: list[dict] = []
    skipped_sources: list[dict] = []
    normalized_sources: list[dict] = []
    chunked_sources: list[dict] = []
    claimed_sources: list[dict] = []
    review_items: list[dict] = []
    error_items: list[dict] = []
    generated_pages: list[dict] = []
    task_id = f"ingest_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    reingested_source_ids: set[str] = set()
    purged_claim_ids: set[str] = set()
    purged_review_ids: set[str] = set()

    # ingest 第一阶段只做来源登记：扫描 raw、生成 source_id、写 sources/state。
    for file_path in collect_files(raw_dir):
        source_hash = file_sha256(file_path)
        relative_path = os.path.relpath(file_path, start=target).replace(os.sep, "/")
        if source_hash in existing_by_hash:
            # 已有相同内容时跳过，避免反复导入导致 state 膨胀。
            skipped_sources.append({
                "path": str(file_path),
                "source_hash": source_hash,
                "reason": "duplicate_hash",
                "source_id": existing_by_hash[source_hash]["source_id"],
            })
            continue

        previous_path_record = latest_source_by_path.get(relative_path)
        if previous_path_record is not None:
            # 同一路径但内容哈希变化时，按“路径级版本更新”处理：
            # 复用原 source_id，清理旧证据链，再让后续 normalized/chunk/claim/wiki
            # 围绕同一个 source_id 重算，而不是无止境地产生平行来源版本。
            source_id = previous_path_record["source_id"]
            version_group = previous_path_record.get("version_group") or build_source_version_group_from_source_path(relative_path)
            previous_source_hash = previous_path_record.get("source_hash")
            previous_normalized_path = previous_path_record.get("normalized_path")

            existing_normalized.pop(source_id, None)
            existing_structured_source_ids.discard(source_id)
            existing_chunked.pop(source_id, None)

            if previous_normalized_path:
                normalized_file_path = target / previous_normalized_path
                if normalized_file_path.exists():
                    normalized_file_path.unlink()

            chunk_file_path = target / "chunks" / f"{source_id}.jsonl"
            if chunk_file_path.exists():
                chunk_file_path.unlink()

            replace_source_scoped_jsonl_records(structure_blocks_path, source_id, [])
            replace_source_scoped_jsonl_records(evidence_blocks_path, source_id, [])
            replace_source_scoped_jsonl_records(knowledge_units_path, source_id, [])

            dirty_claim_ids, deleted_claim_ids = purge_source_from_claims(
                target=target,
                claims_by_id=claims_by_id,
                historical_claims_by_id=historical_claims_by_id,
                source_id=source_id,
            )
            purged_claim_ids.update(deleted_claim_ids)

            dirty_review_ids, deleted_review_ids = purge_deleted_claims_from_reviews(
                reviews_by_id=existing_reviews,
                historical_reviews_by_id=historical_reviews_by_id,
                deleted_claim_ids=deleted_claim_ids,
            )
            purged_review_ids.update(deleted_review_ids)

            claims_by_normalized_text = {
                record["normalized_text"]: record for record in claims_by_id.values()
            }
            claims_by_similarity_bucket = {}
            for record in claims_by_id.values():
                claims_by_similarity_bucket.setdefault(build_similarity_bucket(record["text"]), []).append(record)
            claim_similarity_index = rebuild_claim_similarity_index(list(claims_by_id.values()))

            updated_record = dict(previous_path_record)
            updated_record.update({
                "source_path": relative_path,
                "source_type": infer_source_type(file_path),
                "source_hash": source_hash,
                "dedupe_key": source_hash,
                "version_group": version_group,
                "imported_at": utc_now_iso(),
                "status": "new",
                "normalized_path": None,
                "warnings": [],
            })
            replace_jsonl_record(sources_path, "source_id", source_id, updated_record)
            replace_jsonl_record(
                ingest_state_path,
                "source_id",
                source_id,
                {
                    "task_id": task_id,
                    "source_id": source_id,
                    "state": "new",
                    "last_successful_stage": None,
                    "failed_stage": None,
                    "retry_count": 0,
                    "updated_at": utc_now_iso(),
                },
            )

            sources_by_id[source_id] = updated_record
            latest_source_by_path[relative_path] = updated_record
            existing_by_hash[source_hash] = updated_record
            created_sources.append(updated_record)
            reingested_source_ids.add(source_id)
            continue

        source_id = build_source_id(raw_dir, file_path, source_hash)
        version_group = build_source_version_group(raw_dir, file_path)
        record = {
            "source_id": source_id,
            "source_path": relative_path,
            "source_type": infer_source_type(file_path),
            "source_uri": None,
            "source_hash": source_hash,
            "dedupe_key": source_hash,
            "version_group": version_group,
            "imported_at": utc_now_iso(),
            "status": "new",
            "normalized_path": None,
            "warnings": [],
        }
        append_jsonl(sources_path, record)
        append_jsonl(
            ingest_state_path,
            {
                "task_id": task_id,
                "source_id": source_id,
                "state": "new",
                "last_successful_stage": None,
                "failed_stage": None,
                "retry_count": 0,
                "updated_at": utc_now_iso(),
            },
        )
        created_sources.append(record)
        existing_sources.append(record)
        sources_by_id[source_id] = record
        latest_source_by_path[relative_path] = record

    # 来源登记完成后，第一版先支持 markdown/plain_text 的最小标准化。
    for source_record in load_jsonl(sources_path):
        if source_record["source_id"] in existing_normalized:
            # 已标准化过的记录直接跳过，保证 ingest 可以重复执行而不反复写文件。
            continue
        normalized_record = normalize_source_record(
            target,
            source_record,
            allow_insecure_downloads=not args.disable_insecure_download_retry,
        )
        if normalized_record is None:
            # 暂不支持的文件类型先保留在 sources.jsonl，等待后续转换器接手。
            continue

        append_jsonl(normalized_path, normalized_record)
        normalized_sources.append(normalized_record)

        if normalized_record["extraction_quality"] in {"failed", "poor", "partial"}:
            level = "error" if normalized_record["extraction_quality"] == "failed" else "warning"
            error_items.append(
                append_error_record(
                    error_log_path=error_log_path,
                    task_id=task_id,
                    source_id=normalized_record["source_id"],
                    stage="normalized",
                    level=level,
                    message=f"Normalization finished with quality={normalized_record['extraction_quality']}",
                    details={
                        "warnings": normalized_record.get("warnings", []),
                        "normalized_path": normalized_record["normalized_path"],
                        "source_type": normalized_record["source_type"],
                    },
                )
            )

        updated_source = dict(source_record)
        updated_source["status"] = (
            "failed"
            if normalized_record["extraction_quality"] == "failed"
            else "review_required"
            if normalized_record["extraction_quality"] == "poor"
            else "normalized"
        )
        updated_source["normalized_path"] = normalized_record["normalized_path"]
        # sources.jsonl 记录的是“来源主档案”，状态变化后要同步更新。
        replace_jsonl_record(sources_path, "source_id", source_record["source_id"], updated_source)
        sources_by_id[source_record["source_id"]] = updated_source

        replace_jsonl_record(
            ingest_state_path,
            "source_id",
            source_record["source_id"],
            {
                "task_id": task_id,
                "source_id": source_record["source_id"],
                "state": (
                    "failed"
                    if normalized_record["extraction_quality"] == "failed"
                    else "review_required"
                    if normalized_record["extraction_quality"] == "poor"
                    else "normalized"
                ),
                "last_successful_stage": None if normalized_record["extraction_quality"] == "failed" else "normalized",
                "failed_stage": "normalized" if normalized_record["extraction_quality"] == "failed" else None,
                "retry_count": 0,
                "updated_at": utc_now_iso(),
            },
        )

    normalized_records = load_jsonl(normalized_path)
    if normalized_records and document_analysis_config.enabled:
        run_semantic_batch_task(
            target=target,
            task_name="document_analysis",
            dry_run=False,
        )
        normalized_records = apply_document_analysis_decisions_to_normalized_records(
            target=target,
            normalized_records=normalized_records,
            task_config=document_analysis_config,
        )
        existing_normalized = {record["source_id"]: record for record in normalized_records}

    structured_sources: list[dict] = []
    for normalized_record in normalized_records:
        source_id = normalized_record["source_id"]
        if source_id in existing_structured_source_ids:
            continue
        if normalized_record["extraction_quality"] not in {"good", "partial"}:
            continue

        normalized_file_path = target / normalized_record["normalized_path"]
        if not normalized_file_path.exists():
            continue
        normalized_text = normalized_file_path.read_text(encoding="utf-8")
        compiled_records = compile_structure_knowledge_records(normalized_record, normalized_text)
        replace_source_scoped_jsonl_records(
            structure_blocks_path,
            source_id,
            compiled_records["structure_blocks"],
        )
        replace_source_scoped_jsonl_records(
            evidence_blocks_path,
            source_id,
            compiled_records["evidence_blocks"],
        )
        replace_source_scoped_jsonl_records(
            knowledge_units_path,
            source_id,
            compiled_records["knowledge_units"],
        )
        structured_source = {
            "source_id": source_id,
            "structure_block_count": len(compiled_records["structure_blocks"]),
            "evidence_block_count": len(compiled_records["evidence_blocks"]),
            "knowledge_unit_count": len(compiled_records["knowledge_units"]),
            "updated_at": compiled_records["updated_at"],
        }
        structured_sources.append(structured_source)
        existing_structured_source_ids.add(source_id)

    # 标准化完成后，继续把合格文档切成 chunk，作为 claim / query 的基础证据单元。
    for normalized_record in normalized_records:
        if normalized_record["source_id"] in existing_chunked:
            continue

        chunk_result = chunk_normalized_record(target, normalized_record)
        if chunk_result is None:
            continue

        # 同一 source_id 的 chunk 在原位更新场景下应该整体替换，而不是继续 append。
        replace_jsonl_records_by_filter(
            chunks_path,
            keep_predicate=lambda record, source_id=chunk_result["source_id"]: record.get("source_id") != source_id,
            replacement_records=chunk_result["chunks"],
        )
        chunked_sources.append({
            "source_id": chunk_result["source_id"],
            "chunk_file_path": chunk_result["chunk_file_path"],
            "chunk_count": chunk_result["chunk_count"],
            "updated_at": chunk_result["updated_at"],
        })
        existing_chunked[chunk_result["source_id"]] = {
            "source_id": chunk_result["source_id"],
            "chunk_count": chunk_result["chunk_count"],
        }

        source_record = sources_by_id.get(normalized_record["source_id"])
        if source_record is not None:
            updated_source = dict(source_record)
            updated_source["status"] = "chunked"
            replace_jsonl_record(sources_path, "source_id", source_record["source_id"], updated_source)
            sources_by_id[source_record["source_id"]] = updated_source

        replace_jsonl_record(
            ingest_state_path,
            "source_id",
            normalized_record["source_id"],
            {
                "task_id": task_id,
                "source_id": normalized_record["source_id"],
                "state": "chunked",
                "last_successful_stage": "chunked",
                "failed_stage": None,
                "retry_count": 0,
                "updated_at": utc_now_iso(),
            },
        )

    # Claim 草稿层先做“保守抽取”：
    # 优先从 Knowledge Unit 抽取，缺少 V2 结构账本时再回退到 chunk 文本。
    chunk_records = load_jsonl(chunks_path)
    knowledge_unit_records = load_jsonl(knowledge_units_path)
    chunks_by_source_id_for_claims: dict[str, list[dict]] = {}
    for chunk_record in chunk_records:
        chunks_by_source_id_for_claims.setdefault(chunk_record["source_id"], []).append(chunk_record)
    knowledge_units_by_source_id: dict[str, list[dict]] = {}
    for knowledge_unit_record in knowledge_unit_records:
        if knowledge_unit_record.get("lifecycle_status", "active") != "active":
            continue
        knowledge_units_by_source_id.setdefault(knowledge_unit_record["source_id"], []).append(knowledge_unit_record)
    claims_created_by_source: dict[str, int] = {}
    completed_claim_source_ids = {
        source_id
        for source_id, source_record in sources_by_id.items()
        if source_claim_stage_completed(source_record)
    }
    active_source_ids = choose_active_source_ids(sources_by_id)

    claim_candidate_source_ids = sorted(set(chunks_by_source_id_for_claims) | set(knowledge_units_by_source_id))
    for source_id in claim_candidate_source_ids:
        if source_id in completed_claim_source_ids:
            continue
        source_claim_candidates = build_claim_candidates_for_source(
            source_id=source_id,
            knowledge_units_by_source_id=knowledge_units_by_source_id,
            chunks_by_source_id=chunks_by_source_id_for_claims,
        )
        for claim_record in source_claim_candidates:
            existing_claim = claims_by_normalized_text.get(claim_record["normalized_text"])
            if existing_claim is not None:
                merged_claim = merge_claim_records(existing_claim, claim_record)
                write_claim_file(target, merged_claim)
                replace_jsonl_record(claims_path, "claim_id", merged_claim["claim_id"], merged_claim)
                claims_by_id[merged_claim["claim_id"]] = merged_claim
                claims_by_normalized_text[merged_claim["normalized_text"]] = merged_claim
                continue

            similarity_bucket = build_similarity_bucket(claim_record["text"])
            candidate_claim_ids = collect_claim_review_candidate_ids(
                claim_record=claim_record,
                claims_by_similarity_bucket=claims_by_similarity_bucket,
                claim_similarity_index=claim_similarity_index,
            )
            conflicting_candidates = []
            duplicate_candidates = []
            incoming_has_negation = has_negation(claim_record["text"])

            for candidate_claim_id in sorted(candidate_claim_ids):
                similar_claim = claims_by_id.get(candidate_claim_id)
                if similar_claim is None:
                    continue
                if not claims_are_similar_for_review(claim_record["text"], similar_claim["text"]):
                    continue
                if has_negation(similar_claim["text"]) != incoming_has_negation:
                    conflicting_candidates.append(similar_claim)
                else:
                    duplicate_candidates.append(similar_claim)

            if duplicate_candidates:
                claim_record["duplicate_candidates"] = [item["claim_id"] for item in duplicate_candidates]
                claim_record["review_reason"] = "possible_duplicate_claim"
                claim_record["status"] = "needs_review"

                review_record = build_review_record(
                    kind="claim_duplicate",
                    candidate_claim_ids=sorted([claim_record["claim_id"], *[item["claim_id"] for item in duplicate_candidates]]),
                    reason="Detected highly similar claims that may need merge or archive decisions.",
                    evidence=[
                        {
                            "claim_id": claim_record["claim_id"],
                            "text": claim_record["text"],
                            "source_refs": claim_record["source_refs"],
                        },
                        *[
                            {
                                "claim_id": item["claim_id"],
                                "text": item["text"],
                                "source_refs": item.get("source_refs", []),
                            }
                            for item in duplicate_candidates
                        ],
                    ],
                    recommended_action="merge",
                    signature_parts=[
                        claim_record["normalized_text"],
                        *sorted(item["claim_id"] for item in duplicate_candidates),
                    ],
                )
                if review_record["review_id"] not in existing_reviews:
                    review_file_path = write_review_file(target, review_record)
                    review_record["review_file_path"] = review_file_path
                    append_jsonl(reviews_path, review_record)
                    existing_reviews[review_record["review_id"]] = review_record
                    review_items.append(review_record)

                error_items.append(
                    append_error_record(
                        error_log_path=error_log_path,
                        task_id=task_id,
                        source_id=claim_record["source_ids"][0],
                        stage="claim",
                        level="warning",
                        message="Possible duplicate claims detected",
                        details={
                            "claim_id": claim_record["claim_id"],
                            "duplicate_candidates": claim_record["duplicate_candidates"],
                        },
                    )
                )

            if conflicting_candidates:
                conflict_claim_ids = sorted([claim_record["claim_id"], *[item["claim_id"] for item in conflicting_candidates]])
                conflict_group = hashlib.sha256("|".join(conflict_claim_ids).encode("utf-8")).hexdigest()[:12]
                claim_record["conflict_group"] = f"cfg_{conflict_group}"
                claim_record["review_reason"] = "conflicting_claims_detected"
                claim_record["status"] = "needs_review"

                evidence = [
                    {
                        "claim_id": claim_record["claim_id"],
                        "text": claim_record["text"],
                        "source_refs": claim_record["source_refs"],
                    }
                ]
                evidence.extend({
                    "claim_id": item["claim_id"],
                    "text": item["text"],
                    "source_refs": item.get("source_refs", []),
                } for item in conflicting_candidates)

                review_record = build_review_record(
                    kind="claim_conflict",
                    candidate_claim_ids=conflict_claim_ids,
                    reason="Detected claims with opposite negation pattern but very similar normalized content.",
                    evidence=evidence,
                    recommended_action="keep_both",
                )
                if review_record["review_id"] not in existing_reviews:
                    review_file_path = write_review_file(target, review_record)
                    review_record["review_file_path"] = review_file_path
                    append_jsonl(reviews_path, review_record)
                    existing_reviews[review_record["review_id"]] = review_record
                    review_items.append(review_record)

                error_items.append(
                    append_error_record(
                        error_log_path=error_log_path,
                        task_id=task_id,
                        source_id=claim_record["source_ids"][0],
                        stage="claim",
                        level="warning",
                        message="Conflicting claims detected",
                        details={
                            "claim_id": claim_record["claim_id"],
                            "conflict_group": claim_record["conflict_group"],
                            "candidate_claim_ids": conflict_claim_ids,
                        },
                    )
                )

            claim_file_rel_path = write_claim_file(target, claim_record)
            claim_record["claim_file_path"] = claim_file_rel_path
            append_jsonl(claims_path, claim_record)
            claims_by_id[claim_record["claim_id"]] = claim_record
            claims_by_normalized_text[claim_record["normalized_text"]] = claim_record
            claims_by_similarity_bucket.setdefault(similarity_bucket, []).append(claim_record)
            index_claim_similarity_tokens(claim_similarity_index, claim_record)
            source_id = claim_record["source_ids"][0]
            claims_created_by_source[source_id] = claims_created_by_source.get(source_id, 0) + 1

    for source_id, claim_count in sorted(claims_created_by_source.items()):
        source_claims = [
            record for record in claims_by_id.values()
            if source_id in record.get("source_ids", [])
        ]
        has_review_claim = any(record.get("status") == "needs_review" for record in source_claims)

        claimed_sources.append({
            "source_id": source_id,
            "claim_count": claim_count,
            "needs_review": has_review_claim,
        })

        source_record = sources_by_id.get(source_id)
        if source_record is not None:
            updated_source = dict(source_record)
            updated_source["status"] = "review_required" if has_review_claim else "claimed"
            replace_jsonl_record(sources_path, "source_id", source_id, updated_source)
            sources_by_id[source_id] = updated_source

        replace_jsonl_record(
            ingest_state_path,
            "source_id",
            source_id,
            {
                "task_id": task_id,
                "source_id": source_id,
                "state": "review_required" if has_review_claim else "claimed",
                "last_successful_stage": "claimed",
                "failed_stage": None,
                "retry_count": 0,
                "updated_at": utc_now_iso(),
            },
        )

    # 如果来源被原位重算过，claims/reviews 的内存状态可能已经发生清理或删除，
    # 这里先统一刷回 state 文件，确保后续页面生成和 lint 看到的是同一套事实。
    if reingested_source_ids or purged_claim_ids:
        claim_state_records = build_ordered_claim_state_records(
            live_claims_by_id=claims_by_id,
            historical_claims_by_id=historical_claims_by_id,
        )
        write_jsonl(claims_path, claim_state_records)
        for claim_record in claim_state_records:
            write_claim_file(target, claim_record)

    current_claim_records = [ensure_claim_lifecycle_defaults(record) for record in load_jsonl(claims_path)]
    quality_archived_claim_ids: set[str] = set()
    if current_claim_records and claim_candidate_quality_config.enabled:
        run_semantic_batch_task(
            target=target,
            task_name="claim_candidate_quality",
            dry_run=False,
        )
        current_claim_records, quality_archived_claim_ids, _ = apply_claim_candidate_quality_decisions_to_claim_records(
            target=target,
            claim_records=current_claim_records,
            task_config=claim_candidate_quality_config,
        )
        if quality_archived_claim_ids:
            claims_created_by_source = {
                source_id: count
                for source_id, count in claims_created_by_source.items()
                if count > 0
            }

    semantic_claim_updates_applied = False
    if current_claim_records and claim_role_config.enabled:
        claim_records_before_semantic = [dict(record) for record in current_claim_records]
        run_semantic_batch_task(
            target=target,
            task_name="claim_role",
            dry_run=False,
        )
        current_claim_records = apply_claim_role_decisions_to_claim_records(
            target=target,
            claim_records=current_claim_records,
            task_config=claim_role_config,
        )
        semantic_claim_updates_applied = current_claim_records != claim_records_before_semantic
        live_existing_claims = filter_live_claim_records(current_claim_records)
        historical_existing_claims = [
            record for record in current_claim_records
            if not is_live_claim_record(record)
        ]
        claims_by_id = {record["claim_id"]: record for record in live_existing_claims}
        historical_claims_by_id = {record["claim_id"]: record for record in historical_existing_claims}
        claims_by_normalized_text = {record["normalized_text"]: record for record in live_existing_claims}
        claims_by_similarity_bucket = {}
        for record in live_existing_claims:
            claims_by_similarity_bucket.setdefault(build_similarity_bucket(record["text"]), []).append(record)
        claim_similarity_index = rebuild_claim_similarity_index(live_existing_claims)

    if reingested_source_ids or purged_review_ids:
        review_state_records = build_ordered_review_state_records(
            live_reviews_by_id=existing_reviews,
            historical_reviews_by_id=historical_reviews_by_id,
        )
        write_jsonl(reviews_path, review_state_records)
        for review_record in review_state_records:
            write_review_file(target, review_record)

    # 如果这次 ingest 没有引入任何上游变化，就不必再完整重跑页面生成与索引更新。
    # 这样重复执行 ingest 时，系统会更像“增量编译”而不是“全量重编”。
    can_skip_page_regeneration = workspace_can_skip_page_regeneration(
        sources_by_id=sources_by_id,
        created_sources=created_sources,
        normalized_sources=normalized_sources,
        chunked_sources=chunked_sources,
        claims_created_by_source=claims_created_by_source,
        review_items=review_items,
        semantic_claim_updates_applied=semantic_claim_updates_applied,
    )

    if can_skip_page_regeneration:
        existing_pages = [ensure_page_lifecycle_defaults(record) for record in load_jsonl(pages_path)] if pages_path.exists() else []
        existing_pages = apply_page_alias_overrides_to_records(target, existing_pages)
        live_existing_pages = filter_live_page_records(existing_pages)
        page_records_by_id = {record["page_id"]: record for record in existing_pages}
        all_chunk_records = load_jsonl(chunks_path)
        claims_by_source_id: dict[str, list[dict]] = {}
        for claim_record in claims_by_id.values():
            for source_id in claim_record.get("source_ids", []):
                claims_by_source_id.setdefault(source_id, []).append(claim_record)
        chunks_by_source_id: dict[str, list[dict]] = {}
        for chunk_record in all_chunk_records:
            chunks_by_source_id.setdefault(chunk_record["source_id"], []).append(chunk_record)
        missing_source_page_source_ids = collect_missing_source_page_source_ids(
            active_source_ids=active_source_ids,
            sources_by_id=sources_by_id,
            page_records_by_id=page_records_by_id,
            claims_by_source_id=claims_by_source_id,
            chunks_by_source_id=chunks_by_source_id,
        )
        missing_concept_bucket_keys = collect_missing_concept_bucket_keys(
            claims_by_similarity_bucket=claims_by_similarity_bucket,
            page_records_by_id=page_records_by_id,
        )
        missing_workspace_overview = workspace_overview_page_missing(
            claims_by_similarity_bucket=claims_by_similarity_bucket,
            page_records_by_id=page_records_by_id,
        )
        if (
            missing_source_page_source_ids
            or missing_concept_bucket_keys
            or missing_workspace_overview
        ):
            can_skip_page_regeneration = False

    post_ingest_review_auto_payload = None

    if can_skip_page_regeneration:
        write_jsonl(pages_path, existing_pages)
        alias_index = write_alias_index(target, existing_pages)
        previous_search_index_records = load_search_pages_index(target)
        search_index = {
            "index_path": str(SEARCH_PAGES_INDEX_REL_PATH),
            "record_count": len(previous_search_index_records),
            "rebuilt_count": 0,
            "reused_count": len(previous_search_index_records),
            "index_version": SEARCH_PAGES_INDEX_VERSION,
            "updated_at": utc_now_iso(),
        }
        append_wiki_log(target, task_id, generated_pages)
        payload = {
            "task_id": task_id,
            "workspace": str(target),
            "raw_dir": str(raw_dir),
            "workspace_summary": build_workspace_summary(target, raw_dir),
            "created_sources": created_sources,
            "skipped_sources": skipped_sources,
            "normalized_sources": normalized_sources,
            "structured_sources": structured_sources,
            "chunked_sources": chunked_sources,
            "claimed_sources": claimed_sources,
            "generated_pages": generated_pages,
            "search_index": search_index,
            "alias_index": {
                "index_path": str(ALIAS_INDEX_REL_PATH),
                "canonical_count": len(alias_index.get("canonical_map", {})),
                "alias_key_count": len(alias_index.get("alias_map", {})),
                "conflict_count": len(alias_index.get("conflicts", [])),
                "index_version": alias_index.get("index_version"),
            },
            "review_items": review_items,
            "error_items": error_items,
            "summary": {
                "created_count": len(created_sources),
                "skipped_count": len(skipped_sources),
                "normalized_count": len(normalized_sources),
                "structured_count": len(structured_sources),
                "chunked_count": len(chunked_sources),
                "claimed_count": len(claimed_sources),
                "changed_page_count": 0,
                "review_count": len(review_items),
                "error_count": len([item for item in error_items if item["level"] == "error"]),
                "warning_count": len([item for item in error_items if item["level"] == "warning"]),
                "existing_page_count": len(live_existing_pages),
                "tracked_page_count": len(existing_pages),
            },
        }
        if post_ingest_config.get("review_auto"):
            post_ingest_review_auto_payload = run_post_ingest_review_auto(target)
            payload["post_ingest_review_auto"] = post_ingest_review_auto_payload
        return CommandResult(
            payload=payload,
            message=render_workspace_summary_message(
                "Ingest completed with no upstream changes; wiki regeneration was skipped.",
                target_dir=target,
                raw_dir=raw_dir,
                extra_lines=[
                    f"Task id: {task_id}",
                    (
                        "Ingest: "
                        f"normalized={payload['summary']['normalized_count']}, "
                        f"structured={payload['summary']['structured_count']}, "
                        f"chunks={payload['summary']['chunked_count']}, "
                        f"claims={payload['summary']['claimed_count']}, "
                        f"changed_pages={payload['summary']['changed_page_count']}, "
                        f"reviews_detected={payload['summary']['review_count']}, "
                        f"warnings={payload['summary']['warning_count']}, "
                        f"errors={payload['summary']['error_count']}"
                    ),
                    (
                        render_post_ingest_review_auto_summary(post_ingest_review_auto_payload)
                    )
                    if post_ingest_review_auto_payload is not None else None,
                ],
            ),
        )

    # 先生成来源摘要页，形成 source -> claim -> page 的第一层闭环。
    page_records = [ensure_page_lifecycle_defaults(record) for record in load_jsonl(pages_path)] if pages_path.exists() else []
    page_records = apply_page_alias_overrides_to_records(target, page_records)
    page_records_by_id = {record["page_id"]: record for record in page_records}
    all_claim_records = build_ordered_claim_state_records(
        live_claims_by_id=claims_by_id,
        historical_claims_by_id=historical_claims_by_id,
    )
    all_chunk_records = load_jsonl(chunks_path)
    dirty_claim_ids: set[str] = set()
    dirty_review_ids: set[str] = set()
    existing_overview_page = page_records_by_id.get(build_workspace_overview_page_id())
    if existing_overview_page is not None and existing_overview_page.get("type") == "overview":
        overview_page_id = existing_overview_page["page_id"]
        for claim_record in claims_by_id.values():
            if overview_page_id in claim_record.get("page_ids", []):
                claim_record["page_ids"] = [item for item in claim_record["page_ids"] if item != overview_page_id]
                claim_record["updated_at"] = utc_now_iso()
                dirty_claim_ids.add(claim_record["claim_id"])
        for review_record in existing_reviews.values():
            if overview_page_id in review_record.get("candidate_page_ids", []):
                review_record["candidate_page_ids"] = [
                    item for item in review_record["candidate_page_ids"] if item != overview_page_id
                ]
                dirty_review_ids.add(review_record["review_id"])
    changed_source_ids = {record["source_id"] for record in created_sources}
    changed_source_ids.update(record["source_id"] for record in normalized_sources)
    changed_source_ids.update(record["source_id"] for record in chunked_sources)
    changed_source_ids.update(claims_created_by_source.keys())
    normalized_records_by_source = {
        record["source_id"]: record for record in load_jsonl(normalized_path)
    }
    claims_by_source_id: dict[str, list[dict]] = {}
    for claim_record in filter_live_claim_records(all_claim_records):
        for source_id in claim_record.get("source_ids", []):
            claims_by_source_id.setdefault(source_id, []).append(claim_record)
    chunks_by_source_id: dict[str, list[dict]] = {}
    for chunk_record in all_chunk_records:
        chunks_by_source_id.setdefault(chunk_record["source_id"], []).append(chunk_record)

    changed_source_ids.update(
        collect_missing_source_page_source_ids(
            active_source_ids=active_source_ids,
            sources_by_id=sources_by_id,
            page_records_by_id=page_records_by_id,
            claims_by_source_id=claims_by_source_id,
            chunks_by_source_id=chunks_by_source_id,
        )
    )

    for source_record in load_jsonl(sources_path):
        source_id = source_record["source_id"]
        if source_id not in changed_source_ids:
            continue
        if source_id not in active_source_ids:
            continue
        source_claims = claims_by_source_id.get(source_id, [])
        source_chunks = chunks_by_source_id.get(source_id, [])
        if not source_claims and not source_chunks:
            continue

        # 来源摘要页正文里会展示当前来源状态。
        # 如果这里仍然沿用上一阶段的 claimed/chunked 状态，第一页生成出来的正文
        # 就会和后面真正写回 sources.jsonl 的 generated 状态不一致，导致下一次 ingest
        # 平白再改一遍来源页。这里先在内存里把“即将生成页面”的来源状态提升到 generated，
        # 再用于页面渲染和状态落盘，保证第一次与后续重复运行的页面签名一致。
        source_record_for_page = dict(sources_by_id.get(source_id, source_record))
        source_record_for_page["status"] = "generated"

        normalized_record = normalized_records_by_source.get(source_id)
        page_rel_path = source_summary_page_path(
            source_id,
            normalized_record["title"] if normalized_record else Path(source_record["source_path"]).stem,
        )
        page_text, page_record = build_source_summary_page(
            source_record=source_record_for_page,
            page_rel_path=page_rel_path,
            normalized_record=normalized_record,
            claim_records=source_claims,
            chunk_records=source_chunks,
        )
        page_record = apply_page_alias_overrides(target, page_record)
        page_record["page_path"] = str(page_rel_path)
        stored_page_record, page_changed = upsert_wiki_page(
            target=target,
            page_records_by_id=page_records_by_id,
            page_record=page_record,
            page_text=page_text,
        )
        if page_changed:
            generated_pages.append(stored_page_record)
            dirty_claim_ids.update(
                link_claims_to_page_in_memory(
                    source_claims,
                    stored_page_record["page_id"],
                    claims_by_id,
                )
            )

        replace_jsonl_record(
            ingest_state_path,
            "source_id",
            source_id,
            {
                "task_id": task_id,
                "source_id": source_id,
                "state": "generated",
                "last_successful_stage": "generated",
                "failed_stage": None,
                "retry_count": 0,
                "updated_at": utc_now_iso(),
            },
        )

        replace_jsonl_record(sources_path, "source_id", source_id, source_record_for_page)
        sources_by_id[source_id] = source_record_for_page

    # 再生成概念候选页，让多来源、多 claim 的主题可以在 Wiki 层聚合起来。
    # 这里直接沿用内存中的 claims 映射，保证前面给来源页补上的 page 反链
    # 在概念页生成阶段立刻可见，不必依赖中途先刷回磁盘。
    all_claim_records = list(claims_by_id.values())
    review_records = list(existing_reviews.values())
    concept_claim_groups: dict[str, list[dict]] = {}
    for claim_record in all_claim_records:
        active_claim_source_ids = [
            source_id for source_id in claim_record.get("source_ids", [])
            if source_id in active_source_ids
        ]
        if not active_claim_source_ids:
            continue
        # 概念页聚合与 review 检测共享更稳的 group key，减少“同主题分裂成多页”。
        bucket_key = build_concept_group_key(claim_record)
        concept_claim_groups.setdefault(bucket_key, []).append(claim_record)
    concept_claim_groups = regroup_concept_claims_by_canonical_topic(concept_claim_groups)
    if concept_claim_groups and page_intent_config.enabled:
        run_semantic_batch_task(
            target=target,
            task_name="page_intent",
            dry_run=False,
        )
    page_routes_by_bucket = apply_page_intent_decisions_to_claim_groups(
        target=target,
        concept_claim_groups=concept_claim_groups,
        task_config=page_intent_config,
    )

    changed_bucket_keys = set()
    for claim_record in all_claim_records:
        if any(source_id in changed_source_ids for source_id in claim_record.get("source_ids", [])):
            original_bucket_key = build_concept_group_key(claim_record)
            matching_bucket_key = next(
                (
                    bucket_key
                    for bucket_key, grouped_claims in concept_claim_groups.items()
                    if any(item["claim_id"] == claim_record["claim_id"] for item in grouped_claims)
                ),
                original_bucket_key,
            )
            changed_bucket_keys.add(matching_bucket_key)
    changed_bucket_keys.update(
        collect_missing_concept_bucket_keys(
            claims_by_similarity_bucket=concept_claim_groups,
            page_records_by_id=page_records_by_id,
        )
    )
    if semantic_claim_updates_applied:
        changed_bucket_keys.update(concept_claim_groups.keys())
    for bucket_key, grouped_claims in sorted(concept_claim_groups.items()):
        if bucket_key not in changed_bucket_keys:
            continue
        page_route = page_route_for_bucket(page_routes_by_bucket, bucket_key)
        page_intent = preferred_page_intent_for_claim_group(
            grouped_claims,
            page_route.get("page_intent", "topic"),
        )
        page_route["page_intent"] = page_intent
        page_route["route_target"] = page_intent
        if page_intent == "reject":
            continue
        if page_intent == "concept" and should_generate_concept_page(grouped_claims):
            group_topic_label = choose_group_topic_label(grouped_claims)
            canonical_claim = choose_canonical_claim(grouped_claims, group_topic_label)
            concept_page_id = build_concept_page_id(bucket_key)
            concept_title, concept_title_quality = resolve_concept_title_candidate(
                target=target,
                config=config,
                canonical_claim=canonical_claim,
                claim_records=grouped_claims,
                preferred_section_label=group_topic_label,
            )
            page_rel_path = concept_summary_page_path(
                concept_page_id,
                concept_title,
            )
            page_text, page_record = build_concept_page(
                target=target,
                bucket_key=bucket_key,
                page_rel_path=page_rel_path,
                claim_records=grouped_claims,
                page_records_by_id=page_records_by_id,
                review_records=review_records,
                render_config=readable_concept_render_config,
            )
            page_record = apply_page_route_to_page_record(page_record, page_route)
            page_record = apply_page_alias_overrides(target, page_record)
            page_record["page_path"] = str(page_rel_path)
            stored_page_record, page_changed = upsert_wiki_page(
                target=target,
                page_records_by_id=page_records_by_id,
                page_record=page_record,
                page_text=page_text,
            )
            if page_changed:
                generated_pages.append(stored_page_record)
                dirty_claim_ids.update(
                    link_claims_to_page_in_memory(
                        grouped_claims,
                        stored_page_record["page_id"],
                        claims_by_id,
                    )
                )
                dirty_review_ids.update(
                    link_reviews_to_page_in_memory(
                        review_records=review_records,
                        page_id=stored_page_record["page_id"],
                        claim_ids=stored_page_record["claim_ids"],
                        reviews_by_id=existing_reviews,
                    )
                )

        elif page_intent in {"guide", "example", "topic", "reference", "timeline"}:
            page_id = page_intent_page_id(bucket_key, page_intent)
            page_title_source = choose_group_topic_label(grouped_claims) or choose_canonical_claim(grouped_claims).get("text", "")
            page_rel_path = page_intent_page_path(page_intent, page_id, page_title_source)
            page_text, page_record = build_intent_routed_page(
                target=target,
                config=config,
                bucket_key=bucket_key,
                page_intent=page_intent,
                page_rel_path=page_rel_path,
                claim_records=grouped_claims,
                page_records_by_id=page_records_by_id,
                review_records=review_records,
            )
            page_record = apply_page_route_to_page_record(page_record, page_route)
            page_record = apply_page_alias_overrides(target, page_record)
            page_record["page_path"] = str(page_rel_path)
            stored_page_record, page_changed = upsert_wiki_page(
                target=target,
                page_records_by_id=page_records_by_id,
                page_record=page_record,
                page_text=page_text,
            )
            if page_changed:
                generated_pages.append(stored_page_record)
                dirty_claim_ids.update(
                    link_claims_to_page_in_memory(
                        grouped_claims,
                        stored_page_record["page_id"],
                        claims_by_id,
                    )
                )
                dirty_review_ids.update(
                    link_reviews_to_page_in_memory(
                        review_records=review_records,
                        page_id=stored_page_record["page_id"],
                        claim_ids=stored_page_record["claim_ids"],
                        reviews_by_id=existing_reviews,
                    )
                )

    overview_concept_pages = collect_workspace_overview_concept_pages(
        claims_by_similarity_bucket=concept_claim_groups,
        page_records_by_id=page_records_by_id,
    )
    if should_generate_workspace_overview_page(overview_concept_pages):
        overview_page_rel_path = workspace_overview_page_path()
        overview_page_text, overview_page_record = build_workspace_overview_page(
            target=target,
            page_rel_path=overview_page_rel_path,
            concept_pages=overview_concept_pages,
            page_records_by_id=page_records_by_id,
            claim_records_by_id=claims_by_id,
            render_config=overview_render_config,
        )
        overview_page_record = apply_page_alias_overrides(target, overview_page_record)
        overview_page_record["page_path"] = str(overview_page_rel_path)
        stored_overview_page, overview_page_changed = upsert_wiki_page(
            target=target,
            page_records_by_id=page_records_by_id,
            page_record=overview_page_record,
            page_text=overview_page_text,
        )
        dirty_claim_ids.update(
            link_claims_to_page_in_memory(
                [
                    claims_by_id[claim_id]
                    for claim_id in stored_overview_page["claim_ids"]
                    if claim_id in claims_by_id
                ],
                stored_overview_page["page_id"],
                claims_by_id,
            )
        )
        dirty_review_ids.update(
            link_reviews_to_page_in_memory(
                review_records=review_records,
                page_id=stored_overview_page["page_id"],
                claim_ids=stored_overview_page["claim_ids"],
                reviews_by_id=existing_reviews,
            )
        )
        if overview_page_changed:
            generated_pages.append(stored_overview_page)

    desired_auto_page_ids = {
        expected_source_summary_page_id(source_id)
        for source_id in active_source_ids
        if claims_by_source_id.get(source_id) or chunks_by_source_id.get(source_id)
    }
    forced_stale_page_ids: set[str] = set()
    for bucket_key, grouped_claims in concept_claim_groups.items():
        page_route = page_route_for_bucket(page_routes_by_bucket, bucket_key)
        page_intent = page_route.get("page_intent", "topic")
        forced_stale_page_ids.update(
            {
                build_concept_page_id(bucket_key),
                *{
                    page_intent_page_id(bucket_key, stale_intent)
                    for stale_intent in {"guide", "example", "topic", "reference", "timeline"}
                    if stale_intent != page_intent
                },
            }
        )
        if page_intent == "concept" and should_generate_concept_page(grouped_claims):
            desired_auto_page_ids.add(build_concept_page_id(bucket_key))
        elif page_intent in {"guide", "example", "topic", "reference", "timeline"}:
            desired_auto_page_ids.add(page_intent_page_id(bucket_key, page_intent))
    if should_generate_workspace_overview_page(overview_concept_pages):
        desired_auto_page_ids.add(build_workspace_overview_page_id())

    removed_pages, pruned_claim_ids, pruned_review_ids = prune_stale_auto_pages(
        target=target,
        page_records_by_id=page_records_by_id,
        desired_auto_page_ids=desired_auto_page_ids,
        claims_by_id=claims_by_id,
        reviews_by_id=existing_reviews,
        forced_stale_page_ids=forced_stale_page_ids - desired_auto_page_ids,
    )
    if removed_pages:
        generated_pages.extend(removed_pages)
    dirty_claim_ids.update(pruned_claim_ids)
    dirty_review_ids.update(pruned_review_ids)

    # 页面阶段结束后再统一把内存中的 claims / reviews / pages 落盘，
    # 避免生成大量概念页时频繁整文件重写，明显降低大资料集下的 ingest 耗时。
    if dirty_claim_ids:
        claim_state_records = build_ordered_claim_state_records(
            live_claims_by_id=claims_by_id,
            historical_claims_by_id=historical_claims_by_id,
        )
        write_jsonl(claims_path, claim_state_records)
        for claim_id in dirty_claim_ids:
            write_claim_file(target, claims_by_id[claim_id])

    if dirty_review_ids:
        review_state_records = build_ordered_review_state_records(
            live_reviews_by_id=existing_reviews,
            historical_reviews_by_id=historical_reviews_by_id,
        )
        write_jsonl(reviews_path, review_state_records)
        for review_id in dirty_review_ids:
            write_review_file(target, existing_reviews[review_id])

    # pages.jsonl 这里保留完整页面账本：
    # 在线页面继续参与 query/index，removed 页面则作为历史痕迹留存。
    write_jsonl(pages_path, list(page_records_by_id.values()))

    all_claim_records = list(claims_by_id.values())
    page_records_for_index = list(page_records_by_id.values())
    claim_records_by_id = {record["claim_id"]: record for record in all_claim_records}
    previous_search_index_records = load_search_pages_index(target)
    search_index = write_search_pages_index(
        target=target,
        page_records=page_records_for_index,
        claim_records_by_id=claim_records_by_id,
        previous_records=previous_search_index_records,
    )

    alias_index = write_alias_index(target, page_records_for_index)
    alias_conflict_reviews, _ = build_alias_conflict_reviews(alias_index, existing_reviews)
    for review_record in alias_conflict_reviews:
        review_file_path = write_review_file(target, review_record)
        review_record["review_file_path"] = review_file_path
        append_jsonl(reviews_path, review_record)
        existing_reviews[review_record["review_id"]] = review_record
        review_items.append(review_record)
    rebuild_wiki_index(target, list(page_records_by_id.values()))
    append_wiki_log(target, task_id, generated_pages)

    payload = {
        "task_id": task_id,
        "workspace": str(target),
        "raw_dir": str(raw_dir),
        "workspace_summary": build_workspace_summary(target, raw_dir),
        "created_sources": created_sources,
        "skipped_sources": skipped_sources,
        "normalized_sources": normalized_sources,
        "structured_sources": structured_sources,
        "chunked_sources": chunked_sources,
        "claimed_sources": claimed_sources,
        "generated_pages": generated_pages,
        "search_index": search_index,
        "alias_index": {
            "index_path": str(ALIAS_INDEX_REL_PATH),
            "canonical_count": len(alias_index.get("canonical_map", {})),
            "alias_key_count": len(alias_index.get("alias_map", {})),
            "conflict_count": len(alias_index.get("conflicts", [])),
            "index_version": alias_index.get("index_version"),
        },
        "review_items": review_items,
        "error_items": error_items,
        "summary": {
            "created_count": len(created_sources),
            "skipped_count": len(skipped_sources),
            "normalized_count": len(normalized_sources),
            "structured_count": len(structured_sources),
            "chunked_count": len(chunked_sources),
            "claimed_count": len(claimed_sources),
            "changed_page_count": len(generated_pages),
            "review_count": len(review_items),
            "error_count": len([item for item in error_items if item["level"] == "error"]),
            "warning_count": len([item for item in error_items if item["level"] == "warning"]),
        },
    }
    if post_ingest_config.get("review_auto"):
        post_ingest_review_auto_payload = run_post_ingest_review_auto(target)
        payload["post_ingest_review_auto"] = post_ingest_review_auto_payload
    return CommandResult(
        payload=payload,
        message=render_workspace_summary_message(
            "Ingest registration, normalization, chunking, claim drafting, and wiki generation completed.",
            target_dir=target,
            raw_dir=raw_dir,
            extra_lines=[
                f"Task id: {task_id}",
                (
                    "Ingest: "
                    f"normalized={payload['summary']['normalized_count']}, "
                    f"structured={payload['summary']['structured_count']}, "
                    f"chunks={payload['summary']['chunked_count']}, "
                    f"claims={payload['summary']['claimed_count']}, "
                    f"changed_pages={payload['summary']['changed_page_count']}, "
                    f"reviews_detected={payload['summary']['review_count']}, "
                    f"warnings={payload['summary']['warning_count']}, "
                    f"errors={payload['summary']['error_count']}"
                ),
                (
                    render_post_ingest_review_auto_summary(post_ingest_review_auto_payload)
                )
                if post_ingest_review_auto_payload is not None
                else None,
            ],
        ),
    )


def command_stub(args: argparse.Namespace) -> CommandResult:
    # 预留命令占位，方便先把 CLI 骨架搭起来，再逐步补真实能力。
    return CommandResult(message=f"MyAgentWiki command scaffold: {args.command}")


def command_lint(args: argparse.Namespace) -> CommandResult:
    # lint 在 V1 里承担“结构 + 基础质量巡检”双重职责：
    # - 目录与状态文件是否完整
    # - page / claim / review / alias index 是否彼此对齐
    root = find_project_root()
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else root

    checks = []

    def add_check(name: str, ok: bool, details: str, severity: str = "error") -> None:
        # 统一收集检查项，后面既能给人看，也方便做 JSON 输出。
        checks.append({
            "name": name,
            "ok": ok,
            "severity": severity,
            "details": details,
        })

    if target == root:
        add_check("project_root", True, "Linting repository root scaffold.", severity="info")
        required_paths = [
            "README.md",
            "docs/MyAgentWiki系统详细设计.md",
            "pyproject.toml",
            "config/runtime_manifest.yml",
            "src/myagentwiki/cli.py",
            "templates/project/config/project.yml.tmpl",
        ]
    else:
        schema_guard_payload = workspace_schema_guard_payload(target)
        add_check("workspace_target", True, f"Linting initialized workspace: {target}", severity="info")
        add_check(
            name="workspace_schema_supported",
            ok=schema_guard_payload.get("status") == "supported",
            details=(
                f"workspace.schema_version={schema_guard_payload.get('workspace_schema_version')} "
                f"(expected={schema_guard_payload.get('expected_schema_version')})"
            ),
        )
        required_paths = [
            "wiki/index.md",
            "wiki/log.md",
            "config/project.yml",
            "AGENTS.md",
            "CLAUDE.md",
            str(ALIAS_INDEX_REL_PATH),
        ]

    for rel_path in required_paths:
        # 这里不区分文件还是目录，只要求“该路径存在”。
        path = target / rel_path
        add_check(
            name=f"path_exists:{rel_path}",
            ok=path.exists(),
            details=f"Expected path: {path}",
        )

    if target != root:
        live_claim_records: list[dict] = []
        claim_records_by_id: dict[str, dict] = {}
        config_path = target / "config" / "project.yml"
        raw_dir = resolve_workspace_path(target, load_simple_yaml(config_path)["paths"]["raw"]) if config_path.exists() else target / "raw"
        add_check(
            name="raw_exists",
            ok=raw_dir.exists(),
            details="The raw directory should exist next to the workspace.",
        )
        add_check(
            name="git_initialized",
            ok=(target / ".git").exists(),
            details="Initialized workspace should have a git repository.",
        )
        add_check(
            name="state_sources_exists",
            ok=(target / "state" / "sources.jsonl").exists(),
            details="Workspace should contain state/sources.jsonl.",
        )
        add_check(
            name="state_ingest_state_exists",
            ok=(target / "state" / "ingest_state.jsonl").exists(),
            details="Workspace should contain state/ingest_state.jsonl.",
        )
        add_check(
            name="state_normalized_exists",
            ok=(target / "state" / "normalized.jsonl").exists(),
            details="Workspace should contain state/normalized.jsonl.",
        )
        add_check(
            name="state_structure_blocks_exists",
            ok=(target / STRUCTURE_BLOCKS_REL_PATH).exists(),
            details=f"Workspace should contain {target / STRUCTURE_BLOCKS_REL_PATH}.",
        )
        add_check(
            name="state_evidence_blocks_exists",
            ok=(target / EVIDENCE_BLOCKS_REL_PATH).exists(),
            details=f"Workspace should contain {target / EVIDENCE_BLOCKS_REL_PATH}.",
        )
        add_check(
            name="state_knowledge_units_exists",
            ok=(target / KNOWLEDGE_UNITS_REL_PATH).exists(),
            details=f"Workspace should contain {target / KNOWLEDGE_UNITS_REL_PATH}.",
        )
        add_check(
            name="state_chunks_exists",
            ok=(target / "state" / "chunks.jsonl").exists(),
            details="Workspace should contain state/chunks.jsonl.",
        )
        add_check(
            name="state_claims_exists",
            ok=(target / "state" / "claims.jsonl").exists(),
            details="Workspace should contain state/claims.jsonl.",
        )
        add_check(
            name="state_reviews_exists",
            ok=(target / "state" / "reviews.jsonl").exists(),
            details="Workspace should contain state/reviews.jsonl.",
        )
        add_check(
            name="state_error_log_exists",
            ok=(target / "state" / "error_log.jsonl").exists(),
            details="Workspace should contain state/error_log.jsonl.",
        )
        add_check(
            name="state_pages_exists",
            ok=(target / "state" / "pages.jsonl").exists(),
            details="Workspace should contain state/pages.jsonl.",
        )
        add_check(
            name="index_search_pages_exists",
            ok=(target / SEARCH_PAGES_INDEX_REL_PATH).exists(),
            details=f"Workspace should contain {target / SEARCH_PAGES_INDEX_REL_PATH}.",
        )
        add_check(
            name="index_aliases_exists",
            ok=alias_index_path(target).exists(),
            details=f"Workspace should contain {alias_index_path(target)}.",
        )

        chunk_records = load_jsonl(target / "state" / "chunks.jsonl") if (target / "state" / "chunks.jsonl").exists() else []
        structure_block_records = load_jsonl(target / STRUCTURE_BLOCKS_REL_PATH) if (target / STRUCTURE_BLOCKS_REL_PATH).exists() else []
        if structure_block_records:
            structure_block_ids = [record.get("structure_block_id") for record in structure_block_records]
            add_check(
                name="structure_block_ids_unique",
                ok=len(structure_block_ids) == len(set(structure_block_ids)),
                details=f"All structure_block_id values in {STRUCTURE_BLOCKS_REL_PATH} should be unique.",
            )

        evidence_block_records = load_jsonl(target / EVIDENCE_BLOCKS_REL_PATH) if (target / EVIDENCE_BLOCKS_REL_PATH).exists() else []
        if evidence_block_records:
            evidence_block_ids = [record.get("evidence_block_id") for record in evidence_block_records]
            add_check(
                name="evidence_block_ids_unique",
                ok=len(evidence_block_ids) == len(set(evidence_block_ids)),
                details=f"All evidence_block_id values in {EVIDENCE_BLOCKS_REL_PATH} should be unique.",
            )
            add_check(
                name="evidence_blocks_trace_structure",
                ok=all(record.get("structure_block_ids") for record in evidence_block_records),
                details="Each evidence block should point back to at least one structure block.",
            )

        knowledge_unit_records = load_jsonl(target / KNOWLEDGE_UNITS_REL_PATH) if (target / KNOWLEDGE_UNITS_REL_PATH).exists() else []
        if knowledge_unit_records:
            knowledge_unit_ids = [record.get("knowledge_unit_id") for record in knowledge_unit_records]
            add_check(
                name="knowledge_unit_ids_unique",
                ok=len(knowledge_unit_ids) == len(set(knowledge_unit_ids)),
                details=f"All knowledge_unit_id values in {KNOWLEDGE_UNITS_REL_PATH} should be unique.",
            )
            add_check(
                name="knowledge_units_trace_evidence",
                ok=all(record.get("evidence_block_ids") for record in knowledge_unit_records),
                details="Each knowledge unit should point back to at least one evidence block.",
            )

        if chunk_records:
            chunk_ids = [record.get("chunk_id") for record in chunk_records]
            add_check(
                name="chunk_ids_unique",
                ok=len(chunk_ids) == len(set(chunk_ids)),
                details="All chunk_id values in state/chunks.jsonl should be unique.",
            )

        claim_records = load_jsonl(target / "state" / "claims.jsonl") if (target / "state" / "claims.jsonl").exists() else []
        semantic_decision_records = load_semantic_decisions(target) if semantic_decisions_path(target).exists() else []
        semantic_decisions_by_id = {
            str(record.get("decision_id", "")).strip(): record
            for record in semantic_decision_records
            if str(record.get("decision_id", "")).strip()
        }
        if claim_records:
            claim_records = [ensure_claim_lifecycle_defaults(record) for record in claim_records]
            live_claim_records = filter_live_claim_records(claim_records)
            claim_records_by_id = {record["claim_id"]: record for record in live_claim_records}
            historical_claim_records = [
                record for record in claim_records
                if not is_live_claim_record(record)
            ]
            claim_ids = [record.get("claim_id") for record in claim_records]
            add_check(
                name="claim_ids_unique",
                ok=len(claim_ids) == len(set(claim_ids)),
                details="All claim_id values in state/claims.jsonl should be unique.",
            )
            add_check(
                name="live_claim_source_refs_present",
                ok=all(record.get("source_refs") for record in live_claim_records),
                details="Each live claim should keep at least one source_ref for traceability.",
            )
            add_check(
                name="live_claim_source_ids_present",
                ok=all(record.get("source_ids") for record in live_claim_records),
                details="Each live claim should keep at least one source_id.",
            )
            add_check(
                name="historical_claims_not_live",
                ok=all(record.get("lifecycle_status") in {"superseded", "archived"} for record in historical_claim_records),
                details="Historical claim records should not remain in active lifecycle state.",
            )
            claim_semantic_risk_issues_by_id = {
                record["claim_id"]: claim_semantic_risk_issues(record, semantic_decisions_by_id)
                for record in live_claim_records
            }
            claim_semantic_risk_issues_by_id = {
                claim_id: issues
                for claim_id, issues in claim_semantic_risk_issues_by_id.items()
                if issues
            }
            claim_semantic_risk_preview = ", ".join(
                issues[0]
                for issues in list(claim_semantic_risk_issues_by_id.values())[:8]
            ) or "No live claims carry unreviewed ambiguous semantic decision risk flags."
            add_check(
                name="claim_semantic_risk_flags_reviewed",
                ok=len(claim_semantic_risk_issues_by_id) == 0,
                details=claim_semantic_risk_preview,
                severity="warning",
            )

        review_records = load_jsonl(target / "state" / "reviews.jsonl") if (target / "state" / "reviews.jsonl").exists() else []
        if review_records:
            review_records = [ensure_review_lifecycle_defaults(record) for record in review_records]
            live_review_records = filter_live_review_records(review_records)
            historical_review_records = [
                record for record in review_records
                if not is_live_review_record(record)
            ]
            review_ids = [record.get("review_id") for record in review_records]
            add_check(
                name="review_ids_unique",
                ok=len(review_ids) == len(set(review_ids)),
                details="All review_id values in state/reviews.jsonl should be unique.",
            )
            add_check(
                name="live_review_candidate_claims_present",
                ok=all(
                    record.get("candidate_claim_ids") or record.get("kind") == "alias_conflict"
                    for record in live_review_records
                ),
                details="Claim-oriented live reviews should contain candidate_claim_ids; alias_conflict may use candidate_page_ids only.",
            )
            add_check(
                name="live_review_candidate_pages_present",
                ok=all("candidate_page_ids" in record for record in live_review_records),
                details="Each live review record should contain candidate_page_ids for reverse page lookup.",
            )
            add_check(
                name="historical_reviews_not_live",
                ok=all(record.get("lifecycle_status") in {"superseded", "archived"} for record in historical_review_records),
                details="Historical review records should not remain in active lifecycle state.",
            )

        page_records = load_jsonl(target / "state" / "pages.jsonl") if (target / "state" / "pages.jsonl").exists() else []
        if page_records:
            page_records = [ensure_page_lifecycle_defaults(record) for record in page_records]
            live_page_records = filter_live_page_records(page_records)
            removed_page_records = [
                record for record in page_records
                if record.get("lifecycle_status") == "removed"
            ]
            page_ids = [record.get("page_id") for record in page_records]
            add_check(
                name="page_ids_unique",
                ok=len(page_ids) == len(set(page_ids)),
                details="All page_id values in state/pages.jsonl should be unique.",
            )
            add_check(
                name="page_paths_present",
                ok=all(record.get("page_path") for record in page_records),
                details="Each page record should include page_path.",
            )
            add_check(
                name="page_titles_present",
                ok=all(record.get("title") for record in live_page_records),
                details="Each live page should include title.",
            )
            add_check(
                name="page_types_present",
                ok=all(record.get("type") for record in live_page_records),
                details="Each live page should include type.",
            )
            add_check(
                name="page_canonical_ids_present",
                ok=all(record.get("canonical_id") for record in live_page_records),
                details="Each live page should include canonical_id.",
            )
            add_check(
                name="live_pages_exist_on_disk",
                ok=all((target / record["page_path"]).exists() for record in live_page_records),
                details="Each live page record should have a corresponding wiki markdown file on disk.",
            )
            add_check(
                name="removed_pages_absent_on_disk",
                ok=all(not (target / record["page_path"]).exists() for record in removed_page_records),
                details="Removed page records should keep history in state/pages.jsonl, but their markdown files should already be deleted.",
            )
            add_check(
                name="removed_pages_not_in_live_set",
                ok=all(not is_live_page_record(record) for record in removed_page_records),
                details="Removed page records should not be treated as live pages for index/query rebuilds.",
            )

            canonical_groups: dict[str, list[str]] = {}
            live_page_records_by_id = {record["page_id"]: record for record in live_page_records}
            for record in live_page_records:
                canonical_id = record.get("canonical_id")
                if not canonical_id:
                    continue
                canonical_groups.setdefault(canonical_id, []).append(record.get("type", ""))
            add_check(
                name="canonical_page_family_valid",
                ok=all(
                    len(page_types) == len(set(page_types))
                    and len(page_types) == 1
                    for page_types in canonical_groups.values()
                ),
                details="Each canonical_id should map to at most one live page type.",
            )
            concept_pages = [record for record in live_page_records if record.get("type") == "concept"]
            concept_like_pages = [
                record for record in live_page_records
                if record.get("type") == "concept"
            ]
            add_check(
                name="readable_concept_render_metadata_present",
                ok=all(
                    record.get("render_target") and record.get("render_mode") and record.get("render_status")
                    for record in concept_pages
                ),
                details="Each readable concept page should record render_target, render_mode and render_status for traceability.",
            )
            grounded_concept_issues = {
                record["page_id"]: rendered_page_grounding_issues(
                    target=target,
                    page_record=record,
                    claim_records_by_id=claim_records_by_id,
                    page_records_by_id=live_page_records_by_id,
                )
                for record in concept_pages
            }
            grounded_concept_issues = {
                page_id: issues
                for page_id, issues in grounded_concept_issues.items()
                if issues
            }
            grounded_issue_preview = ", ".join(
                f"{page_id}:{'/'.join(issues[:2])}"
                for page_id, issues in list(grounded_concept_issues.items())[:5]
            ) or "All readable concept pages remain grounded in their linked claims."
            add_check(
                name="readable_concept_pages_grounded",
                ok=len(grounded_concept_issues) == 0,
                details=grounded_issue_preview,
            )
            overview_pages = [record for record in live_page_records if record.get("type") == "overview"]
            add_check(
                name="overview_render_metadata_present",
                ok=all(
                    record.get("render_target") and record.get("render_mode") and record.get("render_status")
                    for record in overview_pages
                ),
                details="Each overview page should record render_target, render_mode and render_status for traceability.",
            )
            grounded_overview_issues = {
                record["page_id"]: rendered_page_grounding_issues(
                    target=target,
                    page_record=record,
                    claim_records_by_id=claim_records_by_id,
                    page_records_by_id=live_page_records_by_id,
                )
                for record in overview_pages
            }
            grounded_overview_issues = {
                page_id: issues
                for page_id, issues in grounded_overview_issues.items()
                if issues
            }
            overview_issue_preview = ", ".join(
                f"{page_id}:{'/'.join(issues[:2])}"
                for page_id, issues in list(grounded_overview_issues.items())[:5]
            ) or "All overview pages remain grounded in their linked concept pages."
            add_check(
                name="overview_pages_grounded",
                ok=len(grounded_overview_issues) == 0,
                details=overview_issue_preview,
            )
            concept_quality_issues = {
                record["page_id"]: concept_page_quality_issues(record, claim_records_by_id)
                for record in concept_like_pages
            }
            concept_quality_issues = {
                page_id: issues
                for page_id, issues in concept_quality_issues.items()
                if issues
            }
            concept_quality_preview = ", ".join(
                f"{page_id}:{'/'.join(issues[:2])}"
                for page_id, issues in list(concept_quality_issues.items())[:8]
            ) or "All concept pages passed title-quality checks."
            add_check(
                name="concept_pages_title_quality",
                ok=len(concept_quality_issues) == 0,
                details=concept_quality_preview,
                severity="warning",
            )
            semantic_consistency_issues = {
                record["page_id"]: page_semantic_consistency_issues(record, claim_records_by_id)
                for record in live_page_records
            }
            semantic_consistency_issues = {
                page_id: issues
                for page_id, issues in semantic_consistency_issues.items()
                if issues
            }
            semantic_consistency_preview = ", ".join(
                f"{page_id}:{'/'.join(issues[:2])}"
                for page_id, issues in list(semantic_consistency_issues.items())[:8]
            ) or "All semantic page types remain consistent with their linked claim roles."
            add_check(
                name="page_semantic_consistency",
                ok=len(semantic_consistency_issues) == 0,
                details=semantic_consistency_preview,
                severity="warning",
            )
            page_brake_issues = {
                record["page_id"]: page_intent_brake_issues(record)
                for record in live_page_records
            }
            page_brake_issues = {
                page_id: issues
                for page_id, issues in page_brake_issues.items()
                if issues
            }
            page_brake_preview = ", ".join(
                f"{page_id}:{'/'.join(issues[:2])}"
                for page_id, issues in list(page_brake_issues.items())[:8]
            ) or "No live pages were routed through semantic page-intent downgrade brakes."
            add_check(
                name="semantic_page_intent_brakes_reviewed",
                ok=len(page_brake_issues) == 0,
                details=page_brake_preview,
                severity="warning",
            )

            alias_index = load_alias_index(target) if alias_index_path(target).exists() else {}
            alias_conflicts = unresolved_alias_conflicts(alias_index) if alias_index else []
            add_check(
                name="alias_conflicts_absent",
                ok=len(alias_conflicts) == 0,
                details="Alias registry should not contain unresolved alias conflicts.",
                severity="warning",
            )

            alias_canonical_ids = set(alias_index.get("canonical_map", {}).keys()) if alias_index else set()
            live_page_canonical_ids = {record.get("canonical_id") for record in live_page_records if record.get("canonical_id")}
            add_check(
                name="alias_index_covers_live_pages",
                ok=live_page_canonical_ids.issubset(alias_canonical_ids),
                details="Alias registry should cover every live page canonical_id.",
            )

            search_index_records = load_search_pages_index(target) if (target / SEARCH_PAGES_INDEX_REL_PATH).exists() else []
            indexed_page_ids = {record.get("page_id") for record in search_index_records}
            expected_live_page_ids = {record.get("page_id") for record in live_page_records}
            add_check(
                name="search_index_covers_live_pages",
                ok=expected_live_page_ids.issubset(indexed_page_ids),
                details="Search index should contain every live page.",
            )

    if target != root:
        report_lines = [
            "# Lint Report",
            "",
            f"- 目标目录: `{target}`",
            f"- 错误数量: `{len([check for check in checks if not check['ok'] and check['severity'] == 'error'])}`",
            f"- 警告数量: `{len([check for check in checks if not check['ok'] and check['severity'] == 'warning'])}`",
            "",
            "## 检查结果 / Checks",
            "",
        ]
        for check in checks:
            marker = "PASS" if check["ok"] else ("WARN" if check["severity"] == "warning" else "FAIL")
            report_lines.append(f"- [{marker}] `{check['name']}`: {check['details']}")
        atomic_write_text(
            target / "reports" / "lint" / "lint_latest.md",
            "\n".join(report_lines).strip() + "\n",
            encoding="utf-8",
        )

    # 先把 error / warning 分开统计，后面扩展更多级别也会比较自然。
    errors = [check for check in checks if not check["ok"] and check["severity"] == "error"]
    warnings = [check for check in checks if not check["ok"] and check["severity"] == "warning"]
    payload = {
        "target": str(target),
        "workspace_summary": build_workspace_summary(target),
        "checks": checks,
        "summary": {
            "errors": len(errors),
            "warnings": len(warnings),
            "ok": len(errors) == 0,
        },
    }
    return CommandResult(
        exit_code=0 if len(errors) == 0 else 1,
        payload=payload,
        message=render_workspace_summary_message(
            "Lint completed." if len(errors) == 0 else "Lint found issues.",
            target_dir=target,
            extra_lines=[
                (
                    "Summary: "
                    f"errors={payload['summary']['errors']}, "
                    f"warnings={payload['summary']['warnings']}, "
                    f"ok={'yes' if payload['summary']['ok'] else 'no'}"
                ),
            ],
        ),
    )


def command_render_page(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    render_target = args.render_target
    render_target_spec = PAGE_RENDER_TARGETS.get(render_target)
    if render_target_spec is None:
        raise KeyError(f"Unknown render_target: {render_target}")

    if render_target_spec.get("rebuild_strategy") == "review_affected_pages":
        live_claims_by_id, _, _ = load_claim_state_maps(target)
        live_reviews_by_id, _, _ = load_review_state_maps(target)
        rebuild_review_affected_pages(
            target=target,
            live_claims_by_id=live_claims_by_id,
            live_reviews_by_id=live_reviews_by_id,
        )

    page_records = [
        ensure_page_lifecycle_defaults(record)
        for record in load_jsonl(target / "state" / "pages.jsonl")
    ]
    live_rendered_pages = live_pages_for_render_target(page_records, render_target)

    selected_pages = live_rendered_pages
    if args.page_id:
        selected_pages = [record for record in selected_pages if record.get("page_id") == args.page_id]
    elif args.canonical_id:
        selected_pages = [record for record in selected_pages if record.get("canonical_id") == args.canonical_id]
    elif args.claim_id:
        selected_pages = [record for record in selected_pages if args.claim_id in record.get("claim_ids", [])]

    if not selected_pages:
        raise KeyError(f"No page matched render_target={render_target} with the requested selector.")

    payload = {
        "workspace": str(target),
        "render_target": render_target,
        "pages": [
            {
                "page_id": record["page_id"],
                "title": record.get("title"),
                "render_target": record.get("render_target") or page_record_render_target(record),
                "canonical_id": record.get("canonical_id"),
                "status": record.get("status"),
                "render_mode": record.get("render_mode"),
                "render_status": record.get("render_status"),
                "page_path": record.get("page_path"),
                "summary": record.get("summary"),
                "claim_ids": record.get("claim_ids", []),
            }
            for record in selected_pages
        ],
        "summary": {
            "page_count": len(selected_pages),
        },
    }
    if len(selected_pages) == 1:
        page_path = target / selected_pages[0]["page_path"]
        payload["page_text"] = page_path.read_text(encoding="utf-8")
    return CommandResult(payload=payload, message=f"Rendered page target: {render_target}")


def command_render_readable_concept(args: argparse.Namespace) -> CommandResult:
    delegated_args = argparse.Namespace(**vars(args))
    delegated_args.render_target = "readable_concept"
    return command_render_page(delegated_args)


def command_semantic_batch(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    payload = run_semantic_batch_task(
        target=target,
        task_name=args.task,
        dry_run=bool(args.dry_run),
    )
    return CommandResult(
        payload=payload,
        message=render_workspace_summary_message(
            f"Semantic batch completed: {args.task}",
            target_dir=target,
            extra_lines=[
                (
                    "Summary: "
                    f"items={payload['summary']['item_count']}, "
                    f"cache_hits={payload['summary']['cache_hits']}, "
                    f"pending_batches={payload['summary']['pending_batch_count']}, "
                    f"written_decisions={payload['summary']['written_decision_count']}, "
                    f"dry_run={'yes' if payload['summary']['dry_run'] else 'no'}"
                ),
            ],
        ),
    )


def command_claim_set_status(args: argparse.Namespace) -> CommandResult:
    target = Path(args.target_dir).expanduser().resolve() if args.target_dir else Path.cwd()
    ensure_workspace_schema_supported(target)
    live_claims_by_id, historical_claims_by_id, _ = load_claim_state_maps(target)
    live_reviews_by_id, _, _ = load_review_state_maps(target)

    claim_record = live_claims_by_id.get(args.claim_id) or historical_claims_by_id.get(args.claim_id)
    if claim_record is None:
        raise KeyError(f"Unknown claim_id: {args.claim_id}")
    if claim_record["claim_id"] not in live_claims_by_id:
        raise ValueError("claim_set_status currently only supports active claims.")

    active_review_ids = sorted(
        review_record["review_id"]
        for review_record in live_reviews_by_id.values()
        if is_actionable_review_record(review_record)
        and args.claim_id in review_record.get("candidate_claim_ids", [])
    )

    updated_claim = dict(claim_record)
    updated_claim["status"] = args.status
    if args.status != "needs_review":
        updated_claim["review_reason"] = None
    updated_claim["updated_at"] = utc_now_iso()
    live_claims_by_id[args.claim_id] = updated_claim

    rebuild_review_affected_pages(
        target=target,
        live_claims_by_id=live_claims_by_id,
        live_reviews_by_id=live_reviews_by_id,
    )

    payload = {
        "workspace": str(target),
        "claim_id": args.claim_id,
        "status": args.status,
        "active_review_ids": active_review_ids,
    }
    return CommandResult(payload=payload, message="Claim status updated.")


def build_parser() -> argparse.ArgumentParser:
    # 所有 CLI 子命令都在这里集中声明，方便一眼看清当前系统有哪些入口。
    parser = argparse.ArgumentParser(
        prog="myagentwiki",
        description="MyAgentWiki CLI scaffold.",
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    # init: 创建一个新的 MyAgentWiki 工作区，并复用/创建其外部 sibling raw 目录。
    init_parser = subparsers.add_parser("init", help="Initialize a new MyAgentWiki workspace.")
    init_parser.add_argument(
        "--source-dir",
        help="Optional path to the sibling raw directory. If omitted, create/use ../raw next to the workspace.",
    )
    init_parser.add_argument("--project-name", required=True, help="Name of the new wiki workspace.")
    init_parser.add_argument("--target-dir", help="Optional explicit target directory.")
    init_parser.add_argument("--json", action="store_true", help="Output JSON.")
    init_parser.set_defaults(handler=command_init)

    # ingest: 扫描 raw，登记来源，并对已支持的文本类型做标准化。
    ingest_parser = subparsers.add_parser("ingest", help="Register raw files into workspace metadata.")
    ingest_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    ingest_parser.add_argument(
        "--disable-insecure-download-retry",
        action="store_true",
        help="Disable the automatic one-time insecure retry for certificate verification failures when downloading remote Markdown images.",
    )
    ingest_parser.add_argument("--json", action="store_true", help="Output JSON.")
    ingest_parser.set_defaults(handler=command_ingest)

    # doctor: 检查当前机器是否满足项目运行要求。
    doctor_parser = subparsers.add_parser("doctor", help="Check runtime environment.")
    doctor_parser.add_argument("--json", action="store_true", help="Output JSON.")
    doctor_parser.set_defaults(handler=command_doctor)

    # bootstrap: 安装本项目声明的 Python 依赖。
    bootstrap_parser = subparsers.add_parser("bootstrap", help="Install or repair Python dependencies.")
    bootstrap_parser.add_argument("--dry-run", action="store_true", help="Show planned install command without running it.")
    bootstrap_parser.add_argument(
        "--extra",
        dest="extras",
        action="append",
        choices=("ocr", "office", "pdf", "dev"),
        default=[],
        help="Optional dependency group to install.",
    )
    bootstrap_parser.add_argument("--json", action="store_true", help="Output JSON.")
    bootstrap_parser.set_defaults(handler=command_bootstrap)

    # lint: 检查仓库骨架或某个已初始化工作区的完整性。
    lint_parser = subparsers.add_parser("lint")
    lint_parser.add_argument("--target-dir", help="Optional workspace directory to lint.")
    lint_parser.add_argument("--json", action="store_true", help="Output JSON.")
    lint_parser.set_defaults(handler=command_lint)

    # query: 在现有 wiki/claim/page 产物上执行多字段 BM25 检索。
    query_parser = subparsers.add_parser("query", help="Search generated wiki pages with weighted BM25 ranking.")
    query_parser.add_argument("text", help="Search query text.")
    query_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    query_parser.add_argument("--limit", type=int, default=5, help="Maximum number of results to return.")
    query_parser.add_argument(
        "--reading-depth",
        choices=tuple(QUERY_READING_DEPTH_LIMITS.keys()),
        default="standard",
        help="Preset reading-pack thickness. `deep` returns more matched claims and chunks per page.",
    )
    query_parser.add_argument(
        "--answer-ready",
        action="store_true",
        help="Render reading_pack as an answer-ready handoff summary for an upper-layer Agent.",
    )
    query_parser.add_argument(
        "--format",
        choices=("summary", "prompt", "messages", "chatml"),
        default="summary",
        help="When used with --answer-ready, choose summary view or direct prompt view.",
    )
    query_parser.add_argument("--claim-limit", type=int, help="Maximum matched claims per page. Overrides reading-depth default.")
    query_parser.add_argument("--chunk-limit", type=int, help="Maximum matched chunks per page. Overrides reading-depth default.")
    query_parser.add_argument("--json", action="store_true", help="Output JSON.")
    query_parser.set_defaults(handler=command_query)

    answer_query_parser = subparsers.add_parser(
        "answer-query",
        help="Return an answer-ready handoff summary derived from query reading_pack.",
    )
    answer_query_parser.add_argument("text", help="Search query text.")
    answer_query_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    answer_query_parser.add_argument("--limit", type=int, default=5, help="Maximum number of results to inspect.")
    answer_query_parser.add_argument(
        "--reading-depth",
        choices=tuple(QUERY_READING_DEPTH_LIMITS.keys()),
        default="standard",
        help="Preset reading-pack thickness. `deep` returns more matched claims and chunks per page.",
    )
    answer_query_parser.add_argument(
        "--format",
        choices=("summary", "prompt", "messages", "chatml"),
        default="summary",
        help="Choose answer-ready summary view or direct prompt view.",
    )
    answer_query_parser.add_argument("--claim-limit", type=int, help="Maximum matched claims per page. Overrides reading-depth default.")
    answer_query_parser.add_argument("--chunk-limit", type=int, help="Maximum matched chunks per page. Overrides reading-depth default.")
    answer_query_parser.add_argument("--json", action="store_true", help="Output JSON.")
    answer_query_parser.set_defaults(handler=command_answer_query)

    render_page_parser = subparsers.add_parser(
        "render-page",
        help="Rebuild and inspect rendered wiki page(s) by render target.",
    )
    render_page_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    render_page_parser.add_argument(
        "--render-target",
        choices=supported_page_render_targets(),
        required=True,
        help="Render target family to rebuild and inspect.",
    )
    render_page_selector_group = render_page_parser.add_mutually_exclusive_group()
    render_page_selector_group.add_argument("--page-id", help="Specific page_id to render.")
    render_page_selector_group.add_argument("--canonical-id", help="Specific canonical_id to render.")
    render_page_selector_group.add_argument("--claim-id", help="Render the page that references this claim.")
    render_page_parser.add_argument("--json", action="store_true", help="Output JSON.")
    render_page_parser.set_defaults(handler=command_render_page)

    render_readable_concept_parser = subparsers.add_parser(
        "render-readable-concept",
        help="Rebuild and inspect readable concept page(s).",
    )
    render_readable_concept_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    render_selector_group = render_readable_concept_parser.add_mutually_exclusive_group()
    render_selector_group.add_argument("--page-id", help="Specific readable concept page_id to render.")
    render_selector_group.add_argument("--canonical-id", help="Specific canonical_id to render.")
    render_selector_group.add_argument("--claim-id", help="Render the readable concept page that references this claim.")
    render_readable_concept_parser.add_argument("--json", action="store_true", help="Output JSON.")
    render_readable_concept_parser.set_defaults(handler=command_render_readable_concept)

    semantic_batch_parser = subparsers.add_parser(
        "semantic-batch",
        help="Run one semantic analysis batch task and persist structured semantic decisions.",
    )
    semantic_batch_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    semantic_batch_parser.add_argument(
        "--task",
        choices=SEMANTIC_TASK_NAMES,
        required=True,
        help="Semantic task family to run.",
    )
    semantic_batch_parser.add_argument("--dry-run", action="store_true", help="Plan batches without writing semantic decisions.")
    semantic_batch_parser.add_argument("--json", action="store_true", help="Output JSON.")
    semantic_batch_parser.set_defaults(handler=command_semantic_batch)

    # claim-set-status: 手工把某条 claim 提升或调整到目标状态，并触发页面重建。
    claim_status_parser = subparsers.add_parser("claim-set-status", help="Update one claim status and rebuild dependent pages.")
    claim_status_parser.add_argument("claim_id", help="Claim id to update.")
    claim_status_parser.add_argument("status", choices=("draft", "stable", "disputed", "needs_review"), help="New claim status.")
    claim_status_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    claim_status_parser.add_argument("--json", action="store_true", help="Output JSON.")
    claim_status_parser.set_defaults(handler=command_claim_set_status)

    # review-list: 查看当前 review 队列及其候选 claim。
    review_list_parser = subparsers.add_parser("review-list", help="List review items and candidate claims.")
    review_list_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    review_list_parser.add_argument("--status", choices=("open", "resolved"), help="Optional review status filter.")
    review_list_parser.add_argument("--json", action="store_true", help="Output JSON.")
    review_list_parser.set_defaults(handler=command_review_list)

    review_auto_parser = subparsers.add_parser(
        "review-auto",
        help="Conservatively auto-resolve high-confidence review items and escalate the rest.",
    )
    review_auto_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    review_auto_parser.add_argument("--dry-run", action="store_true", help="Plan auto actions without mutating workspace state.")
    review_auto_parser.add_argument(
        "--format",
        choices=("summary", "prompt", "messages", "chatml"),
        default="summary",
        help="Choose summary view or direct Agent handoff format.",
    )
    review_auto_parser.add_argument("--json", action="store_true", help="Output JSON.")
    review_auto_parser.set_defaults(handler=command_review_auto)

    # review-apply: 对单条 review 执行人工裁决动作。
    review_apply_parser = subparsers.add_parser("review-apply", help="Apply an action to a review item.")
    review_apply_parser.add_argument("review_id", help="Review id to apply action to.")
    review_apply_parser.add_argument(
        "action",
        choices=("keep_both", "archive_one", "merge", "edit_then_resume", "assign_alias", "remove_alias"),
        help="Decision action to apply.",
    )
    review_apply_parser.add_argument("--primary-claim-id", help="Primary claim id for archive_one / merge.")
    review_apply_parser.add_argument("--secondary-claim-id", help="Secondary claim id for merge.")
    review_apply_parser.add_argument("--primary-page-id", help="Primary page id for alias-conflict assign_alias.")
    review_apply_parser.add_argument("--alias-value", help="Alias value to assign during alias-conflict handling.")
    review_apply_parser.add_argument("--target-dir", help="Workspace directory. Defaults to current directory.")
    review_apply_parser.add_argument("--json", action="store_true", help="Output JSON.")
    review_apply_parser.set_defaults(handler=command_review_apply)

    return parser


def main() -> int:
    # main 保持很薄，只负责“解析参数 -> 调用命令 -> 输出结果”这条主线。
    parser = build_parser()
    args = parser.parse_args()
    result = args.handler(args)
    return print_result(result, as_json=getattr(args, "json", False))


if __name__ == "__main__":
    raise SystemExit(main())
