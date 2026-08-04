from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from openpyxl import Workbook
from myagentwiki.runtime_env import load_simple_yaml


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_FIXTURE_ROOT = REPO_ROOT / "tests" / "fixtures" / "user_project_lab"
DEFAULT_RUNTIME_ROOT = REPO_ROOT / "tests" / "runtime" / "user_project_lab"


def load_manifest(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def cli_environment(*, deterministic: bool = True) -> dict[str, str]:
    environment = {**os.environ, "PYTHONPATH": str(REPO_ROOT / "src")}
    if deterministic:
        environment["MYAGENTWIKI_LLM_MODE"] = "deterministic"
    return environment


def run_cli(*args: str, cwd: Path | None = None) -> dict:
    completed = subprocess.run(
        [sys.executable, "-m", "myagentwiki.cli", *args, "--json"],
        cwd=str(cwd or REPO_ROOT),
        env=cli_environment(),
        capture_output=True,
        text=True,
        check=True,
    )
    return json.loads(completed.stdout)


def validate_default_llm_config(workspace_dir: Path) -> dict:
    config_path = workspace_dir / "config" / "project.yml"
    config = load_simple_yaml(config_path)
    config_text = config_path.read_text(encoding="utf-8")
    llm = config.get("llm", {})
    semantic = config.get("semantic", {})
    automation = config.get("automation", {})
    rendering = config.get("rendering", {})
    checks = {
        "primary_online": llm.get("routing", {}).get("primary") == "online",
        "fallback_cli": llm.get("routing", {}).get("fallback") == "cli",
        "semantic_tasks_llm_assisted": all(
            semantic.get(task_name, {}).get("strategy") == "llm_assisted"
            for task_name in ("document_analysis", "claim_candidate_quality", "claim_role", "page_intent")
        ),
        "automation_tasks_llm_assisted": all(
            automation.get(task_name, {}).get("strategy") == "llm_assisted"
            for task_name in ("review_auto", "image_to_text", "stable_promotion", "concept_candidate_review")
        ),
        "render_tasks_llm_assisted": all(
            rendering.get(task_name, {}).get("mode") == "llm_assisted"
            for task_name in ("readable_concept", "overview")
        ),
        "unimplemented_render_tasks_disabled": all(
            rendering.get(task_name, {}).get("mode") == "disabled"
            for task_name in ("qa_note", "concept_update")
        ),
        "no_task_command": "command:" not in config_text,
        "no_legacy_llm_module_name": "hook" not in config_text.lower(),
    }
    if not all(checks.values()):
        raise AssertionError(f"Generated LLM configuration failed checks: {checks}")
    return checks


def validate_legacy_configuration_message(workspace_dir: Path, runtime_root: Path) -> dict:
    migration_workspace = runtime_root / "migration_workspace"
    if migration_workspace.exists():
        shutil.rmtree(migration_workspace)
    shutil.copytree(workspace_dir, migration_workspace)
    config_path = migration_workspace / "config" / "project.yml"
    config_path.write_text(
        config_path.read_text(encoding="utf-8")
        + "\nsemantic:\n"
        + "  claim_role:\n"
        + '    strategy: "agent_assisted"\n'
        + '    command: ["python3", "-m", "myagentwiki.agent_hook"]\n',
        encoding="utf-8",
    )
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "myagentwiki.cli",
            "semantic-batch",
            "--task",
            "claim_role",
            "--target-dir",
            str(migration_workspace),
            "--json",
        ],
        cwd=str(REPO_ROOT),
        env=cli_environment(),
        capture_output=True,
        text=True,
        check=False,
    )
    payload = json.loads(completed.stdout) if completed.stdout.strip() else {}
    if completed.returncode == 0 or payload.get("error") != "llm_configuration_migration_required":
        raise AssertionError(payload or completed.stderr)
    return {"exit_code": completed.returncode, "payload": payload}


def run_live_llm_check(workspace_dir: Path) -> dict:
    completed = subprocess.run(
        [
            sys.executable,
            str(REPO_ROOT / "scripts" / "debug_llm_routing.py"),
            "live",
            "--workspace",
            str(workspace_dir),
            "--task",
            "claim_stable_promotion",
        ],
        cwd=str(REPO_ROOT),
        env=cli_environment(deterministic=False),
        capture_output=True,
        text=True,
        check=False,
    )
    payload = json.loads(completed.stdout) if completed.stdout.strip() else {"stderr": completed.stderr.strip()}
    return {"exit_code": completed.returncode, "payload": payload}


def ensure_clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def make_png(path: Path, title: str, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (960, 540), color=(250, 248, 240))
    draw = ImageDraw.Draw(image)
    draw.rectangle((40, 40, 920, 500), outline=(70, 70, 70), width=3)
    draw.text((70, 70), title, fill=(20, 20, 20))
    y = 130
    for line in lines:
        draw.text((90, y), f"- {line}", fill=(30, 30, 30))
        y += 54
    image.save(path)


def make_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("状态字段说明", level=1)
    document.add_paragraph("本页说明 state 账本、claim status 和 review 状态之间的关系。")
    table = document.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "字段"
    header[1].text = "英文键"
    header[2].text = "说明"
    rows = [
        ("状态", "status", "表示当前页面或 claim 的可用状态"),
        ("生命周期", "lifecycle_status", "表示对象是否仍在线"),
        ("来源追踪", "source_refs", "表示可回链来源和 chunk"),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = value
    document.save(path)


def make_xlsx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    overview = workbook.active
    overview.title = "概览"
    overview.append(["主题", "Source Count", "说明"])
    overview.append(["Claim", 3, "知识声明层用于承载可追踪结论"])
    overview.append(["Review", 2, "review 闭环负责冲突收口"])
    detail = workbook.create_sheet("明细")
    detail.append(["来源", "状态", "备注"])
    detail.append(["claim-basics.md", "stable", "概念定义"])
    detail.append(["review-workflow.md", "stable", "流程说明"])
    workbook.save(path)


def make_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    # 最简文本 PDF，足够被现有 fallback 抽取页信息与文本片段。
    payload = (
        "%PDF-1.4\n"
        "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] /Contents 4 0 R >> endobj\n"
        "4 0 obj << /Length 74 >> stream\n"
        "BT /F1 12 Tf 50 250 Td (2026-06 v2.0.0 release timeline and workspace overview) Tj ET\n"
        "endstream endobj\n"
        "xref\n0 5\n0000000000 65535 f \n"
        "trailer << /Root 1 0 R /Size 5 >>\nstartxref\n0\n%%EOF\n"
    )
    path.write_bytes(payload.encode("latin-1"))


def seed_raw(raw_dir: Path) -> None:
    make_png(
        raw_dir / "images" / "claim-traceability.png",
        "Claim Traceability",
        [
            "收集 raw 资料",
            "进入 normalized",
            "生成 claims 与 pages",
            "保持 source refs 双向追踪",
        ],
    )
    make_png(
        raw_dir / "images" / "page-links-overview.png",
        "Page Links Overview",
        [
            "概念页链接到综述页",
            "综述页回指来源页",
            "query 使用 related pages 扩展阅读",
            "backlink 支持 incoming/outgoing 检查",
        ],
    )
    make_docx(raw_dir / "reference" / "state-fields.docx")
    make_xlsx(raw_dir / "tables" / "source-coverage.xlsx")
    make_pdf(raw_dir / "timeline" / "release-milestones.pdf")

    write_text(
        raw_dir / "notes" / "claim-basics.md",
        "# Claim 基础\n\n"
        "Claim 是位于 chunk 与 wiki 之间的独立知识声明层。\n\n"
        "Claim 需要保留 source refs、page ids 和 review 关联，方便双向追踪。\n\n"
        "Claim 页面通常会链接到工作区综述页和来源说明页。\n",
    )
    write_text(
        raw_dir / "process" / "review-workflow.md",
        "# Review Workflow\n\n"
        "系统需要把冲突结论送入 review 队列，并允许人工做 merge、keep_both 或 archive_one。\n\n"
        "如果查询想知道如何处理冲突，guide 页面应该优先命中。\n",
    )
    write_text(
        raw_dir / "notes" / "embedded-image-and-table.md",
        "# 混排样例\n\n"
        "## 页面关联\n\n"
        "下图说明 overview 页面、概念页和来源页会互相形成链接。\n\n"
        "![page-links](../images/page-links-overview.png)\n\n"
        "| 字段 | English Key | 说明 |\n"
        "| --- | --- | --- |\n"
        "| 状态 | status | 表示页面或 claim 当前状态 |\n"
        "| 关联页 | related_page_ids | 用于 query 扩展读取相关页面 |\n\n"
        "表格之后的解释段落需要继续进入结构抽取，而不是因为图片或表格而丢失正文。\n",
    )
    write_text(
        raw_dir / "conflicts" / "alias-alpha.md",
        "# 来源追踪\n\n"
        "来源追踪是把 Claim、Chunk、Source 串起来的机制。\n",
    )
    write_text(
        raw_dir / "conflicts" / "alias-beta.md",
        "# 来源追踪\n\n"
        "来源追踪也用于把页面关联回具体证据和 review 判断。\n",
    )
    write_text(
        raw_dir / "bilingual" / "query-expansion.md",
        "# Query Expansion\n\n"
        "Query expansion 需要同时参考中文主题词和少量 English aliases。\n\n"
        "页面之间的 backlinks 可以帮助 query 扩展读取 related pages。\n",
    )


def apply_mutation(raw_dir: Path, mutation: str) -> None:
    if mutation == "update_claim_basics":
        write_text(
            raw_dir / "notes" / "claim-basics.md",
            "# Claim 基础\n\n"
            "Claim 是位于 chunk 与 wiki 之间的独立知识声明层。\n\n"
            "更新后的定义强调：Claim 还需要记录 incoming pages 和 outgoing pages，帮助 query 做相关页面扩展。\n",
        )
        return
    if mutation == "add_new_source":
        write_text(
            raw_dir / "notes" / "workspace-overview.md",
            "# Workspace Overview\n\n"
            "工作区综述页应该汇总主要主题，并把读者引向更细的概念页。\n",
        )
        return
    if mutation == "delete_query_expansion":
        target = raw_dir / "bilingual" / "query-expansion.md"
        if target.exists():
            target.unlink()
        return
    if mutation == "replace_embedded_image":
        make_png(
            raw_dir / "images" / "page-links-overview.png",
            "Updated Page Links Overview",
            [
                "overview 页面成为入口",
                "concept 页面保留 definition",
                "source-summary 保持来源回链",
                "query 通过 linked pages 扩展读取",
            ],
        )
        return
    if mutation == "update_embedded_table":
        write_text(
            raw_dir / "notes" / "embedded-image-and-table.md",
            "# 混排样例\n\n"
            "## 页面关联\n\n"
            "![page-links](../images/page-links-overview.png)\n\n"
            "| 字段 | English Key | 说明 |\n"
            "| --- | --- | --- |\n"
            "| 状态 | status | 表示页面或 claim 当前状态 |\n"
            "| 关联页 | related_page_ids | 用于 query 扩展读取相关页面 |\n"
            "| 反向链接 | incoming_page_ids | 用于回看哪些页面引用了当前页 |\n\n"
            "更新后的表格强调双向链接字段也应进入结构化处理。\n",
        )
        return
    if mutation == "add_conflict_source":
        write_text(
            raw_dir / "conflicts" / "alias-gamma.md",
            "# 来源追踪\n\n"
            "来源追踪页还应该解释 alias 冲突为什么需要 review 收口。\n",
        )
        return
    raise ValueError(f"Unsupported mutation: {mutation}")


def run_queries(manifest: dict, workspace_dir: Path) -> dict:
    results = {}
    for query in manifest.get("queries", []):
        payload = run_cli(
            "query",
            query["text"],
            "--target-dir",
            str(workspace_dir),
            "--intent",
            str(query.get("intent", "lookup")),
            "--link-expansion",
            "auto",
        )
        results[query["name"]] = payload
    return results


def execute_scenario(
    fixture_root: Path,
    runtime_root: Path,
    scenario: str,
    clean: bool,
    *,
    live_llm_check: bool = False,
) -> dict:
    manifest = load_manifest(fixture_root / "fixture_manifest.json")
    if clean:
        ensure_clean_dir(runtime_root)
    else:
        runtime_root.mkdir(parents=True, exist_ok=True)

    raw_dir = runtime_root / "raw"
    assets_dir = runtime_root / "assets"
    workspace_dir = runtime_root / "workspace"
    reports_dir = runtime_root / "reports"
    for path in (raw_dir, assets_dir, reports_dir):
        path.mkdir(parents=True, exist_ok=True)

    seed_raw(raw_dir)
    mutations = manifest.get("scenarios", {}).get(scenario, {}).get("mutations", [])
    for mutation in mutations:
        apply_mutation(raw_dir, mutation)

    doctor = run_cli("doctor")
    bootstrap = run_cli("bootstrap", "--dry-run", "--extra", "dev")
    init_payload = run_cli(
        "init",
        "--source-dir",
        str(raw_dir),
        "--project-name",
        str(manifest.get("project_name", "UserProjectLab")),
        "--target-dir",
        str(workspace_dir),
    )
    llm_config_checks = validate_default_llm_config(workspace_dir)
    migration_check = validate_legacy_configuration_message(workspace_dir, runtime_root)
    ingest_payload = run_cli("ingest", "--target-dir", str(workspace_dir))
    query_payloads = run_queries(manifest, workspace_dir)
    review_payload = run_cli("review-list", "--target-dir", str(workspace_dir))
    lint_payload = run_cli("lint", "--target-dir", str(workspace_dir))

    report = {
        "scenario": scenario,
        "runtime_root": str(runtime_root),
        "doctor": doctor,
        "bootstrap": bootstrap,
        "init": init_payload,
        "llm_config_checks": llm_config_checks,
        "legacy_llm_config_check": migration_check,
        "live_llm_check": run_live_llm_check(workspace_dir) if live_llm_check else {"status": "skipped"},
        "ingest": ingest_payload,
        "queries": query_payloads,
        "review_list": review_payload,
        "lint": lint_payload,
    }
    reports_dir.mkdir(parents=True, exist_ok=True)
    (reports_dir / f"{scenario}.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Run the local user project lab fixture.")
    parser.add_argument("--fixture-root", default=str(DEFAULT_FIXTURE_ROOT))
    parser.add_argument("--runtime-root", default=str(DEFAULT_RUNTIME_ROOT))
    parser.add_argument("--scenario", default="baseline")
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--keep-runtime", action="store_true")
    parser.add_argument("--live-llm-check", action="store_true")
    args = parser.parse_args()

    fixture_root = Path(args.fixture_root).expanduser().resolve()
    runtime_root = Path(args.runtime_root).expanduser().resolve()
    report = execute_scenario(
        fixture_root=fixture_root,
        runtime_root=runtime_root,
        scenario=args.scenario,
        clean=bool(args.clean),
        live_llm_check=bool(args.live_llm_check),
    )
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if not args.keep_runtime:
        return 0
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
